"""
Генератор договора поставки для вакуумных траверс (VacuTec).

ТОЧНАЯ КОПИЯ по образцу:
/home/user/workspace/uploaded_attachments/b2e3741652c9492eaa14bcba105c9cf1/Dogovor_Vakuumnye-traversy.docx

Разделы:
 1. ПРЕДМЕТ ДОГОВОРА
 2. ЦЕНА ПРОДУКЦИИ И ПОРЯДОК ОПЛАТЫ
 3. СРОКИ И ПОРЯДОК ПОСТАВКИ
 4. ПРИЁМКА ПРОДУКЦИИ, КАЧЕСТВО ПРОДУКЦИИ, ГАРАНТИИ
 5. ФОРС-МАЖОРНЫЕ ОБСТОЯТЕЛЬСТВА
 6. ГАРАНТИИ СТОРОН
 7. ПЕРСОНАЛЬНЫЕ ДАННЫЕ
 8. СРОК ДЕЙСТВИЯ ДОГОВОРА
 9. ПРОЧИЕ УСЛОВИЯ
10. ГАРАНТИИ ПОСТАВЩИКА
11. УСЛОВИЯ И ПОРЯДОК ОБМЕНА ЭЛЕКТРОННЫМИ ДОКУМЕНТАМИ
12. ПОЧТОВЫЕ И БАНКОВСКИЕ РЕКВИЗИТЫ  (таблица 2 колонки: Поставщик / Покупатель)

Далее Спецификация № 1: 7 колонок (№, Код, Наименование, Ед, Кол-во, Цена, Сумма) + подписи.
"""
from __future__ import annotations
import io
from dataclasses import dataclass, field
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Поставщики с НДС 22% (если ключ не в списке — без НДС по УСН)
SUPPLIERS_WITH_VAT = {"LKS", "KINEMATIKA"}


# ---------- Утилиты форматирования ----------

def _set_font(run, name="Times New Roman", size=11, bold=False, italic=False):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _para(doc, text, align="just", bold=False, size=11, italic=False, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align == "just":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_font(r, size=size, bold=bold, italic=italic)
    return p


def _heading(doc, text, size=11):
    return _para(doc, text, align="left", bold=True, size=size, space_after=3)


def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:color"), "000000")
        tcBorders.append(elem)
    tcPr.append(tcBorders)


def _fmt_money(v) -> str:
    """Форматирует денежную сумму:
    - 380000    → '380 000'          (целое — без копеек)
    - 380000.5  → '380 000,50'       (есть копейки — всегда 2 знака)
    - 102786.89 → '102 786,89'
    """
    from decimal import Decimal, ROUND_HALF_UP
    if v is None: v = 0
    # Округляем до 2 знаков через Decimal чтобы не было float-глюков
    d = Decimal(str(v)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    # Если копеек нет — показываем целое
    if d == d.to_integral_value():
        return f"{int(d):,}".replace(",", " ")
    # Иначе — целая часть + ",XX"
    integer_part = int(d)
    cents = int((d - integer_part) * 100)
    return f"{integer_part:,}".replace(",", " ") + f",{cents:02d}"


def _calc_totals_vat(lines: list, has_vat: bool):
    """Правильный расчёт через Decimal (без float-ошибок).

    Цены в lines — УЖЕ с НДС (если has_vat=True) или без НДС (если False).
    Сумма позиции = qty * price. Итог = сумма всех позиций.
    Если has_vat: НДС = ИТОГ × 22 / 122 (выделяем в т.ч. НДС).

    Возвращает: (total_decimal, vat_decimal)
    """
    from decimal import Decimal, ROUND_HALF_UP
    total = Decimal("0")
    for ln in lines:
        qty = Decimal(str(ln.get("qty", 0) or 0))
        price = Decimal(str(ln.get("price", 0) or 0))
        line_sum = (qty * price).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        total += line_sum
    total = total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

    if has_vat:
        # НДС = ИТОГ × 22/122
        vat = (total * Decimal("22") / Decimal("122")).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    else:
        vat = Decimal("0")
    return total, vat


# ---------- Основной генератор ----------

def _is_vat_supplier(supplier: dict) -> bool:
    """Поставщик работает с НДС? Без НДС: МОДЕРНИЗАЦИЯ (УСН)."""
    # Смотрим на vat_note — если там "УСН"/"не облагается", то без НДС
    vat_note = str(supplier.get("vat_note", "")).lower()
    if "усн" in vat_note or "не облагается" in vat_note:
        return False
    # Основной тест — по ключам
    for k, v in {"LKS": ["7806586802"], "KINEMATIKA": ["6000001911"],
                 "MODERNIZATSIYA": ["7806601419"]}.items():
        if any(x in str(supplier.get("inn", "")) for x in v):
            return k in SUPPLIERS_WITH_VAT
    return True  # по умолчанию с НДС


def _add_drawings_appendix(doc, drawings: list, contract_number: str, contract_date_str: str):
    """Добавляет Приложение №1 с чертежами (PNG/JPG/PDF), по 2 на страницу."""
    if not drawings:
        return
    import tempfile, os

    # Преобразуем все в список путей к картинкам (PDF → картинки страниц)
    image_paths = []
    tmp_dir = tempfile.mkdtemp(prefix="dogovor_drawings_")

    for dr in drawings:
        # dr = bytes/path/BytesIO
        if isinstance(dr, (bytes, bytearray)):
            # определяем тип по magic bytes
            head = bytes(dr[:8])
            if head.startswith(b"%PDF"):
                # PDF → рендер в пакет PNG
                pdf_path = os.path.join(tmp_dir, f"in_{len(image_paths)}.pdf")
                with open(pdf_path, "wb") as f: f.write(dr)
                try:
                    from pdf2image import convert_from_path
                    pages = convert_from_path(pdf_path, dpi=180)
                    for pi, page in enumerate(pages):
                        p = os.path.join(tmp_dir, f"pdfpage_{len(image_paths)}_{pi}.png")
                        page.save(p, "PNG")
                        image_paths.append(p)
                except Exception:
                    # без pdf2image — пропускаем
                    pass
            else:
                p = os.path.join(tmp_dir, f"img_{len(image_paths)}.png")
                with open(p, "wb") as f: f.write(dr)
                image_paths.append(p)
        elif isinstance(dr, str) and os.path.exists(dr):
            image_paths.append(dr)

    if not image_paths:
        return

    # Старт Приложения
    doc.add_page_break()

    # Шапка Приложения
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Приложение №1 к ")
    _set_font(r, size=10, italic=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(f"Договору поставки № {contract_number} от {contract_date_str}")
    _set_font(r2, size=10, italic=True)

    _para(doc, "Габаритный чертёж оборудования", align="center",
           bold=True, size=12, space_after=12)

    # По 2 картинки на страницу
    for i, img_path in enumerate(image_paths):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        try:
            p.add_run().add_picture(img_path, width=Cm(15))
        except Exception:
            pass
        # если 2 на странице — разрыв после второй (i%2==1) если есть ещё картинки
        if i % 2 == 1 and i < len(image_paths) - 1:
            doc.add_page_break()

    # Подписи в Самом конце Приложения
    _para(doc, "", space_after=12)
    _para(doc, "Подписи сторон", align="center", bold=True, size=11, space_after=12)


def build_dogovor_traversa_docx(
    *,
    lines: list,               # список dict {code, name, unit, qty, price}
                                #   если поставщик с НДС — price уже с НДС 22%
                                #   если МОД (УСН) — price без НДС
    buyer: dict,               # {short, full, address, inn, kpp, ogrn, phone, email, bank, bik, rs, ks, director_position, director_fio_gen, director_fio_short, basis}
    supplier: dict,            # SUPPLIERS[key]
    contract_number: str,      # "27072026-ВД/ЛКС"
    contract_date_str: str,    # "27.07.2026"
    has_vat: bool = None,      # None = авто по поставщику; True/False = принудительно
    prepay_pct: int = 100,
    shipment_days: int = 5,
    warranty_months: int = 12,
    delivery_terms: str = "Самовывоз со склада Поставщика",
    delivery_address: str = "",
    include_stamp: bool = True,
    drawings: list = None,     # список bytes/paths — чертежи в Приложение №1
) -> bytes:
    """Возвращает bytes DOCX-договора траверсы по шаблону 2026."""

    doc = Document()

    # Поля страницы
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Авто-НДС по поставщику — В САМОМ НАЧАЛЕ!
    if has_vat is None:
        has_vat = _is_vat_supplier(supplier)

    # ==================== ШАПКА ====================
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17), WD_ALIGN_PARAGRAPH.RIGHT)
    r1 = p.add_run("Санкт-Петербург")
    _set_font(r1, size=11)
    r2 = p.add_run("\t" + contract_date_str)
    _set_font(r2, size=11)

    _para(doc, f"ДОГОВОР ПОСТАВКИ № {contract_number}",
          align="center", bold=True, size=14, space_after=12)

    # ==================== ПРЕАМБУЛА ====================
    supplier_full = supplier.get("full", supplier["short"])
    if "«" not in supplier_full and "«" in supplier["short"]:
        # добавим в кавычках как в образце
        name_in_quotes = supplier["short"].split("«")[-1].rstrip("»")
        supplier_full = f'{supplier_full} «{name_in_quotes}»'
    supplier_dir_gen = supplier.get("director_fio_gen", "")
    supplier_dir_pos = supplier.get("director_position", "генерального директора").lower()

    buyer_full = buyer.get("full") or buyer.get("short") or "Общество"
    buyer_dir_gen = buyer.get("director_fio_gen", "")
    buyer_dir_pos = (buyer.get("director_position") or "генерального директора").lower()
    buyer_basis = buyer.get("basis") or "Устава"

    preamble = (
        f'{supplier_full}, именуемое в дальнейшем «Поставщик», в лице '
        f'{supplier_dir_pos} {supplier_dir_gen}, действующего на основании '
        f'{supplier.get("director_basis", "Устава")}, с одной стороны, и '
        f'{buyer_full}, именуемое в дальнейшем «Покупатель», в лице '
        f'{buyer_dir_pos} {buyer_dir_gen}, действующего на основании {buyer_basis}, '
        f'с другой стороны, вместе именуемые «Стороны», а по отдельности «Сторона», '
        f'заключили настоящий Договор о нижеследующем:'
    )
    _para(doc, preamble, align="just", size=11, space_after=6)

    # ==================== 1. ПРЕДМЕТ ДОГОВОРА ====================
    _heading(doc, "1. ПРЕДМЕТ ДОГОВОРА")
    _para(doc, "1.1. Поставщик обязуется поставлять Покупателю в течение срока действия Договора заказанную Покупателем продукцию (далее — «Продукция»), а Покупатель обязуется оплачивать и принимать в полном объёме всю поставленную Продукцию в соответствии с условиями настоящего Договора и Спецификациями.")
    _para(doc, "1.2. Номенклатура и количество заказываемой Продукции указываются Покупателем в Спецификации к настоящему договору поставки.")
    _para(doc, "1.3. Номенклатура и количество Продукции по каждой отдельной поставке согласовываются Покупателем и Поставщиком в Спецификации, являющейся неотъемлемой частью Договора.")
    _para(doc, "1.4. В Спецификации указываются: номенклатура и количество Продукции, цена, условия оплаты, условия доставки.")
    _para(doc, "1.5. На основании согласованной Сторонами Спецификации Поставщик выставляет Покупателю счёт.")

    # ==================== 2. ЦЕНА ПРОДУКЦИИ И ПОРЯДОК ОПЛАТЫ ====================
    _heading(doc, "2. ЦЕНА ПРОДУКЦИИ И ПОРЯДОК ОПЛАТЫ")
    vat_clause = "включая НДС 22 %" if has_vat else "без НДС (УСН)"
    _para(doc, f"2.1. Цены на продукцию указываются, {vat_clause}.")
    _para(doc, f"2.2. Покупатель вносит предоплату на расчётный счёт Поставщика, указанный в счёте, в размере {prepay_pct} % от общей стоимости заказа, если иное не предусмотрено Спецификацией.")
    _para(doc, "2.3. При повышении цен Поставщик обязан уведомить Покупателя не менее чем за 3 (три) рабочих дня.")
    _para(doc, "2.4. Поставщик не вправе изменять цены на Продукцию, предоплаченную Покупателем.")
    _para(doc, "2.5. Срок оплаты: Покупатель обязан произвести полную оплату стоимости Продукции в течение 5 (пяти) банковских дней с момента получения уведомления о готовности Продукции к отгрузке.")
    _para(doc, "2.6. В случае частичной оплаты Продукции по Спецификации к настоящему Договору, поставка осуществляется только после полного погашения задолженности в размере 100 % суммы, указанной в Спецификации.")
    _para(doc, "2.7. Подтверждение оплаты: отгрузка Продукции производится исключительно после поступления 100 % оплаты на расчётный счёт Поставщика и подтверждения платежа (выпиской банка или иным документом).")
    _para(doc, "2.8. Ответственность за просрочку оплаты: при непоступлении оплаты в установленный срок Покупатель уплачивает Поставщику пеню в размере 1 % (один процент) от неоплаченной суммы за каждый день просрочки, начиная со дня, следующего за истечением срока оплаты, и до момента фактической оплаты.")
    _para(doc, "2.9. Право Поставщика на односторонний отказ от исполнения Договора: если оплата не поступит в течение 10 (десяти) рабочих дней с момента уведомления о готовности Продукции, Поставщик вправе отказаться от исполнения Договора в одностороннем порядке, уведомив Покупателя, и потребовать возмещения убытков, связанных с неисполнением обязательств Покупателем.")

    # ==================== 3. СРОКИ И ПОРЯДОК ПОСТАВКИ ====================
    _heading(doc, "3. СРОКИ И ПОРЯДОК ПОСТАВКИ")
    _para(doc, "3.1. Уведомление о готовности Продукции: Поставщик обязан уведомить Покупателя в письменной форме (по электронной почте или иным согласованным способом) о готовности Продукции к отгрузке не менее чем за 5 (пять) рабочих дней до планируемой даты отгрузки.")
    _para(doc, "3.2. При задержке отгрузки Продукции на срок более 10 (десяти) рабочих дней Поставщик выплачивает по требованию Покупателя пени в размере 1 % от суммы предоплаты за каждый рабочий день просрочки. Пени начисляются, но не более 10 % от суммы заказа.")
    _para(doc, "3.3. Покупатель обязуется произвести выборку заказа со склада Поставщика в течение 10 (десяти) рабочих дней со дня получения письменного уведомления на электронный адрес о готовности Продукции к отгрузке.")
    _para(doc, "3.4. Хранение заказа свыше 10 (десяти) рабочих дней осуществляется за дополнительную плату. Стоимость хранения составляет 1 % от суммы заказа за каждый день хранения. Сроки хранения должны быть согласованы Сторонами дополнительно.")
    _para(doc, "3.5. Доставка Продукции осуществляется транспортом Покупателя на условиях самовывоза или доставки до транспортных компаний на выбор Покупателя — ООО «Байкал-Сервис», ООО «Деловые линии», ООО «Первая экспедиционная компания», ООО «КИТ», ООО «ПЭК». Расходы по транспортировке несёт Покупатель.")
    _para(doc, "3.6. Поставщик обязуется сообщить Покупателю о наличии Продукции и готовности её к отгрузке. При реализации Продукции Поставщик предоставляет Покупателю (представителю по доверенности) сопроводительные документы: 1) универсальный передаточный документ по форме, рекомендованной Письмом ФНС России от 21.10.2013 № ММВ-20-3/96@; 2) паспорта или иные документы, подтверждающие качество Продукции производителем; 3) транспортные накладные (при необходимости).")
    _para(doc, "3.7. При отправке Продукции до склада Покупателя силами и средствами Поставщика (транспортной компанией или привлечением стороннего перевозчика) товаросопроводительные документы оформляются в 2 (двух) экземплярах.")
    _para(doc, "3.8. Поставщик считается выполнившим своё обязательство по передаче Продукции Покупателю: а) в случае доставки Продукции Поставщиком — с момента сдачи Продукции на склад Покупателя (Грузополучателя); б) в случае самовывоза Продукции Покупателем — с момента получения Продукции представителем Покупателя (Грузополучателя) на складе Поставщика; в) в случае доставки Продукции сторонним перевозчиком — с момента вручения Поставщиком Продукции перевозчику для его доставки.")
    _para(doc, "3.9. Право собственности на купленную Продукцию и риск случайной гибели или случайной порчи, утраты или повреждения переходит к Покупателю с момента, когда Поставщик в соответствии с условиями настоящего Договора считается выполнившим своё обязательство по передаче Продукции Покупателю.")
    _para(doc, "3.10. Место поставки по каждой поставке в рамках настоящего Договора указывается в Спецификации и/или счёте. Если Стороны не предусмотрели место поставки в Спецификации и/или счёте, таковым считается: 1) при доставке Продукции силами Поставщика — фактический адрес Покупателя; 2) при доставке Продукции силами транспортной компании — терминал транспортной компании в городе, в котором расположено место нахождения Покупателя.")
    _para(doc, "3.11. В случае организации доставки Продукции силами транспортной компании Покупатель, подписывая настоящий Договор, выражает своё согласие на получение Поставщиком от транспортной компании копий документов, подтверждающих факт доставки, и признаёт их надлежащим доказательством исполнения Поставщиком своих обязательств.")
    _para(doc, "3.12. Дополнительные условия отгрузки и транспортировки, включая особые требования к упаковке, маркировке и документации, определяются в Спецификации для каждой поставки отдельно.")

    # ==================== 4. ПРИЁМКА ПРОДУКЦИИ ====================
    _heading(doc, "4. ПРИЁМКА ПРОДУКЦИИ, КАЧЕСТВО ПРОДУКЦИИ, ГАРАНТИИ")
    _para(doc, "4.1. Поставщик обязуется поставить Продукцию строго в соответствии со Спецификацией.")
    _para(doc, "4.2. Право собственности на Продукцию переходит от Поставщика к Покупателю: 4.2.1. в момент передачи Продукции Покупателю на складе Покупателя, если доставка осуществляется за счёт Поставщика; 4.2.2. в момент передачи Продукции Покупателю или его представителю на складе Поставщика, если доставка осуществляется силами Покупателя; 4.2.3. в момент передачи Продукции транспортной компании, если доставка Продукции производится с привлечением сторонней транспортной компании.")
    _para(doc, "4.3. В случае самовывоза со склада Поставщика Продукция отгружается только при наличии правильно заполненной доверенности, оформленной на лицо, получающее Продукцию.")
    _para(doc, "4.4. В случае доставки Продукции за счёт Поставщика Покупатель при получении Продукции обязан подписать универсальный передаточный документ (УПД), заверить его круглой печатью Покупателя либо доверенным лицом Покупателя, действующим на основании доверенности.")
    _para(doc, "4.5. В случае доставки Продукции через транспортную компанию Покупатель обязан непосредственно после приёмки Продукции на своём складе подписать УПД и отправить его Поставщику: копию — в день приёмки по электронной почте; оригинал — курьерской доставкой или почтой в течение 3-х рабочих дней.")
    _para(doc, "4.6. В целях упрощения и ускорения документооборота Стороны могут осуществлять электронный обмен документами в соответствии с ФЗ от 06.04.2011 № 63-ФЗ «Об электронной подписи».")
    _para(doc, "4.7. Покупатель обязуется произвести приёмку Продукции на своём складе по количеству и по качеству в течение 5 рабочих дней с момента получения Продукции.")
    _para(doc, "4.8. При обнаружении несоответствия качества или расхождения по количеству во время приёмки Продукции Покупатель заполняет Бланк рекламации установленной формы и направляет Поставщику по электронной почте.")
    _para(doc, "4.9. Рекламацию по качеству Продукции Покупатель вправе предъявить в течение гарантийного срока в установленном порядке по электронной почте zakaz@rolls-kran.ru.")

    # ==================== 5. ФОРС-МАЖОР ====================
    _heading(doc, "5. ФОРС-МАЖОРНЫЕ ОБСТОЯТЕЛЬСТВА")
    _para(doc, "5.1. Стороны освобождаются от ответственности за частичное или полное неисполнение обязательств по Договору, если надлежащее исполнение оказалось невозможным вследствие обстоятельств непреодолимой силы, возникших после заключения Договора в результате событий чрезвычайного характера, которые Сторона не могла ни предвидеть, ни предотвратить разумными мерами.")
    _para(doc, "5.2. При наступлении обстоятельств непреодолимой силы, указанных в пункте 5.1 Договора, Сторона должна без промедления известить о них в письменной форме другую Сторону. Извещение должно содержать данные о характере обстоятельств, а также оценку их влияния на возможность исполнения Стороной обязательств по Договору и срок исполнения обязательств.")

    # ==================== 6. ГАРАНТИИ СТОРОН ====================
    _heading(doc, "6. ГАРАНТИИ СТОРОН")
    _para(doc, "6.1. Стороны гарантируют на момент заключения и в период действия настоящего договора как Сторонами договора, так и их работниками или посредниками:")
    _para(doc, "6.1.1. полный отказ от операций и действий, связанных с легализацией (отмыванием) доходов, полученных преступным путём, а также соблюдение требований Федерального закона № 115-ФЗ от 07.08.2001 «О противодействии легализации (отмыванию) доходов, полученных преступным путём, и финансированию терроризма»;")
    _para(doc, "6.1.2. неприменение и пресечение коррупционных действий в отношениях со второй Стороной договора и третьими лицами в соответствии с Федеральным законом № 273-ФЗ от 25.12.2008 «О противодействии коррупции»;")
    _para(doc, "6.1.3. сдачу полной и достоверной налоговой отчётности в соответствии с НК РФ, предоставление запрашиваемых МИФНС документов в случае встречной или иной проверки.")
    _para(doc, "6.2. Стороны гарантируют достоверность информации и заверений, представленных при заключении и исполнении договора, относящихся к правоспособности Сторон, полномочиям на его заключение, а также своему финансовому и имущественному положению.")
    _para(doc, "6.3. В случае нарушения одной из Сторон гарантий, предусмотренных разделом 6 настоящего договора, вторая Сторона вправе в одностороннем порядке отказаться от настоящего договора путём направления виновной Стороне соответствующего уведомления.")

    # ==================== 7. ПЕРСОНАЛЬНЫЕ ДАННЫЕ ====================
    _heading(doc, "7. ПЕРСОНАЛЬНЫЕ ДАННЫЕ")
    _para(doc, "7.1. Подписанием настоящего договора Стороны дают согласие на обработку персональных данных своих работников, переданных второй Стороне с целью надлежащего исполнения настоящего договора, включая персональные данные подписантов, представителей, работников бухгалтерии.")
    _para(doc, "7.2. Стороны обязуются соблюдать правила обработки персональных данных, предусмотренные Законом № 152-ФЗ, соблюдать конфиденциальность персональных данных и обеспечивать безопасность персональных данных при их обработке.")
    _para(doc, "7.3. Стороны вправе осуществлять следующие действия по обработке персональных данных: сбор, запись, систематизация, накопление, хранение, уточнение, извлечение, использование, передачу, блокирование, удаление, уничтожение персональных данных.")
    _para(doc, "7.4. Стороны вправе осуществлять действия по обработке персональных данных в течение всего срока действия настоящего договора и пяти лет с даты прекращения его действия.")
    _para(doc, "7.5. Стороны вправе в любое время отозвать согласие на обработку персональных данных путём направления второй Стороне письменного уведомления. Сторона, получившая такое уведомление, обязуется прекратить обработку персональных данных в срок не позднее 30 дней с даты получения уведомления.")

    # ==================== 8. СРОК ДЕЙСТВИЯ ДОГОВОРА ====================
    _heading(doc, "8. СРОК ДЕЙСТВИЯ ДОГОВОРА")
    _para(doc, "8.1. Договор вступает в силу с момента подписания его Сторонами и действует до 31.12.2026 г. При этом в части неисполненных обязательств Договор действует до полного их исполнения.")
    _para(doc, "8.2. Настоящий Договор подлежит автоматической пролонгации на каждый следующий год, если ни одна из Сторон договора за один календарный месяц не заявит о своём намерении его прекратить.")
    _para(doc, "8.3. Договор может быть расторгнут по взаимному согласию Сторон.")
    _para(doc, "8.4. Договор может быть расторгнут Стороной в одностороннем порядке только при условии отсутствия задолженности перед другой Стороной.")

    # ==================== 9. ПРОЧИЕ УСЛОВИЯ ====================
    _heading(doc, "9. ПРОЧИЕ УСЛОВИЯ")
    _para(doc, "9.1. Договор составлен в двух экземплярах, по одному экземпляру для каждой из Сторон. Оба экземпляра имеют одинаковую юридическую силу.")
    _para(doc, "9.2. Договор может быть дополнен или изменён по соглашению Сторон, путём подписания дополнительных соглашений к Договору.")
    _para(doc, "9.3. Все спорные вопросы Стороны стараются разрешить путём переговоров. Если спорные вопросы не могут быть разрешены путём переговоров, они передаются на рассмотрение арбитражного суда по местонахождению истца.")
    _para(doc, "9.4. Юридическую силу имеют оригиналы документов. Электронные и факсимильные копии подлежат замене на оригиналы.")
    _para(doc, f"9.5. Стороны могут вести почтовую, факсимильную и электронную переписку только с адресов и на адреса, указанные в настоящем Договоре. Электронная почта Поставщика: {supplier.get('email', 'zakaz@rolls-kran.ru')}. Почтовый адрес Поставщика: {supplier.get('address', '')}.")
    _para(doc, "9.6. В случае изменения реквизитов или адресов для переписки Стороны письменно информируют друг друга в срок не более 3 рабочих дней.")
    _para(doc, "9.7. Все ранее подписанные между Сторонами договоры и соглашения утрачивают свою силу с момента подписания настоящего Договора.")

    # ==================== 10. ГАРАНТИИ ПОСТАВЩИКА ====================
    _heading(doc, "10. ГАРАНТИИ ПОСТАВЩИКА")
    _para(doc, "10.1. Поставщик заверяет и гарантирует следующее:")
    _para(doc, "— Поставщик является надлежащим образом зарегистрированной организацией;")
    _para(doc, "— все сведения о Поставщике в ЕГРЮЛ достоверны на момент подписания договора; если в дальнейшем в ЕГРЮЛ появится запись о недостоверности данных о Поставщике, он обязуется в течение месяца внести в ЕГРЮЛ уточнённые сведения;")
    _para(doc, "— Поставщик располагает необходимыми ресурсами для исполнения настоящего Договора;")
    _para(doc, "— Поставщик отразит все операции по настоящему Договору, включая полученные от Покупателя авансы и реализацию Продукции Покупателю, в бухгалтерской и налоговой отчётности;")
    _para(doc, "— в случае получения Поставщиком требований налогового органа о представлении документов, относящихся к сделке с Покупателем, Поставщик обязуется исполнить требование в течение пяти рабочих дней со дня получения требования;")
    _para(doc, "— Поставщик обязуется выставлять Покупателю правильно оформленные счета-фактуры и первичные бухгалтерские документы в соответствии с требованиями действующего законодательства;")
    _para(doc, "— Поставщик самостоятельно выполняет обязательства по настоящему договору.")

    # ==================== 11. ЭДО ====================
    _heading(doc, "11. УСЛОВИЯ И ПОРЯДОК ОБМЕНА ЭЛЕКТРОННЫМИ ДОКУМЕНТАМИ")
    _para(doc, "11.1.1. Стороны договора признают юридическую силу электронных документов, подписанных усиленной электронной подписью (УЭП), равнозначной бумажным документам, подписанным собственноручной подписью.")
    _para(doc, f"11.1.2. {supplier['short']} (ИНН {supplier['inn']}) осуществляет обмен электронными документами через систему электронного документооборота (ЭДО) {supplier.get('edo_provider', 'СБИС')}, идентификатор участника: {supplier.get('edo_id', '')}.")
    _para(doc, "11.1.3. Электронные документы, направляемые между Сторонами, должны соответствовать требованиям, установленным законодательством РФ, а также настоящим договором.")
    _para(doc, "11.2. Порядок обмена: электронные документы направляются через систему ЭДО СБИС. Каждая Сторона обязана обеспечить конфиденциальность и целостность электронных документов при их передаче и хранении.")
    _para(doc, "11.3. Перечень электронных документов: договоры; универсальные передаточные документы (УПД); акты выполненных работ/оказанных услуг; приложения к договорам; акты сверок; чертежи; документы, связанные с исполнением обязательств.")
    _para(doc, "11.4. Требования: электронные документы должны быть подписаны усиленной электронной подписью уполномоченного представителя Стороны. Формат электронных документов должен соответствовать требованиям, установленным законодательством РФ.")
    _para(doc, "11.5. Сторона, направившая электронный документ, несёт ответственность за его достоверность и соответствие требованиям законодательства. В случае возникновения сомнений в подлинности электронного документа Сторона вправе запросить бумажный оригинал.")
    _para(doc, "11.6. Стороны обязуются хранить электронные документы в течение срока, установленного законодательством РФ или настоящим договором. По требованию одной из Сторон электронные документы должны быть предоставлены в бумажном виде.")

    # ==================== 12. РЕКВИЗИТЫ (табл 2 колонки) ====================
    _heading(doc, "12. ПОЧТОВЫЕ И БАНКОВСКИЕ РЕКВИЗИТЫ")

    tbl_req = doc.add_table(rows=1, cols=2)
    tbl_req.autofit = False
    tbl_req.columns[0].width = Cm(8.5)
    tbl_req.columns[1].width = Cm(8.5)

    def _fill_req_cell(cell, title, org):
        for para in cell.paragraphs:
            para._element.getparent().remove(para._element)
        p = cell.add_paragraph()
        r = p.add_run(title); _set_font(r, size=11, bold=True)
        p2 = cell.add_paragraph()
        _set_font(p2.add_run(""), size=11)
        p3 = cell.add_paragraph()
        _set_font(p3.add_run(org.get("full") or org.get("short", "")), size=11, bold=True)
        _set_font(cell.add_paragraph().add_run(f"Юр. адрес / факт. адрес: {org.get('address', '')}"), size=10)
        if org.get("inn"): _set_font(cell.add_paragraph().add_run(f"ИНН {org['inn']}"), size=10)
        if org.get("kpp"): _set_font(cell.add_paragraph().add_run(f"КПП {org['kpp']}"), size=10)
        if org.get("ogrn"): _set_font(cell.add_paragraph().add_run(f"ОГРН {org['ogrn']}"), size=10)
        if org.get("rs"):   _set_font(cell.add_paragraph().add_run(f"р/с {org['rs']}"), size=10)
        if org.get("bank"): _set_font(cell.add_paragraph().add_run(f"в банке {org['bank']}"), size=10)
        if org.get("ks"):   _set_font(cell.add_paragraph().add_run(f"к/с {org['ks']}"), size=10)
        if org.get("bik"):  _set_font(cell.add_paragraph().add_run(f"БИК {org['bik']}"), size=10)
        phone = org.get("phone") or org.get("phone_short") or org.get("phone_direct", "")
        if phone: _set_font(cell.add_paragraph().add_run(f"Тел.: {phone}"), size=10)
        if org.get("email"): _set_font(cell.add_paragraph().add_run(f"E-mail: {org['email']}"), size=10)
        _set_font(cell.add_paragraph().add_run(""), size=10)
        pos = org.get("director_position", "Генеральный директор")
        _set_font(cell.add_paragraph().add_run(f"{pos} {org.get('short', '')}"), size=11)
        _set_font(cell.add_paragraph().add_run(""), size=11)
        sig = cell.add_paragraph()
        _set_font(sig.add_run(f"_____________________ / {org.get('director_fio_short', '')} /"), size=11)
        mp = cell.add_paragraph()
        _set_font(mp.add_run("М.П."), size=11, bold=True)
        _set_cell_border(cell)

    _fill_req_cell(tbl_req.rows[0].cells[0], "ПОСТАВЩИК", supplier)
    _fill_req_cell(tbl_req.rows[0].cells[1], "ПОКУПАТЕЛЬ", buyer)

    # =================================================================
    # СПЕЦИФИКАЦИЯ (страница 2)
    # =================================================================
    doc.add_page_break()

    spec_no = 1
    _para(doc, f"Спецификация № {spec_no} от {contract_date_str}",
          align="center", bold=True, size=12, space_after=3)
    _para(doc, f"к договору поставки № {contract_number}",
          align="center", size=11, space_after=12)

    _para(doc, preamble, align="just", size=11, space_after=6)

    _heading(doc, "Наименование, количество и стоимость поставляемого Товара:")

    # Правильный расчёт через Decimal
    # Цены в lines УЖЕ с НДС (если has_vat=True) — берём как есть из КП
    total, vat_amount = _calc_totals_vat(lines, has_vat)

    # Таблица спецификации: 7 колонок
    tbl_spec = doc.add_table(rows=1 + len(lines), cols=7)
    tbl_spec.autofit = False
    widths = [Cm(0.8), Cm(2.5), Cm(6.5), Cm(1.2), Cm(1.5), Cm(2.5), Cm(2.5)]
    for i, w in enumerate(widths):
        tbl_spec.columns[i].width = w

    headers = ["№", "Код", "Наименование Товара", "Ед.", "Кол-во", "Цена, ₽", "Сумма, ₽"]
    for i, h in enumerate(headers):
        cell = tbl_spec.rows[0].cells[i]
        for para in cell.paragraphs:
            para._element.getparent().remove(para._element)
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); _set_font(r, size=10, bold=True)
        _set_cell_border(cell)

    for idx, ln in enumerate(lines, 1):
        qty = float(ln.get("qty", 0))
        price = float(ln.get("price", 0))
        summ = qty * price
        row = tbl_spec.rows[idx].cells
        values = [
            str(idx),
            str(ln.get("code", "") or ""),
            str(ln.get("name", "") or ""),
            str(ln.get("unit", "шт") or "шт"),
            (f"{qty:.0f}" if qty == int(qty) else f"{qty:.2f}"),
            _fmt_money(price),
            _fmt_money(summ),
        ]
        for i, v in enumerate(values):
            cell = row[i]
            for para in cell.paragraphs:
                para._element.getparent().remove(para._element)
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 3, 4) else \
                          WD_ALIGN_PARAGRAPH.RIGHT if i in (5, 6) else \
                          WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(v); _set_font(r, size=10)
            _set_cell_border(cell)

    _para(doc, "", space_after=6)

    if has_vat:
        total_txt = f"Общая цена Товара по настоящей Спецификации к Договору составляет: {_fmt_money(total)} ₽, в том числе НДС 22 % — {_fmt_money(vat_amount)} ₽ (НДС включён в стоимость)."
    else:
        total_txt = f"Общая цена Товара по настоящей Спецификации к Договору составляет: {_fmt_money(total)} ₽, без НДС (УСН)."
    _para(doc, total_txt, align="just", bold=True, space_after=6)

    _para(doc, "Комплект документов к поставке: паспорт ГОСТ, руководство по эксплуатации, сертификат.",
          align="just", bold=True, space_after=6)
    _para(doc, f"Порядок оплаты: Покупатель обязуется произвести оплату Товара в следующем порядке: — {prepay_pct} % от общей цены Спецификации — предоплата после подписания Договора и Спецификации.",
          align="just", bold=True, space_after=6)
    _para(doc, f"Срок отгрузки продукции: {shipment_days} рабочих дней",
          align="just", bold=True, space_after=6)

    delivery_line = delivery_terms
    if delivery_address:
        delivery_line = f"{delivery_terms}, адрес доставки: {delivery_address}"
    _para(doc, f"Условия доставки: {delivery_line}.",
          align="just", bold=True, space_after=6)
    _para(doc, f"Гарантия: {warranty_months} месяцев с даты поставки.",
          align="just", bold=True, space_after=12)

    # Подписи
    tbl_sig = doc.add_table(rows=1, cols=2)
    tbl_sig.autofit = False
    tbl_sig.columns[0].width = Cm(8.5)
    tbl_sig.columns[1].width = Cm(8.5)

    def _sig_cell(cell, title, org):
        for para in cell.paragraphs:
            para._element.getparent().remove(para._element)
        p = cell.add_paragraph()
        r = p.add_run(f"{title}:"); _set_font(r, size=11, bold=True)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(org.get("short", "")); _set_font(r2, size=11)
        for _ in range(4):
            cell.add_paragraph()
        p_sig = cell.add_paragraph()
        _set_font(p_sig.add_run(f"___________________________/ {org.get('director_fio_short', '')} /"), size=11)
        p_mp = cell.add_paragraph()
        _set_font(p_mp.add_run("М.П."), size=11, bold=True)
        p_pos = cell.add_paragraph()
        _set_font(p_pos.add_run(f"{org.get('director_position', 'Генеральный директор')} "), size=11)
        _set_cell_border(cell)

    _sig_cell(tbl_sig.rows[0].cells[0], "ПОСТАВЩИК", supplier)
    _sig_cell(tbl_sig.rows[0].cells[1], "ПОКУПАТЕЛЬ", buyer)

    # ==================== ПРИЛОЖЕНИЕ №1 (чертежи) ====================
    if drawings:
        _add_drawings_appendix(doc, drawings, contract_number, contract_date_str)

    # Возвращаем bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ==================== АДАПТЕР для обратной совместимости с build_dogovor_docx ====================

def build_dogovor_universal(
    q,                              # QuoteData — старый стиль
    buyer: dict,
    contract_number: str,
    contract_date_str: str,
    *,
    prepay_pct: int = 70,
    delivery_terms: str = "Самовывоз со склада Поставщика",
    include_stamp: bool = True,
    shipment_term=None,             # строка или число дней
    warranty_text=None,
    delivery_address: str = "",
    drawings: list = None,
    supplier: dict = None,          # если не задан, берём из app.SUPPLIER
    has_vat: bool = None,           # None = из q.include_vat или авто
    **_ignored_kwargs,              # съедаем любые старые аргументы
) -> bytes:
    """
    Адаптер: принимает QuoteData с .lines (SpecLine) и вызывает единую болванку.

    Позволяет заменить build_dogovor_docx() без переписывания всех вызовов.
    """
    # Поставщик: берём из аргумента или из app.SUPPLIER
    if supplier is None:
        import sys as _sys
        app_mod = _sys.modules.get("__main__") or _sys.modules.get("app")
        if app_mod is None:
            import app as app_mod
        supplier = getattr(app_mod, "SUPPLIER", None)
        if supplier is None:
            import suppliers
            supplier = suppliers.get_supplier("LKS")

    # Преобразуем q.lines (SpecLine) → list of dict
    lines_raw = getattr(q, "lines", []) or []
    lines = []
    for ln in lines_raw:
        if hasattr(ln, "code"):  # SpecLine
            lines.append({
                "code": getattr(ln, "code", "") or "",
                "name": getattr(ln, "name", "") or "",
                "unit": getattr(ln, "unit", "шт") or "шт",
                "qty":  float(getattr(ln, "qty", 0) or 0),
                "price": float(getattr(ln, "price", 0) or 0),
            })
        elif isinstance(ln, dict):
            lines.append({
                "code": str(ln.get("code", "") or ""),
                "name": str(ln.get("name", "") or ""),
                "unit": str(ln.get("unit", "шт") or "шт"),
                "qty":  float(ln.get("qty", 0) or 0),
                "price": float(ln.get("price", 0) or 0),
            })

    # has_vat: если не задан, берём из q.include_vat, иначе авто
    if has_vat is None:
        if hasattr(q, "include_vat"):
            has_vat = bool(q.include_vat)
        else:
            has_vat = _is_vat_supplier(supplier)

    # shipment_term → shipment_days (вытаскиваем число)
    shipment_days = 20
    if shipment_term:
        try:
            import re
            m = re.search(r"(\d+)", str(shipment_term))
            if m: shipment_days = int(m.group(1))
        except Exception: pass

    # warranty_text → warranty_months
    warranty_months = 12
    if warranty_text:
        try:
            import re
            m = re.search(r"(\d+)\s*мес", str(warranty_text))
            if m: warranty_months = int(m.group(1))
        except Exception: pass

    return build_dogovor_traversa_docx(
        lines=lines, buyer=buyer, supplier=supplier,
        contract_number=contract_number,
        contract_date_str=contract_date_str,
        has_vat=has_vat,
        prepay_pct=int(prepay_pct or 100),
        shipment_days=shipment_days,
        warranty_months=warranty_months,
        delivery_terms=str(delivery_terms or "Самовывоз со склада Поставщика"),
        delivery_address=delivery_address,
        include_stamp=include_stamp,
        drawings=drawings,
    )
