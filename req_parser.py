# -*- coding: utf-8 -*-
"""Парсер реквизитов покупателя из свободного текста или DOCX.

Извлекает: наименование (краткое и полное), ИНН, КПП, ОГРН, юр. и почтовый
адрес, телефон, e-mail, банк, БИК, р/с, к/с, ФИО и должность директора.
"""
from __future__ import annotations

import io
import re
from typing import Optional


# --------------------------------------------------------
# Извлечение текста из DOCX (без external tool - python-docx)
# --------------------------------------------------------
def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    # тексты из таблиц (карточки партнёра часто оформлены таблицей)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    parts.append(txt)
    return "\n".join(parts)


# --------------------------------------------------------
# Универсальные регэкспы
# --------------------------------------------------------
_INN_RE = re.compile(r"\bИНН\s*[:\-]?\s*(\d{10}|\d{12})\b")
_KPP_RE = re.compile(r"\bКПП\s*[:\-]?\s*(\d{9})\b")
_OGRN_RE = re.compile(r"\bОГРН(?:ИП)?\s*[:\-]?\s*(\d{13,15})\b")
_BIK_RE = re.compile(r"\bБИК\s*[:\-]?\s*(\d{9})\b")
_RS_RE = re.compile(
    r"(?:р/с|расч[её]тный\s+сч[её]т|расч\.\s*сч[её]т)"
    r"\s*[№:\-]?\s*([\d\s]{20,30})",
    re.IGNORECASE,
)
_KS_RE = re.compile(
    r"(?:к/с|кор+еспондентский\s+сч[её]т|корр\.\s*сч[её]т)"
    r"\s*[№:\-]?\s*([\d\s]{20,30})",
    re.IGNORECASE,
)
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
# Телефоны — только там, где явно есть разделители (скобки/дефисы/пробелы),
# чтобы не подхватывать номера счетов/ИНН/ОГРН.
_PHONE_RE = re.compile(
    r"(?:\+?\s*7|8)"
    r"[\s\-\.]*\(?\d{3,4}\)?[\s\-\.]+"
    r"\d{2,3}[\s\-\.]+\d{2}[\s\-\.]+\d{2}"
)

# Наименование: ищем «ООО ...», «АО ...», «ПАО ...», «ИП ...», «ЗАО ...»
_ORG_SHORT_RE = re.compile(
    r"(?<![А-ЯЁа-яё])((?:ООО|АО|ПАО|ЗАО|НАО|ИП|ОАО|ГК|АНО)\s+"
    r"(?:«[^»]+»|\"[^\"]+\"|[«\"][^«»\"]+[»\"]|[А-ЯЁA-Z][А-ЯЁA-Zа-яёa-z0-9\-\s]{1,60}))",
    re.UNICODE,
)
_ORG_FULL_RE = re.compile(
    r"(Общество\s+с\s+ограниченной\s+ответственностью"
    r"|Акционерное\s+общество"
    r"|Публичное\s+акционерное\s+общество"
    r"|Индивидуальный\s+предприниматель"
    r"|Закрытое\s+акционерное\s+общество)",
    re.IGNORECASE,
)


def _clean_digits(s: str) -> str:
    return re.sub(r"\s+", "", s)


def _find_addresses(text: str) -> tuple[str, str]:
    """Возвращает (юр. адрес, почтовый адрес)."""
    ur, pos = "", ""

    m = re.search(
        r"(?:юр(?:идический)?\.?\s*(?:/\s*факт(?:ический)?\.?)?\s*адрес"
        r"|адрес(?:\s+регистрации)?)\s*[:\-]?\s*"
        r"(\d{6},[^\n]+)",
        text, re.IGNORECASE,
    )
    if m:
        ur = m.group(1).strip().rstrip(".;")

    m = re.search(
        r"почтовый\s*адрес\s*[:\-]?\s*(\d{6},[^\n]+)",
        text, re.IGNORECASE,
    )
    if m:
        pos = m.group(1).strip().rstrip(".;")

    # Если юр. адрес не нашёлся — берём первый индекс+адрес в тексте
    if not ur:
        m = re.search(r"(\d{6},[^\n]{10,200})", text)
        if m:
            ur = m.group(1).strip().rstrip(".;")

    return ur, pos


def _find_bank(text: str) -> str:
    m = re.search(
        r"(?:в\s+банке|банк(?:овские\s+реквизиты)?|наименование\s+банка)"
        r"\s*[:\-]?\s*([^\n]+)",
        text, re.IGNORECASE,
    )
    if not m:
        return ""
    bank = m.group(1).strip()
    # обрезаем до следующего лейбла
    bank = re.split(
        r"\s*(?:р/с|к/с|БИК|ИНН|КПП|ОГРН|Расч\.\s*сч|Кор\.\s*сч|тел|email|e-mail)",
        bank, maxsplit=1, flags=re.IGNORECASE,
    )[0]
    return bank.strip(" ,;:.-\"'«»")


def _find_director(text: str) -> tuple[str, str, str]:
    """Возвращает (position, fio_gen, fio_short).

    Ищет 'в лице генерального директора Иванова Ивана Ивановича' /
    'директор Петров П.П.' и т.п.
    """
    position = ""
    fio_gen = ""
    fio_short = ""

    # Вариант "в лице <должность в Р.п.> ФИО Р.п."
    m = re.search(
        r"в\s+лице\s+([А-ЯЁа-яё\s]{5,50}?)"
        r"\s+([А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?"
        r"\s+[А-ЯЁ][а-яё]+(?:а|я|ы|и)"
        r"\s+[А-ЯЁ][а-яё]+(?:ича|евича|овича|инича|ича))",
        text,
    )
    if m:
        position = m.group(1).strip()
        fio_gen = m.group(2).strip()
    else:
        # Вариант простой "Директор: Иванов И.И." / "Ген. директор Иванов И.И."
        # Фамилия может быть двойной (Бонч-Бруевич).
        surname = r"[А-ЯЁ][а-яё]+(?:-[А-ЯЁ][а-яё]+)?"
        m = re.search(
            r"(генеральн(?:ый|ого)\s+директор(?:а)?"
            r"|директор(?:а)?"
            r"|ИП|индивидуальн(?:ый|ого)\s+предпринимател(?:ь|я)"
            r"|управляющ(?:ий|его))"
            r"\s*[:\-]?\s*"
            r"(" + surname + r"(?:\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.?"
            r"|\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+))",
            text, re.IGNORECASE,
        )
        if m:
            position = m.group(1).strip()
            fio_raw = m.group(2).strip()
            # если это ФИО полностью — построим родительный
            if re.search(r"[А-ЯЁ][а-яё]{3,}\s+[А-ЯЁ][а-яё]{2,}\s+[А-ЯЁ][а-яё]{2,}", fio_raw):
                fio_gen = _to_genitive(fio_raw)
            else:
                fio_gen = fio_raw
    # Нормализуем должность
    position = position.replace("генерального", "генеральный").replace(
        "директора", "директор").replace("управляющего", "управляющий").strip()
    position = re.sub(r"^ип$", "Индивидуальный предприниматель",
                      position, flags=re.IGNORECASE)
    if position:
        position = position[0].upper() + position[1:]

    # Короткое ФИО
    if fio_gen:
        parts = fio_gen.split()
        if len(parts) >= 3:
            surname_nom = _to_nominative_surname(parts[0])
            fio_short = f"{surname_nom} {parts[1][0].upper()}.{parts[2][0].upper()}."
            # Если вторая/третья часть в именительном — построим генитив
            if not (parts[2].endswith("ича") or parts[2].endswith("евича")
                    or parts[2].endswith("овича") or parts[2].endswith("иничны")):
                fio_gen = _to_genitive(fio_gen)
        elif len(parts) == 2:
            fio_short = fio_gen
    return position, fio_gen, fio_short


# --- очень простые правила для мужских ФИО ---
def _to_genitive(fio_nom: str) -> str:
    parts = fio_nom.split()
    if len(parts) != 3:
        return fio_nom
    s, n, p = parts

    def _surname(x: str) -> str:
        if x.endswith("ов") or x.endswith("ев") or x.endswith("ин") or x.endswith("ын"):
            return x + "а"
        if x.endswith("ский") or x.endswith("цкий"):
            return x[:-2] + "ого"
        if x.endswith("а"):
            return x[:-1] + "ы"
        return x

    def _name(x: str) -> str:
        if x.endswith("й"):
            return x[:-1] + "я"
        if x.endswith("а"):
            return x[:-1] + "ы"
        if x.endswith("я"):
            return x[:-1] + "и"
        if x.endswith("ь"):
            return x[:-1] + "я"
        return x + "а"

    def _patr(x: str) -> str:
        if x.endswith("ич"):
            return x + "а"
        if x.endswith("вна"):
            return x[:-1] + "ы"
        return x

    return f"{_surname(s)} {_name(n)} {_patr(p)}"


def _to_nominative_surname(surname_gen: str) -> str:
    # двойная фамилия
    if "-" in surname_gen:
        parts = surname_gen.split("-")
        return "-".join(_to_nominative_surname(p) for p in parts)
    if surname_gen.endswith("ова") or surname_gen.endswith("ева"):
        return surname_gen[:-1]  # Иванова → Иванов
    if surname_gen.endswith("ина") or surname_gen.endswith("ына"):
        return surname_gen[:-1]
    if surname_gen.endswith("ого"):
        return surname_gen[:-3] + "ий"
    return surname_gen


# --------------------------------------------------------
# Главная функция
# --------------------------------------------------------
def parse_requisites(text: str) -> dict:
    """Возвращает словарь полей покупателя, готовый к передаче в build_dogovor_docx."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    result: dict = {
        "short": "", "full": "", "address": "", "post_address": "",
        "inn": "", "kpp": "", "ogrn": "",
        "phone": "", "email": "",
        "bank": "", "bik": "", "rs": "", "ks": "",
        "director_position": "", "director_fio_gen": "", "director_fio_short": "",
        "basis": "Устава",
    }

    m = _INN_RE.search(text)
    if m:
        result["inn"] = m.group(1)
    m = _KPP_RE.search(text)
    if m:
        result["kpp"] = m.group(1)
    m = _OGRN_RE.search(text)
    if m:
        result["ogrn"] = m.group(1)
    m = _BIK_RE.search(text)
    if m:
        result["bik"] = m.group(1)
    m = _RS_RE.search(text)
    if m:
        result["rs"] = _clean_digits(m.group(1))[:20]
    m = _KS_RE.search(text)
    if m:
        result["ks"] = _clean_digits(m.group(1))[:20]
    m = _EMAIL_RE.search(text)
    if m:
        result["email"] = m.group(0)
    phones = _PHONE_RE.findall(text)
    if phones:
        cleaned = []
        for p in phones[:3]:
            # очищаем номера счетов (без разделителей)
            if re.search(r"[\s\-\.\(\)]", p):
                cleaned.append(re.sub(r"\s+", " ", p).strip())
        # дедупликация
        seen = set()
        unique = []
        for c in cleaned:
            digits = re.sub(r"\D", "", c)[-10:]
            if digits and digits not in seen:
                seen.add(digits)
                unique.append(c)
        result["phone"] = ", ".join(unique[:2])

    # Наименование
    m = _ORG_SHORT_RE.search(text)
    if m:
        result["short"] = re.sub(r"\s+", " ", m.group(1)).strip()

    # Полное наименование — из «Общество с ограниченной ответственностью «X»»
    fm = _ORG_FULL_RE.search(text)
    if fm:
        # ищем после этого куска название в кавычках
        m2 = re.search(
            fm.group(1) + r"\s*[«\"]?([^«»\"\n,]{2,80})[»\"]?",
            text, re.IGNORECASE,
        )
        if m2:
            name = m2.group(1).strip().rstrip(",;.")
            form = fm.group(1).strip()
            result["full"] = f"{form} «{name}»"

    # Авто-расшифровка правовой формы: если в тексте только ООО/АО/пр.,
    # разворачиваем в полное наименование.
    LEGAL_FORM_MAP = {
        "ООО": "Общество с ограниченной ответственностью",
        "АО": "Акционерное общество",
        "ПАО": "Публичное акционерное общество",
        "ЗАО": "Закрытое акционерное общество",
        "НАО": "Непубличное акционерное общество",
        "ОАО": "Открытое акционерное общество",
        "ИП": "Индивидуальный предприниматель",
        "ГК": "Группа компаний",
        "АНО": "Автономная некоммерческая организация",
    }
    if not result["full"] and result["short"]:
        # Разбираем "ООО «Ромашка»" в аббревиатуру + название
        m_form = re.match(r"^(ООО|АО|ПАО|ЗАО|НАО|ОАО|ИП|ГК|АНО)\s+(.+)$",
                           result["short"])
        if m_form and m_form.group(1) in LEGAL_FORM_MAP:
            form_full = LEGAL_FORM_MAP[m_form.group(1)]
            name_part = m_form.group(2).strip()
            if not name_part.startswith("«") and not name_part.startswith("\""):
                name_part = f"«{name_part}»"
            result["full"] = f"{form_full} {name_part}"
        else:
            result["full"] = result["short"]

    # Обратно: если есть full, но нет short — сворачиваем
    if result["full"] and not result["short"]:
        for abbr, long_form in LEGAL_FORM_MAP.items():
            if result["full"].lower().startswith(long_form.lower()):
                rest = result["full"][len(long_form):].strip()
                if not rest.startswith("«"):
                    rest = f"«{rest}»"
                result["short"] = f"{abbr} {rest}"
                break

    # адреса
    ur, pos_addr = _find_addresses(text)
    result["address"] = ur
    result["post_address"] = pos_addr

    # банк
    result["bank"] = _find_bank(text)

    # директор
    pos, fio_gen, fio_short = _find_director(text)
    result["director_position"] = pos
    result["director_fio_gen"] = fio_gen
    result["director_fio_short"] = fio_short

    return result
