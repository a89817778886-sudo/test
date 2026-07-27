# -*- coding: utf-8 -*-
"""Парсер стороннего КП (PDF/DOCX) в список строк спецификации + извлечение
чертежа (PDF/JPG/PNG) в PNG-байты для вставки в Приложение к договору.

Использует: (1) табличный парсинг для DOCX с явными таблицами (с fallback на
текстовый режим), (2) текстовый построчный regex-парсинг для PDF — устойчив
к переносу колонки "Код" на несколько строк (типично для КП rolls-kran/РОЛТЭК
и аналогичных шаблонов).
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Optional


@dataclass
class ExternalSpecLine:
    code: str
    name: str
    unit: str
    qty: float
    price: float

    @property
    def total(self) -> float:
        return round(self.qty * self.price, 2)


def _to_float(s: str) -> Optional[float]:
    if not s:
        return None
    s = s.strip().replace("\xa0", " ")
    s = re.sub(r"\s+", "", s)
    s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def _clean_cell(s: Optional[str]) -> str:
    return (s or "").strip().replace("\n", " ")


_STOP_WORDS = ("итого", "всего к оплате", "сумма к оплате", "сумма:",
               "всего:", "итог:")

# Строка позиции КП: "<код> <наименование ...> <ед.> <кол-во> <цена> <сумма>"
# Пример: "130.2 Направляющая RC35 (40х35х2,5) 2 м, оцинкованная шт 4 1 765,34 7 061,36"
_LINE_RE = re.compile(
    r"^(?P<code>[A-Za-zА-Яа-я0-9][A-Za-zА-Яа-я0-9._\-]{0,40})\s+"
    r"(?P<name>.+?)\s+"
    r"(?P<unit>шт|м|компл\.?|усл\.?|кг|м²|м2|комплект)\.?\s+"
    r"(?P<qty>\d+(?:[.,]\d+)?)\s+"
    r"(?P<price>\d[\d\s]*,\d{2})\s+"
    r"(?P<sum>\d[\d\s]*,\d{2})\s*$"
)

# Более гибкий вариант без строгого якоря по коду — для позиций, где код
# был разбит на несколько строк исходного PDF и не попал в ту же строку.
_LINE_RE_LOOSE = re.compile(
    r"(?P<name>[A-ZА-Я][^\d]{4,}?)\s+"
    r"(?P<unit>шт|м|компл\.?|усл\.?|кг|м²|м2|комплект)\.?\s+"
    r"(?P<qty>\d+(?:[.,]\d+)?)\s+"
    r"(?P<price>\d[\d\s]*,\d{2})\s+"
    r"(?P<sum>\d[\d\s]*,\d{2})\s*$"
)


def _merge_wrapped_lines(raw_lines: list[str]) -> list[str]:
    """В типовых КП каждая товарная позиция целиком помещается на одну
    строку извлечённого текста (даже если код визуально был перенесён
    внутри ячейки PDF). Возвращаем непустые строки — дальнейшая regex
    фильтрация отсеет служебные/неполные строки."""
    return [re.sub(r"\s+", " ", line.strip()) for line in raw_lines if line.strip()]


def parse_pdf_text(file_bytes: bytes) -> list[ExternalSpecLine]:
    import pdfplumber
    all_lines: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            all_lines.extend(text.split("\n"))

    merged = _merge_wrapped_lines(all_lines)
    result: list[ExternalSpecLine] = []
    for line in merged:
        low = line.lower()
        if any(sw in low for sw in _STOP_WORDS):
            continue

        m = _LINE_RE.match(line)
        if m:
            code = m.group("code")
            name = m.group("name").strip()
            unit = m.group("unit")
            qty = _to_float(m.group("qty"))
            price = _to_float(m.group("price"))
            if qty is None or price is None:
                continue
            if len(code) < 1 or len(name) < 3:
                continue
            result.append(ExternalSpecLine(code=code, name=name, unit=unit,
                                           qty=qty, price=price))
            continue

        m2 = _LINE_RE_LOOSE.search(line)
        if m2:
            name = m2.group("name").strip()
            unit = m2.group("unit")
            qty = _to_float(m2.group("qty"))
            price = _to_float(m2.group("price"))
            if qty is None or price is None:
                continue
            if len(name) < 3:
                continue
            result.append(ExternalSpecLine(code="", name=name, unit=unit,
                                           qty=qty, price=price))
    return result


def parse_docx_tables(file_bytes: bytes) -> list[ExternalSpecLine]:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    lines: list[ExternalSpecLine] = []

    for table in doc.tables:
        rows = [[_clean_cell(c.text) for c in r.cells] for r in table.rows]
        if not rows:
            continue
        header = [h.lower() for h in rows[0]]

        def find_col(*keywords):
            for i, h in enumerate(header):
                if any(k in h for k in keywords):
                    return i
            return None

        col_name = find_col("наименован", "товар", "продукц")
        col_unit = find_col("ед.", "единиц")
        col_qty = find_col("кол-во", "количество", "кол.")
        col_price = find_col("цена")
        col_code = find_col("код", "артикул")

        if col_name is None or col_qty is None or col_price is None:
            continue

        for row in rows[1:]:
            if len(row) <= max(col_name, col_qty, col_price):
                continue
            name = row[col_name]
            if not name or name.lower().startswith(("итого", "всего", "сумма к оплате")):
                continue
            qty = _to_float(row[col_qty]) if col_qty is not None else None
            price = _to_float(row[col_price]) if col_price is not None else None
            if qty is None or price is None:
                continue
            unit = row[col_unit] if col_unit is not None and col_unit < len(row) else "шт"
            code = row[col_code] if col_code is not None and col_code < len(row) else ""
            lines.append(ExternalSpecLine(code=code, name=name, unit=unit or "шт",
                                          qty=qty, price=price))

    if not lines:
        full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        merged = _merge_wrapped_lines(full_text)
        for line in merged:
            low = line.lower()
            if any(sw in low for sw in _STOP_WORDS):
                continue
            m = _LINE_RE.match(line)
            if not m:
                continue
            qty = _to_float(m.group("qty"))
            price = _to_float(m.group("price"))
            if qty is None or price is None:
                continue
            lines.append(ExternalSpecLine(code=m.group("code"),
                                          name=m.group("name").strip(),
                                          unit=m.group("unit"),
                                          qty=qty, price=price))
    return lines


def parse_external_kp(file_bytes: bytes, filename: str) -> list[ExternalSpecLine]:
    """Определяет тип файла по расширению и парсит таблицу товаров."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return parse_docx_tables(file_bytes)
    if lower.endswith(".pdf"):
        return parse_pdf_text(file_bytes)
    raise ValueError("Поддерживаются только файлы .pdf и .docx")


def extract_drawing_image(file_bytes: bytes, filename: str, dpi: int = 200) -> bytes:
    """Возвращает PNG-байты чертежа. PDF — рендерит первую страницу;
    JPG/PNG — перекодирует в PNG для унификации."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page = doc[0]
        pix = page.get_pixmap(dpi=dpi)
        return pix.tobytes("png")
    else:
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        out = io.BytesIO()
        img.save(out, format="PNG")
        return out.getvalue()


# ================================================================
# АВТОЗАПОЛНЕНИЕ РЕКВИЗИТОВ ИЗ ДОКУМЕНТА
# ================================================================

@dataclass
class ExtractedRequisites:
    """Реквизиты, извлечённые из внешнего КП (все поля опциональные)."""
    company_full: str = ""     # ООО «Ромашка» / полное наименование
    company_short: str = ""    # краткое
    inn: str = ""
    kpp: str = ""
    ogrn: str = ""
    address: str = ""          # юр. или почтовый адрес
    bank_name: str = ""
    bank_bik: str = ""
    bank_account: str = ""     # р/с
    corr_account: str = ""     # к/с
    phone: str = ""
    email: str = ""
    director_short: str = ""   # "Иванов И.И."
    director_gen: str = ""     # "Иванова Ивана Ивановича" — если сможем определить
    director_title: str = ""   # "Генеральный директор"
    site: str = ""

    def is_empty(self) -> bool:
        """True если ничего не извлечено."""
        return not any([self.inn, self.kpp, self.ogrn, self.company_full,
                        self.company_short, self.bank_account, self.phone,
                        self.email])


# --- регулярки ---
# Специально обрабатываем формат «ИНН/КПП 1234567890/123456789»
_RE_INN_KPP_COMBO = re.compile(r"ИНН[\s/\\.,:]*КПП[\s:]*(\d{10})\s*/\s*(\d{9})", re.IGNORECASE)
# ИНН отдельно (после «ИНН» может быть пробел, двоеточие, точка, №)
_RE_INN = re.compile(r"\bИНН[\s:.№#]*(\d{10}|\d{12})\b", re.IGNORECASE)
_RE_KPP = re.compile(r"\bКПП[\s:.№#]*(\d{9})\b", re.IGNORECASE)
_RE_OGRN = re.compile(r"\bОГРН(?:ИП)?[\s:]*[№#]?\s*(\d{13,15})\b", re.IGNORECASE)
_RE_BIK = re.compile(r"\bБИК[\s:]*[№#]?\s*(\d{9})\b", re.IGNORECASE)
# Р/с — 20-значный счёт
_RE_RS = re.compile(
    r"(?:расчет[\.а-я]*\s*счет|р/?\s*с|р\.\s*/?\s*с|расч[её]тный счет)"
    r"[\s№:.]*(\d{20})", re.IGNORECASE)
# К/с
_RE_KS = re.compile(
    r"(?:корр[\.а-я]*\s*счет|к/?\s*с|к\.\s*/?\s*с|корресп)"
    r"[\s№:.]*(\d{20})", re.IGNORECASE)
# Просто 20-значное число рядом со словом "р/с" или "счёт"
_RE_ACC_ANY = re.compile(r"\b(\d{20})\b")

# Телефон — требуем чтобы начинался как +7 или 8, а не был частью длинного числа
_RE_PHONE = re.compile(
    r"(?<![\d])("
    r"\+?7[\s\-\(]+\d{3}[\s\-\)]+\d{3}[\s\-]*\d{2}[\s\-]*\d{2}|"
    r"8[\s\-\(]+\d{3}[\s\-\)]+\d{3}[\s\-]*\d{2}[\s\-]*\d{2}"
    r")(?![\d])")
_RE_EMAIL = re.compile(
    r"\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b")
_RE_SITE = re.compile(
    r"\b((?:https?://)?(?:www\.)?[a-zA-Z0-9-]+\.(?:ru|com|net|рф)(?:/\S*)?)\b",
    re.IGNORECASE)

# Наименование ЮЛ: ищем ООО / ОАО / ЗАО / АО / ИП «...»
# Аббревиатура — отдельное слово (не цепляет ОО в «ОБЩЕСТВО»). Только Uppercase.
_RE_COMPANY = re.compile(
    r"(?<![А-Яё])(ООО|ОАО|ЗАО|АО|ПАО|ИП|ТОО|НКО|ФКП|ФГУП|МУП|ГУП|НПФ|АНО)(?![А-Яё])\s*[«\"']([^»\"']{3,120})[»\"']")
# Полное наименование
_RE_COMPANY_FULL = re.compile(
    r"(Общество с ограниченной ответственностью|Открытое акционерное общество|"
    r"Закрытое акционерное общество|Публичное акционерное общество|"
    r"Акционерное общество|Индивидуальный предприниматель|"
    r"Федеральное казенное предприятие|Федеральное государственное унитарное предприятие|"
    r"Муниципальное унитарное предприятие|Автономная некоммерческая организация)"
    r"\s*[«\"']([^»\"']{3,150})[»\"']", re.IGNORECASE)

# Адрес: «Адрес: 123456, г. Москва, ул. …» или «Юр. адрес»
# Поддерживаем многострочный адрес (часто в карточках адрес на 2-3 строки)
_RE_ADDRESS = re.compile(
    r"(?:юр(?:идического|идический)?\.?\s*адрес(?:\s+юр\.?\s*лица|\s+регистрации)?|"
    r"почтовый\s*адрес|фактический\s*адрес|местонахождение|адрес)"
    r"[\s:]*(\d{6}[,\s]+.{5,500}?)(?=\s*(?:ИНН|КПП|ОГРН|БИК|Р/?с|К/?с|Расч[её]т|Корр|Тел|E-?mail|email|Банк|Директор|Почтовый|Фактич|$))",
    re.IGNORECASE | re.DOTALL)
# Fallback: любая строка с индексом 6 цифр
_RE_ADDRESS_FALLBACK = re.compile(
    r"\b(\d{6}[,\s]+(?:г\.|город|обл\.?|область|РФ|Россия).{5,400}?)(?=\s*(?:ИНН|КПП|ОГРН|БИК|Тел|E-?mail|email|Директор|Банк|$))",
    re.IGNORECASE | re.DOTALL)

# Банк: строка типа «в ПАО СБЕРБАНК» или «Банк: Филиал ...».
# Поддерживает банк на несколько строк (ПАО/АО/Филиал с переносами).
# Банк: читаем всё от ключевого слова (ПАО/АО/Филиал) до стоп-слова
_RE_BANK = re.compile(
    r"(?:в\s+|банк[\s:]+|наименование\s+банка[\s:]+)"
    r"((?:ПАО|АО|ООО|ФАКБ|ФКБ|Филиал|Банк)[^\n]*?(?:\n[^\n]*(?:АО|ПАО|БАНК|Bank)[^\n]*)?)\s*(?=БИК|Р/?с|Расч[её]тный|К/?с|Корр|ИНН|КПП|ОГРН|Адрес|Юр\.|Тел|E-mail|email|Директор)",
    re.IGNORECASE)

# Директор. Исключаем «Заместитель», «Помощник», «Исполняющий».
_RE_DIRECTOR = re.compile(
    r"(?<![а-яёА-ЯЁ])(?:генеральный\s+директор|директор|руководитель|президент|глава|ИП)"
    r"(?:\s*\([^)]*\))?"  # возможно (на основании Устава)
    r"[\s:]*([А-ЯЁ][а-яё\-]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.)",
    re.IGNORECASE)
_RE_DIRECTOR_FULL = re.compile(
    r"(?<![а-яёА-ЯЁ])(?:генеральный\s+директор|директор|руководитель|президент|глава|ИП)"
    r"(?:\s*\([^)]*\))?"
    r"[\s:]*([А-ЯЁ][а-яё\-]+\s+[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+)?)",
    re.IGNORECASE)


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Все страницы PDF в единый текст."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
    return "\n".join(text_parts)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Весь текст DOCX (параграфы + ячейки таблиц + текстовые боксы).

    Обычный python-docx игнорирует текст в textbox/shape и в вложенных таблицах.
    Карточки организаций часто верстают в таких блоках — добавляем вытягивание всего <w:t> через XML.
    """
    from docx import Document
    parts: list[str] = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        # Рекурсивно проходим таблицы (в т.ч. вложенные)
        def _walk_table(tbl):
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            parts.append(p.text)
                    for nested in cell.tables:
                        _walk_table(nested)
        for tbl in doc.tables:
            _walk_table(tbl)
    except Exception:
        pass

    # Fallback — вытягиваем весь текст из word/document.xml включая textbox’ы
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            for name in ["word/document.xml", "word/header1.xml", "word/footer1.xml"]:
                if name not in z.namelist():
                    continue
                xml = z.read(name)
                root = ET.fromstring(xml)
                # Найти все <w:t> везде в дереве
                ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                for t in root.iter(f"{ns}t"):
                    if t.text and t.text.strip():
                        parts.append(t.text)
    except Exception:
        pass

    # Дедупликация последовательных одинаковых строк (чтобы регексы не видели дубли)
    dedup: list[str] = []
    _last = None
    for line in parts:
        if line != _last:
            dedup.append(line)
            _last = line
    return "\n".join(dedup)


def _decline_word_gen(word: str, is_female: bool) -> str:
    """Склонить отдельное слово (имя/отчество/фамилию) в родительный падеж."""
    if not word:
        return word
    w = word.strip()
    wl = w.lower()

    # Фамилии и имена женские
    if is_female:
        # Фамилии-прилагательные -ская/-цкая/-кая (Гладкая, Достоевская)
        if wl.endswith(("ская", "цкая", "ческая")):
            return w[:-2] + "ой"
        if wl.endswith("кая") and len(wl) > 5:
            return w[:-2] + "ой"
        # Фамилии -ова, -ева, -ина, -ына
        if wl.endswith(("ова", "ева", "ина", "ына")):
            return w[:-1] + "ой"
        # Женское имя -ия (Анастасия → Анастасии)
        if wl.endswith("ия"):
            return w[:-1] + "и"
        if wl.endswith("ья"):
            return w[:-1] + "и"
        # Женское имя/отчество на -га/-ка/-ха (Ольга → Ольги, Маринка → Маринки)
        if wl.endswith(("га", "ка", "ха")):
            return w[:-1] + "и"
        # Общее правило — женское на -а → -ы (Мария → Марии выше, Светлана → Светланы)
        if wl.endswith("а"):
            return w[:-1] + "ы"
        if wl.endswith("я"):
            return w[:-1] + "и"
        # Женские на согласный (Любовь) в родительном не меняются
        return w

    # Мужские фамилии
    if wl.endswith(("ов", "ев", "ёв", "ин", "ын")):
        return w + "а"
    if wl.endswith(("ий", "ый", "ой")):
        return w[:-2] + "ого"
    if wl.endswith(("ский", "цкий")):
        return w[:-2] + "ого"
    # Мужские имена и отчества
    if wl.endswith("ий"):
        return w[:-1] + "я"     # Николай → Николая, Геннадий → Геннадия
    if wl.endswith("а"):
        return w[:-1] + "ы"     # Никита → Никиты
    if wl.endswith("я"):
        return w[:-1] + "и"     # Илья → Ильи
    if wl.endswith('ь'):
        return w[:-1] + "я"     # Игорь → Игоря
    if wl.endswith('й'):
        return w[:-1] + "я"     # Алексей → Алексея, Сергей → Сергея
    if wl.endswith("о") or wl.endswith("е"):
        return w              # Петровское не склоняем
    # Мужское имя на согласный (Иван, Пётр)
    if wl and wl[-1] in "бвгджзклмнпрстфхцчшщ":
        # Пётр → Петра, Иван → Ивана
        # Проверяем неблагозвучное чередование: Пётр → Петр (егда с подвижной ё)
        if "ё" in w and w.endswith("тр"):
            w = w.replace("ё", "е")
        return w + "а"
    return w


def _is_female_by_patronymic(patronymic: str) -> bool:
    """Определяет пол по отчеству: оканчивается на -на = женщина."""
    p = patronymic.strip().lower()
    return p.endswith(("овна", "евна", "ична", "инична", "ынична"))


def _fio_short_to_gen(fio_short: str) -> str:
    """Преобразует ФИО в родительный падеж.

    Поддерживает:
      • Полное ФИО: «Сурина Анастасия Витальевна» → «Суриной Анастасии Витальевны»
      • Мужское: «Иванов Иван Иванович» → «Иванова Ивана Ивановича»
      • Сокращённое: «Иванов И.И.» → «Иванова И.И.»
    """
    fio_short = fio_short.strip()
    parts = fio_short.split()
    if not parts:
        return ""

    # Определяем пол: есть ли полное отчество?
    is_female = False
    if len(parts) >= 3:
        patronymic = parts[2]
        if len(patronymic) > 4 and not patronymic.endswith("."):
            is_female = _is_female_by_patronymic(patronymic)

    # Склоняем каждую часть
    result_parts = []
    for i, part in enumerate(parts):
        # Инициалы (И.И.) не склоняем
        if len(part) <= 3 and part.endswith("."):
            result_parts.append(part)
            continue
        result_parts.append(_decline_word_gen(part, is_female))

    return " ".join(result_parts)


def extract_requisites_from_text(text: str) -> ExtractedRequisites:
    """Извлекает реквизиты компании из произвольного текста (скопированного/введённого).

    Находит: ИНН, КПП, ОГРН, банк, БИК, р/с, к/с, адрес, телефон, email,
    наименование ЮЛ, ФИО директора. Все поля опциональные.
    """
    if not text or not text.strip():
        return ExtractedRequisites()
    return _extract_requisites_impl(text)


def extract_requisites(file_bytes: bytes, filename: str) -> ExtractedRequisites:
    """Извлекает реквизиты компании из PDF/DOCX файла.

    Обертка над extract_requisites_from_text — сначала вытягивает текст из файла.
    """
    lower = filename.lower()
    if lower.endswith(".docx"):
        text = _extract_text_from_docx(file_bytes)
    elif lower.endswith(".pdf"):
        text = _extract_text_from_pdf(file_bytes)
    else:
        return ExtractedRequisites()

    if not text:
        return ExtractedRequisites()
    return _extract_requisites_impl(text)


def _extract_requisites_impl(text: str) -> ExtractedRequisites:
    """Внутренняя: выполняет регексы по готовому тексту."""

    # Нормализуем пробелы
    text = text.replace("\xa0", " ")

    req = ExtractedRequisites()

    # Сначала пробуем комбинацию «ИНН/КПП X/Y»
    m = _RE_INN_KPP_COMBO.search(text)
    if m:
        req.inn = m.group(1)
        req.kpp = m.group(2)
    else:
        # ИНН
        m = _RE_INN.search(text)
        if m:
            req.inn = m.group(1)
        # КПП
        m = _RE_KPP.search(text)
        if m:
            req.kpp = m.group(1)

    # ОГРН
    m = _RE_OGRN.search(text)
    if m:
        req.ogrn = m.group(1)

    # БИК — основной регекс
    m = _RE_BIK.search(text)
    if m:
        req.bank_bik = m.group(1)
    # Fallback: метка «БИК» на одной строке, число на следующей (МВМС-формат)
    if not req.bank_bik:
        _lines_bik = [l.strip() for l in text.split("\n")]
        for i, L in enumerate(_lines_bik):
            if L.upper().rstrip(":.") in ("БИК", "БИК БАНКА"):
                for j in range(i + 1, min(i + 3, len(_lines_bik))):
                    _cand = _lines_bik[j].strip()
                    if re.fullmatch(r"\d{9}", _cand):
                        req.bank_bik = _cand
                        break
                if req.bank_bik:
                    break

    # Р/с
    m = _RE_RS.search(text)
    if m:
        req.bank_account = m.group(1)

    # К/с — основной регекс
    m = _RE_KS.search(text)
    if m:
        req.corr_account = m.group(1)
    # Fallback: метка «Корреспондентский счет» / «К/с» / «Корр. счет» на одной строке, число в следующих
    if not req.corr_account:
        _lines_ks = [l.strip() for l in text.split("\n")]
        for i, L in enumerate(_lines_ks):
            Lup = L.upper().rstrip(":.")
            if Lup in ("КОРРЕСПОНДЕНТСКИЙ СЧЕТ", "КОРРЕСПОНДЕНТСКИЙ", "К/С", "К-С", "КОРР. СЧЕТ", "КОРР.СЧЕТ"):
                for j in range(i + 1, min(i + 4, len(_lines_ks))):
                    _cand = _lines_ks[j].strip()
                    if re.fullmatch(r"\d{20}", _cand):
                        req.corr_account = _cand
                        break
                if req.corr_account:
                    break

    # Если не нашли р/с явно, пытаемся эвристикой:
    # Ищем все 20-значные числа. Первое обычно р/с, второе к/с.
    if not req.bank_account or not req.corr_account:
        all_accs = _RE_ACC_ANY.findall(text)
        # Отфильтруем ИНН (12-значные) и другие короткие
        accs = [a for a in all_accs if len(a) == 20]
        if not req.bank_account and accs:
            req.bank_account = accs[0]
        if not req.corr_account and len(accs) >= 2:
            req.corr_account = accs[1]

    # Полное наименование
    m = _RE_COMPANY_FULL.search(text)
    if m:
        form_full = m.group(1).strip()
        name_body = m.group(2).strip()
        req.company_full = f"{form_full} «{name_body}»"
        # Автогенерация краткого из полного
        _short_map = {
            "общество с ограниченной": "ООО",
            "открытое акционерное": "ОАО",
            "закрытое акционерное": "ЗАО",
            "публичное акционерное": "ПАО",
            "акционерное общество": "АО",
            "индивидуальный предприниматель": "ИП",
            "федеральное казенное предприятие": "ФКП",
            "федеральное государственное унитарное предприятие": "ФГУП",
            "муниципальное унитарное предприятие": "МУП",
            "автономная некоммерческая организация": "АНО",
        }
        # Схлопываем переносы/табы в имени
        name_body = re.sub(r"[\s\n\t]+", " ", name_body)
        req.company_full = re.sub(r"[\s\n\t]+", " ", req.company_full)
        for full_form, short_form in _short_map.items():
            if form_full.lower().startswith(full_form):
                req.company_short = f"{short_form} «{name_body}»"
                break

    # Краткое наименование (если ещё не найдено)
    if not req.company_short:
        m = _RE_COMPANY.search(text)
        if m:
            form_short = m.group(1).upper()
            name_body = re.sub(r"[\s\n\t]+", " ", m.group(2).strip())
            req.company_short = f"{form_short} «{name_body}»"
            # Если полное не нашли — используем короткое как fallback
            if not req.company_full:
                expand = {
                    "ООО": "Общество с ограниченной ответственностью",
                    "ОАО": "Открытое акционерное общество",
                    "ЗАО": "Закрытое акционерное общество",
                    "АО": "Акционерное общество",
                    "ПАО": "Публичное акционерное общество",
                    "ИП": "Индивидуальный предприниматель",
                }.get(form_short, form_short)
                req.company_full = f"{expand} «{name_body}»"

    # Адрес
    m = _RE_ADDRESS.search(text)
    if m:
        req.address = re.sub(r"\s+", " ", m.group(1)).strip()
    else:
        m = _RE_ADDRESS_FALLBACK.search(text)
        if m:
            req.address = re.sub(r"\s+", " ", m.group(1)).strip()

    # Банк — несколько стратегий подбора
    m = _RE_BANK.search(text)
    if m:
        bank = m.group(1).strip().rstrip(",;.")
        bank = re.sub(r"[\s\t\n]+", " ", bank).strip().rstrip(",;.")
        if len(bank) > 5 and not bank.lower().startswith(("россии", "рф", "россия")):
            req.bank_name = bank
    # Фолбэк: если не нашли — ищем любую строку с ФИЛИАЛ/ПАО/АО и словом БАНК
    if not req.bank_name:
        for line in text.split("\n"):
            L = line.strip().strip(",;.").strip('"«» ')
            L_up = L.upper()
            if not L or len(L) > 200 or len(L) < 6:
                continue
            # Строка содержит ФОРМУ банка (ПАО/АО/Филиал) И слово БАНК
            has_form = any(k in L_up for k in ["ПАО ", " ПАО", "АО ", " АО", "ФИЛИАЛ"])
            has_bank = "БАНК" in L_up or "СБЕРБАНК" in L_up or "ВТБ" in L_up or "АЛЬФА" in L_up or "ПСБ" in L_up or "ТИНЬКОФФ" in L_up or "РОССЕЛЬХОЗ" in L_up
            if has_form and has_bank:
                # Отсеиваем строки-метки (без конкретного банка) вроде «Наименование банка»
                if L_up.startswith("НАИМЕНОВАНИЕ"): continue
                # Исключаем если в строке есть ключевые слова компании-клиента (чтобы не взять КП)
                if any(k in L_up for k in ["ООО ", "ОГРН ", "ИНН "]): continue
                # Отсеиваем если в строке есть слова перевозчика/покупателя
                if any(k in L_up for k in ["ГРУЗОПОЛУЧАТЕЛЬ", "ПОКУПАТЕЛЬ", "ПОСТАВЩИК", "ПЛАТЕЛЬЩИК"]): continue
                req.bank_name = L
                break
    # Ещё один вариант: строка «Банк» / «Наименование банка» только метка, а валуе в следующих строках
    if not req.bank_name:
        lines = [L.strip() for L in text.split("\n")]
        for i, L in enumerate(lines):
            L_up = L.upper()
            if L_up in ("БАНК", "БАНК:", "НАИМЕНОВАНИЕ БАНКА", "НАИМЕНОВАНИЕ БАНКА:"):
                # Начиная со следующей строки собираем непустые до БИК/Р-с/ИНН/Корр.
                collected = []
                for j in range(i + 1, min(i + 5, len(lines))):
                    nx = lines[j].strip()
                    if not nx:
                        continue
                    nx_up = nx.upper()
                    if any(nx_up.startswith(s) for s in ("БИК", "Р/С", "Р/СЧ", "РАСЧЕТН", "К/С", "КОРР", "ИНН", "КПП", "ОГРН", "АДРЕС", "ТЕЛ", "E-MAIL", "EMAIL", "ДИРЕКТОР")):
                        break
                    collected.append(nx)
                    if len(" ".join(collected)) > 30:
                        break
                if collected:
                    cand = " ".join(collected).strip().rstrip(",;.")
                    cand = re.sub(r"[\s\t\n]+", " ", cand)
                    if len(cand) > 5:
                        req.bank_name = cand
                        break

    # Телефон — основной регекс
    m = _RE_PHONE.search(text)
    if m:
        req.phone = m.group(1).strip()
    # Fallback: метка «Телефон» / «Тел» на одной строке, число на следующей (МВМС-формат)
    if not req.phone:
        _lines_ph = [l.strip() for l in text.split("\n")]
        for i, L in enumerate(_lines_ph):
            Lup = L.upper().rstrip(":.")
            if Lup in ("ТЕЛЕФОН", "ТЕЛ", "ТЕЛ.", "КОНТАКТНЫЙ ТЕЛЕФОН", "Телефон/Факс"):
                for j in range(i + 1, min(i + 3, len(_lines_ph))):
                    _cand = _lines_ph[j].strip()
                    # Простой тест: цифр ≥ 10
                    _digits = re.sub(r"\D", "", _cand)
                    if len(_digits) >= 10 and len(_digits) <= 15:
                        req.phone = _cand
                        break
                if req.phone:
                    break
    # Нормализация: 84950250077 → +7 (495) 025-00-77
    if req.phone:
        _d = re.sub(r"\D", "", req.phone)
        if len(_d) == 11 and _d.startswith(("7", "8")):
            req.phone = f"+7 ({_d[1:4]}) {_d[4:7]}-{_d[7:9]}-{_d[9:11]}"

    # Email — отфильтровываем @rolls-kran.ru и другие собственные почты поставщика
    for m in _RE_EMAIL.finditer(text):
        email = m.group(1)
        if "rolls-kran" in email.lower():
            continue
        req.email = email
        break
    # Fallback: email мог быть разорван по строкам: info \n @ \n mwms.ru
    if not req.email:
        # Схлопываем только в сегментах вокруг символа @ (чтобы не склеивать с соседними словами)
        for _at_pos in [m.start() for m in re.finditer(r"@", text)]:
            # Захватываем 60 символов вокруг @ и склеиваем внутри пробелы
            _seg = text[max(0, _at_pos - 40):_at_pos + 40]
            _seg_c = re.sub(r"\s+", "", _seg)
            # Ограниченный email-регекс: домен макс 20 символов, только ASCII без кириллицы
            for m in re.finditer(r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9-]+\.[a-zA-Z]{2,4})(?![a-zA-Z0-9])", _seg_c):
                email = m.group(1)
                if "rolls-kran" in email.lower():
                    continue
                # Адрес валидный: точка в домене ближе к концу, локал не пустой
                if "@" in email and email.count("@") == 1:
                    req.email = email
                    break
            if req.email:
                break

    # Сайт
    for m in _RE_SITE.finditer(text):
        site = m.group(1)
        if "rolls-kran" in site.lower():
            continue
        if any(x in site.lower() for x in ["mail.", "@"]):
            continue
        req.site = site
        break

    # Директор — сначала подбор по строкам: метка на одной строке, ФИО на следующей
    _lines_dir = [l.strip() for l in text.split("\n")]
    for i, L in enumerate(_lines_dir):
        L_up = L.upper().rstrip(":.")
        # Метка должности (только «главный» директор, без Заместителя/Помощника)
        if L_up in ("ГЕНЕРАЛЬНЫЙ ДИРЕКТОР", "ДИРЕКТОР", "РУКОВОДИТЕЛЬ",
                    "ПРЕЗИДЕНТ", "ГЛАВА", "ИНДИВИДУАЛЬНЫЙ ПРЕДПРИНИМАТЕЛЬ",
                    "ФИО ДИРЕКТОРА", "Ф.И.О. ДИРЕКТОРА"):
            # Ищем ФИО в следующих 3-4 строках
            for j in range(i + 1, min(i + 5, len(_lines_dir))):
                cand = _lines_dir[j].strip().strip(",;.").strip()
                if not cand:
                    continue
                # Завершаем если наткнулись на следующую метку
                if cand.upper().startswith(("ОСНОВАНИЕ", "СВИДЕТЕЛЬ", "УСТАВ", "ОГРН",
                        "ИНН", "АДРЕС", "ТЕЛ", "E-MAIL", "EMAIL", "БАНК", "БИК",
                        "Р/С", "К/С", "КОРР", "8.", "9.", "10.")):
                    break
                # Проверяем: похоже на ФИО? 2-4 слова с заглавными буквами
                cand_parts = cand.split()
                if 2 <= len(cand_parts) <= 4 and all(
                    len(p) >= 2 and p[0].isupper() and any(c.isalpha() for c in p)
                    for p in cand_parts
                ):
                    req.director_short = cand
                    break
            if req.director_short:
                break

    # Директор
    m = _RE_DIRECTOR_FULL.search(text)
    if m:
        fio_full = m.group(1).strip()
        # Если это уже сокращённый вид (Иванов И.И.) — вычислим полный отдельно
        if "." in fio_full:
            req.director_short = fio_full
        else:
            # Полное ФИО типа "Иванов Иван Иванович"
            parts = fio_full.split()
            if len(parts) >= 2:
                surname = parts[0]
                name_ini = parts[1][0] + "."
                patr_ini = parts[2][0] + "." if len(parts) >= 3 else ""
                req.director_short = f"{surname} {name_ini}{patr_ini}"
                # Родительный падеж полного ФИО
                req.director_gen = _fio_full_to_gen(fio_full)
    else:
        m = _RE_DIRECTOR.search(text)
        if m:
            req.director_short = m.group(1).strip()

    # Если ПОЛНОЕ ФИО (без точек — значит имя-отчество полностью) — вычисляем сокращённое и род. падеж
    if req.director_short and "." not in req.director_short:
        parts = req.director_short.split()
        if len(parts) >= 2:
            surname = parts[0]
            name_ini = parts[1][0] + "." if len(parts[1]) > 0 else ""
            patr_ini = parts[2][0] + "." if len(parts) >= 3 and len(parts[2]) > 0 else ""
            req.director_gen = _fio_full_to_gen(req.director_short)
            req.director_short = f"{surname} {name_ini}{patr_ini}".strip()

    # Если получили только сокращённое ФИО и не сгенерили род. падеж
    if req.director_short and not req.director_gen:
        req.director_gen = _fio_short_to_gen(req.director_short)

    # Должность директора — по умолчанию Генеральный директор
    if req.director_short or req.director_gen:
        if re.search(r"генеральн\w*\s+директор", text, re.IGNORECASE):
            req.director_title = "Генеральный директор"
        elif re.search(r"\bдиректор", text, re.IGNORECASE):
            req.director_title = "Директор"
        elif re.search(r"\bруководитель", text, re.IGNORECASE):
            req.director_title = "Руководитель"
        else:
            req.director_title = "Генеральный директор"

    return req


def _fio_full_to_gen(fio: str) -> str:
    """Полное ФИО в родительный падеж (делегируем на уточнённый алгоритм)."""
    return _fio_short_to_gen(fio)


def _fio_full_to_gen_OLD(fio: str) -> str:
    """[deprecated] Старая примитивная эвристика."""
    parts = fio.strip().split()
    if len(parts) < 2:
        return fio
    surname = parts[0]
    name = parts[1] if len(parts) >= 2 else ""
    patronymic = parts[2] if len(parts) >= 3 else ""

    # Фамилия
    if surname.endswith(("ов", "ев", "ин", "ын")):
        surname_gen = surname + "а"
    elif surname.endswith(("ий", "ый")):
        surname_gen = surname[:-2] + "ого"
    elif surname.endswith("а") and len(surname) > 3:
        surname_gen = surname[:-1] + "ой"
    else:
        surname_gen = surname

    # Имя (мужское)
    male_names_map = {
        "Александр": "Александра", "Алексей": "Алексея", "Андрей": "Андрея",
        "Антон": "Антона", "Артем": "Артема", "Артём": "Артёма",
        "Борис": "Бориса", "Вадим": "Вадима", "Валерий": "Валерия",
        "Виктор": "Виктора", "Виталий": "Виталия", "Владимир": "Владимира",
        "Владислав": "Владислава", "Вячеслав": "Вячеслава", "Геннадий": "Геннадия",
        "Григорий": "Григория", "Дмитрий": "Дмитрия", "Евгений": "Евгения",
        "Егор": "Егора", "Игорь": "Игоря", "Илья": "Ильи", "Иван": "Ивана",
        "Кирилл": "Кирилла", "Константин": "Константина", "Максим": "Максима",
        "Михаил": "Михаила", "Никита": "Никиты", "Николай": "Николая",
        "Олег": "Олега", "Павел": "Павла", "Пётр": "Петра", "Петр": "Петра",
        "Роман": "Романа", "Сергей": "Сергея", "Станислав": "Станислава",
        "Юрий": "Юрия",
    }
    name_gen = male_names_map.get(name, name + "а" if not name.endswith("а") else name[:-1] + "ы")

    # Отчество
    if patronymic.endswith("вич"):
        patr_gen = patronymic + "а"
    elif patronymic.endswith("ич"):
        patr_gen = patronymic + "а"
    else:
        patr_gen = patronymic

    return f"{surname_gen} {name_gen} {patr_gen}".strip()
