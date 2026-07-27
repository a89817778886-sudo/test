# -*- coding: utf-8 -*-
"""Генератор договора поставки по стороннему КП.

Не связан с QuoteData (кран/траверса) — принимает уже готовую спецификацию
и реквизиты покупателя. Формирует DOCX и PDF версии с Приложением № 1
(чертёж).
"""
from __future__ import annotations
import io
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass, field
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    BaseDocTemplate, Frame, Image, KeepTogether, PageBreak,
    PageTemplate, Paragraph, Spacer, Table, TableStyle,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from money_words import money_to_words, money_full


MEDIA_DIR = Path(__file__).parent / "media"


@dataclass
class ExternalContractData:
    """Все входные данные для внешнего договора."""
    # Реквизиты покупателя
    buyer_full: str = ""                # ООО «Ромашка»
    buyer_short: str = ""
    buyer_inn: str = ""
    buyer_kpp: str = ""
    buyer_ogrn: str = ""
    buyer_address: str = ""
    buyer_bank: str = ""
    buyer_bik: str = ""
    buyer_account: str = ""
    buyer_corr_account: str = ""
    buyer_director_short: str = ""      # Иванов И.И.
    buyer_director_gen: str = ""        # Иванова Ивана Ивановича
    buyer_director_title: str = "Генеральный директор"
    buyer_basis: str = "Устава"         # действующего на основании…
    buyer_phone: str = ""
    buyer_email: str = ""
    buyer_type: str = "OOO"             # "OOO" | "IP" — адаптация преамбулы

    # Спецификация (уже распарсенные позиции)
    items: List[Dict[str, Any]] = field(default_factory=list)
    # каждая позиция: {code, name, unit, qty, price}

    # Параметры
    has_vat: bool = True                # True → цена включает НДС 22%, False → без НДС
    prepay_pct: int = 70                # % предоплаты
    delivery_type: str = "Самовывоз"    # или "Доставка"
    delivery_address: str = ""          # если Доставка — куда; если Самовывоз — откуда
    delivery_cost: float = 0.0
    contract_number: str = ""
    contract_date: str = ""             # 22.07.2026
    supplier_key: str = "LKS"           # LKS / MODERNIZATSIYA / KINEMATIKA
    include_stamp: bool = True
    drawing_bytes: Optional[bytes] = None  # PNG байты чертежа
    drawing_caption: str = "Габаритный чертёж оборудования"
    comment: str = ""                   # необязательный комментарий

    # Условия договора
    shipment_term_days: int = 20        # срок изготовления
    warranty_months: int = 12
    contract_valid_days: int = 14
    pickup_days: int = 10
    storage_free_days: int = 10
    storage_penalty_pct: float = 1.0
    shipment_penalty_pct: float = 1.0

    @property
    def total_amount(self) -> float:
        """Сумма всех позиций + доставка (доставка идёт отдельной строкой в договоре)."""
        return sum(float(it.get("qty", 0)) * float(it.get("price", 0))
                   for it in self.items) + float(self.delivery_cost or 0)

    @property
    def spec_amount(self) -> float:
        """Только позиции (без доставки)."""
        return sum(float(it.get("qty", 0)) * float(it.get("price", 0))
                   for it in self.items)


# --- DOCX генератор ---


def _sup_get(sup: Any, *keys, default: str = "") -> str:
    """Универсальное получение поля из supplier (dict или объект)."""
    for k in keys:
        if isinstance(sup, dict):
            if k in sup and sup[k]:
                return str(sup[k])
        else:
            v = getattr(sup, k, None)
            if v:
                return str(v)
    return default

def _fmt_money(x: float) -> str:
    """1234567.89 → '1 234 567,89'"""
    return f"{x:,.2f}".replace(",", " ").replace(".", ",")


def build_external_contract_docx(data: ExternalContractData,
                                 supplier: Any) -> bytes:
    """Строит DOCX-договор поставки по внешнему КП."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(1.5)

    # Заголовок
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"ДОГОВОР ПОСТАВКИ № {data.contract_number}")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"г. Санкт-Петербург                                          {data.contract_date}")

    doc.add_paragraph()

    # Преамбула
    supplier_full = _sup_get(supplier, "label", "short", "full", default="ООО «ЛКС»")
    supplier_short = _sup_get(supplier, "short", "label", default="ООО «ЛКС»")
    supplier_director_gen = _sup_get(supplier, "director_fio_gen")
    supplier_director_short = _sup_get(supplier, "director_fio_short")
    supplier_director_title = _sup_get(supplier, "director_position", "director_title", default="Генеральный директор")
    supplier_basis = _sup_get(supplier, "director_basis", "basis", default="Устава")

    _is_ip = str(getattr(data, "buyer_type", "OOO") or "OOO").upper() == "IP"
    _buyer_form = "именуемый" if _is_ip else "именуемое"
    if _is_ip:
        _buyer_person_clause = "действующего на основании Свидетельства о государственной регистрации в качестве ИП"
    else:
        _buyer_person_clause = (f"в лице {data.buyer_director_title.lower()} {data.buyer_director_gen or '___'}, "
                                f"действующего на основании {data.buyer_basis}")
    preamble = (
        f"{supplier_full}, именуемое в дальнейшем «Поставщик», в лице "
        f"{supplier_director_title.lower()} {supplier_director_gen}, "
        f"действующего на основании {supplier_basis}, с одной стороны, "
        f"и {data.buyer_full}, {_buyer_form} в дальнейшем «Покупатель», "
        f"{_buyer_person_clause}, с другой стороны, "
        f"совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:"
    )
    p = doc.add_paragraph(preamble)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(8)

    def h(text):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        p.paragraph_format.space_before = Pt(8)

    def para(text, indent=True):
        p = doc.add_paragraph(text)
        if indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(4)

    total = data.total_amount
    vat_text = "с НДС 22 %" if data.has_vat else "без НДС"

    # 1. Предмет
    h("1. ПРЕДМЕТ ДОГОВОРА")
    para("1.1. Поставщик обязуется передать в собственность Покупателю, "
         "а Покупатель обязуется принять и оплатить оборудование "
         "(далее — Товар), номенклатура, количество и цена которого "
         "указаны в Спецификации, являющейся неотъемлемой частью настоящего Договора.")
    if data.drawing_bytes:
        para("1.2. Габаритный чертёж поставляемого Товара приведён в "
             "Приложении № 1 к настоящему Договору.")

    # 2. Цена и порядок расчётов
    h("2. ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ")
    para(f"2.1. Общая цена настоящего Договора составляет "
         f"{_fmt_money(total)} руб. ({money_to_words(total)}), {vat_text}.")

    prepay_sum = total * data.prepay_pct / 100
    remainder_sum = total - prepay_sum
    remainder_pct = round(100 - data.prepay_pct)
    if data.prepay_pct == 100:
        para(f"2.2. Оплата производится в размере 100 % — "
             f"{_fmt_money(total)} руб. ({money_to_words(total)}) — "
             f"единовременно в течение 5 (пяти) рабочих дней с момента "
             f"подписания настоящего Договора и Спецификации.")
    else:
        para(f"2.2. Оплата производится в следующем порядке:")
        para(f"— Предоплата {data.prepay_pct} % — "
             f"{_fmt_money(prepay_sum)} руб. ({money_to_words(prepay_sum)}) — "
             f"в течение 5 (пяти) рабочих дней после подписания настоящего "
             f"Договора и Спецификации;", indent=False)
        para(f"— Окончательный расчёт {remainder_pct} % — "
             f"{_fmt_money(remainder_sum)} руб. ({money_to_words(remainder_sum)}) — "
             f"в течение 5 (пяти) рабочих дней после уведомления Поставщиком "
             f"Покупателя о готовности Товара к отгрузке.", indent=False)
    para("2.3. Все расчёты производятся в безналичной форме на расчётный счёт Поставщика.")

    # 3. Срок и условия поставки
    h("3. СРОК И УСЛОВИЯ ПОСТАВКИ")
    para(f"3.1. Срок изготовления Товара составляет до {data.shipment_term_days} "
         f"({int_to_words_days(data.shipment_term_days)}) рабочих дней с момента "
         f"поступления предоплаты на расчётный счёт Поставщика.")
    if data.delivery_type == "Самовывоз":
        para(f"3.2. Условия поставки: Самовывоз со склада Поставщика "
             f"по адресу: {data.delivery_address or 'г. Санкт-Петербург, склад Поставщика'}.")
        para(f"3.3. Покупатель обязан вывезти Товар со склада Поставщика "
             f"в течение {data.pickup_days} ({int_to_words_days(data.pickup_days)}) "
             f"рабочих дней с даты уведомления о готовности к отгрузке. "
             f"При нарушении этого срока Покупатель уплачивает Поставщику "
             f"пени в размере {data.shipment_penalty_pct:.1f} % от стоимости "
             f"неотгруженного Товара за каждый день просрочки.")
        para(f"3.4. Хранение Товара свыше {data.storage_free_days} "
             f"({int_to_words_days(data.storage_free_days)}) рабочих дней "
             f"после уведомления оплачивается Покупателем дополнительно в "
             f"размере {data.storage_penalty_pct:.1f} % от стоимости Товара "
             f"за каждый день хранения.")
    else:
        para(f"3.2. Условия поставки: Доставка силами Поставщика по адресу: "
             f"{data.delivery_address or '___'}. Стоимость доставки составляет "
             f"{_fmt_money(data.delivery_cost)} руб. и включена в общую цену Договора.")

    # 4. Гарантия
    h("4. КАЧЕСТВО И ГАРАНТИЯ")
    para(f"4.1. Поставщик гарантирует качество Товара в течение "
         f"{data.warranty_months} ({int_to_words_months(data.warranty_months)}) "
         f"месяцев с даты передачи Товара Покупателю.")
    para("4.2. Гарантия не распространяется на дефекты, возникшие вследствие "
         "нарушения Покупателем правил эксплуатации, транспортировки или хранения.")

    # 5. Ответственность
    h("5. ОТВЕТСТВЕННОСТЬ СТОРОН")
    para("5.1. За нарушение обязательств по настоящему Договору Стороны несут "
         "ответственность в соответствии с действующим законодательством РФ.")
    para("5.2. Стороны освобождаются от ответственности за частичное или "
         "полное неисполнение обязательств, если это неисполнение является "
         "следствием обстоятельств непреодолимой силы (форс-мажор).")

    # 6. Прочие условия
    h("6. ПРОЧИЕ УСЛОВИЯ")
    para(f"6.1. Настоящий Договор вступает в силу с момента его подписания "
         f"обеими Сторонами и действует до полного исполнения Сторонами "
         f"своих обязательств. Настоящий Договор действителен для акцепта "
         f"в течение {data.contract_valid_days} "
         f"({int_to_words_days(data.contract_valid_days)}) рабочих дней "
         f"с даты его составления.")
    para("6.2. Все споры и разногласия решаются путём переговоров. При "
         "невозможности урегулирования — в Арбитражном суде по месту "
         "нахождения Ответчика.")
    para("6.3. Договор составлен в двух экземплярах, имеющих равную "
         "юридическую силу, по одному для каждой из Сторон.")

    if data.comment:
        h("7. ДОПОЛНИТЕЛЬНЫЕ УСЛОВИЯ")
        para(data.comment)

    # Реквизиты сторон — 2-колоночная таблица
    doc.add_paragraph()
    h("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(8.5)
    tbl.columns[1].width = Cm(8.5)
    cell_l = tbl.cell(0, 0)
    cell_r = tbl.cell(0, 1)

    def _fill_party(cell, is_supplier: bool):
        if is_supplier:
            lines = [
                ("Поставщик:", True),
                (supplier_full, True),
                (f"ИНН {_sup_get(supplier, "inn")} / КПП {_sup_get(supplier, "kpp")}", False),
                (f"ОГРН {_sup_get(supplier, "ogrn")}", False),
                (f"Адрес: {_sup_get(supplier, "address")}", False),
                (f"р/с {_sup_get(supplier, "rs", "bank_account")}", False),
                (f"в {_sup_get(supplier, "bank", "bank_name")}", False),
                (f"БИК {_sup_get(supplier, "bik")}", False),
                (f"к/с {_sup_get(supplier, "ks", "corr_account")}", False),
                (f"Тел.: {_sup_get(supplier, "phone", "phone_short")}", False),
                (f"E-mail: {_sup_get(supplier, "email")}", False),
                ("", False),
                (f"{supplier_director_title}", False),
                (f"{supplier_director_short} ___________________", False),
                ("М.П.", False),
            ]
        else:
            lines = [
                ("Покупатель:", True),
                (data.buyer_full or "ООО «___»", True),
                (f"ИНН {data.buyer_inn} / КПП {data.buyer_kpp}", False),
                (f"ОГРН {data.buyer_ogrn}", False),
                (f"Адрес: {data.buyer_address}", False),
                (f"р/с {data.buyer_account}", False),
                (f"в {data.buyer_bank}", False),
                (f"БИК {data.buyer_bik}", False),
                (f"к/с {data.buyer_corr_account}", False),
                (f"Тел.: {data.buyer_phone}", False),
                (f"E-mail: {data.buyer_email}", False),
                ("", False),
                (f"{data.buyer_director_title}", False),
                (f"{data.buyer_director_short or '___'} ___________________", False),
                ("М.П.", False),
            ]
        # Очистим стандартный параграф
        cell.text = ""
        for text, bold in lines:
            p = cell.add_paragraph()
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.bold = bold

    _fill_party(cell_l, is_supplier=True)
    _fill_party(cell_r, is_supplier=False)

    # ============= СПЕЦИФИКАЦИЯ (новая страница) =============
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"СПЕЦИФИКАЦИЯ № 1")
    r.bold = True
    r.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"к Договору поставки № {data.contract_number} от {data.contract_date}")
    r.font.size = Pt(11)

    doc.add_paragraph()

    spec_tbl = doc.add_table(rows=1, cols=6)
    spec_tbl.style = "Light Grid Accent 1"
    hdr = spec_tbl.rows[0].cells
    for i, h_text in enumerate(["№", "Код", "Наименование", "Ед.", "Кол-во", "Цена, ₽"]):
        hdr[i].text = h_text
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)

    for idx, item in enumerate(data.items, 1):
        row = spec_tbl.add_row().cells
        row[0].text = str(idx)
        row[1].text = str(item.get("code", ""))
        row[2].text = str(item.get("name", ""))
        row[3].text = str(item.get("unit", "шт"))
        row[4].text = f"{float(item.get('qty', 0)):.1f}".rstrip("0").rstrip(".")
        row[5].text = _fmt_money(float(item.get("price", 0)))
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # Строка "доставка" отдельно, если стоимость > 0
    if data.delivery_cost > 0:
        row = spec_tbl.add_row().cells
        row[0].text = str(len(data.items) + 1)
        row[1].text = "—"
        row[2].text = f"Доставка ({data.delivery_type})"
        row[3].text = "усл."
        row[4].text = "1"
        row[5].text = _fmt_money(data.delivery_cost)

    # Итого
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"ИТОГО: {_fmt_money(total)} руб., {vat_text}")
    r.bold = True
    r.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Прописью: {money_to_words(total)}")
    r.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph()

    # Подписи под спецификацией
    tbl2 = doc.add_table(rows=1, cols=2)
    c_l = tbl2.cell(0, 0)
    c_r = tbl2.cell(0, 1)
    c_l.text = ""
    c_l.add_paragraph().add_run("Поставщик:").bold = True
    c_l.add_paragraph(f"{supplier_director_title}")
    c_l.add_paragraph(f"{supplier_director_short}  _________________")
    c_l.add_paragraph("М.П.")
    c_r.text = ""
    c_r.add_paragraph().add_run("Покупатель:").bold = True
    c_r.add_paragraph(f"{data.buyer_director_title}")
    c_r.add_paragraph(f"{data.buyer_director_short or '___'}  _________________")
    c_r.add_paragraph("М.П.")

    # ============= ПРИЛОЖЕНИЕ № 1 (чертёж) =============
    if data.drawing_bytes:
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("ПРИЛОЖЕНИЕ № 1")
        r.bold = True
        r.font.size = Pt(13)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"к Договору поставки № {data.contract_number} от {data.contract_date}")
        r.font.size = Pt(11)
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(data.drawing_caption)
        r.bold = True
        r.font.size = Pt(11)
        # Сохраним временно чтобы вставить в docx
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(data.drawing_bytes)
            tmp_path = tmp.name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(tmp_path, width=Cm(16))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def int_to_words_days(n: int) -> str:
    """Число прописью для дней (мужской род)."""
    from money_words import int_to_words_ru
    return int_to_words_ru(n, currency_gender="male")


def int_to_words_months(n: int) -> str:
    from money_words import int_to_words_ru
    return int_to_words_ru(n, currency_gender="male")


# --- PDF генератор ---

def _register_font(name: str, file: str) -> str:
    """Регистрирует шрифт. Ищет в корне lks-kp-app, затем в fonts/. Fallback — Helvetica."""
    base = Path(__file__).parent
    for candidate in (base / file, base / "fonts" / file):
        if candidate.exists():
            try:
                pdfmetrics.registerFont(TTFont(name, str(candidate)))
                return name
            except Exception:
                pass
    return "Helvetica"


def build_external_contract_pdf(data: ExternalContractData,
                                supplier: Any) -> bytes:
    """Строит PDF-договор для просмотра/печати. Логика та же что в DOCX."""
    # Шрифты: DejaVuSans-Fallback — единственный в проекте шрифт,
    # который поддерживает все нужные глифы (кириллица + №). Ставим его
    # везде, полужирный вариант — SF UI Display-Medium (он поддерживает
    # кириллицу, а № в заголовках не так критичен).
    reg = _register_font("ExtCtRegular", "DejaVuSans-Fallback.ttf")
    bold = _register_font("ExtCtBold", "DejaVuSans-Fallback.ttf")
    # Зарегистрируем семейство чтобы <b>…</b> внутри Paragraph работал
    if reg != "Helvetica":
        try:
            pdfmetrics.registerFontFamily(reg, normal=reg, bold=bold,
                                          italic=reg, boldItalic=bold)
        except Exception:
            pass

    buf = io.BytesIO()
    doc = BaseDocTemplate(buf, pagesize=A4,
                          leftMargin=20*mm, rightMargin=15*mm,
                          topMargin=20*mm, bottomMargin=15*mm)
    frame = Frame(doc.leftMargin, doc.bottomMargin,
                  doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="default", frames=frame)])

    styles = {
        "title": ParagraphStyle("title", fontName=bold, fontSize=13,
                                alignment=1, spaceAfter=6, leading=15),
        "subtitle": ParagraphStyle("subtitle", fontName=reg, fontSize=10,
                                   alignment=1, spaceAfter=8, leading=12),
        "h": ParagraphStyle("h", fontName=bold, fontSize=10.5, alignment=0,
                            spaceBefore=8, spaceAfter=4, leading=13),
        "p": ParagraphStyle("p", fontName=reg, fontSize=10, alignment=4,
                            firstLineIndent=25, spaceAfter=3, leading=13),
        "party_bold": ParagraphStyle("pb", fontName=bold, fontSize=9,
                                     alignment=0, leading=11, spaceAfter=1),
        "party": ParagraphStyle("pr", fontName=reg, fontSize=8.5,
                                alignment=0, leading=10, spaceAfter=0),
        "spec_cell": ParagraphStyle("sc", fontName=reg, fontSize=8.5,
                                    alignment=0, leading=10),
        "spec_hdr": ParagraphStyle("sh", fontName=bold, fontSize=9,
                                   alignment=1, leading=11, textColor=colors.white),
    }

    story = []
    supplier_full = _sup_get(supplier, "label", "short", "full", default="ООО «ЛКС»")
    supplier_director_gen = _sup_get(supplier, "director_fio_gen")
    supplier_director_short = _sup_get(supplier, "director_fio_short")
    supplier_director_title = _sup_get(supplier, "director_position", "director_title", default="Генеральный директор")
    supplier_basis = _sup_get(supplier, "director_basis", "basis", default="Устава")

    story.append(Paragraph(f"ДОГОВОР ПОСТАВКИ № {data.contract_number}", styles["title"]))
    story.append(Paragraph(f"г. Санкт-Петербург &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; {data.contract_date}",
                          styles["subtitle"]))

    _is_ip_pdf = str(getattr(data, "buyer_type", "OOO") or "OOO").upper() == "IP"
    _buyer_form_pdf = "именуемый" if _is_ip_pdf else "именуемое"
    if _is_ip_pdf:
        _person_pdf = "действующего на основании Свидетельства о государственной регистрации в качестве ИП"
    else:
        _person_pdf = (f"в лице {data.buyer_director_title.lower()} {data.buyer_director_gen or '___'}, "
                       f"действующего на основании {data.buyer_basis}")
    preamble = (
        f"{supplier_full}, именуемое в дальнейшем «Поставщик», в лице "
        f"{supplier_director_title.lower()} {supplier_director_gen}, "
        f"действующего на основании {supplier_basis}, с одной стороны, "
        f"и {data.buyer_full or 'ООО «___»'}, {_buyer_form_pdf} в дальнейшем «Покупатель», "
        f"{_person_pdf}, с другой стороны, "
        f"совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:"
    )
    story.append(Paragraph(preamble, styles["p"]))

    total = data.total_amount
    vat_text = "с НДС 22 %" if data.has_vat else "без НДС"

    story.append(Paragraph("1. ПРЕДМЕТ ДОГОВОРА", styles["h"]))
    story.append(Paragraph(
        "1.1. Поставщик обязуется передать в собственность Покупателю, "
        "а Покупатель обязуется принять и оплатить оборудование (далее — Товар), "
        "номенклатура, количество и цена которого указаны в Спецификации, "
        "являющейся неотъемлемой частью настоящего Договора.", styles["p"]))
    if data.drawing_bytes:
        story.append(Paragraph(
            "1.2. Габаритный чертёж поставляемого Товара приведён в Приложении № 1 "
            "к настоящему Договору.", styles["p"]))

    story.append(Paragraph("2. ЦЕНА ДОГОВОРА И ПОРЯДОК РАСЧЁТОВ", styles["h"]))
    story.append(Paragraph(
        f"2.1. Общая цена настоящего Договора составляет "
        f"<b>{_fmt_money(total)} руб.</b> ({money_to_words(total)}), {vat_text}.",
        styles["p"]))

    prepay_sum = total * data.prepay_pct / 100
    remainder_sum = total - prepay_sum
    remainder_pct = round(100 - data.prepay_pct)

    if data.prepay_pct == 100:
        story.append(Paragraph(
            f"2.2. Оплата производится в размере 100 % — "
            f"<b>{_fmt_money(total)} руб.</b> ({money_to_words(total)}) — "
            f"единовременно в течение 5 (пяти) рабочих дней с момента "
            f"подписания настоящего Договора и Спецификации.", styles["p"]))
    else:
        story.append(Paragraph("2.2. Оплата производится в следующем порядке:", styles["p"]))
        story.append(Paragraph(
            f"— Предоплата <b>{data.prepay_pct} %</b> — "
            f"<b>{_fmt_money(prepay_sum)} руб.</b> ({money_to_words(prepay_sum)}) — "
            f"в течение 5 (пяти) рабочих дней после подписания Договора и Спецификации;",
            styles["p"]))
        story.append(Paragraph(
            f"— Окончательный расчёт <b>{remainder_pct} %</b> — "
            f"<b>{_fmt_money(remainder_sum)} руб.</b> ({money_to_words(remainder_sum)}) — "
            f"в течение 5 (пяти) рабочих дней после уведомления о готовности к отгрузке.",
            styles["p"]))
    story.append(Paragraph("2.3. Все расчёты производятся в безналичной форме "
                          "на расчётный счёт Поставщика.", styles["p"]))

    story.append(Paragraph("3. СРОК И УСЛОВИЯ ПОСТАВКИ", styles["h"]))
    story.append(Paragraph(
        f"3.1. Срок изготовления Товара составляет до "
        f"<b>{data.shipment_term_days}</b> ({int_to_words_days(data.shipment_term_days)}) "
        f"рабочих дней с момента поступления предоплаты на расчётный счёт Поставщика.",
        styles["p"]))
    if data.delivery_type == "Самовывоз":
        story.append(Paragraph(
            f"3.2. Условия поставки: <b>Самовывоз</b> со склада Поставщика "
            f"по адресу: {data.delivery_address or 'г. Санкт-Петербург, склад Поставщика'}.",
            styles["p"]))
        story.append(Paragraph(
            f"3.3. Покупатель обязан вывезти Товар со склада в течение "
            f"{data.pickup_days} ({int_to_words_days(data.pickup_days)}) рабочих дней "
            f"с даты уведомления о готовности. При нарушении — пени "
            f"{data.shipment_penalty_pct:.1f} % от стоимости неотгруженного Товара "
            f"за каждый день просрочки.", styles["p"]))
        story.append(Paragraph(
            f"3.4. Хранение свыше {data.storage_free_days} "
            f"({int_to_words_days(data.storage_free_days)}) рабочих дней "
            f"оплачивается дополнительно {data.storage_penalty_pct:.1f} % от стоимости "
            f"Товара за каждый день хранения.", styles["p"]))
    else:
        story.append(Paragraph(
            f"3.2. Условия поставки: <b>Доставка силами Поставщика</b> по адресу: "
            f"{data.delivery_address or '___'}. Стоимость доставки — "
            f"{_fmt_money(data.delivery_cost)} руб., включена в цену Договора.",
            styles["p"]))

    story.append(Paragraph("4. КАЧЕСТВО И ГАРАНТИЯ", styles["h"]))
    story.append(Paragraph(
        f"4.1. Поставщик гарантирует качество Товара в течение "
        f"<b>{data.warranty_months}</b> ({int_to_words_months(data.warranty_months)}) "
        f"месяцев с даты передачи Товара Покупателю.", styles["p"]))
    story.append(Paragraph(
        "4.2. Гарантия не распространяется на дефекты, возникшие вследствие нарушения "
        "Покупателем правил эксплуатации, транспортировки или хранения.", styles["p"]))

    story.append(Paragraph("5. ОТВЕТСТВЕННОСТЬ СТОРОН", styles["h"]))
    story.append(Paragraph(
        "5.1. За нарушение обязательств по настоящему Договору Стороны несут "
        "ответственность в соответствии с действующим законодательством РФ.", styles["p"]))
    story.append(Paragraph(
        "5.2. Стороны освобождаются от ответственности за неисполнение обязательств "
        "вследствие обстоятельств непреодолимой силы (форс-мажор).", styles["p"]))

    story.append(Paragraph("6. ПРОЧИЕ УСЛОВИЯ", styles["h"]))
    story.append(Paragraph(
        f"6.1. Настоящий Договор действителен для акцепта в течение "
        f"<b>{data.contract_valid_days}</b> ({int_to_words_days(data.contract_valid_days)}) "
        f"рабочих дней с даты его составления.", styles["p"]))
    story.append(Paragraph(
        "6.2. Все споры и разногласия решаются путём переговоров. При невозможности "
        "урегулирования — в Арбитражном суде по месту нахождения Ответчика.", styles["p"]))
    story.append(Paragraph(
        "6.3. Договор составлен в двух экземплярах, имеющих равную юридическую силу, "
        "по одному для каждой из Сторон.", styles["p"]))

    if data.comment:
        story.append(Paragraph("7. ДОПОЛНИТЕЛЬНЫЕ УСЛОВИЯ", styles["h"]))
        story.append(Paragraph(data.comment, styles["p"]))

    # Реквизиты
    story.append(Spacer(1, 10))
    story.append(Paragraph("РЕКВИЗИТЫ И ПОДПИСИ СТОРОН", styles["h"]))

    def _party_cell(is_supplier: bool):
        if is_supplier:
            block = [
                Paragraph("Поставщик:", styles["party_bold"]),
                Paragraph(supplier_full, styles["party_bold"]),
                Paragraph(f"ИНН {_sup_get(supplier, "inn")} / КПП {_sup_get(supplier, "kpp")}", styles["party"]),
                Paragraph(f"ОГРН {_sup_get(supplier, "ogrn")}", styles["party"]),
                Paragraph(f"Адрес: {_sup_get(supplier, "address")}", styles["party"]),
                Paragraph(f"р/с {_sup_get(supplier, "rs", "bank_account")}", styles["party"]),
                Paragraph(f"в {_sup_get(supplier, "bank", "bank_name")}", styles["party"]),
                Paragraph(f"БИК {_sup_get(supplier, "bik")}", styles["party"]),
                Paragraph(f"к/с {_sup_get(supplier, "ks", "corr_account")}", styles["party"]),
                Paragraph(f"Тел.: {_sup_get(supplier, "phone", "phone_short")}", styles["party"]),
                Paragraph(f"E-mail: {_sup_get(supplier, "email")}", styles["party"]),
                Spacer(1, 6),
                Paragraph(supplier_director_title, styles["party"]),
                Paragraph(f"{supplier_director_short} ___________________", styles["party"]),
                Paragraph("М.П.", styles["party"]),
            ]
        else:
            block = [
                Paragraph("Покупатель:", styles["party_bold"]),
                Paragraph(data.buyer_full or "ООО «___»", styles["party_bold"]),
                Paragraph(f"ИНН {data.buyer_inn} / КПП {data.buyer_kpp}", styles["party"]),
                Paragraph(f"ОГРН {data.buyer_ogrn}", styles["party"]),
                Paragraph(f"Адрес: {data.buyer_address}", styles["party"]),
                Paragraph(f"р/с {data.buyer_account}", styles["party"]),
                Paragraph(f"в {data.buyer_bank}", styles["party"]),
                Paragraph(f"БИК {data.buyer_bik}", styles["party"]),
                Paragraph(f"к/с {data.buyer_corr_account}", styles["party"]),
                Paragraph(f"Тел.: {data.buyer_phone}", styles["party"]),
                Paragraph(f"E-mail: {data.buyer_email}", styles["party"]),
                Spacer(1, 6),
                Paragraph(data.buyer_director_title, styles["party"]),
                Paragraph(f"{data.buyer_director_short or '___'} ___________________", styles["party"]),
                Paragraph("М.П.", styles["party"]),
            ]
        return block

    parties_tbl = Table([[_party_cell(True), _party_cell(False)]],
                        colWidths=[87*mm, 87*mm])
    parties_tbl.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("LINEBETWEEN", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(parties_tbl)

    # ---- СПЕЦИФИКАЦИЯ ----
    story.append(PageBreak())
    story.append(Paragraph(f"СПЕЦИФИКАЦИЯ № 1", styles["title"]))
    story.append(Paragraph(
        f"к Договору поставки № {data.contract_number} от {data.contract_date}",
        styles["subtitle"]))
    story.append(Spacer(1, 6))

    # Таблица позиций
    spec_data = [[
        Paragraph("№", styles["spec_hdr"]),
        Paragraph("Код", styles["spec_hdr"]),
        Paragraph("Наименование", styles["spec_hdr"]),
        Paragraph("Ед.", styles["spec_hdr"]),
        Paragraph("Кол-во", styles["spec_hdr"]),
        Paragraph("Цена, ₽", styles["spec_hdr"]),
        Paragraph("Сумма, ₽", styles["spec_hdr"]),
    ]]
    for idx, item in enumerate(data.items, 1):
        qty = float(item.get("qty", 0))
        price = float(item.get("price", 0))
        spec_data.append([
            Paragraph(str(idx), styles["spec_cell"]),
            Paragraph(str(item.get("code", "")), styles["spec_cell"]),
            Paragraph(str(item.get("name", "")), styles["spec_cell"]),
            Paragraph(str(item.get("unit", "шт")), styles["spec_cell"]),
            Paragraph(f"{qty:g}", styles["spec_cell"]),
            Paragraph(_fmt_money(price), styles["spec_cell"]),
            Paragraph(_fmt_money(qty * price), styles["spec_cell"]),
        ])
    if data.delivery_cost > 0:
        spec_data.append([
            Paragraph(str(len(data.items) + 1), styles["spec_cell"]),
            Paragraph("—", styles["spec_cell"]),
            Paragraph(f"Доставка ({data.delivery_type})", styles["spec_cell"]),
            Paragraph("усл.", styles["spec_cell"]),
            Paragraph("1", styles["spec_cell"]),
            Paragraph(_fmt_money(data.delivery_cost), styles["spec_cell"]),
            Paragraph(_fmt_money(data.delivery_cost), styles["spec_cell"]),
        ])
    spec_tbl = Table(spec_data, colWidths=[10*mm, 22*mm, 62*mm, 12*mm, 16*mm, 25*mm, 27*mm],
                     repeatRows=1)
    spec_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F97316")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(spec_tbl)

    story.append(Spacer(1, 8))
    story.append(Paragraph(
        f"<para align='right'><b>ИТОГО: {_fmt_money(total)} руб., {vat_text}</b></para>",
        ParagraphStyle("total", fontName=bold, fontSize=11, alignment=2)))
    story.append(Paragraph(
        f"<para align='right'><i>Прописью: {money_to_words(total)}</i></para>",
        ParagraphStyle("totalw", fontName=reg, fontSize=9.5, alignment=2)))

    story.append(Spacer(1, 16))
    story.append(Table([[
        [Paragraph("Поставщик:", styles["party_bold"]),
         Paragraph(supplier_director_title, styles["party"]),
         Paragraph(f"{supplier_director_short} _________________", styles["party"]),
         Paragraph("М.П.", styles["party"])],
        [Paragraph("Покупатель:", styles["party_bold"]),
         Paragraph(data.buyer_director_title, styles["party"]),
         Paragraph(f"{data.buyer_director_short or '___'} _________________", styles["party"]),
         Paragraph("М.П.", styles["party"])],
    ]], colWidths=[87*mm, 87*mm]))

    # ---- ПРИЛОЖЕНИЕ № 1 ----
    if data.drawing_bytes:
        story.append(PageBreak())
        story.append(Paragraph("ПРИЛОЖЕНИЕ № 1", styles["title"]))
        story.append(Paragraph(
            f"к Договору поставки № {data.contract_number} от {data.contract_date}",
            styles["subtitle"]))
        story.append(Paragraph(data.drawing_caption,
                              ParagraphStyle("cap", fontName=bold, fontSize=11,
                                             alignment=1, spaceAfter=8)))
        img = Image(io.BytesIO(data.drawing_bytes), width=170*mm, height=200*mm,
                    kind="proportional")
        img.hAlign = "CENTER"
        story.append(img)

    doc.build(story)
    return buf.getvalue()
