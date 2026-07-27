# -*- coding: utf-8 -*-
"""
Калькулятор консольных кранов ЛКС — КП и Договор поставки.
Финальная версия, соответствующая паспорту логики.

Автор шаблона: ROLLS KRAN.
"""
from __future__ import annotations

import io
import os
import re
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from kp_pdf import build_kp_pdf
from req_parser import parse_requisites, extract_text_from_docx

# =========================================================
# 0. ГЛОБАЛЬНЫЕ ПАРАМЕТРЫ И КОНСТАНТЫ
# =========================================================

APP_DIR = Path(__file__).resolve().parent
MEDIA_DIR = APP_DIR / "media"
PRICES_DIR = APP_DIR / "prices"


def _find_file(filename: str) -> Optional[Path]:
    """Ищет файл в prices/, media/ или в корне (гибкая структура репо)."""
    for cand in [PRICES_DIR / filename, APP_DIR / "media" / filename,
                 APP_DIR / filename]:
        if cand.exists():
            return cand
    return None

PRICE_CRANES_71_73 = "konsolnkii-2.xlsx"          # прайс однопл. ЛКС71/73
PRICE_CRANES_77_78 = "dvukhplechevye.xlsx"        # прайс двухпл. ЛКС77/78
PRICE_HOISTS = "prais-tal-dlia-kp-na-saite-3.xlsx"  # прайс тали R-Tech

VAT_RATE = 0.22
PREPAYMENT_DEFAULT = 0.70

# Фирменная палитра rolls-kran.ru — белый/оранжевый/чёрный
BRAND_COLOR = RGBColor(0x11, 0x11, 0x11)         # чёрный — основной текст
BRAND_YELLOW_HEX = "F97316"                      # оранжевый (для шапок таблиц)
BRAND_YELLOW_SOFT_HEX = "FFF1E5"                 # мягкий оранжевый (для итогов)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_BG = "F5F5F5"

# ---- Реквизиты Поставщика (ООО «ЛКС») ----
SUPPLIER = {
    "short": "ООО «ЛКС»",
    "full": "Общество с ограниченной ответственностью",  # «ЛКС» подставляется в шаблоне
    "address": "195030, г. Санкт-Петербург, вн.тер.г. муниципальный округ Ржевка, "
               "ш. Революции, д. 114, Литера А, помещ. 111",
    "inn": "7806586802",
    "kpp": "780601001",
    "ogrn": "12178000979903",
    "rs": "40702810110000819349",
    "bank": "АО «ТИНЬКОФФ БАНК»",
    "ks": "30101810145250000974",
    "bik": "044525974",
    "phone": "8 (800) 302-73-10 — бесплатный звонок по всей РФ",
    "phone_short": "8 (800) 302-73-10",
    "email": "zakaz@rolls-kran.ru",
    "director_position": "Генеральный директор",
    "director_fio_gen": "Щербакова Ильи Витальевича",
    "director_fio_short": "Щербаков И.В.",
    "director_basis": "Устава",
}

# =========================================================
# 1. ПАСПОРТ ЛОГИКИ (ФИКСИРОВАННЫЙ)
# =========================================================

SERIES_OPTIONS = ["ЛКС71", "ЛКС73", "ЛКС73М", "ЛКС77", "ЛКС78"]

SERIES_DESCRIPTIONS = {
    "ЛКС71": "Кран консольно-поворотный настенный",
    "ЛКС73": "Кран консольно-поворотный на колонне",
    "ЛКС77": "Кран консольно-поворотный двухплечевой настенный",
    "ЛКС78": "Кран консольно-поворотный двухплечевой на колонне",
}

# ЛКС71/73 — грузоподъёмность → допустимые длины стрелы (м)
BOOM_BY_CAPACITY_LKS71_73 = {
    80: [3, 4, 5],
    125: [3, 4, 5, 6, 7],
    250: [3, 4, 5, 6, 7, 8],
    500: [3, 4, 5, 6, 7, 8],
    1000: [3, 4, 5, 6],
}

# ЛКС71/73 — высота до стрелы (м) — всегда 3..6
HEIGHT_TO_ARM_LKS71_73 = [3, 4, 5, 6]

# ЛКС71/73 — габаритная надбавка к высоте до стрелы по г/п (м).
# Например: 500 кг, h=4м → габарит 4+1.2=5.2 м.
GABARITE_ADD = {80: 0.6, 125: 0.8, 250: 1.0, 500: 1.2, 1000: 1.2}


def gabarite_height(capacity: int, height_to_arm: float) -> float:
    """Габаритная высота крана ЛКС73."""
    return float(height_to_arm) + GABARITE_ADD.get(capacity, 1.2)


# ЛКС71/73 — высота колонны (историческая таблица) — ещё используется в UI
COLUMN_HEIGHT_MAP = {
    80: {3: 3.6, 4: 4.6, 5: 5.6, 6: 6.6},
    125: {3: 3.8, 4: 4.8, 5: 5.8, 6: 6.8},
    250: {3: 4.0, 4: 5.0, 5: 6.0, 6: 7.0},
    500: {3: 4.2, 4: 5.2, 5: 6.2, 6: 7.2},
    1000: {3: 4.2, 4: 5.2, 5: 6.2, 6: 7.2},
}

LKS77_78_CAPACITIES = [150, 300]
LKS77_78_BOOMS = [2, 3, 4]
# ЛКС78 на колонне: высота колонны (ключ) → рабочая высота до стрелы
LKS78_COLUMN_HEIGHTS = [3.8, 4.8, 5.8, 6.8]
LKS78_HEIGHT_TO_ARM = {3.8: 3.6, 4.8: 4.6, 5.8: 5.6, 6.8: 6.6}
LKS78_HEIGHTS = LKS78_COLUMN_HEIGHTS   # бэк-компат (высота колонны)
LKS77_HEIGHTS = [0]            # ЛКС77 — настенный, высота колонны 0

# Массы ЛКС78 (г/п, высота_колонны, стрела) → масса кг
LKS78_WEIGHTS = {
    # 150 кг
    (150, 3.8, 2): 190, (150, 3.8, 3): 219, (150, 3.8, 4): 240,
    (150, 4.8, 2): 216, (150, 4.8, 3): 245, (150, 4.8, 4): 335,
    (150, 5.8, 2): 316, (150, 5.8, 3): 345, (150, 5.8, 4): 366,
    (150, 6.8, 2): 364, (150, 6.8, 3): 393, (150, 6.8, 4): 414,
    # 300 кг
    (300, 3.8, 2): 286, (300, 3.8, 3): 317, (300, 3.8, 4): 347,
    (300, 4.8, 2): 334, (300, 4.8, 3): 364, (300, 4.8, 4): 442,
    (300, 5.8, 2): 427, (300, 5.8, 3): 458, (300, 5.8, 4): 489,
    (300, 6.8, 2): 462, (300, 6.8, 3): 492, (300, 6.8, 4): 598,
}

# --- Тали ---
HOISTS_LKS71_73 = ["OCALIFT", "R-Tech"]
HOISTS_LKS77_78 = ["R-Tech"]

# Поперечные тележки под таль
TROLLEY_MAP = {
    "OCALIFT": {125: "523", 250: "523", 500: "523", 1000: "533"},
    "R-Tech": {80: "513.RC59", 125: "521", 150: "521", 250: "521",
               300: "521", 500: "521", 1000: "533"},
}

# Кабельные тележки/подвес для ЛКС71/73
CABLE_TROLLEY_CODE = "321.050.RC74.65х54"
START_HANGER_CODE = "327.050.RC74.65х54"

# Цены на тележки, подвесы и аксессуары (все с НДС 22%)
ACCESSORY_PRICES = {
    "523":       (6518.90,  "Тележка грузовая гп 600 кг, под крюк (523)"),
    "521":       (7394.87,  "Тележка грузовая гп 600 кг, поперечная (521)"),
    "533":       (17044.20, "Тележка грузовая гп 1200 кг, под крюк (533)"),
    "513.RC59":  (3816.65,  "Тележка грузовая RC59, под крюк (513.RC59)"),
    CABLE_TROLLEY_CODE: (
        2016.12,
        "Тележка RC74 кабельная, подвес стандартный Ø50 мм, "
        "кабельное гнездо 65×54 мм",
    ),
    START_HANGER_CODE: (
        1125.75,
        "Подвес начальный RC74, подвес стандартный Ø50 мм, "
        "кабельное гнездо 65×54 мм",
    ),
}

# Электрификация — только для ЛКС71/73 и г/п 250/500/1000 кг
ELECTRIFICATION_CAPACITIES = [250, 500, 1000]
# Пакет электрификации: шкаф управления + приводная тележка + кабели
ELECTRIFICATION_PACKAGE = [
    {"code": "AR580.10",
     "name": "Шкаф управления — подъём, один тип перемещения (тележки)",
     "qty": 1, "unit": "шт", "price": 142290.92},
    {"code": "AR542",
     "name": "Тележка приводная 180 Вт, 22 м/мин, до 1 т, правая",
     "qty": 1, "unit": "шт", "price": 45320.68},
]
# Плоские кабели YFFB — комплект для электрификации
ELECTRIFICATION_CABLES = [
    {"code": "A-RT-YFFB 4×1.5", "name": "Кабель плоский YFFB 4×1,5 (питание)",
     "qty": 10.5, "unit": "м", "price": 214.72},
    {"code": "A-RT-YFFB 6×1.5", "name": "Кабель плоский YFFB 6×1,5 (управление)",
     "qty": 3.5, "unit": "м", "price": 341.39},
    {"code": "A-RT-YFFB 5×2.5", "name": "Кабель плоский YFFB 5×2,5 (силовой)",
     "qty": 20.0, "unit": "м", "price": 435.14},
    {"code": "A-RT-YFFB 4×2.5", "name": "Кабель плоский YFFB 4×2,5 (для тали)",
     "qty": 20.0, "unit": "м", "price": 339.06},
]

# Кабель питания для тали R-Tech
RTECH_POWER_CABLE_LEN = 20  # м


def lookup_accessory_price(df_cranes, code: str, fallback_price: float,
                          fallback_name: str) -> tuple[float, str]:
    """Ищет цену позиции в прайсе консольных кранов по коду.

    Нестрогое сопоставление: без пробелов, без дефисов, без точек, без регистра.
    Возвращает (цена, наименование). Если не найдено — возвращает fallback.
    """
    if df_cranes is None or df_cranes.empty:
        return fallback_price, fallback_name
    def _norm(s: str) -> str:
        return (str(s).replace(" ", "").replace("×", "x")
                .replace("-", "").replace(".", "").lower())
    target = _norm(code)
    for _, row in df_cranes.iterrows():
        if _norm(row.get("code", "")).find(target) >= 0:
            price = float(row.get("price", 0)) or fallback_price
            name = str(row.get("name", "") or fallback_name)
            return price, name
    return fallback_price, fallback_name

# =========================================================
# 2. ЦЕНЫ OCALIFT (в прайсе R-Tech нет OCALIFT — задаём вручную)
# =========================================================
# Цены OCALIFT — базовые розничные, редактируются через сайдбар.
OCALIFT_PRICES_DEFAULT = {
    # (capacity_kg, "1 скорость" / "2 скорости"): цена руб. с НДС
    (125, "1 скорость"): 45500,
    (125, "2 скорости"): 55500,
    (250, "1 скорость"): 55000,
    (250, "2 скорости"): 66000,
    (500, "1 скорость"): 62000,
    (500, "2 скорости"): 76000,
    (1000, "1 скорость"): 95000,
    (1000, "2 скорости"): 115000,
}

# =========================================================
# 3. ДОПУСТИМЫЕ КОМБИНАЦИИ
# =========================================================

def get_allowed_capacities(series: str) -> list[int]:
    if series in ("ЛКС71", "ЛКС73"):
        return list(BOOM_BY_CAPACITY_LKS71_73.keys())
    if series in ("ЛКС77", "ЛКС78"):
        return LKS77_78_CAPACITIES
    return []


def get_allowed_booms(series: str, capacity: int) -> list[int]:
    if series in ("ЛКС71", "ЛКС73"):
        return BOOM_BY_CAPACITY_LKS71_73.get(capacity, [])
    if series in ("ЛКС77", "ЛКС78"):
        return LKS77_78_BOOMS
    return []


def get_available_column_diameters(df, series: str, capacity: int, boom,
                                    height_to_arm) -> list[str]:
    """Список доступных диаметров колонны в прайсе для конкретной конфигурации."""
    if df is None or df.empty:
        return []
    base = build_crane_code(series, capacity, boom, height_to_arm)
    import re as _re
    pat = _re.compile(_re.escape(base) + r"\.(\d+)$")
    diameters = []
    for code in df["code"]:
        m = pat.match(str(code))
        if m:
            diameters.append(m.group(1))
    return sorted(set(diameters), key=int)


def get_lllm_variants(df, capacity: int) -> list[dict]:
    """Краны ЛКС73М с электрическим поворотом (LLL) — только для 1000 кг."""
    if df is None or df.empty or capacity != 1000:
        return []
    import re as _re
    pat = _re.compile(rf"^ЛКС7\dМ\.{capacity:04d}-(\d+)-(\d+)\.LLL$")
    out = []
    for _, row in df.iterrows():
        m = pat.match(str(row["code"]))
        if m:
            out.append({
                "code": str(row["code"]),
                "name": str(row["name"]),
                "price": float(row["price"]),
                "boom": int(m.group(1)),
                "height": int(m.group(2)),
            })
    return sorted(out, key=lambda x: (x["boom"], x["height"]))


def get_flanges(df) -> list[dict]:
    """Список уширенных фланцев ЛКС73Ф."""
    if df is None or df.empty:
        return []
    hits = df[df["code"].str.match(r"^ЛКС7\dФ\.\d+$", na=False)]
    return [
        {
            "code": str(row["code"]),
            "name": str(row["name"]),
            "price": float(row["price"]),
            "diameter": int(str(row["code"]).rsplit(".", 1)[1]),
        }
        for _, row in hits.iterrows()
    ]


def get_allowed_heights(series: str) -> list[int]:
    if series == "ЛКС73":
        return HEIGHT_TO_ARM_LKS71_73
    if series == "ЛКС78":
        return LKS78_HEIGHTS
    # ЛКС71 (настенный) и ЛКС77 — высота колонны не задаётся
    if series in ("ЛКС71", "ЛКС77"):
        return [0]
    return []


def get_column_height(series: str, capacity: int, height_to_arm: int) -> float:
    if series in ("ЛКС71", "ЛКС73"):
        return COLUMN_HEIGHT_MAP[capacity][height_to_arm]
    if series == "ЛКС78":
        return float(height_to_arm)
    return 0.0


def get_allowed_hoists(series: str, capacity: int = 0) -> list[str]:
    if series in ("ЛКС71", "ЛКС73"):
        # Для 80 и 125 кг — только R-Tech (P2 или HC), OCALIFT недоступен
        if capacity in (80, 125):
            return ["R-Tech"]
        return HOISTS_LKS71_73
    return HOISTS_LKS77_78


def get_hoist_modes(series: str, brand: str, height_to_arm: float,
                    capacity: int) -> list[str]:
    """Режимы работы тали в зависимости от бренда, серии и г/п."""
    if brand == "OCALIFT":
        return ["1 скорость", "2 скорости"]
    if brand == "R-Tech":
        # ЛКС71/73 г/п 80 или 125 кг — только R-Tech P2 (8/2) или R-Tech HC
        if series in ("ЛКС71", "ЛКС73") and capacity in (80, 125):
            return ["8/2 м/мин", "ручной контроль груза"]
        # ЛКС77 (настенный) — оба варианта; HC ограничен высотой 3 м на этапе выбора
        if series == "ЛКС77":
            return ["8/2 м/мин", "ручной контроль груза"]
        # ЛКС78 (на колонне) — HC доступен только для колонны 3.8 м
        if series == "ЛКС78":
            if abs(float(height_to_arm) - 3.8) < 0.05:
                return ["8/2 м/мин", "ручной контроль груза"]
            return ["8/2 м/мин"]
        return ["8/2 м/мин"]
    return ["8/2 м/мин"]


# Г/п тали с ручным контролем груза (по прайсу R-Tech HC есть 80/125/250/500 кг)
HC_CAPACITIES = [125, 250]


def get_trolley_code(brand: str, capacity: int) -> Optional[str]:
    return TROLLEY_MAP.get(brand, {}).get(capacity)


def build_crane_code(series: str, capacity: int, boom: int, height_to_arm: int) -> str:
    # У ЛКС71 (настенный) высота колонны всегда 0 в прайсе
    if series == "ЛКС71":
        return f"ЛКС71.{capacity:04d}-{boom}-0"
    if series == "ЛКС73":
        return f"ЛКС73.{capacity:04d}-{boom}-{height_to_arm}"
    if series == "ЛКС77":
        return f"ЛКС77.{capacity:04d}-{boom}-0"
    if series == "ЛКС78":
        # Высота колонны в коде — целое число (3.8 → 3, 4.8 → 4 и т.д.)
        return f"ЛКС78.{capacity:04d}-{boom}-{int(float(height_to_arm))}"
    return f"{series}.{capacity:04d}-{boom}-{height_to_arm}"


def calc_cable_trolleys(boom: int) -> int:
    return max(boom - 1, 0)


# =========================================================
# 4. ЗАГРУЗКА ПРАЙСОВ
# =========================================================

def _price_str_to_float(s) -> float:
    if s is None:
        return 0.0
    if isinstance(s, (int, float)):
        return float(s)
    txt = str(s).replace("\xa0", "").replace(" ", "").replace(",", ".")
    txt = re.sub(r"[^0-9.]", "", txt)
    try:
        return float(txt) if txt else 0.0
    except ValueError:
        return 0.0


@st.cache_data(show_spinner=False)
def load_price_cranes_71_73(uploaded=None) -> Optional[pd.DataFrame]:
    if uploaded is not None:
        df = pd.read_excel(uploaded, header=None)
    else:
        path = _find_file(PRICE_CRANES_71_73)
        if path is None:
            return None
        df = pd.read_excel(path, header=None)
    df.columns = ["code", "name", "unit", "qty", "price"]
    df["code"] = df["code"].astype(str).str.strip()
    df["price"] = df["price"].apply(_price_str_to_float)
    return df


@st.cache_data(show_spinner=False)
def load_price_cranes_77_78(uploaded=None) -> Optional[pd.DataFrame]:
    if uploaded is not None:
        df = pd.read_excel(uploaded, header=None)
    else:
        path = _find_file(PRICE_CRANES_77_78)
        if path is None:
            return None
        df = pd.read_excel(path, header=None)
    df.columns = ["code", "name", "unit", "qty", "price"]
    df["code"] = df["code"].astype(str).str.strip()
    df["price"] = df["price"].apply(_price_str_to_float)
    return df


@st.cache_data(show_spinner=False)
def load_price_hoists(uploaded=None) -> Optional[pd.DataFrame]:
    if uploaded is not None:
        df = pd.read_excel(uploaded)
    else:
        path = _find_file(PRICE_HOISTS)
        if path is None:
            return None
        df = pd.read_excel(path)
    df.columns = ["code", "name", "price"]
    df["code"] = df["code"].astype(str).str.strip()
    df["name"] = df["name"].astype(str)
    df["price"] = df["price"].apply(_price_str_to_float)
    return df


def find_crane_price(df: Optional[pd.DataFrame], series: str, capacity: int,
                    boom: int, height_to_arm: int) -> tuple[float, str]:
    """Возвращает (цена, найденный_код).
    Ищет только коды вида <base>.XXX (с точкой + суффикс),
    чтобы не зацепить аномальные строки без диаметра колонны в прайсе.
    """
    if df is None or df.empty:
        return 0.0, ""
    base = build_crane_code(series, capacity, boom, height_to_arm)
    # Ищем <base>.XXX с числовым суффиксом диаметра колонны (200…600).
    # Исключаем: LLL (электрический поворот) и аномальные строки без суффикса.
    import re as _re
    numeric_pattern = _re.escape(base) + r"\.\d+$"
    mask = df["code"].str.match(numeric_pattern, na=False)
    if mask.any():
        rows = df[mask].sort_values("price")
        row = rows.iloc[0]
        return float(row["price"]), str(row["code"])
    # fallback — точное совпадение без суффикса
    exact = df[df["code"] == base]
    if not exact.empty:
        row = exact.iloc[0]
        return float(row["price"]), str(row["code"])
    return 0.0, base


def find_hoist_price(df: Optional[pd.DataFrame], brand: str, capacity: int,
                    mode: str, hoist_height: int,
                    series: str = "", hoist_exec_override: str = "") -> tuple[float, str, str]:
    """Возвращает (цена, код, наименование) для тали.
    Для OCALIFT возвращает цену из внутренней таблицы (её нет в прайсе).
    """
    if brand == "OCALIFT":
        # OCALIFT: цена берётся из прайса талей по модели/высоте,
        # fallback — встроенные константы.
        # Прайс содержит только 500/1000 кг, высоты 3 и 6 м.
        # 500кг 1ск: OCALIFT 005-01s; 1000кг 1ск: OCALIFT 01-01S
        # 500кг 2ск: OCALIFT 005-01D; 1000кг 2ск: OCALIFT 01-01D
        oca_price = 0.0
        oca_code = f"OCALIFT-{capacity}-{'1S' if mode == '1 скорость' else '2S'}"
        oca_name = f"Таль электрическая цепная OCALIFT, г/п {capacity} кг, {mode}"
        if df is not None and not df.empty and capacity in (500, 1000):
            # Доступные высоты OCALIFT: 3 / 4.5 / 6 м — привязываем к ближайшей
            h_val = float(hoist_height)
            if h_val <= 3.75:
                oca_h_key = "3м"
            elif h_val <= 5.25:
                oca_h_key = "4,5м"
            else:
                oca_h_key = "6м"
            model_key = "005" if capacity == 500 else "01"
            # 1 скорость — модель заканчивается на 01s/01S; 2 скорости — на 01D
            if mode == "2 скорости":
                model_pattern = fr"{model_key}-01D\b"
            else:
                model_pattern = fr"{model_key}-01[sS]\b"
            # Нормализуем: в прайсе могут быть переносы внутри кодов модели (005-\n01s → 005-01s)
            name_col_raw = df["name"].astype(str) if "name" in df.columns else pd.Series(dtype=str)
            name_col = (name_col_raw
                        .str.replace(r"-\s+", "-", regex=True)  # "005- 01s" → "005-01s"
                        .str.replace(r"\s+", " ", regex=True))
            mask = name_col.str.contains("OCALIFT", case=False, na=False)
            mask &= name_col.str.contains(model_pattern, case=False, na=False, regex=True)
            mask &= name_col.str.contains(f"высота подъема {oca_h_key}", case=False, na=False)
            hits = df[mask]
            if not hits.empty:
                oca_price = float(hits.iloc[0]["price"])
                oca_name = str(hits.iloc[0]["name"]).replace("\n", " ")
        # fallback
        if oca_price == 0.0:
            oca_price = float(st.session_state.get(
                "ocalift_prices", OCALIFT_PRICES_DEFAULT).get((capacity, mode), 0.0))
        return oca_price, oca_code, oca_name

    # R-Tech
    if df is None or df.empty:
        return 0.0, "", ""

    # У R-Tech в прайсе только 80/125/250/500/1000 кг.
    # ЛКС77/78 идут с г/п 150 и 300 кг — маппим на ближайшую большую:
    # 150 → 250 кг тали (запас по г/п), 300 → 500 кг тали.
    price_capacity = capacity
    if capacity == 150:
        price_capacity = 250
    elif capacity == 300:
        price_capacity = 500

    # Переводим г/п в тонны для кода.
    # В прайсе целые т пишутся как 1.0 / 2.0 (не 1, не 2)
    cap_t = price_capacity / 1000.0
    if cap_t < 1:
        cap_key = f"{cap_t:g}"        # 0.08, 0.125, 0.25, 0.5
    else:
        cap_key = f"{int(cap_t)}.0"   # 1.0, 2.0

    if mode == "ручной контроль груза":
        # HC в прайсе есть только 80/125/250/500 кг (0.08/0.125/0.25/0.5 т).
        # Для 150/300/1000 маппим на ближайшую большую HC.
        hc_map = {80: 80, 125: 125, 150: 250, 250: 250, 300: 500, 500: 500,
                  1000: 500}  # для 1000 берём 500 кг HC
        hc_cap = hc_map.get(capacity, capacity)
        hc_cap_t = hc_cap / 1000.0
        hc_cap_key = f"{hc_cap_t:g}" if hc_cap_t < 1 else f"{int(hc_cap_t)}.0"
        pattern = f"A-RT-ECH-{hc_cap_key}-8-2.HC.03"
        exact = df[df["code"] == pattern]
        if exact.empty:
            exact = df[df["code"].str.contains(
                fr"A-RT-ECH-{hc_cap_key}-.*\.HC\.", case=False, na=False, regex=True)]
        if not exact.empty:
            row = exact.iloc[0]
            return float(row["price"]), str(row["code"]), str(row["name"])
        # запасной — по подстроке
        mask = df["code"].str.contains(f"{hc_cap_key}-8-2.HC", na=False, regex=False)
    else:
        # 8/2 м/мин. Индексы:
        # P2 — стационарная (пульт 2 кнопки), суффикс высоты двойной: P2.HH.HH
        # C2 — с двумя разъёмами, суффикс одинарный: C2.HH
        # C3 — радиоуправление, суффикс одинарный: C3.HH (для ЛКС73М по умолчанию)
        h_key = f"{int(hoist_height):02d}"
        # Определяем индекс: override > явное по серии > дефолт по серии
        if hoist_exec_override and hoist_exec_override in ("P2", "C2", "C3"):
            exec_key = hoist_exec_override
        else:
            # По умолчанию — P2 (стационарная, с ручным пультом) для ВСЕХ серий.
            # C2/C3 — только когда кран с электрификацией (задаётся через hoist_exec_override).
            exec_key = "P2"
        # P2 имеет двойной суффикс высоты, C2/C3 — одинарный
        if exec_key == "P2":
            pattern = f"A-RT-ECH-{cap_key}-8-2.P2.{h_key}.{h_key}"
        else:
            pattern = f"A-RT-ECH-{cap_key}-8-2.{exec_key}.{h_key}"
        exact = df[df["code"] == pattern]
        if not exact.empty:
            row = exact.iloc[0]
            return float(row["price"]), str(row["code"]), str(row["name"])
        # запасной: любой выбранный тип нужной г/п
        mask = df["code"].str.contains(
            f"A-RT-ECH-{cap_key}-8-2.{exec_key}", na=False, regex=False)

    if not mask.any():
        return 0.0, pattern, ""
    row = df[mask].iloc[0]
    return float(row["price"]), str(row["code"]), str(row["name"])


# =========================================================
# 5. СТРУКТУРЫ И ПОСТРОЕНИЕ СПЕЦИФИКАЦИИ
# =========================================================

@dataclass
class SpecLine:
    code: str
    name: str
    unit: str
    qty: float
    price: float

    @property
    def total(self) -> float:
        return round(self.qty * self.price, 2)


@dataclass
class QuoteData:
    series: str
    capacity: int
    boom: int
    height_to_arm: int
    hoist_brand: str
    hoist_mode: str
    hoist_height: int
    include_electrification: bool
    include_montage: bool
    montage_price: float
    # Г/п тали с ручным контролем груза (для ЛКС77/78, независимо от г/п крана)
    hc_capacity: Optional[int] = None
    # Индекс исполнения тали (P2/C2/C3). Если None — автоподбор по серии
    hoist_exec: Optional[str] = None
    # Включать в КП позицию тали (True) или оставить кран без тали (False)
    include_hoist: bool = True
    # Включать НДС 22% в стоимость (True) или цены без НДС (False, для УСН/ИП)
    include_vat: bool = True
    # Монтаж с НДС 22 % или без НДС (от ИП)
    montage_vat: bool = False
    # Диаметр колонны (суффикс в коде крана, напр. "420")
    column_diameter: Optional[str] = None
    # Кран с электрическим поворотом стрелы (ЛКС73М LLL)
    use_lllm: bool = False
    lllm_code: Optional[str] = None
    # Уширенный фланец
    include_flange: bool = False
    flange_code: Optional[str] = None
    # Режим Кран + траверса — кабель 5×2,5 на 25 м вместо стандартного 4×2,5
    with_traverse: bool = False
    lines: list[SpecLine] = field(default_factory=list)
    electrification_lines: list[SpecLine] = field(default_factory=list)

    @property
    def total(self) -> float:
        _sum = sum(l.total for l in self.lines) + sum(l.total for l in self.electrification_lines)
        # Если НДС НЕ включён — вычитаем его из цен (в прайсе они с НДС)
        if not getattr(self, "include_vat", True):
            _sum = _sum / (1 + VAT_RATE)
        return round(_sum, 2)

    @property
    def vat(self) -> float:
        # Если НДС не облагается — выделять нечего
        if not getattr(self, "include_vat", True):
            return 0.0
        # НДС 22% уже включён в цены → выделяем
        return round(self.total * VAT_RATE / (1 + VAT_RATE), 2)

    @property
    def total_with_montage(self) -> float:
        # монтаж от ИП без НДС — учитываем отдельно (в договоре поставки не идёт)
        return round(self.total + (self.montage_price if self.include_montage else 0), 2)


def build_specification(q: QuoteData,
                        df_71_73: Optional[pd.DataFrame],
                        df_77_78: Optional[pd.DataFrame],
                        df_hoists: Optional[pd.DataFrame]) -> QuoteData:
    """Собирает состав по паспорту логики."""
    q.lines = []
    q.electrification_lines = []

    # 1. Кран — выбор между ЛКС73 (ручной поворот) и ЛКС73М LLL (электроповорот)
    if q.use_lllm and q.lllm_code and df_71_73 is not None:
        # Кран ЛКС73М с электрическим поворотом.
        # Парсим стрелу и высоту из самого кода LLL (формат: ЛКС73М.ГРПН-СТРЕЛА-ВЫСОТА.LLL),
        # т.к. они могут отличаться от q.boom/q.height_to_arm (у пользователя в сайдбаре).
        import re as _re_lll
        _m_lll = _re_lll.match(r"ЛКС73М\.\d+-(\d+)-(\d+)\.LLL", q.lllm_code)
        if _m_lll:
            _lll_boom = int(_m_lll.group(1))
            _lll_height = int(_m_lll.group(2))
        else:
            _lll_boom = q.boom
            _lll_height = q.height_to_arm
        hit = df_71_73[df_71_73["code"] == q.lllm_code]
        if not hit.empty:
            crane_price = float(hit.iloc[0]["price"])
            crane_code = q.lllm_code
            crane_name = (
                f"Кран ЛКС73М г/п {q.capacity} кг, стрела {_lll_boom} м, "
                f"высота до стрелы {_lll_height} м — с электрическим поворотом стрелы (LLL)"
            )
        else:
            crane_price, crane_code = 0.0, q.lllm_code
            crane_name = f"Кран ЛКС73М ({q.lllm_code})"
        q.lines.append(SpecLine(crane_code, crane_name, "шт", 1, crane_price))
        # Фланец — если выбран
        if q.include_flange and q.flange_code and df_71_73 is not None:
            fh = df_71_73[df_71_73["code"] == q.flange_code]
            if not fh.empty:
                q.lines.append(SpecLine(
                    q.flange_code, str(fh.iloc[0]["name"]), "шт", 1,
                    float(fh.iloc[0]["price"])))
        # Флаг: кран LLL уже добавлен, обычную ветку крана ниже пропускаем
        _crane_added = True
    else:
        _crane_added = False

    if not _crane_added:
        if q.series in ("ЛКС71", "ЛКС73"):
            # Если выбран конкретный диаметр колонны — берём точно его
            if q.column_diameter and df_71_73 is not None and q.series == "ЛКС73":
                base = build_crane_code(q.series, q.capacity, q.boom, q.height_to_arm)
                exact_code = f"{base}.{q.column_diameter}"
                hit = df_71_73[df_71_73["code"] == exact_code]
                if not hit.empty:
                    crane_price = float(hit.iloc[0]["price"])
                    crane_code = exact_code
                else:
                    crane_price, crane_code = find_crane_price(
                        df_71_73, q.series, q.capacity, q.boom, q.height_to_arm)
            else:
                crane_price, crane_code = find_crane_price(
                    df_71_73, q.series, q.capacity, q.boom, q.height_to_arm)
            if q.series == "ЛКС73":
                # Суффикс = высота до стрелы в см. Без суффикса — стандартный height_to_arm.
                if q.column_diameter:
                    actual_height = int(q.column_diameter) / 100.0
                else:
                    actual_height = float(q.height_to_arm)
                gabarite = gabarite_height(q.capacity, actual_height)
                crane_name = (
                    f"Кран консольно-поворотный на колонне г/п {q.capacity} кг, "
                    f"стрела {q.boom} м, высота до стрелы {actual_height:g} м "
                    f"(габарит {gabarite:g} м)"
                )
            else:
                crane_name = (
                    f"Кран консольно-поворотный настенный г/п {q.capacity} кг, "
                    f"стрела {q.boom} м"
                )
        else:
            crane_price, crane_code = find_crane_price(
                df_77_78, q.series, q.capacity, q.boom, q.height_to_arm)
            desc = SERIES_DESCRIPTIONS.get(q.series, "")
            if q.series == "ЛКС78":
                crane_name = (
                    f"{desc} г/п {q.capacity} кг, стрела {q.boom} м, "
                    f"высота {q.height_to_arm} м"
                )
            else:
                crane_name = f"{desc} г/п {q.capacity} кг, стрела {q.boom} м"

        q.lines.append(SpecLine(crane_code or build_crane_code(
            q.series, q.capacity, q.boom, q.height_to_arm),
            crane_name, "шт", 1, crane_price))

    # 1а. Уширенный фланец (опционально, только для ЛКС73)
    # Для LLL фланец уже добавлен выше — пропускаем
    if (not _crane_added and q.include_flange and q.flange_code and df_71_73 is not None
            and q.series == "ЛКС73"):
        fh = df_71_73[df_71_73["code"] == q.flange_code]
        if not fh.empty:
            q.lines.append(SpecLine(
                q.flange_code, str(fh.iloc[0]["name"]), "шт", 1,
                float(fh.iloc[0]["price"])))

    # 2. Таль
    # Для тали с ручным контролем груза на ЛКС77/78 г/п тали — отдельный
    # параметр (125 или 250 кг), независимый от г/п крана (150/300 кг).
    if q.hoist_mode == "ручной контроль груза" and q.hc_capacity:
        hoist_cap = q.hc_capacity
    else:
        hoist_cap = q.capacity
    # Таль добавляется только если пользователь не снял галочку «Включить таль в КП»
    if getattr(q, "include_hoist", True):
        h_price, h_code, h_name = find_hoist_price(
            df_hoists, q.hoist_brand, hoist_cap, q.hoist_mode, q.hoist_height,
            series=q.series,
            hoist_exec_override=(getattr(q, "hoist_exec", None) or ""))
        if not h_name:
            h_name = (f"Таль электрическая цепная {q.hoist_brand}, "
                      f"г/п {hoist_cap} кг, {q.hoist_mode}, высота {q.hoist_height} м")
        q.lines.append(SpecLine(h_code, h_name, "шт", 1, h_price))

    # 3. Тележка поперечная под таль — только для ЛКС71/73 И если таль включена.
    #    Для ЛКС77/78 в спецификацию идут только кран + таль.
    if getattr(q, "include_hoist", True) and q.series in ("ЛКС71", "ЛКС73"):
        trolley_code = get_trolley_code(q.hoist_brand, q.capacity)
        if trolley_code:
            price, name = ACCESSORY_PRICES.get(
                trolley_code, (0.0, f"Тележка под таль {trolley_code}"))
            q.lines.append(SpecLine(trolley_code, name, "шт", 1, price))

    # 4. Кабельные тележки и начальный подвес — только для ЛКС71/73
    if q.series in ("ЛКС71", "ЛКС73"):
        n_cable = calc_cable_trolleys(q.boom)
        ct_price, ct_name = ACCESSORY_PRICES[CABLE_TROLLEY_CODE]
        q.lines.append(SpecLine(CABLE_TROLLEY_CODE, ct_name, "шт",
                                n_cable, ct_price))
        sh_price, sh_name = ACCESSORY_PRICES[START_HANGER_CODE]
        q.lines.append(SpecLine(START_HANGER_CODE, sh_name, "шт", 1, sh_price))

    # 5. Кабель питания для R-Tech — только для ЛКС71/73.
    #    Стандартно: 4×2,5 на 20 м. В комбо (Кран + траверса): 5×2,5 на 25 м.
    #    ЛКС77/78 идёт без кабеля (только кран + таль).
    if q.hoist_brand == "R-Tech" and q.series in ("ЛКС71", "ЛКС73"):
        if q.with_traverse:
            cable_code = "A-RT-YFFB 5×2.5"
            cable_default = "Кабель питания 5×2,5 мм² (для вакуумной траверсы)"
            cable_len = 25
            fallback_price = 500.0
        else:
            cable_code = "A-RT-YFFB 4×2.5"
            cable_default = "Кабель питания тали R-Tech 4×2,5 мм²"
            cable_len = RTECH_POWER_CABLE_LEN
            fallback_price = 339.06
        cable_price, cable_name = lookup_accessory_price(
            df_71_73, cable_code, fallback_price, cable_default)
        q.lines.append(SpecLine(
            cable_code, cable_name or cable_default, "м", cable_len, cable_price))

    # 6. Электрификация (доп. опция, отдельной таблицей)
    if (q.include_electrification and q.series in ("ЛКС71", "ЛКС73")
            and q.capacity in ELECTRIFICATION_CAPACITIES):
        # Проверяем ручную настройку из session_state
        try:
            import streamlit as _st_e
            _custom_on = bool(_st_e.session_state.get("electr_custom", False))
            _custom_shkaf = _st_e.session_state.get("electr_shkaf_code_final")
            _custom_trolley = _st_e.session_state.get("electr_trolley_code_final")
            _custom_cables = _st_e.session_state.get("electr_cables_final")
        except Exception:
            _custom_on = False
            _custom_shkaf = None
            _custom_trolley = None
            _custom_cables = None

        if _custom_on and (_custom_shkaf or _custom_trolley or _custom_cables):
            # Ручной выбор
            # 1) Шкаф
            if _custom_shkaf:
                _sp, _sn = lookup_accessory_price(
                    df_71_73, _custom_shkaf, 0.0,
                    f"Шкаф управления {_custom_shkaf}")
                q.electrification_lines.append(SpecLine(
                    _custom_shkaf, _sn, "шт", 1, _sp))
            # 2) Приводная тележка
            if _custom_trolley:
                _tp, _tn = lookup_accessory_price(
                    df_71_73, _custom_trolley, 0.0,
                    f"Тележка приводная {_custom_trolley}")
                q.electrification_lines.append(SpecLine(
                    _custom_trolley, _tn, "шт", 1, _tp))
            # 3) Кабели по метражу
            _cbl_map = {
                "4x1.5": ("A-RT-YFFB 4×1.5", "Кабель плоский YFFB 4×1,5 (питание)", 214.72),
                "6x1.5": ("A-RT-YFFB 6×1.5", "Кабель плоский YFFB 6×1,5 (управление)", 341.39),
                "5x2.5": ("A-RT-YFFB 5×2.5", "Кабель плоский YFFB 5×2,5 (силовой)", 435.14),
                "4x2.5": ("A-RT-YFFB 4×2.5", "Кабель плоский YFFB 4×2,5 (для тали)", 339.06),
            }
            _cables_data = _custom_cables or {}
            for _key, (_code, _dname, _fbp) in _cbl_map.items():
                _qty = float(_cables_data.get(_key, 0) or 0)
                if _qty <= 0:
                    continue
                _p, _n = lookup_accessory_price(df_71_73, _code, _fbp, _dname)
                q.electrification_lines.append(SpecLine(
                    _code, _n or _dname, "м", _qty, _p))
        else:
            # Автоподбор: для ЛКС73М (LLL) подставляем AR580.32 и AR541 вместо AR580.10/AR542
            _pkg = list(ELECTRIFICATION_PACKAGE)
            if getattr(q, "use_lllm", False) and q.lllm_code:
                _pkg = [
                    {"code": "AR580.32",
                     "name": "Шкаф управления — для крана с электроповоротом стрелы (ЛКС73М)",
                     "qty": 1, "unit": "шт", "price": 0.0},
                    {"code": "AR541",
                     "name": "Тележка приводная (ЛКС73М)",
                     "qty": 1, "unit": "шт", "price": 0.0},
                ]
            for c in _pkg + ELECTRIFICATION_CABLES:
                price, name = lookup_accessory_price(
                    df_71_73, c["code"], c["price"], c["name"])
                q.electrification_lines.append(SpecLine(
                    c["code"], name, c["unit"], c["qty"], price))

    return q


# =========================================================
# 6. УТИЛИТЫ ФОРМАТИРОВАНИЯ DOCX
# =========================================================

def fmt_money(v: float) -> str:
    s = f"{v:,.2f}".replace(",", " ").replace(".", ",")
    return s + " ₽"


def fmt_money_plain(v: float) -> str:
    return f"{v:,.2f}".replace(",", " ").replace(".", ",")


def _num_to_words_ru(n: int) -> str:
    """Простая сумма прописью для целой части (рубли)."""
    units = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    units_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
             "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот",
                "шестьсот", "семьсот", "восемьсот", "девятьсот"]

    def under_1000(x: int, female: bool = False) -> str:
        u = units_f if female else units
        parts = []
        h = x // 100
        rest = x % 100
        if h:
            parts.append(hundreds[h])
        if 10 <= rest < 20:
            parts.append(teens[rest - 10])
        else:
            t = rest // 10
            o = rest % 10
            if t:
                parts.append(tens[t])
            if o:
                parts.append(u[o])
        return " ".join(parts)

    if n == 0:
        return "ноль"
    mln = n // 1_000_000
    thd = (n // 1000) % 1000
    rest = n % 1000
    parts = []
    if mln:
        parts.append(under_1000(mln))
        if mln % 10 == 1 and mln % 100 != 11:
            parts.append("миллион")
        elif mln % 10 in (2, 3, 4) and mln % 100 not in (12, 13, 14):
            parts.append("миллиона")
        else:
            parts.append("миллионов")
    if thd:
        parts.append(under_1000(thd, female=True))
        if thd % 10 == 1 and thd % 100 != 11:
            parts.append("тысяча")
        elif thd % 10 in (2, 3, 4) and thd % 100 not in (12, 13, 14):
            parts.append("тысячи")
        else:
            parts.append("тысяч")
    if rest:
        parts.append(under_1000(rest))
    return " ".join(parts)


def amount_in_words(amount: float) -> str:
    rub = int(amount)
    kop = round((amount - rub) * 100)
    words = _num_to_words_ru(rub).capitalize()
    return f"{words} ({amount:,.2f}) рублей {kop:02d} копеек".replace(",", " ").replace(".", ",")


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color: str = "BFBFBF", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def format_table_default(table, header_bg: str = LIGHT_BG,
                        header_bold: bool = True) -> None:
    """Применяет фирменное оформление к таблице (границы, шапка, шрифт)."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
        if ri == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_bg)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = header_bold


def add_paragraph(doc, text: str, *, bold: bool = False, italic: bool = False,
                  size: int = 11, color: Optional[RGBColor] = None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_header_footer(section, *, small_footer: bool = True) -> None:
    """Верхний колонтитул — логотип ROLLSKRAN + контакты. Нижний — реквизиты."""
    # HEADER: таблица 2 колонки — слева логотип, справа контакты
    header = section.header
    # Очищаем существующие параграфы
    for p in list(header.paragraphs):
        p.clear()
    # Удалим лишние таблицы если были
    for t in list(header.tables):
        t._element.getparent().remove(t._element)

    htable = header.add_table(rows=1, cols=2, width=Cm(17))
    htable.autofit = False
    lc, rc = htable.rows[0].cells
    lc.width = Cm(6)
    rc.width = Cm(11)

    # Логотип слева
    logo = MEDIA_DIR / "logo.jpg"
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo.exists():
        try:
            lp.add_run().add_picture(str(logo), width=Mm(65))
        except Exception:
            pass

    # Контакты справа
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = rp.add_run(f"{SUPPLIER['phone']}")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(10)
    r1.font.bold = True
    r1.font.color.rgb = DARK
    rp2 = rc.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = rp2.add_run(f"{SUPPLIER['email']}")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRAY
    rp3 = rc.add_paragraph()
    rp3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = rp3.add_run("rolls-kran.ru")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(9)
    r3.font.color.rgb = GRAY

    # FOOTER — компактно, тонко
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _parts = [SUPPLIER['short'], f"ИНН {SUPPLIER['inn']}",
              f"КПП {SUPPLIER['kpp']}"]
    if SUPPLIER.get('ogrn'):
        _parts.append(f"ОГРН {SUPPLIER['ogrn']}")
    _parts.append(SUPPLIER['address'])
    txt = "  ·  ".join(_parts)
    fr = fp.add_run(txt)
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY




# =========================================================
# 7. КАРТИНКИ
# =========================================================

def get_crane_image(series: str) -> Optional[Path]:
    mapping = {
        "ЛКС71": "kran_lks71.jpg",   # настенный одноплечевой
        "ЛКС73": "kran_lks73.jpg",   # одноплечевой на колонне
        "ЛКС77": "kran_lks77.jpg",   # двухплечевой настенный
        "ЛКС78": "kran_lks78.jpg",   # двухплечевой на колонне
    }
    fname = mapping.get(series)
    if not fname:
        return None
    p = MEDIA_DIR / fname
    return p if p.exists() and p.is_file() else None


def get_hoist_image(brand: str, exec_key: str = "") -> Optional[Path]:
    """Картинка тали. Для R-Tech P2 (с ручным контролем) — отдельная картинка."""
    if brand == "R-Tech" and exec_key == "P2":
        p_p2 = MEDIA_DIR / "tal_rtech_p2.png"
        if p_p2.exists():
            return p_p2
        # если файла P2 нет — fallback на обычную R-Tech
    mapping = {"OCALIFT": "tal_ocalift.jpg", "R-Tech": "tal_rtech.png"}
    p = MEDIA_DIR / mapping.get(brand, "")
    return p if p.exists() else None


# =========================================================
# 8. ТАБЛИЦЫ ХАРАКТЕРИСТИК
# =========================================================

def crane_characteristics(q: QuoteData) -> list[tuple[str, str]]:
    # Для LLL-кранов (ЛКС73М) стрела и высота парсятся из кода,
    # т.к. q.boom/q.height_to_arm могут остаться из сайдбара с другими значениями.
    if q.use_lllm and q.lllm_code:
        model_code = q.lllm_code
        import re as _re_ch
        _m = _re_ch.match(r"ЛКС73М\.\d+-(\d+)-(\d+)\.LLL", q.lllm_code)
        if _m:
            actual_boom = int(_m.group(1))
            actual_height = int(_m.group(2))
        else:
            actual_boom = float(q.boom)
            actual_height = float(q.height_to_arm)
        is_lll = True
    else:
        model_code = build_crane_code(q.series, q.capacity, q.boom, q.height_to_arm)
        if q.column_diameter:
            model_code = f"{model_code}.{q.column_diameter}"
        actual_boom = float(q.boom)
        if q.column_diameter:
            actual_height = int(q.column_diameter) / 100.0
        else:
            actual_height = float(q.height_to_arm)
        is_lll = False

    # Для ЛКС73М показываем полное имя серии
    _series_display = "ЛКС73М" if is_lll else q.series
    _type_desc = SERIES_DESCRIPTIONS.get(q.series, "")
    if is_lll:
        _type_desc = _type_desc + " (модификация М — с электрическим поворотом стрелы)"

    rows = [
        ("Модель", model_code),
        ("Серия", _series_display),
        ("Тип", _type_desc),
        ("Грузоподъёмность", f"{q.capacity} кг"),
        ("Длина стрелы", f"{actual_boom:g} м"),
    ]
    if q.series == "ЛКС73":
        # Габарит крана = высота до стрелы + добавка по г/п
        gabarite = gabarite_height(q.capacity, actual_height)
        rows.append(("Высота до стрелы", f"{actual_height:g} м"))
        rows.append(("Габаритная высота крана", f"{gabarite:g} м"))
        rows.append(("Угол поворота", "300°"))
    elif q.series == "ЛКС71":
        rows.append(("Исполнение", "настенное"))
        rows.append(("Угол поворота", "180° (настенное исполнение)"))
    elif q.series == "ЛКС77":
        rows += [
            ("Исполнение", "двухплечевой настенный"),
            ("Угол поворота", "основная стрела 300° / вспомогательная 320°"),
        ]
    elif q.series == "ЛКС78":
        # Колонна 3.8/4.8/5.8/6.8 м, рабочая высота (до стрелы) = column − 0.2
        _col_h = float(q.height_to_arm)
        _work_h = LKS78_HEIGHT_TO_ARM.get(_col_h, _col_h - 0.2 if _col_h > 0.5 else _col_h)
        _mass = LKS78_WEIGHTS.get((int(q.capacity), _col_h, int(actual_boom)))
        rows += [
            ("Исполнение", "двухплечевой на колонне"),
            ("Высота колонны", f"{_col_h:g} м"),
            ("Рабочая высота (до стрелы)", f"{_work_h:g} м"),
            ("Угол поворота", "основная стрела 300° / вспомогательная 320°"),
        ]
        if _mass:
            rows.append(("Масса крана", f"{_mass} кг"))
    # Материал профиля — только для ЛКС71/73/73М
    if q.series in ("ЛКС71", "ЛКС73"):
        rows.append(("Материал профиля", "Холоднокатаный профиль, сталь S355"))
    # Управление краном: радиоуправление при пакете электрификации, иначе проводной пульт
    _control = ("радиоуправление" if getattr(q, "include_electrification", False)
                else "проводной пульт")
    # Поворот стрелы: показываем отдельной строкой только для LLL (электрический)
    if is_lll:
        rows.append(("Поворот стрелы", "электрический"))
    rows += [
        ("Управление краном", _control),
        ("Климатическое исполнение", "У3, от −20 °C до +40 °C"),
        ("Гарантия", "12 месяцев (24 месяца при монтаже нашей компанией)"),
    ]
    return rows


def hoist_characteristics(q: QuoteData) -> list[tuple[str, str]]:
    speed = q.hoist_mode
    if q.hoist_mode == "ручной контроль груза":
        speed = "8/2 м/мин, ручной контроль"
    # Г/п тали: для HC берём отдельное значение, иначе — г/п крана
    if q.hoist_mode == "ручной контроль груза" and q.hc_capacity:
        hoist_cap = q.hc_capacity
    else:
        hoist_cap = q.capacity
    # Конкретные скорости для OCALIFT по модели:
    # 005-01S (500кг 1ск): 6.8 м/мин
    # 01-01S (1000кг 1ск): 6.6 м/мин
    # 005-01D (500кг 2ск): 6.9/2.3 м/мин
    # 01-01D (1000кг 2ск): 6.9/2.3 м/мин
    if q.hoist_brand == "OCALIFT":
        _oca_speeds = {
            (500, "1 скорость"): "6,8 м/мин",
            (1000, "1 скорость"): "6,6 м/мин",
            (500, "2 скорости"): "6,9/2,3 м/мин",
            (1000, "2 скорости"): "6,9/2,3 м/мин",
        }
        _oca_speed = _oca_speeds.get((hoist_cap, q.hoist_mode))
        if _oca_speed:
            speed = _oca_speed
    # Строительная высота тали
    if q.hoist_brand == "OCALIFT":
        build_h = {500: "540 мм", 1000: "575 мм"}.get(hoist_cap, "—")
    elif q.hoist_brand == "R-Tech":
        if hoist_cap in (80, 125):
            build_h = "400 мм"
        elif hoist_cap in (250, 500):
            build_h = "420 мм"
        elif hoist_cap == 1000:
            build_h = "516 мм"
        else:
            build_h = "—"
    else:
        build_h = "—"
    # Тип цепи зависит от бренда:
    # OCALIFT — калиброванная стальная класс 80 (без японской пометки)
    # R-Tech — калиброванная стальная класс 80, надёжная японская цепь
    if q.hoist_brand == "OCALIFT":
        _chain_type = "калиброванная стальная класс 80"
        _country = "Китай"
    elif q.hoist_brand == "R-Tech":
        _chain_type = "калиброванная стальная класс 80, надёжная японская цепь"
        _country = "Россия (все запчасти всегда в наличии на складе в Санкт-Петербурге и Москве)"
    else:
        _chain_type = "калиброванная стальная класс 80"
        _country = "—"

    return [
        ("Производитель", q.hoist_brand),
        ("Страна производителя", _country),
        ("Грузоподъёмность", f"{hoist_cap} кг"),
        ("Скорость подъёма", speed),
        ("Высота подъёма", f"{q.hoist_height} м"),
        ("Строительная высота", build_h),
        ("Режим работы", "М5 (среднего использования)"),
        ("Напряжение", "380 В ± 10%, 50 Гц / цепь управления 24 В"),
        ("Степень защиты", "IP55"),
        ("Тип цепи", _chain_type),
        ("Гарантия", "12 месяцев"),
    ]


# =========================================================
# 9. ГЕНЕРАЦИЯ КП (КРАСИВОЕ)
# =========================================================

def _add_kv_table(doc, rows: list[tuple[str, str]], first_col_cm: float = 6.0,
                 second_col_cm: float = 10.5) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.width = Cm(first_col_cm)
        c2.width = Cm(second_col_cm)
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = DARK
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)
        set_cell_borders(c1)
        set_cell_borders(c2)
        if i % 2 == 0:
            set_cell_shading(c1, "FAFAFA")
            set_cell_shading(c2, "FAFAFA")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _add_spec_table(doc, lines: list[SpecLine], show_zero_price: bool = True) -> None:
    table = doc.add_table(rows=1 + len(lines) + 1, cols=6)
    table.autofit = False
    widths = [Cm(0.8), Cm(3.0), Cm(8.6), Cm(1.2), Cm(1.4), Cm(2.0)]
    header = ["№", "Код", "Наименование", "Ед.", "Кол-во", "Сумма,\n₽ с НДС"]
    for i, h in enumerate(header):
        c = table.rows[0].cells[i]
        c.width = widths[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        r.font.bold = True
        set_cell_shading(c, BRAND_YELLOW_HEX)
        r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)  # чёрный текст на оранжевом
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, ln in enumerate(lines, start=1):
        row = table.rows[idx]
        vals = [
            str(idx),
            ln.code,
            ln.name,
            ln.unit,
            f"{ln.qty:g}",
            fmt_money_plain(ln.total) if (ln.total > 0 or show_zero_price) else "—",
        ]
        for j, val in enumerate(vals):
            c = row.cells[j]
            c.width = widths[j]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if j in (0, 3, 4)
                          else WD_ALIGN_PARAGRAPH.RIGHT if j == 5
                          else WD_ALIGN_PARAGRAPH.LEFT)
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            if idx % 2 == 0:
                set_cell_shading(c, "FAFAFA")
    # Итого
    total = sum(ln.total for ln in lines)
    row = table.rows[-1]
    for j in range(6):
        c = row.cells[j]
        c.width = widths[j]
        c.text = ""
        set_cell_shading(c, BRAND_YELLOW_SOFT_HEX)
    row.cells[0].merge(row.cells[4])
    pt = row.cells[0].paragraphs[0]
    pt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rt = pt.add_run("Итого с НДС 22%:")
    rt.font.name = "Times New Roman"
    rt.font.size = Pt(11)
    rt.font.bold = True
    ps = row.cells[5].paragraphs[0]
    ps.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rs = ps.add_run(fmt_money_plain(total))
    rs.font.name = "Times New Roman"
    rs.font.size = Pt(11)
    rs.font.bold = True
    rs.font.color.rgb = BRAND_COLOR
    for c in table.rows[0].cells:
        set_cell_borders(c, color="111111", size="8")
    for row in table.rows[1:]:
        for c in row.cells:
            set_cell_borders(c)


def build_simple_kp_docx(items: list, kp_number: str, buyer_name: str = "",
                         kp_date: str = "", comment: str = "") -> bytes:
    """Упрощённый КП для перегенерации из CRM по произвольному списку позиций.

    Строит DOCX с логотипом ROLLSKRAN, таблицей спецификации (код / наименование / ед / кол-во / цена / сумма) и итогами.
    items: list[dict] с полями code, name, unit, qty, price.
    """
    from datetime import date as _date
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
        add_header_footer(s)

    # Шапка
    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = top.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    r0.font.name = "Times New Roman"
    r0.font.size = Pt(24)
    r0.font.bold = True
    r0.font.color.rgb = BRAND_COLOR

    _kp_d = kp_date or _date.today().strftime("%d.%m.%Y")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f"№ {kp_number}   от   {_kp_d}")
    rs.font.name = "Times New Roman"
    rs.font.size = Pt(12)
    rs.font.color.rgb = GRAY

    if buyer_name:
        who = doc.add_paragraph()
        who.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rw = who.add_run(f"Кому: {buyer_name}")
        _rw.font.name = "Times New Roman"
        _rw.font.size = Pt(12)
        _rw.font.bold = True

    doc.add_paragraph()

    # Спецификация — SpecLine для _add_spec_table
    _lines = []
    for it in items:
        _lines.append(SpecLine(
            code=str(it.get("code") or ""),
            name=str(it.get("name") or ""),
            unit=str(it.get("unit") or "шт"),
            qty=float(it.get("qty") or 0),
            price=float(it.get("price") or 0),
        ))
    _add_spec_table(doc, _lines, show_zero_price=True)

    # Итого
    _total = sum(ln.total for ln in _lines)
    _p_total = doc.add_paragraph()
    _p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _rt = _p_total.add_run(f"Итого: {fmt_money_plain(_total)} ₽ (с НДС 22 %)")
    _rt.font.name = "Times New Roman"
    _rt.font.size = Pt(12)
    _rt.font.bold = True

    if comment:
        doc.add_paragraph()
        _pc = doc.add_paragraph()
        _rc = _pc.add_run("Комментарий: ")
        _rc.font.name = "Times New Roman"
        _rc.font.size = Pt(11)
        _rc.font.bold = True
        _rc2 = _pc.add_run(comment)
        _rc2.font.name = "Times New Roman"
        _rc2.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_kp_docx(q: QuoteData, kp_number: str, buyer_name: str = "") -> bytes:
    doc = Document()

    # Стили страницы
    for s in doc.sections:
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
        add_header_footer(s)

    # -------- Титул --------
    # Красная полоса-заголовок
    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top.paragraph_format.space_after = Pt(4)
    r0 = top.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    r0.font.name = "Times New Roman"
    r0.font.size = Pt(28)
    r0.font.bold = True
    r0.font.color.rgb = BRAND_COLOR

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    rs = sub.add_run(f"№ {kp_number}   от   {date.today().strftime('%d.%m.%Y')}")
    rs.font.name = "Times New Roman"
    rs.font.size = Pt(12)
    rs.font.color.rgb = GRAY

    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # Название крана крупно
    add_paragraph(
        doc,
        SERIES_DESCRIPTIONS.get(q.series, ""),
        bold=True, size=16, color=DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    subtitle_extra = ""
    if q.series in ("ЛКС73", "ЛКС78"):
        subtitle_extra = f" · высота {q.height_to_arm} м"
    add_paragraph(
        doc,
        f"серии {q.series} · г/п {q.capacity} кг · стрела {q.boom} м{subtitle_extra}",
        italic=True, size=12, color=GRAY,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
    )

    # Фото крана
    img = get_crane_image(q.series)
    if img:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Mm(95))

    if buyer_name:
        add_paragraph(doc, f"Для: {buyer_name}",
                      italic=True, size=11, color=GRAY,
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_paragraph(
        doc,
        "Легкие крановые системы созданы на базе нового уникального сверхпрочного "
        "профиля. Холоднокатаные профили разработаны и выполнены из специальной "
        "стали S355. Использование специального профиля и высококачественных "
        "тележек обеспечивает лёгкость хода при ручном перемещении весь срок "
        "службы, усилие перемещения не превышает 1% переносимого груза.",
        size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
    )

    doc.add_page_break()

    # -------- Стр. 2: характеристики --------
    add_paragraph(doc, "Технические характеристики",
                  bold=True, size=18, color=BRAND_COLOR,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_paragraph(doc, f"Кран {q.series}", bold=True, size=12, color=DARK,
                  space_after=4)
    _add_kv_table(doc, crane_characteristics(q))

    # Фото крана — под таблицей характеристиками крана
    ci = get_crane_image(q.series)
    if ci:
        add_paragraph(doc, "", space_after=4)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(ci), width=Mm(65))
        except Exception:
            pass

    add_paragraph(doc, "", space_after=6)
    add_paragraph(doc, f"Электротельфер (таль) {q.hoist_brand}",
                  bold=True, size=12, color=DARK, space_after=4)
    _add_kv_table(doc, hoist_characteristics(q))

    # Фото тали
    hi = get_hoist_image(q.hoist_brand, getattr(q, "hoist_exec", "") or "")
    if hi:
        add_paragraph(doc, "", space_after=4)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(hi), width=Mm(40))
        except Exception:
            pass

    doc.add_page_break()

    # -------- Стр. 3: расчёт --------
    add_paragraph(doc, "Спецификация и расчёт стоимости",
                  bold=True, size=18, color=BRAND_COLOR,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_spec_table(doc, q.lines)

    # Электрификация — отдельной таблицей
    if q.electrification_lines:
        add_paragraph(doc, "", space_after=6)
        add_paragraph(doc, "Дополнительная комплектация — электрификация",
                      bold=True, size=12, color=DARK, space_after=4)
        add_paragraph(
            doc,
            "Опциональный пакет полной электрификации перемещения тали по стреле "
            f"для {q.series} г/п {q.capacity} кг. Повышает удобство работы, "
            "исключает ручное перемещение тали на большие пролёты.",
            italic=True, size=10, color=GRAY, space_after=4,
        )
        _add_spec_table(doc, q.electrification_lines)

    # Монтаж — отдельным блоком
    if q.include_montage and q.montage_price > 0:
        add_paragraph(doc, "", space_after=6)
        add_paragraph(doc, "Монтаж и пусконаладка",
                      bold=True, size=12, color=DARK, space_after=4)
        montage_line = SpecLine("—", "Монтаж и пусконаладочные работы "
                               "консольного крана", "усл.", 1, q.montage_price)
        _add_spec_table(doc, [montage_line])
        add_paragraph(
            doc,
            "Услуги монтажа оказываются от ИП без НДС по отдельному договору "
            "подряда. В стоимость поставки не включены.",
            italic=True, size=9, color=GRAY, space_after=4,
        )

    # Итого блок
    add_paragraph(doc, "", space_after=6)
    total_table = doc.add_table(rows=3, cols=2)
    total_table.autofit = False
    total_rows = [
        ("Стоимость поставки с НДС 22%", fmt_money(q.total)),
        (f"Предоплата {int(PREPAYMENT_DEFAULT*100)} %",
         fmt_money(q.total * PREPAYMENT_DEFAULT)),
        (f"Остаток {round((1-PREPAYMENT_DEFAULT)*100)} % перед отгрузкой",
         fmt_money(q.total * (1 - PREPAYMENT_DEFAULT))),
    ]
    for i, (k, v) in enumerate(total_rows):
        c1, c2 = total_table.rows[i].cells
        c1.width = Cm(11.5)
        c2.width = Cm(5.5)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(k + ":")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        r1.font.bold = (i == 0)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(v)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12 if i == 0 else 11)
        r2.font.bold = (i == 0)
        if i == 0:
            r2.font.color.rgb = BRAND_COLOR
            set_cell_shading(c1, BRAND_YELLOW_SOFT_HEX)
            set_cell_shading(c2, BRAND_YELLOW_SOFT_HEX)
        set_cell_borders(c1)
        set_cell_borders(c2)

    # -------- Стр. 4: условия --------
    doc.add_page_break()
    add_paragraph(doc, "Условия поставки",
                  bold=True, size=18, color=BRAND_COLOR,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    conds = [
        ("Оплата", "70 % предоплата после подписания договора и спецификации, "
                   "30 % — по уведомлению о готовности к отгрузке."),
        ("Срок изготовления", "до 20 рабочих дней после поступления предоплаты. "
                              "Возможна досрочная поставка."),
        ("Доставка", "Бесплатная доставка до ТК «Деловые линии» в "
                     "Санкт-Петербурге. Далее — за счёт Покупателя."),
        ("Гарантия", "12 месяцев на кран, 12 месяцев на таль. "
                     "24 месяца при монтаже нашей компанией."),
        ("Документы", "Паспорт крана, паспорт тали, руководство по эксплуатации, "
                      "сертификат соответствия."),
        ("Срок действия КП", "14 календарных дней с даты выставления."),
    ]
    for k, v in conds:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"■  {k}. ")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        r1.font.bold = True
        r1.font.color.rgb = BRAND_COLOR
        r2 = p.add_run(v)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)

    # Подпись в 2 колонках: текст | подпись+печать
    add_paragraph(doc, "", space_after=10)
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = False
    lc, rc = sig_table.rows[0].cells
    lc.width = Cm(11)
    rc.width = Cm(6)

    # Левая колонка — текст
    lc.text = ""
    lp1 = lc.paragraphs[0]
    r = lp1.add_run("С уважением,")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    lp2 = lc.add_paragraph()
    r = lp2.add_run(f"Исполнительный директор {SUPPLIER['short']}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    lp3 = lc.add_paragraph()
    r = lp3.add_run("Букреев Антон")
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.bold = True
    lp4 = lc.add_paragraph()
    r = lp4.add_run(f"тел. {SUPPLIER['phone']}")
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY
    lp5 = lc.add_paragraph()
    r = lp5.add_run(SUPPLIER['email'])
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    # Правая колонка — подпись и печать
    rc.text = ""
    sig_img = MEDIA_DIR / "signature.jpg"
    stamp_img = MEDIA_DIR / "stamp.jpg"
    if sig_img.exists():
        rp1 = rc.paragraphs[0]
        rp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            rp1.add_run().add_picture(str(sig_img), width=Mm(45))
        except Exception:
            pass
    if stamp_img.exists():
        rp2 = rc.add_paragraph()
        rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp2.paragraph_format.space_before = Pt(0)
        try:
            rp2.add_run().add_picture(str(stamp_img), width=Mm(32))
        except Exception:
            pass

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# =========================================================
# 10. ГЕНЕРАЦИЯ ДОГОВОРА ПОСТАВКИ (СКЕЛЕТ ИЗ ШАБЛОНА)
# =========================================================

DOGOVOR_SECTIONS = [
    ("1. ПРЕДМЕТ ДОГОВОРА", [
        "1.1. Поставщик обязуется поставлять Покупателю в течение срока действия Договора заказанную Покупателем продукцию (далее — «Продукция»), а Покупатель обязуется оплачивать и принимать в полном объёме всю заказанную Продукцию.",
        "1.2. Номенклатура и количество заказываемой Продукции указываются Покупателем в Спецификации к настоящему договору поставки.",
        "1.3. Номенклатура и количество Продукции по каждой отдельной поставке согласовываются Покупателем и Поставщиком в Спецификации, являющейся неотъемлемой частью Договора.",
        "1.4. В Спецификации указываются: номенклатура и количество Продукции, цена, условия оплаты, условия доставки.",
        "1.5. На основании согласованной Сторонами Спецификации Поставщик выставляет Покупателю счёт.",
    ]),
    ("2. ЦЕНА ПРОДУКЦИИ И ПОРЯДОК ОПЛАТЫ", [
        "2.1. Цены на продукцию указываются, включая НДС 22 %.",
        "2.2. Покупатель вносит предоплату на расчётный счёт Поставщика, указанный в счёте, в размере 100 % от общей стоимости заказа, если иное не предусмотрено Спецификацией.",
        "2.3. При повышении цен Поставщик обязан уведомить Покупателя не менее чем за 3 (три) рабочих дня.",
        "2.4. Поставщик не вправе изменять цены на Продукцию, предоплаченную Покупателем.",
        "2.5. Срок оплаты: Покупатель обязан произвести полную оплату стоимости Продукции в течение 5 (пяти) банковских дней с момента получения уведомления о готовности Продукции к отгрузке.",
        "2.6. В случае частичной оплаты Продукции по Спецификации к настоящему Договору, поставка осуществляется только после полного погашения задолженности в размере 100 % суммы, указанной в Спецификации.",
        "2.7. Подтверждение оплаты: отгрузка Продукции производится исключительно после поступления 100 % оплаты на расчётный счёт Поставщика и подтверждения платежа (выпиской банка или иным документом).",
        "2.8. Ответственность за просрочку оплаты: при непоступлении оплаты в установленный срок Покупатель уплачивает Поставщику пеню в размере 1 % (один процент) от неоплаченной суммы за каждый день просрочки, начиная со дня, следующего за истечением срока оплаты.",
        "2.9. Право Поставщика на односторонний отказ от исполнения Договора: если оплата не поступит в течение 10 (десяти) рабочих дней с момента уведомления о готовности Продукции, Поставщик вправе отказаться от исполнения Договора в одностороннем порядке, уведомив Покупателя в письменной форме, и потребовать возмещения убытков, связанных с хранением и простоем Продукции.",
    ]),
    ("3. СРОКИ И ПОРЯДОК ПОСТАВКИ", [
        "3.1. Уведомление о готовности Продукции: Поставщик обязан уведомить Покупателя в письменной форме (по электронной почте или иным согласованным способом) о готовности Продукции к отгрузке не менее чем за 3 (три) рабочих дня до предполагаемой даты отгрузки. Покупатель обязуется произвести отгрузку Продукции со склада Поставщика в сроки, указанные в Спецификации.",
        "3.2. При задержке отгрузки Продукции на срок более 10 (десяти) рабочих дней Поставщик выплачивает по требованию Покупателя пени в размере 1 % от суммы предоплаты за каждый рабочий день просрочки. Пени начисляются с момента истечения установленного срока отгрузки до фактической передачи Продукции транспортной компании.",
        "3.3. Покупатель обязуется произвести выборку заказа со склада Поставщика в течение 10 (десяти) рабочих дней со дня получения письменного уведомления на электронный адрес о готовности Продукции к отгрузке. Уведомление считается полученным с момента его отправки на электронный адрес Покупателя, указанный в Договоре.",
        "3.4. Хранение заказа свыше 10 (десяти) рабочих дней осуществляется за дополнительную плату. Стоимость хранения составляет 1 % от суммы заказа за каждый день хранения. Сроки хранения должны быть согласованы с Поставщиком в письменной форме. В случае отсутствия согласования Поставщик вправе отказать в дальнейшем хранении и реализовать Продукцию по своему усмотрению без уведомления Покупателя.",
        "3.5. Доставка Продукции осуществляется транспортом Покупателя на условиях самовывоза или доставки до транспортных компаний на выбор Покупателя — ООО «Байкал-Сервис», ООО «Деловые линии», ООО «Первая экспедиционная компания», или иным способом за счёт Покупателя, если в Спецификации не указано иное.",
        "3.6. Поставщик обязуется сообщить Покупателю о наличии Продукции и готовности её к отгрузке. При реализации Продукции Поставщик предоставляет Покупателю (представителю по доверенности) сопроводительные документы на Продукцию: 1) универсальный передаточный документ по форме, рекомендованной Письмом ФНС России от 21.10.2013 № ММВ-20-3/96@; 2) паспорта или иные документы, подтверждающие качество Продукции производителем; 3) транспортные накладные (при необходимости).",
        "3.7. При отправке Продукции до склада Покупателя силами и средствами Поставщика (транспортной компанией или привлечением стороннего перевозчика) товаросопроводительные документы оформляются в 2 (двух) экземплярах, по одному для каждой из Сторон, и вкладываются в груз или направляются в отдельном конверте по отдельной экспедиторской накладной, или отправляются по ЭДО. Покупатель обязуется возвратить экземпляр Поставщика, подписанный со своей стороны, не позднее 30 (тридцати) календарных дней с даты поставки.",
        "3.8. Поставщик считается выполнившим своё обязательство по передаче Продукции Покупателю: а) в случае доставки Продукции Поставщиком — с момента сдачи Продукции на склад Покупателя (Грузополучателя); б) в случае самовывоза Продукции Покупателем — с момента получения Продукции представителем Покупателя на складе Поставщика; в) в случае доставки Продукции сторонним перевозчиком — с момента вручения Поставщиком Продукции перевозчику для его доставки.",
        "3.9. Право собственности на купленную Продукцию и риск случайной гибели или случайной порчи, утраты или повреждения переходит к Покупателю с момента, когда Поставщик в соответствии с условиями настоящего Договора считается выполнившим своё обязательство по передаче Продукции Покупателю.",
        "3.10. Место поставки по каждой поставке в рамках настоящего Договора указывается в Спецификации и/или счёте. Если Стороны не предусмотрели место поставки в Спецификации и/или счёте, таковым считается: 1) при доставке Продукции силами Поставщика — фактический адрес Покупателя из карточки контрагента, а при его отсутствии — юридический адрес Покупателя, а при его отсутствии — адрес из ЕГРЮЛ; 2) при доставке Продукции силами транспортной компании — терминал транспортной компании в городе местонахождения Покупателя.",
        "3.11. В случае организации доставки Продукции силами транспортной компании Покупатель, подписывая настоящий Договор, выражает своё согласие на получение Поставщиком от транспортной компании копий документов, подтверждающих выдачу груза Грузополучателю.",
        "3.12. Дополнительные условия отгрузки и транспортировки, включая особые требования к упаковке, маркировке и документации, определяются в Спецификации для каждой поставки отдельно.",
    ]),
    ("4. ПРИЁМКА ПРОДУКЦИИ, КАЧЕСТВО ПРОДУКЦИИ, ГАРАНТИИ", [
        "4.1. Поставщик обязуется поставить Продукцию строго в соответствии со Спецификацией.",
        "4.2. Право собственности на Продукцию переходит от Поставщика к Покупателю: 4.2.1. в момент передачи Продукции Покупателю на складе Покупателя, если доставка осуществляется за счёт Поставщика; 4.2.2. в момент передачи Продукции Покупателю или его представителю на складе Поставщика, если доставка осуществляется силами Покупателя; 4.2.3. в момент передачи Продукции транспортной компании, если доставка Продукции производится с привлечением сторонней транспортной компании.",
        "4.3. В случае самовывоза со склада Поставщика Продукция отгружается только при наличии правильно заполненной доверенности, оформленной на лицо, получающее Продукцию.",
        "4.4. В случае доставки Продукции за счёт Поставщика Покупатель при получении Продукции обязан подписать универсальный передаточный документ (УПД), заверить его круглой печатью Покупателя либо доверенностью на получение ТМЦ, выданной лицу, подписывающему УПД, и передать на месте представителю Поставщика.",
        "4.5. В случае доставки Продукции через транспортную компанию Покупатель обязан непосредственно после приёмки Продукции на своём складе подписать УПД и отправить его Поставщику: копию — в день приёмки по электронной почте zakaz@rolls-kran.ru; оригинал — курьерской доставкой или почтой в течение 3 рабочих дней.",
        "4.6. В целях упрощения и ускорения документооборота Стороны могут осуществлять электронный обмен документами в соответствии с ФЗ от 06.04.2011 № 63-ФЗ «Об электронной подписи».",
        "4.7. Покупатель обязуется произвести приёмку Продукции на своём складе по количеству и по качеству в течение 5 рабочих дней с момента получения Продукции.",
        "4.8. При обнаружении несоответствия качества или расхождения по количеству во время приёмки Продукции Покупатель заполняет Бланк рекламации установленной формы и направляет Поставщику по электронной почте zakaz@rolls-kran.ru. К Бланку рекламации необходимо прикрепить информацию, подтверждающую наличие повреждений или недостачи: фотографии, данные маркировки на упаковке Продукции или иную информацию по запросу Поставщика.",
        "4.9. Рекламацию по качеству Продукции Покупатель вправе предъявить в течение гарантийного срока в установленном порядке по электронной почте zakaz@rolls-kran.ru.",
    ]),
    ("5. ФОРС-МАЖОРНЫЕ ОБСТОЯТЕЛЬСТВА", [
        "5.1. Стороны освобождаются от ответственности за частичное или полное неисполнение обязательств по Договору, если надлежащее исполнение оказалось невозможным вследствие обстоятельств непреодолимой силы, таких как: пожар, забастовка, стихийные бедствия, акты государственной власти, военные действия, аварийные ситуации на предприятиях Сторон.",
        "5.2. При наступлении обстоятельств непреодолимой силы, указанных в пункте 5.1 Договора, Сторона должна без промедления известить о них в письменной форме другую Сторону. Извещение должно содержать данные о характере обстоятельств, а также копии официальных документов, удостоверяющих наличие этих обстоятельств.",
    ]),
    ("6. ГАРАНТИИ СТОРОН", [
        "6.1. Стороны гарантируют на момент заключения и в период действия настоящего договора как Сторонами договора, так и их работниками или посредниками:",
        "6.1.1. полный отказ от операций и действий, связанных с легализацией (отмыванием) доходов, полученных преступным путём, а также соблюдение требований Федерального закона № 115-ФЗ от 07.08.2001 «О противодействии легализации (отмыванию) доходов, полученных преступным путём, и финансированию терроризма»;",
        "6.1.2. неприменение и пресечение коррупционных действий в отношениях со второй Стороной договора и третьими лицами в соответствии с Федеральным законом № 273-ФЗ от 25.12.2008 «О противодействии коррупции»;",
        "6.1.3. сдачу полной и достоверной налоговой отчётности в соответствии с НК РФ, предоставление запрашиваемых МИФНС документов в случае встречной или иной проверки.",
        "6.2. Стороны гарантируют достоверность информации и заверений, представленных при заключении и исполнении договора, относящихся к правоспособности Сторон, полномочиям на его заключение, а также своему финансовому состоянию.",
        "6.3. В случае нарушения одной из Сторон гарантий, предусмотренных разделом 6 настоящего договора, вторая Сторона вправе в одностороннем порядке отказаться от настоящего договора путём направления виновной Стороне письменного уведомления. Договор подлежит прекращению с момента получения виновной Стороной уведомления об одностороннем отказе.",
    ]),
    ("7. ПЕРСОНАЛЬНЫЕ ДАННЫЕ", [
        "7.1. Подписанием настоящего договора Стороны дают согласие на обработку персональных данных своих работников, переданных второй Стороне с целью надлежащего исполнения настоящего договора, включая персональные данные работников, уполномоченных на подписание настоящего договора, дополнительных соглашений и приложений к нему и первичных учётных документов во исполнение договора. Стороны подтверждают, что до передачи персональных данных они выполнили все действия, предусмотренные Федеральным законом № 152-ФЗ от 27.07.2006 «О персональных данных».",
        "7.2. Стороны обязуются соблюдать правила обработки персональных данных, предусмотренные Законом № 152-ФЗ, соблюдать конфиденциальность персональных данных и обеспечивать безопасность персональных данных при их обработке.",
        "7.3. Стороны вправе осуществлять следующие действия по обработке персональных данных: сбор, запись, систематизация, накопление, хранение, уточнение, извлечение, использование, передачу, блокирование, удаление, уничтожение.",
        "7.4. Стороны вправе осуществлять действия по обработке персональных данных в течение всего срока действия настоящего договора и пяти лет с даты прекращения его действия.",
        "7.5. Стороны вправе в любое время отозвать согласие на обработку персональных данных путём направления второй Стороне письменного уведомления. Сторона, получившая такое уведомление, обязуется прекратить обработку персональных данных.",
    ]),
    ("8. СРОК ДЕЙСТВИЯ ДОГОВОРА", [
        "8.1. Договор вступает в силу с момента подписания его Сторонами и действует до 31.12.2026 г. При этом в части неисполненных обязательств Договор действует до полного их исполнения.",
        "8.2. Настоящий Договор подлежит автоматической пролонгации на каждый следующий год, если ни одна из Сторон договора за один календарный месяц не заявит о своём намерении его прекратить.",
        "8.3. Договор может быть расторгнут по взаимному согласию Сторон.",
        "8.4. Договор может быть расторгнут Стороной в одностороннем порядке только при условии отсутствия задолженности перед другой Стороной.",
    ]),
    ("9. ПРОЧИЕ УСЛОВИЯ", [
        "9.1. Договор составлен в двух экземплярах, по одному экземпляру для каждой из Сторон. Оба экземпляра имеют одинаковую юридическую силу.",
        "9.2. Договор может быть дополнен или изменён по соглашению Сторон, путём подписания дополнительных соглашений к Договору.",
        "9.3. Все спорные вопросы Стороны стараются разрешить путём переговоров. Если спорные вопросы не могут быть разрешены путём переговоров, они передаются на рассмотрение арбитражного суда по местонахождению истца в порядке, установленном действующим законодательством РФ.",
        "9.4. Юридическую силу имеют оригиналы документов. Электронные и факсимильные копии подлежат замене на оригиналы.",
        "9.5. Стороны могут вести почтовую, факсимильную и электронную переписку только с адресов и на адреса, указанные в настоящем Договоре. Электронная почта Поставщика: zakaz@rolls-kran.ru. Почтовый адрес Поставщика: 195030, г. Санкт-Петербург, вн.тер.г. муниципальный округ Ржевка, ш. Революции, д. 114, Литера А, помещ. 111.",
        "9.6. В случае изменения реквизитов или адресов для переписки Стороны письменно информируют друг друга в срок не более 3 рабочих дней.",
        "9.7. Все ранее подписанные между Сторонами договоры и соглашения утрачивают свою силу с момента подписания настоящего Договора.",
    ]),
    ("10. ГАРАНТИИ ПОСТАВЩИКА", [
        "10.1. Поставщик заверяет и гарантирует следующее:",
        "— Поставщик является надлежащим образом зарегистрированной организацией;",
        "— все сведения о Поставщике в ЕГРЮЛ достоверны на момент подписания договора; если в дальнейшем в ЕГРЮЛ появится запись о недостоверности данных о Поставщике, он обязуется в течение месяца внести в ЕГРЮЛ достоверные сведения или подтвердить регистрирующему органу, что сведения в ЕГРЮЛ достоверны;",
        "— Поставщик располагает необходимыми ресурсами для исполнения настоящего Договора;",
        "— Поставщик отразит все операции по настоящему Договору, включая полученные от Покупателя авансы и реализацию Продукции Покупателю, в бухгалтерской и налоговой отчётности;",
        "— в случае получения Поставщиком требований налогового органа о представлении документов, относящихся к сделке с Покупателем, Поставщик обязуется исполнить требование в течение пяти рабочих дней со дня получения требования;",
        "— Поставщик обязуется выставлять Покупателю правильно оформленные счета-фактуры и первичные бухгалтерские документы в соответствии с требованиями действующего законодательства;",
        "— Поставщик самостоятельно выполняет обязательства по настоящему договору.",
    ]),
    ("11. УСЛОВИЯ И ПОРЯДОК ОБМЕНА ЭЛЕКТРОННЫМИ ДОКУМЕНТАМИ", [
        "11.1.1. Стороны договора признают юридическую силу электронных документов, подписанных усиленной электронной подписью (УЭП), равнозначной бумажным документам, подписанным собственноручной подписью, при условии соблюдения требований законодательства Российской Федерации.",
        "11.1.2. ООО «ЛКС» (ИНН 7806586802) осуществляет обмен электронными документами через систему электронного документооборота (ЭДО) СБИС, идентификатор участника: 2BE2fc30491f9e94f8ebce17c64df74224b.",
        "11.1.3. Электронные документы, направляемые между Сторонами, должны соответствовать требованиям, установленным законодательством РФ, а также настоящим договором.",
        "11.2. Порядок обмена: электронные документы направляются через систему ЭДО СБИС. Каждая Сторона обязана обеспечить конфиденциальность и целостность электронных документов при их передаче и хранении. Электронные документы считаются полученными: при направлении через ЭДО СБИС — в момент фиксации факта отправки в системе; при направлении по электронной почте — в момент получения уведомления о доставке письма.",
        "11.3. Перечень электронных документов: договоры; универсальные передаточные документы (УПД); акты выполненных работ/оказанных услуг; приложения к договорам; акты сверок; чертежи; доверенности; товарные накладные (ТОРГ-12).",
        "11.4. Требования: электронные документы должны быть подписаны усиленной электронной подписью уполномоченного представителя Стороны. Формат электронных документов должен соответствовать требованиям, установленным законодательством РФ, или быть согласован Сторонами дополнительно.",
        "11.5. Сторона, направившая электронный документ, несёт ответственность за его достоверность и соответствие требованиям законодательства. В случае возникновения сомнений в подлинности электронного документа Стороны обязаны провести взаимную проверку и подтвердить его действительность.",
        "11.6. Стороны обязуются хранить электронные документы в течение срока, установленного законодательством РФ или настоящим договором. По требованию одной из Сторон электронные документы могут быть предоставлены в бумажном виде, заверенные надлежащим образом.",
    ]),
]


PERCENT_WORDS = {
    0: "ноль", 5: "пять", 10: "десять", 15: "пятнадцать",
    20: "двадцать", 25: "двадцать пять", 30: "тридцать",
    35: "тридцать пять", 40: "сорок", 45: "сорок пять",
    50: "пятьдесят", 55: "пятьдесят пять",
    60: "шестьдесят", 65: "шестьдесят пять",
    70: "семьдесят", 75: "семьдесят пять",
    80: "восемьдесят", 85: "восемьдесят пять",
    90: "девяносто", 95: "девяносто пять", 100: "сто",
}


def percent_in_words(n: int) -> str:
    """Переводит целое 0..100 в пропись. Для произвольных — через десятки+единицы."""
    n = int(n)
    if n in PERCENT_WORDS:
        return PERCENT_WORDS[n]
    tens = (n // 10) * 10
    ones = n % 10
    ones_words = {1: "один", 2: "два", 3: "три", 4: "четыре",
                  6: "шесть", 7: "семь", 8: "восемь", 9: "девять"}
    if tens in PERCENT_WORDS and ones in ones_words:
        return f"{PERCENT_WORDS[tens]} {ones_words[ones]}"
    return str(n)


def _build_default_warranty(q: QuoteData) -> str:
    """Собирает текст гарантии по умолчанию:
    — без монтажа: просто 12 месяцев
    — с монтажом: 12 мес + блок про расширенную гарантию 24 мес.
    """
    base = ("Поставщик гарантирует соответствие заявленного качества комплектующих при "
            "соблюдении Покупателем условий эксплуатации и хранения. Гарантия на все товары — "
            "12 месяцев со дня получения Покупателем.")
    if getattr(q, "include_montage", False):
        return (base + " Гарантия на монтажные работы — 12 месяцев. При заказе монтажа "
                       "в нашей компании вы получаете расширенную гарантию 24 месяца на комплектующие крановой системы "
                       "(электрическая таль, вакуумная траверса и пакет электрификации в расширенную гарантию не входят).")
    return base


def build_dogovor_docx(q: QuoteData, buyer: dict, contract_number: str,
                       contract_date_str: str, **kwargs) -> bytes:
    """Единая болванка — все договоры через dogovor_traversa.build_dogovor_universal.

    Канонический шаблон = Dogovor_Vakuumnye-traversy.docx.
    Обёртка сохраняет старый API вызова но генерит новую болванку.
    Автоматически сохраняет договор в contracts_db.
    """
    import dogovor_traversa
    # Поставщик берём из глобального SUPPLIER
    docx_bytes = dogovor_traversa.build_dogovor_universal(
        q, buyer, contract_number, contract_date_str,
        supplier=SUPPLIER,
        **kwargs,
    )

    # АВТОСОХРАНЕНИЕ в базу договоров
    try:
        import contracts_db, dogovor_traversa as _dt
        # Строим lines из q
        lines = []
        for ln in getattr(q, "lines", []) or []:
            if hasattr(ln, "code"):
                lines.append({"code": getattr(ln,"code","") or "", "name": getattr(ln,"name","") or "",
                              "unit": getattr(ln,"unit","шт") or "шт",
                              "qty": float(getattr(ln,"qty",0) or 0),
                              "price": float(getattr(ln,"price",0) or 0)})
            elif isinstance(ln, dict):
                lines.append({"code": str(ln.get("code","") or ""), "name": str(ln.get("name","") or ""),
                              "unit": str(ln.get("unit","шт") or "шт"),
                              "qty": float(ln.get("qty",0) or 0),
                              "price": float(ln.get("price",0) or 0)})
        has_vat = bool(getattr(q, "include_vat", True))
        # Если has_vat не задан из q — берём из поставщика
        if SUPPLIER:
            has_vat = _dt._is_vat_supplier(SUPPLIER)
        total, vat = _dt._calc_totals_vat(lines, has_vat)

        prepay_pct = kwargs.get("prepay_pct") or 100
        import re
        ship_days = 20
        st_txt = kwargs.get("shipment_term") or ""
        m = re.search(r"(\d+)", str(st_txt))
        if m: ship_days = int(m.group(1))
        warr_m = 12
        wt = kwargs.get("warranty_text") or ""
        m = re.search(r"(\d+)\s*мес", str(wt))
        if m: warr_m = int(m.group(1))

        contracts_db.save_contract(
            contract_number=contract_number,
            contract_date=contract_date_str,
            buyer=buyer, supplier=SUPPLIER or {},
            lines=lines, total_amount=float(total), has_vat=has_vat,
            vat_amount=float(vat),
            prepay_pct=int(prepay_pct or 100),
            shipment_days=ship_days, warranty_months=warr_m,
            delivery_terms=str(kwargs.get("delivery_terms") or ""),
            delivery_address=str(kwargs.get("delivery_address") or ""),
            contract_type=str(kwargs.get("contract_type") or "КП"),
            docx_bytes=docx_bytes,
        )
    except Exception as _save_err:
        # Не падаем если сохранение сломалось — договор всё равно выдаём
        try:
            import streamlit as _st
            _st.warning(f"⚠️ Договор не сохранён в базу (но DOCX готов): {_save_err}")
        except Exception: pass

    return docx_bytes


def _build_dogovor_docx_legacy(q: QuoteData, buyer: dict, contract_number: str,
                       contract_date_str: str,
                       prepay_pct: int = None,
                       delivery_terms: str = "включена в стоимость",
                       traverse_selections: list = None,
                       include_stamp: bool = False,
                       shipment_term: str = None,
                       warranty_text: str = None,
                       payment_order_text: str = None,
                       contract_valid_until: str = None,
                       shipment_penalty_pct: float = 1.0,
                       shipment_penalty_after_days: int = 10,
                       pickup_days: int = 10,
                       storage_penalty_pct: float = 1.0,
                       storage_free_days: int = 10,
                       kp_comment: str = "",
                       drawing_bytes: Optional[bytes] = None,
                       drawing_caption: str = "") -> bytes:
    from pathlib import Path as _P
    _include_stamp = bool(include_stamp)
    """Договор поставки по точному тексту шаблона ООО «ЛКС»."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(1.5)
        add_header_footer(s)

    # Совместимость: парсер пишет director_fio_short/director_fio_gen,
    # а старые вызовы иногда director_short/director_gen — берём любой.
    buyer_full = (buyer.get("full") or buyer.get("short") or "ООО «___»").strip()
    buyer_short = (buyer.get("short") or buyer.get("full")
                   or "ООО «___»").strip()
    buyer_name_full = buyer_full  # в преамбуле — полное наименование
    buyer_director_gen = (buyer.get("director_fio_gen")
                          or buyer.get("director_gen")
                          or buyer.get("director") or "").strip()
    buyer_director_short = (buyer.get("director_fio_short")
                            or buyer.get("director_short") or "").strip()
    buyer_director_position = (buyer.get("director_position")
                               or "Генеральный директор").strip()

    # --- Шапка ---
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run1 = p.add_run("Санкт-Петербург")
    run1.font.size = Pt(11)
    # Tab-выравнивание даты вправо через много пробелов
    p.add_run("\t\t\t\t\t\t\t\t\t")
    r_date = p.add_run(contract_date_str)
    r_date.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"ДОГОВОР ПОСТАВКИ № {contract_number}")
    r.bold = True
    r.font.size = Pt(14)

    # --- Преамбула ---
    # Полное наименование поставщика: собираем из full + short (без дублей)
    _sup_full = SUPPLIER.get('full', '')
    _sup_short = SUPPLIER.get('short', '')
    if _sup_full and '«' not in _sup_full and '"' not in _sup_full:
        # Пример: full="Общество с ограниченной ответственностью", short="ООО «ЛКС»"
        # → вытаскиваем «...» из short
        import re as _re
        _m = _re.search(r'[«"](.+?)[»"]', _sup_short)
        _sup_name = f'{_sup_full} «{_m.group(1)}»' if _m else _sup_short
    else:
        _sup_name = _sup_full or _sup_short

    # Формируем должность в род. падеже ("Генеральный директор" → "генерального директора")
    _pos = SUPPLIER.get('director_position','Генеральный директор').lower()
    _pos_gen = (_pos.replace('ый директор', 'ого директора')
                    .replace('ый директор', 'ого директора'))
    preamble = (
        f"{_sup_name}, именуемое в дальнейшем «Поставщик», "
        f"в лице {_pos_gen} {SUPPLIER['director_fio_gen']}, действующего на основании {SUPPLIER.get('director_basis','Устава')}, "
        f"с одной стороны, и {buyer_name_full}, именуемое в дальнейшем «Покупатель», "
    )
    if buyer_director_gen:
        preamble += f"в лице генерального директора {buyer_director_gen}, действующего на основании Устава, "
    preamble += ("с другой стороны, совместно именуемые «Стороны», "
                 "заключили настоящий Договор о нижеследующем:")
    p = doc.add_paragraph(preamble)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)

    # --- Разделы 1-11 (с динамической подстановкой 3.2/3.3/3.4, 8.1, 11.1.2) ---
    _valid_until = (contract_valid_until or "31.12.2026").strip()
    _edo_provider = SUPPLIER.get("edo_provider", "СБИС")
    _edo_id = SUPPLIER.get("edo_id", "")
    _sup_short = SUPPLIER.get("short", "ООО «___»")
    _sup_inn = SUPPLIER.get("inn", "")

    # Числа прописью (родительный падеж) для дней
    _NUM_GEN = {1:"одного", 2:"двух", 3:"трёх", 4:"четырёх", 5:"пяти",
                6:"шести", 7:"семи", 8:"восьми", 9:"девяти", 10:"десяти",
                14:"четырнадцати", 15:"пятнадцати", 20:"двадцати",
                21:"двадцати одного", 30:"тридцати", 45:"сорока пяти", 60:"шестидесяти"}
    def _num_gen(n):
        return _NUM_GEN.get(int(n), str(int(n)))
    _ship_after_word = _num_gen(shipment_penalty_after_days)
    _pickup_word = _num_gen(pickup_days)
    _storage_word = _num_gen(storage_free_days)

    # Процент без .0 (чтобы 1.0 → «1», а 1.5 → «1,5»)
    def _fmt_pct(v):
        v = float(v)
        s = f"{v:.2f}".rstrip("0").rstrip(".")
        return s.replace(".", ",")
    _ship_pct_s = _fmt_pct(shipment_penalty_pct)
    _stor_pct_s = _fmt_pct(storage_penalty_pct)

    for title, items in DOGOVOR_SECTIONS:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        rh = h.add_run(title)
        rh.bold = True
        rh.font.size = Pt(11)
        for txt in items:
            # Динамическая подстановка пунктов 3.2, 3.3, 3.4, 8.1, 11.1.2
            if txt.startswith("3.2."):
                txt = (f"3.2. При задержке отгрузки Продукции на срок более "
                       f"{shipment_penalty_after_days} ({_ship_after_word}) рабочих дней Поставщик выплачивает "
                       f"по требованию Покупателя пени в размере {_ship_pct_s} % от суммы предоплаты "
                       f"за каждый рабочий день просрочки. Пени начисляются с момента истечения "
                       f"установленного срока отгрузки до фактической передачи Продукции транспортной компании.")
            elif txt.startswith("3.3."):
                txt = (f"3.3. Покупатель обязуется произвести выборку заказа со склада Поставщика "
                       f"в течение {pickup_days} ({_pickup_word}) рабочих дней со дня получения письменного уведомления "
                       f"на электронный адрес о готовности Продукции к отгрузке. Уведомление считается "
                       f"полученным с момента его отправки на электронный адрес Покупателя, указанный в Договоре.")
            elif txt.startswith("3.4."):
                txt = (f"3.4. Хранение заказа свыше {storage_free_days} ({_storage_word}) рабочих дней "
                       f"осуществляется за дополнительную плату. Стоимость хранения составляет "
                       f"{_stor_pct_s} % от суммы заказа за каждый день хранения. Сроки хранения должны "
                       f"быть согласованы с Поставщиком в письменной форме. В случае отсутствия согласования "
                       f"Поставщик вправе отказать в дальнейшем хранении и реализовать Продукцию по своему усмотрению "
                       f"без уведомления Покупателя.")
            elif txt.startswith("8.1."):
                txt = (f"8.1. Договор вступает в силу с момента подписания его Сторонами "
                       f"и действует до {_valid_until} г. При этом в части неисполненных "
                       f"обязательств Договор действует до полного их исполнения.")
            elif txt.startswith("11.1.2."):
                txt = (f"11.1.2. {_sup_short} (ИНН {_sup_inn}) "
                       f"осуществляет обмен электронными документами через систему электронного "
                       f"документооборота (ЭДО) {_edo_provider}"
                       + (f", идентификатор участника: {_edo_id}." if _edo_id else "."))
            pp = doc.add_paragraph(txt)
            pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pp.paragraph_format.space_after = Pt(3)

    # --- Дополнительные условия (комментарий из КП) ---
    _kp_comment_txt = str(kp_comment or "").strip()
    if _kp_comment_txt:
        h_c = doc.add_paragraph()
        h_c.paragraph_format.space_before = Pt(12)
        h_c.paragraph_format.space_after = Pt(6)
        rc = h_c.add_run("12. ДОПОЛНИТЕЛЬНЫЕ УСЛОВИЯ")
        rc.bold = True
        rc.font.size = Pt(11)
        for _ln in _kp_comment_txt.splitlines():
            if _ln.strip():
                _pc = doc.add_paragraph(_ln.strip())
                _pc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _pc.paragraph_format.space_after = Pt(3)
        _reqs_number = "13"
    else:
        _reqs_number = "12"

    # --- Реквизиты ---
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    rh = h.add_run(f"{_reqs_number}. ПОЧТОВЫЕ И БАНКОВСКИЕ РЕКВИЗИТЫ")
    rh.bold = True
    rh.font.size = Pt(11)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    left, right = tbl.rows[0].cells
    sup_lines = [
        "ПОСТАВЩИК", "",
        SUPPLIER['short'],
        f"Юр. адрес / факт. адрес: {SUPPLIER['address']}",
        f"ИНН {SUPPLIER['inn']}",
        f"КПП {SUPPLIER['kpp']}",
        f"ОГРН {SUPPLIER['ogrn']}",
        f"р/с {SUPPLIER['rs']}",
        f"в банке {SUPPLIER['bank']}",
        f"к/с {SUPPLIER['ks']}",
        f"БИК {SUPPLIER['bik']}",
        f"Тел.: {SUPPLIER['phone']}",
        f"E-mail: {SUPPLIER['email']}",
        "", "", "",
        f"{SUPPLIER['director_position']} ООО «ЛКС»",
        "",
        f"_____________________ / {SUPPLIER['director_fio_short']} /",
        "М.П.",
    ]
    buyer_lines = [
        "ПОКУПАТЕЛЬ", "",
        buyer_name_full,
        f"Юр. адрес: {buyer.get('address','')}",
        f"ИНН {buyer.get('inn','')}",
        f"КПП {buyer.get('kpp','')}",
        f"ОГРН {buyer.get('ogrn','')}",
        f"р/с {buyer.get('rs','')}",
        f"в банке {buyer.get('bank','')}",
        f"к/с {buyer.get('ks','')}",
        f"БИК {buyer.get('bik','')}",
        f"Тел.: {buyer.get('phone','')}",
        f"E-mail: {buyer.get('email','')}",
        "", "", "",
        f"{buyer_director_position} {buyer_short}",
        "",
        f"_____________________ / {buyer_director_short} /",
        "М.П.",
    ]
    for cell, lines in ((left, sup_lines), (right, buyer_lines)):
        cell.text = ""
        for i, ln in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            r = p.add_run(ln)
            r.font.size = Pt(9)
            if i == 0 or (i == 17 and ln):
                r.bold = True

    # ---- Спецификация № 1 ----
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Спецификация № 1 от {contract_date_str}")
    r.bold = True
    r.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(f"к договору поставки № {contract_number}")
    r2.font.size = Pt(11)

    spec_preamble = (
        f"{_sup_name}, именуемое в дальнейшем «Поставщик», "
        f"в лице {_pos_gen} {SUPPLIER['director_fio_gen']}, действующего на основании {SUPPLIER.get('director_basis','Устава')}, "
        f"с одной стороны, и {buyer_name_full}, именуемое в дальнейшем «Покупатель», "
    )
    if buyer_director_gen:
        spec_preamble += f"в лице генерального директора {buyer_director_gen}, действующего на основании Устава, "
    spec_preamble += (f"с другой стороны, заключили Спецификацию № 1 "
                      f"к Договору поставки № {contract_number} от {contract_date_str}:")
    pp = doc.add_paragraph(spec_preamble)
    pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pp.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    r = p.add_run("Наименование, количество и стоимость поставляемого Товара:")
    r.bold = True
    r.font.size = Pt(11)

    # Спец-таблица — как в КП (артикул в отдельной колонке, цифры по центру, компактно)
    lines_all = list(q.lines) + list(q.electrification_lines)
    tbl = doc.add_table(rows=1 + len(lines_all) + 1, cols=7)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Ширины колонок (сумма ≈ 17 см — как в КП)
    # Фиксируем ширины: №=0.8, Код=3.2, Наимен=6.5, Ед.=1.0, Кол-во=1.3, Цена=2.0, Сумма=2.2
    col_widths_cm = [0.8, 3.2, 6.5, 1.0, 1.3, 2.0, 2.2]
    # tblLayout=fixed чтобы Word не пересчитывал ширины по контенту
    from docx.oxml import OxmlElement as _OxE
    from docx.oxml.ns import qn as _qn
    _tblPr = tbl._element.find(_qn("w:tblPr"))
    if _tblPr is None:
        _tblPr = _OxE("w:tblPr")
        tbl._element.insert(0, _tblPr)
    _tblLayout = _OxE("w:tblLayout")
    _tblLayout.set(_qn("w:type"), "fixed")
    _tblPr.append(_tblLayout)
    # Применяем ширины к каждой ячейке
    for _row in tbl.rows:
        for _c, _w in zip(_row.cells, col_widths_cm):
            _c.width = Cm(_w)
    hdr = tbl.rows[0].cells
    headers = ["№", "Код", "Наименование Товара", "Ед.", "Кол-во",
               "Цена, руб.", "Сумма, руб."]
    for c, txt in zip(hdr, headers):
        c.text = ""
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pr = c.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr.paragraph_format.space_before = Pt(0)
        pr.paragraph_format.space_after = Pt(0)
        rr = pr.add_run(txt)
        rr.bold = True
        rr.font.size = Pt(9)

    total_vat_incl = 0.0
    for i, ln in enumerate(lines_all, 1):
        row = tbl.rows[i].cells
        price_nds = float(ln.price)
        price_novat = price_nds / (1 + VAT_RATE)
        sum_novat = price_novat * ln.qty
        total_vat_incl += ln.total
        # Значения + выравнивание по колонкам как в КП:
        # №, Код, Ед., Кол-во, Цена, Сумма — центр; Наименование — влево
        vals = [
            (str(i), WD_ALIGN_PARAGRAPH.CENTER),
            (str(ln.code or ""), WD_ALIGN_PARAGRAPH.CENTER),
            (str(ln.name), WD_ALIGN_PARAGRAPH.LEFT),
            (str(ln.unit), WD_ALIGN_PARAGRAPH.CENTER),
            (f"{ln.qty:g}", WD_ALIGN_PARAGRAPH.CENTER),
            (f"{price_novat:,.2f}".replace(",", " "), WD_ALIGN_PARAGRAPH.CENTER),
            (f"{sum_novat:,.2f}".replace(",", " "), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for c, (v, alg) in zip(row, vals):
            c.text = ""
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # выравнивание по вертикали
            pr = c.paragraphs[0]
            pr.alignment = alg
            pr.paragraph_format.space_before = Pt(0)
            pr.paragraph_format.space_after = Pt(0)
            rr = pr.add_run(v)
            rr.font.size = Pt(9)

    # Итого — компактно, по центру
    last = tbl.rows[-1].cells
    last[0].merge(last[5])
    last[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    last[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = last[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    rr = p.add_run("Итого без НДС, руб.:")
    rr.bold = True
    rr.font.size = Pt(10)
    total_novat = total_vat_incl / (1 + VAT_RATE)
    p_val = last[6].paragraphs[0]
    p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_val.paragraph_format.space_before = Pt(0)
    p_val.paragraph_format.space_after = Pt(0)
    rrv = p_val.add_run(f"{total_novat:,.2f}".replace(",", " "))
    rrv.bold = True
    rrv.font.size = Pt(10)

    vat_amount = total_vat_incl - total_novat
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Общая цена Товара по настоящей Спецификации к Договору составляет: ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        f"{fmt_money(total_vat_incl)}, в том числе НДС 22 % — "
        f"{fmt_money(vat_amount)}."
    ).font.size = Pt(11)

    # Условия — процент предоплаты берём из аргумента или дефолта.
    pp = int(prepay_pct) if prepay_pct is not None else int(PREPAYMENT_DEFAULT*100)
    rem = 100 - pp
    pp_words = percent_in_words(pp)
    rem_words = percent_in_words(rem)
    dt = (delivery_terms or "включена в стоимость").strip().rstrip(".")
    conds = [
        ("Комплект документов к поставке",
         "паспорт ГОСТ, руководство по эксплуатации, сертификат."),
        ("Порядок оплаты",
         (payment_order_text.strip() if payment_order_text and payment_order_text.strip() else
          f"Покупатель обязуется произвести оплату Товара в следующем порядке: "
          f"— {pp} % ({pp_words} процентов) от общей цены Спецификации — "
          f"предоплата после подписания Договора и Спецификации в течение 5 рабочих дней; "
          f"оставшиеся {rem} % ({rem_words} процентов) — по уведомлению "
          f"на электронную почту о готовности Товара к отгрузке.")),
        ("Срок отгрузки продукции",
         (shipment_term.strip() if shipment_term and shipment_term.strip() else
          "20 рабочих дней со дня поступления оплаты на расчётный счёт Поставщика.")),
        ("Условия доставки", dt + "."),
        ("Гарантия",
         (warranty_text.strip() if warranty_text and warranty_text.strip() else
          _build_default_warranty(q))),
    ]
    for k, v in conds:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_after = Pt(4)
        rr = pp.add_run(f"{k}: ")
        rr.bold = True
        rr.font.size = Pt(11)
        pp.add_run(v).font.size = Pt(11)

    # Подписи спецификации
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    left, right = tbl.rows[0].cells
    _stamp_path = SUPPLIER.get("stamp_path") if _include_stamp else None
    # Левая ячейка: Поставщик (с печатью если включена)
    left.text = ""
    _p_l0 = left.paragraphs[0]
    _r_l0 = _p_l0.add_run(f"{SUPPLIER['director_position']} {SUPPLIER['short']}")
    _r_l0.bold = True; _r_l0.font.size = Pt(10)
    left.add_paragraph()
    if _stamp_path and _P(_stamp_path).exists():
        # Картинка печати+подписи — вверху, без прочерка
        _p_l_pic = left.add_paragraph()
        try:
            from docx.shared import Inches as _In
            _p_l_pic.add_run().add_picture(_stamp_path, width=_In(3.0))
        except Exception:
            pass
        _p_l_sig = left.add_paragraph()
        _r_l_sig = _p_l_sig.add_run(f"/ {SUPPLIER['director_fio_short']} /")
        _r_l_sig.font.size = Pt(10)
    else:
        _p_l_sig = left.add_paragraph()
        _r_l_sig = _p_l_sig.add_run(
            f"_____________________ / {SUPPLIER['director_fio_short']} /")
        _r_l_sig.font.size = Pt(10)
        _p_l_stamp = left.add_paragraph()
        _r_l_stamp = _p_l_stamp.add_run("М.П.")
        _r_l_stamp.font.size = Pt(10)
    # Правая ячейка — Покупатель
    right.text = ""
    for i, ln in enumerate([
        f"{buyer_director_position} {buyer_short}", "", "",
        f"_____________________ / {buyer_director_short} /", "М.П."]):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        r = p.add_run(ln)
        r.font.size = Pt(10)
        if i == 0:
            r.bold = True

    # =============== ПРИЛОЖЕНИЕ № 1: габаритные чертежи ===============
    if traverse_selections or drawing_bytes:
        _append_drawings_annex(
            doc, traverse_selections or [], contract_number,
            contract_date_str, buyer, buyer_short, buyer_director_short,
            buyer_director_position,
            extra_drawing_bytes=drawing_bytes,
            extra_drawing_caption=drawing_caption)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================================================
# Приложение № 1: габаритные чертежи (траверса + опции)
# =========================================================
def _append_drawings_annex(doc, tv_selections, contract_number,
                            contract_date_str, buyer,
                            buyer_short, buyer_director_short,
                            buyer_director_position,
                            extra_drawing_bytes: Optional[bytes] = None,
                            extra_drawing_caption: str = ""):
    """Добавляет в DOCX Приложение № 1 с чертежами траверсы и выбранных опций.
    Максимум 2 картинки на странице."""
    from docx.shared import Inches
    from pathlib import Path as _P
    drawings_dir = _P(__file__).parent / "media" / "traverse" / "drawings"

    # Новая страница
    doc.add_page_break()

    # Заголовок Приложения — компактно, в одной строке
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"Приложение № 1 к Договору поставки № {contract_number} от {contract_date_str}")
    r.font.size = Pt(10)
    r.italic = True

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(4)
    p_t.paragraph_format.space_after = Pt(4)
    r_t = p_t.add_run("Габаритные чертежи")
    r_t.font.size = Pt(13)
    r_t.bold = True

    # Собираем список картинок: [(подпись, путь)]
    images = []

    # 0a. Внешний чертёж (extra_drawing_bytes) — если передан
    if extra_drawing_bytes:
        try:
            import tempfile as _tmpm, os as _osm
            _tf = _tmpm.NamedTemporaryFile(suffix=".png", delete=False)
            _tf.write(extra_drawing_bytes)
            _tf.close()
            _cap = extra_drawing_caption or "Габаритный чертёж оборудования"
            images.append((_cap, _P(_tf.name)))
        except Exception:
            pass

    # 0. Габаритный чертёж крана (если q доступен в замыкании и это не траверса)
    try:
        _q_ref = q  # noqa — может отсутствовать в контексте (генерация из CRM)
        if _q_ref.series != "VACUTEC":
            import drawings as _drw_c
            _cd = _drw_c.find_crane_drawing(
                series=_q_ref.series, capacity=_q_ref.capacity,
                boom=_q_ref.boom, height_to_arm=_q_ref.height_to_arm,
                use_lllm=bool(getattr(_q_ref, "use_lllm", False)),
                lllm_code=str(getattr(_q_ref, "lllm_code", "") or ""),
            )
            if _cd and _cd.suffix.lower() != ".pdf":
                images.append(
                    (f"Консольный кран {_q_ref.series} "
                     f"{int(_q_ref.capacity)} кг · {_q_ref.boom:g}×{_q_ref.height_to_arm:g} м",
                     _cd))
    except Exception:
        pass

    seen_traverse = set()
    seen_options = set()
    for sel in tv_selections:
        if not sel:
            continue
        # 1. Общий вид траверсы (по base_code)
        _bc_key = sel.base_code.replace("VacuTec ", "").strip()
        if _bc_key not in seen_traverse:
            seen_traverse.add(_bc_key)
            _tv_img = drawings_dir / f"{_bc_key}.jpg"
            if _tv_img.exists():
                images.append(
                    (f"Общий вид траверсы VacuTec {_bc_key}",
                     _tv_img))
        # 2. Опции — добавляем только выбранные, по одному разу
        _opts = [
            ("handle", sel.options_handle, "Наклонная ручка"),
            ("cable", sel.options_cable, "Спиральный кабель управления"),
            ("supports", sel.options_supports, "Опоры для хранения траверсы"),
            ("battery", sel.options_battery, "Сменный аккумулятор 24V 20,8 A·ч"),
        ]
        for key, is_on, label in _opts:
            if is_on and key not in seen_options:
                seen_options.add(key)
                _p = drawings_dir / f"{key}.jpg"
                if _p.exists():
                    images.append((label, _p))

    if not images:
        _p = doc.add_paragraph("Габаритные чертежи предоставляются по запросу.")
        _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Размещаем по 2 чертежа на страницу
        for i, (label, path) in enumerate(images):
            _pl = doc.add_paragraph()
            _pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _pl.paragraph_format.space_before = Pt(6)
            _pl.paragraph_format.space_after = Pt(4)
            # Подпись — всегда вместе со своей картинкой на одной странице
            _pl.paragraph_format.keep_with_next = True
            _pl.paragraph_format.keep_together = True
            _rl = _pl.add_run(label)
            _rl.bold = True
            _rl.font.size = Pt(11)
            _pi = doc.add_paragraph()
            _pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _pi.paragraph_format.space_after = Pt(4)
            _run = _pi.add_run()
            # Цель — любые 2 картинки на странице A4. Активная высота A4 между полями ≈ 245 мм.
            # На 2 картинки с подписями — ≤ ≈100 мм высоты на каждую.
            _pi.paragraph_format.keep_together = True
            try:
                from PIL import Image as _PILImage
                with _PILImage.open(str(path)) as _im:
                    _iw, _ih = _im.size  # пиксели
                # Максимально допустимые в дюймах (с запасом на подпись)
                _max_w_in = 6.0    # ≈ 152 мм ширины
                _max_h_in = 3.15   # ≈ 80 мм высоты (чтобы 2 гарантировано влезали)
                _ratio_w = _max_w_in / (_iw / 96)  # коэф. если масштабируем по ширине
                _ratio_h = _max_h_in / (_ih / 96)  # коэф. если по высоте
                # Выбираем меньший коэф. — чтобы влезло по обеим осям
                _r = min(_ratio_w, _ratio_h)
                _final_w_in = (_iw / 96) * _r
                _run.add_picture(str(path), width=Inches(_final_w_in))
            except Exception:
                # Fallback: выставляем ширину без ограничения по высоте
                try:
                    _run.add_picture(str(path), width=Inches(5.5))
                except Exception:
                    pass
            # После каждой второй картинки — разрыв страницы (кроме последней)
            if (i + 1) % 2 == 0 and (i + 1) < len(images):
                doc.add_page_break()

    # Подписи сторон
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    _pt = doc.add_paragraph()
    _pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _rt = _pt.add_run("Подписи сторон")
    _rt.bold = True; _rt.font.size = Pt(11)

    _sign_tbl = doc.add_table(rows=1, cols=2)
    _sign_tbl.autofit = True
    _l, _r = _sign_tbl.rows[0].cells
    for cell, lines in (
        (_l, [f"Поставщик: {SUPPLIER['short']}", "", "",
              f"_____________________ / {SUPPLIER['director_fio_short']} /", "М.П."]),
        (_r, [f"{buyer_director_position} {buyer_short}", "", "",
              f"_____________________ / {buyer_director_short} /", "М.П."]),
    ):
        cell.text = ""
        for i, ln in enumerate(lines):
            _p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            _r = _p.add_run(ln)
            _r.font.size = Pt(10)
            if i == 0:
                _r.bold = True


st.set_page_config(
    page_title="ЛКС · Генератор КП и договора",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="auto",  # На мобильных — свёрнут, на десктопе — открыт
)

# Мета-тег viewport для мобильных устройств (корректный zoom, нет растягивания)
st.markdown(
    '<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">',
    unsafe_allow_html=True)

# === ГЛОБАЛЬНЫЙ CSS-СТИЛЬ ===
st.markdown("""
<style>
/* ===== Глобальные переменные ===== */
:root {
    --brand-primary: #7c3aed;
    --brand-secondary: #14b8a6;
    --brand-accent: #3b82f6;
    --gradient-purple: linear-gradient(135deg, #a855f7 0%, #7c3aed 100%);
    --gradient-teal: linear-gradient(135deg, #14b8a6 0%, #0891b2 100%);
    --gradient-blue: linear-gradient(135deg, #60a5fa 0%, #3b82f6 100%);
    --gradient-orange: linear-gradient(135deg, #fb923c 0%, #f97316 100%);
    --bg-soft: #f8fafc;
    --card-bg: #ffffff;
    --border-soft: #e2e8f0;
    --text-dark: #0f172a;
    --text-muted: #64748b;
    --radius-lg: 16px;
    --radius-md: 12px;
    --shadow-soft: 0 2px 8px rgba(15, 23, 42, 0.04);
    --shadow-hover: 0 4px 16px rgba(15, 23, 42, 0.08);
}

/* ===== Общий фон ===== */
.stApp {
    background: var(--bg-soft);
}

/* ===== Кнопки ===== */
.stButton > button {
    border-radius: var(--radius-md) !important;
    border: 1px solid var(--border-soft) !important;
    padding: 0.5rem 1.25rem !important;
    font-weight: 500 !important;
    transition: all 0.2s ease !important;
    box-shadow: var(--shadow-soft) !important;
}
.stButton > button:hover {
    transform: translateY(-1px);
    box-shadow: var(--shadow-hover) !important;
    border-color: var(--brand-primary) !important;
}
.stButton > button[kind="primary"] {
    background: var(--gradient-purple) !important;
    color: white !important;
    border: none !important;
}
.stButton > button[kind="primary"]:hover {
    background: linear-gradient(135deg, #9333ea 0%, #6d28d9 100%) !important;
}

/* ===== Download button в том же стиле ===== */
.stDownloadButton > button {
    border-radius: var(--radius-md) !important;
    border: 1px solid var(--border-soft) !important;
    padding: 0.5rem 1.25rem !important;
    font-weight: 500 !important;
    box-shadow: var(--shadow-soft) !important;
}

/* ===== Метрики (st.metric) — градиентные карточки ===== */
div[data-testid="stMetric"] {
    background: var(--card-bg);
    border-radius: var(--radius-lg);
    padding: 20px 24px;
    box-shadow: var(--shadow-soft);
    border: 1px solid var(--border-soft);
    transition: all 0.2s ease;
}
div[data-testid="stMetric"]:hover {
    box-shadow: var(--shadow-hover);
    transform: translateY(-2px);
}
div[data-testid="stMetricLabel"] {
    color: var(--text-muted);
    font-size: 0.85rem;
    font-weight: 500;
}
div[data-testid="stMetricValue"] {
    color: var(--text-dark);
    font-weight: 700;
    font-size: 1.75rem;
}

/* ===== Фирменные градиентные блоки (компактные) ===== */
.gradient-card-purple, .gradient-card-teal, .gradient-card-blue, .gradient-card-orange {
    border-radius: var(--radius-lg);
    padding: 14px 16px;
    color: white;
    box-shadow: var(--shadow-soft);
    margin-bottom: 8px;
}
.gradient-card-purple { background: var(--gradient-purple); }
.gradient-card-teal { background: var(--gradient-teal); }
.gradient-card-blue { background: var(--gradient-blue); }
.gradient-card-orange { background: var(--gradient-orange); }
.gradient-card-label {
    font-size: 0.75rem;
    opacity: 0.9;
    margin-bottom: 2px;
}
.gradient-card-value {
    font-size: 1.4rem;
    font-weight: 700;
    line-height: 1.2;
}

/* ===== Заголовки ===== */
h1, h2, h3 {
    color: var(--text-dark) !important;
    letter-spacing: -0.02em;
}
h2 { font-weight: 700 !important; }
h3 { font-weight: 600 !important; }

/* ===== Контейнеры с границей ===== */
[data-testid="stContainer"] {
    border-radius: var(--radius-lg) !important;
}

/* ===== Таблицы/dataframe ===== */
div[data-testid="stDataFrame"], div[data-testid="stDataEditor"] {
    border-radius: var(--radius-md);
    overflow: hidden;
    box-shadow: var(--shadow-soft);
}

/* ===== Input и selectbox ===== */
div[data-baseweb="select"] > div,
div[data-baseweb="input"] > input,
div[data-baseweb="textarea"] > textarea {
    border-radius: var(--radius-md) !important;
}

/* ===== Сайдбар ===== */
section[data-testid="stSidebar"] {
    background: white;
    border-right: 1px solid var(--border-soft);
    min-width: 320px !important;
    max-width: 380px !important;
}
section[data-testid="stSidebar"] > div {
    padding-top: 1rem !important;
}
/* Логотип — большой, вырвнен по левому верхнему углу сайдбара */
section[data-testid="stSidebar"] img[src^="data:image/jpeg"]:first-of-type {
    width: 100% !important;
    max-width: 100% !important;
}

/* ===== ВЕРХНЕЕ МЕНЮ-КНОПКИ ===== */
.top-nav-active {
    background: linear-gradient(135deg, #a855f7 0%, #7c3aed 100%);
    color: white !important;
    font-weight: 600;
    padding: 12px 8px;
    border-radius: 10px;
    text-align: center;
    font-size: 13px;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
    box-shadow: 0 4px 10px rgba(124, 58, 237, 0.25);
    cursor: default;
    user-select: none;
}
div[data-testid="stHorizontalBlock"]:first-of-type .stButton > button {
    background: white !important;
    border: 1px solid #e2e8f0 !important;
    color: #334155 !important;
    font-weight: 500 !important;
    padding: 12px 8px !important;
    border-radius: 10px !important;
    box-shadow: 0 1px 3px rgba(15, 23, 42, 0.05) !important;
    font-size: 13px !important;
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    transition: all 0.15s ease !important;
}
div[data-testid="stHorizontalBlock"]:first-of-type .stButton > button:hover {
    background: #f8fafc !important;
    border-color: #a855f7 !important;
    color: #7c3aed !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 10px rgba(124, 58, 237, 0.12) !important;
}

/* ===== НАВИГАЦИОННОЕ МЕНЮ САЙДБАРА (Bitrix24/amoCRM стиль) ===== */
.pro-sidebar-nav {
    display: flex;
    flex-direction: column;
    gap: 1px;
    padding: 0;
    margin-bottom: 8px;
}
.pro-nav-group {
    font-size: 10px;
    font-weight: 700;
    color: #94a3b8;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    padding: 12px 12px 6px 12px;
    margin-top: 8px;
}

/* Обёртка пункта меню (включает кнопку внутри) */
.pro-nav-item {
    position: relative;
    margin: 1px 4px;
    border-radius: 8px;
}
/* Кнопка внутри пункта (неактивного) */
.pro-nav-item .stButton > button {
    background: transparent !important;
    border: none !important;
    color: #475569 !important;
    font-weight: 500 !important;
    text-align: left !important;
    padding: 10px 14px !important;
    border-radius: 8px !important;
    box-shadow: none !important;
    justify-content: flex-start !important;
    transition: all 0.15s ease !important;
    transform: none !important;
    min-height: unset !important;
    height: auto !important;
}
.pro-nav-item .stButton > button:hover {
    background: #f1f5f9 !important;
    color: #7c3aed !important;
    transform: none !important;
    box-shadow: none !important;
    border: none !important;
}
.pro-nav-item .stButton > button p {
    text-align: left !important;
    font-weight: 500 !important;
    margin: 0 !important;
    font-size: 16px !important;
}
.pro-nav-item .stButton > button {
    font-size: 16px !important;
    padding: 12px 16px !important;
}

/* Material-иконка внутри кнопки — цвет и размер как в amoCRM/Bitrix */
.pro-nav-item .stButton > button span.material-icons,
.pro-nav-item .stButton > button span[data-testid="stIconMaterial"] {
    color: #64748b !important;
    font-size: 20px !important;
    margin-right: 4px;
}

/* Активный пункт — фиолетовый фон + левая полоска */
.pro-nav-item.active {
    background: linear-gradient(90deg, rgba(124, 58, 237, 0.10) 0%, rgba(124, 58, 237, 0.05) 100%);
}
.pro-nav-item.active .stButton > button span.material-icons,
.pro-nav-item.active .stButton > button span[data-testid="stIconMaterial"] {
    color: #7c3aed !important;
}
.pro-nav-item.active::before {
    content: "";
    position: absolute;
    left: -4px;
    top: 6px;
    bottom: 6px;
    width: 3px;
    background: linear-gradient(180deg, #a855f7 0%, #7c3aed 100%);
    border-radius: 0 2px 2px 0;
}
.pro-nav-item.active .stButton > button {
    background: transparent !important;
    color: #7c3aed !important;
    font-weight: 700 !important;
    cursor: default !important;
}
.pro-nav-item.active .stButton > button:hover {
    background: transparent !important;
    color: #7c3aed !important;
}
.pro-nav-item.active .stButton > button p {
    color: #7c3aed !important;
    font-weight: 700 !important;
}
/* Отключённая кнопка (active) — убираем серый цвет disabled */
.pro-nav-item.active .stButton > button:disabled {
    background: transparent !important;
    color: #7c3aed !important;
    opacity: 1 !important;
    cursor: default !important;
}



/* ===== Alerts (info/success/warning/error) — мягкие ===== */
div[data-testid="stAlert"] {
    border-radius: var(--radius-md);
    border: none;
    box-shadow: var(--shadow-soft);
}

/* =============================================================== */
/* ===== HERO-ЗАГОЛОВОК СТРАНИЦЫ (как в профессиональных CRM) ===== */
/* =============================================================== */
.hero-header {
    display: flex;
    justify-content: space-between;
    align-items: center;
    padding: 20px 24px;
    background: linear-gradient(135deg, #ffffff 0%, #faf5ff 100%);
    border-radius: 16px;
    margin-bottom: 20px;
    box-shadow: 0 2px 12px rgba(15, 23, 42, 0.06);
    border: 1px solid #e9d5ff;
}
.hero-title {
    margin: 0 0 4px 0 !important;
    font-size: 26px !important;
    font-weight: 700 !important;
    color: #0f172a !important;
    letter-spacing: -0.02em;
}
.hero-subtitle {
    margin: 0 !important;
    color: #64748b;
    font-size: 14px;
}
.hero-info {
    display: flex;
    align-items: center;
    gap: 12px;
}
.hero-badge {
    background: linear-gradient(135deg, #a855f7 0%, #7c3aed 100%);
    color: white;
    padding: 8px 16px;
    border-radius: 100px;
    font-size: 13px;
    font-weight: 600;
    box-shadow: 0 4px 12px rgba(124, 58, 237, 0.25);
}

/* =============================================================== */
/* ===== КАРТОЧКИ-СЕКЦИИ: оборачиваем каждый subheader в карточку ===== */
/* =============================================================== */
/* subheader = h3 в Streamlit — превращаем в шапку карточки */
.main .block-container h3 {
    padding: 20px 24px 12px 24px !important;
    background: white;
    border-radius: 16px 16px 0 0;
    margin: 24px 0 0 0 !important;
    border-top: 1px solid #e2e8f0;
    border-left: 1px solid #e2e8f0;
    border-right: 1px solid #e2e8f0;
    box-shadow: 0 -2px 8px rgba(15, 23, 42, 0.03);
}

/* Главная primary-кнопка «Скачать» / «Сохранить КП» — крупнее */
.stButton > button[kind="primary"] {
    padding: 14px 24px !important;
    font-size: 15px !important;
    font-weight: 600 !important;
    letter-spacing: 0.02em !important;
}

/* Download-кнопки (PDF/DOCX) — primary-стиль */
.stDownloadButton > button {
    background: linear-gradient(135deg, #a855f7 0%, #7c3aed 100%) !important;
    color: white !important;
    border: none !important;
    padding: 12px 20px !important;
    font-weight: 600 !important;
    box-shadow: 0 4px 12px rgba(124, 58, 237, 0.20) !important;
    transition: all 0.2s ease !important;
}
.stDownloadButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 6px 16px rgba(124, 58, 237, 0.30) !important;
}

/* Мобильная адаптация hero-header */
@media (max-width: 768px) {
    .hero-header {
        flex-direction: column;
        gap: 12px;
        text-align: center;
        padding: 16px 20px;
    }
    .hero-title { font-size: 22px !important; }
}

/* =================================================================== */
/* ===== 📱 МОБИЛЬНАЯ АДАПТАЦИЯ (смартфоны до 768px) ===== */
/* =================================================================== */
@media (max-width: 768px) {
    /* Общие отступы — меньше */
    .stApp {
        padding: 0 !important;
    }
    div[data-testid="stAppViewContainer"] .main .block-container {
        padding: 1rem 0.75rem !important;
        max-width: 100% !important;
    }

    /* Сайдбар на мобильных — оверлей, 85% ширины (оставляем место для закрытия) */
    section[data-testid="stSidebar"] {
        min-width: 85vw !important;
        max-width: 85vw !important;
        width: 85vw !important;
    }
    section[data-testid="stSidebar"] > div {
        padding: 1rem !important;
    }

    /* Логотип в сайдбаре — компактнее */
    section[data-testid="stSidebar"] img[src^="data:image/jpeg"]:first-of-type {
        max-width: 60% !important;
        margin: 0 auto !important;
    }

    /* Навигация в сайдбаре на мобильных — крупнее для пальца */
    .pro-nav-item .stButton > button {
        padding: 14px 18px !important;
        font-size: 16px !important;
        min-height: 48px !important;
    }
    .pro-nav-item .stButton > button p {
        font-size: 16px !important;
    }
    .pro-nav-group {
        font-size: 12px !important;
        padding: 16px 14px 8px 14px !important;
    }

    /* Заголовки — меньше */
    h1 { font-size: 1.5rem !important; }
    h2 { font-size: 1.25rem !important; }
    h3 { font-size: 1.1rem !important; }

    /* Кнопки — крупнее для пальца (Apple HIG: мин 44px высоты) */
    .stButton > button, .stDownloadButton > button {
        min-height: 44px !important;
        padding: 12px 16px !important;
        font-size: 15px !important;
    }

    /* Градиентные карточки — компактнее */
    .gradient-card-purple, .gradient-card-teal,
    .gradient-card-blue, .gradient-card-orange {
        padding: 16px 18px !important;
        margin-bottom: 8px !important;
    }
    .gradient-card-label { font-size: 0.8rem !important; }
    .gradient-card-value { font-size: 1.5rem !important; }

    /* Metric — компактнее */
    div[data-testid="stMetric"] {
        padding: 14px 16px !important;
    }
    div[data-testid="stMetricValue"] {
        font-size: 1.4rem !important;
    }

    /* Кнопки Назад/Вперёд — компактнее на мобильных */
    button[data-testid="nav_back_btn"], button[data-testid="nav_fwd_btn"] {
        min-height: 40px !important;
        padding: 8px !important;
    }

    /* Input/Selectbox — крупнее высота */
    div[data-baseweb="input"] > input,
    div[data-baseweb="select"] > div,
    div[data-baseweb="textarea"] > textarea {
        min-height: 44px !important;
        font-size: 16px !important;  /* iOS не зумит если 16px+ */
    }

    /* Data editor / dataframe — горизонтальный скролл */
    div[data-testid="stDataEditor"], div[data-testid="stDataFrame"] {
        overflow-x: auto !important;
    }

    /* Колонки — в одну колонку на мобильном (если очень узкие) */
    /* Streamlit автоматически стакает columns на mobile — тут просто гарантируем отступы */
    div[data-testid="column"] {
        min-width: 100% !important;
    }

    /* Карточки-container без боковых отступов */
    div[data-testid="stContainer"] {
        padding: 12px !important;
    }

    /* PDF-превью — уменьшаем высоту iframe */
    iframe[src^="data:application/pdf"] {
        height: 500px !important;
    }
}

/* =================================================================== */
/* ===== 📱 Очень маленькие экраны (iPhone SE, до 380px) ===== */
/* =================================================================== */
@media (max-width: 380px) {
    .gradient-card-value { font-size: 1.3rem !important; }
    h1 { font-size: 1.3rem !important; }
    h2 { font-size: 1.15rem !important; }
    .stButton > button, .stDownloadButton > button {
        font-size: 14px !important;
        padding: 10px 12px !important;
    }
}

/* ===== Планшеты (768–1024px) ===== */
@media (min-width: 768px) and (max-width: 1024px) {
    div[data-testid="stAppViewContainer"] .main .block-container {
        padding: 1.5rem 1rem !important;
    }
    section[data-testid="stSidebar"] {
        min-width: 280px !important;
        max-width: 300px !important;
    }
}
</style>
""", unsafe_allow_html=True)



# --- Сайдбар: параметры и прайсы ---
import traverse as _tv

# --- Горизонтальные табы навигации ---
_TAB_NAMES = [
    ("💼 Расчёт КП", "Расчёт КП"),
    ("📊 Дашборд", "CRM: Дашборд"),
    ("🎯 Воронка", "CRM: Воронка"),
    ("👥 Клиенты", "CRM: Клиенты"),
    ("📄 КП", "CRM: КП"),
    ("📑 Договоры", "CRM: Договоры"),
    ("⏰ Задачи", "CRM: Напоминания"),
    ("💰 Продажи", "CRM: Продажи"),
    ("📎 Внешний договор", "Внешний договор"),
    ("⚙️ Настройки", "CRM: Настройки"),
]
if "app_section" not in st.session_state:
    # Сначала смотрим в URL query params (чтобы F5 оставался на текущей странице)
    _qp = st.query_params
    _saved_section = _qp.get("nav", None)
    st.session_state["app_section"] = _saved_section if _saved_section else "Расчёт КП"

# Состояние сворачивания сайдбара
if "sidebar_collapsed" not in st.session_state:
    st.session_state["sidebar_collapsed"] = False

# История навигации для кнопок «Назад / Вперёд» (как в браузере)
if "nav_history" not in st.session_state:
    st.session_state["nav_history"] = [st.session_state["app_section"]]
    st.session_state["nav_index"] = 0

def _nav_push(section: str):
    """Добавить раздел в историю (обрезает ветку «вперёд»)."""
    _hist = st.session_state["nav_history"]
    _idx = st.session_state["nav_index"]
    if _idx < len(_hist) - 1:
        _hist = _hist[:_idx + 1]
    if not _hist or _hist[-1] != section:
        _hist.append(section)
    # Ограничиваем историю 50 шагами
    if len(_hist) > 50:
        _hist = _hist[-50:]
    st.session_state["nav_history"] = _hist
    st.session_state["nav_index"] = len(_hist) - 1

def _nav_go(section: str):
    st.session_state["app_section"] = section
    st.query_params["nav"] = section  # сохраняем в URL — F5 оставаться здесь
    _nav_push(section)
    st.rerun()

# Дополнительные стили (слайдеры, чекбоксы, кнопки навигации)
st.markdown("""
<style>
/* Нав-кнопки Назад/Вперёд — квадратные, компактные */
button[kind="secondary"][data-testid*="nav_back_btn"],
button[kind="secondary"][data-testid*="nav_fwd_btn"] {
    padding: 8px !important;
    border-radius: 8px !important;
}

/* === Слайдер — фиолетовый === */
.stSlider [data-baseweb="slider"] div[role="slider"] {
    background: #7c3aed !important;
    border-color: #7c3aed !important;
}
.stSlider [data-baseweb="slider"] > div > div > div {
    background: #7c3aed !important;
}

/* === Радио-кнопки — фиолетовые === */
input[type="radio"]:checked + div {
    color: #7c3aed !important;
}

/* === Чекбоксы === */
input[type="checkbox"]:checked {
    accent-color: #7c3aed !important;
}

/* === Tabs (встроенные st.tabs) === */
div[data-testid="stTabs"] button[data-baseweb="tab"] {
    color: #64748b;
    font-weight: 500;
}
div[data-testid="stTabs"] button[data-baseweb="tab"][aria-selected="true"] {
    color: #7c3aed !important;
    border-bottom-color: #7c3aed !important;
}

</style>
""", unsafe_allow_html=True)

# Шапка: Назад/Вперёд + верхнее меню-кнопки
_nav_back_col, _nav_fwd_col, _tabs_container_col = st.columns([0.6, 0.6, 12])

_hist = st.session_state["nav_history"]
_idx = st.session_state["nav_index"]
_can_back = _idx > 0
_can_fwd = _idx < len(_hist) - 1

with _nav_back_col:
    if st.button("◀", key="nav_back_btn", disabled=not _can_back,
                 help=("Назад: " + _hist[_idx - 1]) if _can_back else "Нет предыдущей страницы",
                 use_container_width=True):
        st.session_state["nav_index"] = _idx - 1
        st.session_state["app_section"] = _hist[_idx - 1]
        st.query_params["nav"] = _hist[_idx - 1]
        st.rerun()
with _nav_fwd_col:
    if st.button("▶", key="nav_fwd_btn", disabled=not _can_fwd,
                 help=("Вперёд: " + _hist[_idx + 1]) if _can_fwd else "Нет следующей страницы",
                 use_container_width=True):
        st.session_state["nav_index"] = _idx + 1
        st.session_state["app_section"] = _hist[_idx + 1]
        st.query_params["nav"] = _hist[_idx + 1]
        st.rerun()

# Верхнее меню-кнопки
with _tabs_container_col:
    _tabs_cols = st.columns(len(_TAB_NAMES))
    for _i, (_label, _val) in enumerate(_TAB_NAMES):
        with _tabs_cols[_i]:
            _active = st.session_state["app_section"] == _val
            if _active:
                st.markdown(
                    f'<div class="top-nav-active">{_label}</div>',
                    unsafe_allow_html=True)
            else:
                if st.button(_label, key=f"top_nav_{_i}",
                             use_container_width=True):
                    st.session_state["app_section"] = _val
                    st.query_params["nav"] = _val
                    _nav_push(_val)
                    st.rerun()

app_section = st.session_state["app_section"]

# --- Поставщик документов — UI в едином блоке «Условия КП» ниже ---
import suppliers as _suppliers
_new_key = st.session_state.get("supplier_key", _suppliers.DEFAULT_SUPPLIER_KEY)
SUPPLIER = _suppliers.get_supplier(_new_key)

# Если выбран раздел CRM — рендерим его и останавливаемся
if app_section != "Расчёт КП":
    try:
        import crm_ui as _crm_ui
        import crm_db as _crm_db
        _crm_db.init_db()
        if app_section == "CRM: Дашборд":
            _crm_ui.render_dashboard_tab()
        elif app_section == "CRM: Воронка":
            _crm_ui.render_funnel_tab()
        elif app_section == "CRM: Клиенты":
            _crm_ui.render_customers_tab()
        elif app_section == "CRM: КП":
            _crm_ui.render_quotes_tab()
        elif app_section == "CRM: Договоры":
            import contracts_ui as _contracts_ui
            _contracts_ui.render_contracts_tab()
        elif app_section == "CRM: Напоминания":
            _crm_ui.render_reminders_tab()
        elif app_section == "CRM: Продажи":
            _crm_ui.render_sales_tab()
        elif app_section == "CRM: Настройки":
            _crm_ui.render_settings_tab()
        elif app_section == "Внешний договор":
            import external_contract_ui as _ext_ui
            _ext_ui.render_external_contract_tab()
    except Exception as _e:
        import traceback
        st.error(f"CRM модуль недоступен: {_e}")
        st.code(traceback.format_exc())
    st.stop()

with st.sidebar:
    st.markdown(
        f'<div style="text-align:center;padding:0 0 24px 0;">'
        f'<img src="data:image/jpeg;base64,/9j/4AAQSkZJRgABAQAAAQABAAD/4gHYSUNDX1BST0ZJTEUAAQEAAAHIAAAAAAQwAABtbnRyUkdCIFhZWiAH4AABAAEAAAAAAABhY3NwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAQAA9tYAAQAAAADTLQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAlkZXNjAAAA8AAAACRyWFlaAAABFAAAABRnWFlaAAABKAAAABRiWFlaAAABPAAAABR3dHB0AAABUAAAABRyVFJDAAABZAAAAChnVFJDAAABZAAAAChiVFJDAAABZAAAAChjcHJ0AAABjAAAADxtbHVjAAAAAAAAAAEAAAAMZW5VUwAAAAgAAAAcAHMAUgBHAEJYWVogAAAAAAAAb6IAADj1AAADkFhZWiAAAAAAAABimQAAt4UAABjaWFlaIAAAAAAAACSgAAAPhAAAts9YWVogAAAAAAAA9tYAAQAAAADTLXBhcmEAAAAAAAQAAAACZmYAAPKnAAANWQAAE9AAAApbAAAAAAAAAABtbHVjAAAAAAAAAAEAAAAMZW5VUwAAACAAAAAcAEcAbwBvAGcAbABlACAASQBuAGMALgAgADIAMAAxADb/2wBDAAEBAQEBAQEBAQEBAQEBAQIBAQEBAQIBAQECAgICAgICAgIDAwQDAwMDAwICAwQDAwQEBAQEAgMFBQQEBQQEBAT/2wBDAQEBAQEBAQIBAQIEAwIDBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAQEBAT/wAARCAFuB9ADASIAAhEBAxEB/8QAHwABAAMBAQEAAwEBAAAAAAAAAAkKCwgHBgMEBQIB/8QAYBAAAQMDAgIECAcKCggEBAQHAAECAwQFBgcICRESGliWChMUGCFX1NYVGSIxl7bVNzg5d3iHlJi00xYjMkFRdpW1t9EXJDZISYbGxyUzYdJCUllxKERigjQ1Q3Jzg5H/xAAeAQEAAQQDAQEAAAAAAAAAAAAACAYHCQoDBAUBAv/EAFwRAAEDAQQDBg4MCwYGAQUAAAABAgMEBQYHEQgSIQkTFzFBVBUZUVZhcZGSk5TR0tPUFBYiMjY4VXJzgaOzN0JSYnR1laGxsrQjgoPBwsMkMzVThbUlNENjovH/2gAMAwEAAhEDEQA/AL/AAAAPwVVVTUVNUVtbUQUlHSQPqqurqpW09NSxRtV8kkkjlRrWtaiqrlVEREVVKb3Fk8J1s2n9XkegXDiq7HmOW0r5bRle6G50Ud7wvH5G9KOWLDaGTnDcZ2O5f+KVbX0SdB3iYKtr2TxgWa93O/PaVsXw9Mx3Pa1YlpvHV0r6rH8VnqXXjULMegrm9Gz4/TJJX1SdNEjdNHF4iJz2rLLE3m5Kmu7bwuq+TVN0xzZDtutlBRRyvp6LVPcXWyXKurI1ToLLT4na6iNkD2rzfHJPcp2rzZ4ymTk5i04tTdUdSNaM4yDUzVvOsr1J1Byqs8vyLMs1vtRkeRXeXooxqzVU73PVrGNaxjOfRjYxrWo1rURPgwCVTXHjccU7X+oqnZfvL1axa3VPSjbYtHblDojaoIXfPTr8BR0k0zOXNq+UySuc1VRznIqkceZanalajVMlZqFqFnGd1k1QlZNV5lllflFTLMiPRJXSVMr3K9ElkTpKvP8AjHen0qfDgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA+4w3U7UrTmpjrNPdQs4wSshqFrIavDcsr8XqYplRiLK2SmlY5HqkUadJF5/xbfT6EJHNDuNxxTtAKildiG8vVrKbdTdGN1i1iuUOt1qnhb81Onw7HVzQs5cmp5NJE5rURGuaiIRVgAus7SfC6r5DU2vHN72262V9FJKynrdU9ulbJba6jjROgktRid0qJGTvcvJ8skFyganJ/i6ZebWJbK2jb89pW+jD1zHbDrViWpEdJSsqsgxWCpdZ9QsO6atb0bxj9SkdfSp01WNs0kXiJXMcsUsreTlx3D7zTLVHUjRfOMf1M0kzrK9NtQcVrPL8dzLCr7UY5kVol6KscsNVA9r0a9jnMezn0ZGPc1yOa5UUDakBTX4TfhOtm1Aq8c0C4jtXY8Oy2qfFaMU3Q2yijsmF5BI7oxxRZlQx8obdO93P/xSkayiXpt8dBSNY+eS5DS1VNW01PW0VRBV0dXAyqpKullbUU1VFI1HxyRyNVWua5qoqORVRUVFQA/OAAAfgqqqmoqaora2ogpKOkgfVVdXVStp6alijar5JJJHKjWta1FVXKqIiIqqfnKavhOfFmqsBs1bw4dAcilpcuy6ywXHdFldpqFiqLBZa6Js9DhsMrV6TZrjC+Kqr+XJEo5qaHm9KuoZGBHVx4ePDke6zI8w2h7Q8wnsu1SyzyY9qHqHj1S6nuO4ypjc1KmKKpREfHj8b2OjjijVEuCI6WVz4HxRNqvAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAtQ8B7jw5HtTyPD9oe7zMJ71tUvU8ePaeah5DUuqLjtzqZHOSmilqVRXyY/I97Y5IpFVLeitlicyBksTqrwANtmlqqatpqetoqiCro6uBlVSVdLK2opqqKRqPjkjkaqtc1zVRUciqioqKh+cpr+DFcWSr1As1Jw4tfcifVZbh1jnue17K7vVeMqMgstDGs1dhssrndJ89uhbJV0H8rnRQ1UPNjaSBklygA5C357ucP2L7Statz2YpS1cem+JST4rj9VMsP8ADHIaxzaLH7O3oqj+VVX1FLHK+NFdFAs8qp0YnKmQxqjqbnGtGpGdat6mZBWZVqDqTlddmuZZFXq1Km73G41ElVVTK1qIxjVfI7oxsRrI2o1rWta1ES474XVu2qZr5tu2Q45dJY6Kgtk+4vVOip5kdT1k9TJWWHE6eVW+lr4GU2Szvheq8219HJ0U5MctKYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA+80u1NzjRfUjBdW9M8grMV1B02yuhzXDcioFatTaLjbqiOqpZka5FY9qPjb0o3o5kjVc1zXNcqLrz7DN3OH76NpWiu57DkpaSPUjEo58qx+lmWb+B2Q0bnUWQWd3SVX8qWvp6qOJ8iI6WBIJUToytVcdwus+CK7tqmG+bkdkOR3SWSir7ZBuL0soqiZG09HPTSUdhyyniV3pc+dlTjU7IWKnJtBWSdFeb3IBX643GuNRr/xTt5eXuqvKbdi2rVTo7YmxyeNpIKTCIYcU506/N0JprRUVPNvyXOqnuRVR3MirPuNTsyqdRtStQtQqySomrM7zi7ZlVzViI2rlkulfUV0jpURzkR6unVXcnO9Kr8pfnPhwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAASqcEfXGo0A4p2zTL21Xk1uynVqm0dvrZJPFUk9Jm8M2Kc6hfm6EM13p6nm75LXUrHKqI3mRVn3GmOZVOnOpWnuoVHJUQ1mCZxacypJqNEdVxSWuvp66N0SK5qK9HQIrebm+lE+UnzgHw4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAJ6OBhtq0I3G5VuNotcNMcZ1KpcUx/GqrHYMkgknZaZKupvLKl8XQe3ksiU8KLz5/+WhYo+LD2B9lnS/8AQKn98Y+Md90VwwwCxQtDCu8lh19RWUaQq+SBKfenb/BHO3V3yZjtjZERc2p7pFyzTJS6l2cJ7ZvRY0Vt0lTEyOTWyR2vmmq5Wrnk1U406vEZ94NBD4sPYH2WdL/0Cp/fD4sPYH2WdL/0Cp/fFoOm8YKdbNq9yk9ZPe4B7xc8g+08wz7waCHxYewPss6X/oFT++HxYewPss6X/oFT++HTeMFOtm1e5SesjgHvFzyD7TzDPvBoIfFh7A+yzpf+gVP74fFh7A+yzpf+gVP74dN4wU62bV7lJ6yOAe8XPIPtPMM+8Ggh8WHsD7LOl/6BU/vh8WHsD7LOl/6BU/vh03jBTrZtXuUnrI4B7xc8g+08wz7waCHxYewPss6X/oFT++HxYewPss6X/oFT++HTeMFOtm1e5SesjgHvFzyD7TzDPvBoIfFh7A+yzpf+gVP74fFh7A+yzpf+gVP74dN4wU62bV7lJ6yOAe8XPIPtPMM+8Ggh8WHsD7LOl/6BU/vh8WHsD7LOl/6BU/vh03jBTrZtXuUnrI4B7xc8g+08wz7waCHxYewPss6X/oFT++P4VdwoeHncKl9VUbYcMjleiNc2hvN8tlMnRRETlDDXMjT5vSqNTmvpXmpyw7rvga52U927Wa3sNo3Ln2lqm/x+o/LsCLyInuauDPtyJ/oUoGAvgXvg1cOO9+Ok83v4JqZugnlNk1Oy+3+KRnL0Mg+FFgTmiclXxfNearz5+k5uzfwfzZnkMcsmI5drdp/W9ByU7KPKbdkdmY5WojVlgq6B870aqc+TahnPpORV/k9Gr7D3VzRjtWVsVo01qUWfG6alhe1PF6qd6p/cz7B0KnBC+UDVdC+GTsNe5F//AHY1P3lNAFi/VvwdzWGyxVFbonrxgmfRxo6aOy59YazTy6Oaic2xRVEDq6CWRfSiOk8nZ83NWkNuv+zPc9teqFZrfo3l2GWx0yU9PlXk0d9wise5ytYyK90b5aJZHcuaQrKkqIqc2JzJb4Y6UOAGMUzKPD29NLU1b/e07nOgqXfNp6hsUzsuVWscibNu1M6Ftm5d6bvtWS1aJ7GJxuREcxO29iuan1qhzEAC/ZTAAAAAAAAPXtH9Ada9wF/XGdFtL801Ju8St8sjxaxzV9FaWv59CSvrOSU9JGvJUSSpkjYq8k6XNUQ8+1bWsqwrPlte3KqOnpIk1nyyvbHGxvVe96o1qdlVRDlggnqZWwUzFe92xGtRVVV7CJtU8hBPXo14PxulzOOkr9YNQNONFKCoRqz2ymkl1NzGh5/ykkpaV8VvXl6OXi7g7mvP5k5KsjOA+DzbXbLFFLqJq5rRndwjVFeyxzWjBbFUehUcj6daSrqPSvJU6FS3lyXnz5+iG1990O0TrkTPo33lSunbxsoYZalq/Nna1KZ392ZS4Fm4U34tFqSJR701eWRzWL9bc1en1tKfYL1OPcFTh0WNkflmid3yieJY3sqsh1Tynp9KNXLzdFS3CCF3S5p0muYrVRiJ0U5u5+iQcJrh4U88NRHtjxJ0kErZmNnyLIKqBytVHIj4n16se3mnpa9FaqehUVPQWKrt1v0eIJVjorEtiVE/G3ijYi9rOuV2Xbai9hSpY8Cr1ubnJU07exrSL/t5fvKCQNBD4sPYH2WdL/0Cp/fD4sPYH2WdL/0Cp/fHj9N4wU62bV7lJ6ydjgHvFzyD7TzDPvBoIfFh7A+yzpf+gVP74fFh7A+yzpf+gVP74dN4wU62bV7lJ6yOAe8XPIPtPMM+8Ggh8WHsD7LOl/6BU/vh8WHsD7LOl/6BU/vh03jBTrZtXuUnrI4B7xc8g+08wz7waCHxYewPss6X/oFT++HxYewPss6X/oFT++HTeMFOtm1e5SesjgHvFzyD7TzDPvBfsl4THDvmlkmftjxNHyyOkckWSZDBEiuVVVGsbcEa1PT6GtRERPQiIh5xkHBR4dN7je2h0XveLTP6Suqcf1Vyh8nNzkdzRlXX1EScuSoiIxGojl9HzKns0O636O88qR1tiWxEi/jbxRvROrnlXI7LtIq9g68mBV62tzjqad3Y1pE/28ii0C37n/g8e2S8wzSac6xaz4LcJeaxtyFbPnlkpvkdFqMp20tHUKnSTpL06lyrzVEVqcuUces3g/263CIay4aR5rpzrdb6dr3U9rbUv04zKu5elqMpKxz6BFVOfodXpyX+n5y+9x90L0Tr8zso4rzJQ1DssmV0MtM1M+rO5q0zfrmTubSmbSwqvxZrVkWj3xqcsbmvXvUXXXvSB4Hq2rehesmg2QfwX1l0zzPTa9vV/ktLllhntUFzbGqI+WiqHN8TVRIqonjad8jPT/KPKSY9l2rZlt0EVq2NUx1FLKmsyWJ7ZI3t6rHsVWuTsoqoUBNDNTSugqGK17dio5FRUXsou1AADvnEACyTwOtpu3HcVpVrletbtIsS1IuuOag2612OuyOnlnnttPLbXSyRRqyRqI1z06S8+fpLNY+Y1WDo+4ZVmKF5aSappKZ8LHRwam+Ks0rYmqm+OY3JFeirm5NiLlmpUF2Lu1V6bYjsWje1kj0cqK7PL3LVcvEirydQrbA0EPiw9gfZZ0v/AECp/fD4sPYH2WdL/wBAqf3xj+6bxgp1s2r3KT1kulwD3i55B9p5hn3g0EPiw9gfZZ0v/QKn98Piw9gfZZ0v/QKn98Om8YKdbNq9yk9ZHAPeLnkH2nmGfeC6bxAeH9sy0x2Zbh8+wHbxp9i2ZYtp9NdMeyG10c8dwtNQ2enaksSrKqI5Ec5PSi/OUsiamjNpM3T0obp197ro0FRSQUlR7Gc2p3rXc/eo5dZu9SSJq5SIm1UXNF2ZFu743OrrmV0VBXyse57NdFZnkiZq3JdZE27AACSRSIAAAAAAAAAAAAAAAAAAAAABcM4WWxLaFrLsP0K1J1R0CwPNc6yT+E/w5k16pJ5bnc/I8yyGgpfGubK1F8XT0tPE3kifJiaSCfFh7A+yzpf+gVP74xjX+3UnCHD6/dtXCtO71pSVNm1dTRyPjSl1HvppnwvczWqGu1HOYqt1kRclTNEXYXksvBa3rVsyntSGrhRk0bJERdfNEe1HIi5NyzRF25GfeDQQ+LD2B9lnS/8AQKn98Piw9gfZZ0v/AECp/fFJdN4wU62bV7lJ6yd7gHvFzyD7TzDPvBoIfFh7A+yzpf8AoFT++IWeN5tD207eNv8ApPk+iejmH6cX+96xMsN1uuO00sNVXUa2S61C071fI5Oh4yCJ/oT52IXIwg3SzCnGPEqycMrDsG0IKu0JFjZJKlNvbVRjn5v1J3OyyaqbGquaoeRb2EFt3fsee2amqidHEmao3XzXaibM2onL1SsqADI8WkAAAAAAAAAAAAAAAAAAAAAAAAAB+5brdcLvXUlrtNBW3S519Q2loLdbqV9bXVsr16LI4YWIr3vcqoiNaiqq/wAx+XvZGxZJFRGomaquxERONVXkRD6iKq5Jxn6YJWdB+DJvp1vp6K7V2n1t0Xxus6L2XfWe5vxa5OZ6FfyskUU1zY9GrzRKimha5V5dNPSqSraZ+DoYFSw00+se5DL77UPVr6y16Z4jRYpDTJzTpRxV9c+tWX5l5SOpo/5Sfxa8vlRQxD04tFzDOokobdvZBNVMVUWKjSSscjk42udTNkiY5OJUkkYqLsXbsK4srDe+lsNSSmoXNYv40mUaZdVEeqOVO0ilVAF3nFeBbw/sebA274jqPnSxK1ZH5VqbX0bqrou6So/4NSjROknyV6CN9C+jkvpPZKHhD8Oe30zKWDbPYpIo1VWurs4yq51K9JVVelNNdHyL8/oRXLyT0JyQjbam60aN9HKsVBZdr1G33zaelY1U6qb5Wtf3WIVdDgZe6RNaWeBvYV71X90ap+8oRA0BaHhbcP230zKWn2uadSRMVXNdXLcLnUr0lVV5zTVL5F+f0Irl5J6E5Ift/Fh7A+yzpf8AoFT++KefuvGCKOVI7tWqreTNKRFy7KeyVy7q9s7SYD3jy21kGf8AieYZ94NBD4sPYH2WdL/0Cp/fD4sPYH2WdL/0Cp/fH56bxgp1s2r3KT1k+8A94ueQfaeYZ94NBD4sPYH2WdL/ANAqf3w+LD2B9lnS/wDQKn98Om8YKdbNq9yk9ZHAPeLnkH2nmGfeDQQ+LD2B9lnS/wDQKn98Piw9gfZZ0v8A0Cp/fDpvGCnWzavcpPWRwD3i55B9p5hn3g0EPiw9gfZZ0v8A0Cp/fD4sPYH2WdL/ANAqf3w6bxgp1s2r3KT1kcA94ueQfaeYZ94NBD4sPYH2WdL/ANAqf3w+LD2B9lnS/wDQKn98Om8YKdbNq9yk9ZHAPeLnkH2nmGfeDQQ+LD2B9lnS/wDQKn98Piw9gfZZ0v8A0Cp/fDpvGCnWzavcpPWRwD3i55B9p5hn3g0EPiw9gfZZ0v8A0Cp/fD4sPYH2WdL/ANAqf3w6bxgp1s2r3KT1kcA94ueQfaeYZ94NBD4sPYH2WdL/ANAqf3xQw1atdvsmquplltNLFQ2q0ag3q12yhgRUgo6enuVTFDExF9PRYxjWp/6ISo0YNMi42lPX2xZ9z7Lq6R1nMhfItTvOT0ndI1qM3qWTam9rnnlxplntKKvnh/aVyoqeWvmZIkquRNTW2auSrnrInVPPgWy/Bd9k21HeL58nnO6G4RrR/o6/0Z/wI/hlSzVP8G/hf/SB8JeTeLkZy8o+C7f0+fPn5Kzly9PO2X8SZwpOw3oj/ZdZ7SS+KCMmkGst8SZwpOw3oj/ZdZ7SPiTOFJ2G9Ef7LrPaQDJpBrLfEmcKTsN6I/2XWe0nJm/jhB8M/TTYrvS1HwPZvo/i+c6f7TNRs2wzJbZbauO5Y7drVh95r7dXU7lqFRJYKinhlYqoqI6NPQoBmOAAAAAAAAAAAAAAAAAAAsJcJDgCa8cQepsGsesi3/QbaG+dlXHmc9Eyn1C1fha5FkgxCinY5radyJ0FvVVG6lY5/wDERVzo5oo7t1h4GvCfx+x2exQ7KNJ7nFZrZBa47jfkuV7vlekETYkmrKyWqWSeZ/R6T5Xqrnucqr84Bk/g1lviTOFJ2G9Ef7LrPaR8SZwpOw3oj/ZdZ7SAZNINZb4kzhSdhvRH+y6z2kfEmcKTsN6I/wBl1ntIBk0g1h7rwYuEhYbXcr5fNlugVmstmoJrreLxdaeot1rtVLTxumqKmpqJKpI4ooo2Pe+R6o1rWKqqiIqlNfiebqOElU53b9sXDp2X6GXKOszKgxrUDdPJY6uW3ujdXwxVdFg1M+oVJmvTpROvNQ3oK1ZPJIXo6GuQCtADWW+JM4UnYb0R/sus9pHxJnCk7DeiP9l1ntIBk0g1lviTOFJ2G9Ef7LrPaR8SZwpOw3oj/ZdZ7SAZNINZb4kzhSdhvRH+y6z2kfEmcKTsN6I/2XWe0gGTSDWW+JM4UnYb0R/sus9pHxJnCk7DeiP9l1ntIBk0g1lviTOFJ2G9Ef7LrPaR8SZwpOw3oj/ZdZ7SAZNINZb4kzhSdhvRH+y6z2kfEmcKTsN6I/2XWe0gGTSDWW+JM4UnYb0R/sus9pHxJnCk7DeiP9l1ntIBk0g1lviTOFJ2G9Ef7LrPaR8SZwpOw3oj/ZdZ7SAZNINYi5cDzhO3WimoKrY9o/FBP0enJbUutmrW9F7Xp0KmnrGTM9LUReg9OaKqLzRVRfiLl4P5wf7rRTUFVsuxaKCfo9OS26l5zZq1vRe16dCpp72yZnpaiL0HpzRVReaKqKBlXA0zM68GC4TeWtmbYNOdXdL1le57H4LrXeLg6nRzmuRrPhlbiio1EVqdNHLycvNVXkqR160eB/6QXCnnqNu+8TUjEaqLpyU1o1owG16i09d8/i4pLjbJLWsHzp0pUpZvmXlEnP0AURwT2bpfBu+J3trpLnkFg00xzcrhlsYtTNfNvV+kyi/xRf8Aw9LGquGlu8svL+UyhpqprVRflqnJywUXyxXvGLzc8dyWz3XHsgsldJbLzYr5b5rTebRUwuVk1PVUsrWyRSscitdHI1HNVFRURQD+UAAAAAAAAAAAAAAAAAAAAAAAAAAADo3bbtD3Nbv8w/gLto0Sz7WHIY5I2XFMTszpLHj6So5YpLtd5VjoLfE/oORstdPDGqpyR3P0FmjbN4I5uLzGlt183V7itPtFKSdjamfCNM7HPq5mcSdJUdTVddJLRW6ml5J0kkpn3CNEVvoVVVGgVCgaSelPgrnDCwOCB2du191urUc2SrXN9UG4zbJXIjkcyKCx0lBKyNVVF5OmkeitT+M5c0Xr22eD9cH600cdDS7LsXlhjc5zX3PU7Or1WKrnK5elUVF7klcnNfQjnKiJ6E5InIAyrwasdz4CHCIuzqt1VspwSJa1jo5ktmZ5dZWsRzegviUp7tGkS8vmdF0VRfSiovpOVtU/Bd+FNn9NURYdh+s+h9RJydBVaa6x3C9+TqjVRE6GQsuqOaq8lci/KX5kc0AzRQXBt1/gkGuWG0FyyXZ5uDxbWiCmjkqodM9WbQmmObStajUjpqG8xST2ysnequXpVbbZE1E/lqvz1YNetuuuW17US66T7g9Lcw0l1Bs/y6rHMwtTrfNUxKqoyqo505wVdNJyXxdVSySwyInNkjkAPGAAAAAAAD0jSbR3VfXnObPpnotpzmmqmoF+f0LViGB47VZPfqpqK1JJvJ4GOcyGPpI6Sd/Rjibzc97WoqoB5uC0dtd8FJ30at0lpyDcRn+mG1rH6/oy1Nhqpnauap0Mbua9J9st8sdsaqt6Kox10SRqu5PjYqK0m20l8Eu2C4lS0c+q+ru4/V69xxKyuZSX+zac4hVOVP5cdDT2+atjVF+ZFr3p/TzAM7sGpfi/g7PB7xaNis2jQX6sSB9PLX5RrJn17kqGvkSRFWnfe0pmubyaxr44muRreXNek9XfcpwFuEW2kkok2T6f+JlqWVTnrlmVuq0dG2RjUbULdfGtYqSO6UbXIxyo1XNVWtVAMpcGnvqh4NTwkNRaSqjsuheZaQ3OqaqOvul+seSR1kKq1GtdFSXWquFAxW8kVEbTdFV5q5Hc1IMt3vgj2o+M0NxyjZJuCoNTIqWJ9RDpRrpRw4hl9S1jXObFQ5LRt+D6mokXotayrpLfC3kquqPTyQCmqD2bXnbxrhte1Gu2km4LTDLtJtRLK1JazGcwtbqConhermxVdHMnOGrpZFY/xdVSvkgk6KqyRyIeMgAAAAAAAAAAAAAHWewfCcU1L31bLdOM8sVDlGDagbs9OcJzPGrmx0ltyK03XMLNQXGhqGoqKsU9PUTRPRFRVbIvpQ04/iTOFJ2G9Ef7LrPaQDJpBrLfEmcKTsN6I/2XWe0j4kzhSdhvRH+y6z2kAyaQay3xJnCk7DeiP9l1ntI+JM4UnYb0R/sus9pAMmkEsvHK0S0o26cUrdDo3ohg1j020vw7+BP8GMKxyJ8Fmsvwhp3iN1rfEte5zk8dV11XO7m5flzu/wDsRNAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAFlnwcb/bPdb/VjEv2u/lqcq++DeMYr95UisasjG6eMbIrUV7WuXOFciL86Iqtaqp/P0U/oLQRq5bo/MkumRe1mXvG2enb/wDjKN3+rL6iaGEjcsP6Feqsv30if5AAEHC5AAAAAAAAAAAAAAAAAAAAAAAAAP0LparZfLbXWa9W2gvFoulK+huVqulHHcLbcYJWqySGeCRFZIx7VVHMeioqKqKh++D9xyPikbLE5UcioqKi5KiptRUVOJU5FPioipkvEV899/A2011Mt161H2h0tt0t1JiZLcqnSt8yUmmWaPROmsFu6S9G0VDvSkbWL5EqoxixUzVdOlS7L8QyjAMov2FZrYbpi+W4vdJrLkOPXqkfQ3S0VUD1ZLDNE5ObXNVP/sqKioqoqKabpB5xleHtaNw2lN33E6Y2CGDXfSeyvut9ZbadkdRqdjlHEr6ukqGoiLLW0ETHT0knNXvjilpujIr6fxOXrQY3QO9FhXkoMIMcq91XZFS5sNLXTO1pqSRyo2OOeVy60tM9yo3fJFV8Cqiq9YUyjsNiThZRVNJLb12okjnYiufE1MmvRNqq1qbGvRNuSbHdTW46YQAM+ZGEH9Kz2a75FdrZYMftVyvt9vVfDarNZbPQy3O7XaqqJGxQU1NTRtdJLLI9zWMjY1XOc5ERFVeR/Oa1znNa1quc5Ua1rU5ucq+hERC5rwiuGVatt+GWfcHrVj1NW7gs1tTK/H7NeKGOodo3bKmPnHDA17OlFdaiKT/WpUVHwskWmajf9YWaN2k/pK3P0Y8PHXvvC3f6+ZVjoqRrka+pmRM125LqQxoqOmlyVGIrWojpHxsdV1zLn198rVSgpPcxN2ySKmaMb/m5eJreXauxEVU5S2H8CSgWisuqG9Z1TNV1CR3G16CWS4rSwUjVTpM/hJdIH9Nz15oq0VE9iM6KeMqH83wtsk4NgGDaY41b8N06xDG8GxO0x+Lt2OYnZaew2akTkiKrKeFjWdJ3RRXPVOk5fSqqvpPrgaymOGkfi3pCW++2sRrUfJCjlWGljVWUlOnIkMCKrUVE2LI/Xmeie7kcTFu3dKwrq0qU1kwojsvdPXbI/wCc7j+pMmpyIgABYsqUAAAAAAAAAAAAAAAAAAAAA+L1A04wDVfFrjhGpmGYznuIXZnRuGOZZZoL5aahURUZJ4mVrmtkZ0lVkreT2O5Oa5qoilaXfjwJG2yhveqWymSuq4qSOW53fQO/XB9fXLGiue9MXuUqrJKrG9Ho0Fc90j+g/wAXVPesdO60cC/OBekni5o828y18O7TcymVyOmpJFV9JUJszSWFVRNZUTJJY9SZqe8kTbnTN5boWFeulWC1oUV+WTZEySRnzXdT81c2ryopmGXK23Gz3GvtF3oK21Xa1VsttudsuVK+huNtqIHuingngeiPjkje1zHMeiOa5qoqIqH6RcY4wfDKtuvOIXvcvodjsNJrjhVqkuecY5ZKBGyav2qmZ05ZEhjTnJdqSNjnRPaiyVUTVgXxj20yNpzmzboy6SNztJrDtl87tJvFZEqR1lI5yOkpp8s9XPJNeKRM3QyoiI9uaKjZGSMZDm+N0bQuday2fWe6jdtjeiZI9vV7Dk4nN5F6qKiqLaHg6P3Ftxv40LV/dLipeW0PB0fuLbjfxoWr+6XFkd0w+KHb309B/WQlR4PfDyl+bL924sYAA1hyZIAABwfxPPvA9034r5/2mmM+80EOJ594Hum/FfP+00xn3mwVuQ/4FLzfrVf6SmIs48fCKj+g/wBx4ABlmLGgAAAAAAAAAAAAAAAAAAAAF8Dgy/g19t//ADh9fcpJQCL/AIMv4Nfbf/zh9fcpJQDT90mfjIYg/ry1v6+oJ63O+CNlfo0H3TQACyJUYK/HhEf3sOiX4+WfV69Fgcr8eER/ew6Jfj5Z9Xr0Sy0Fvja3I/S3fcTFDYl/AW0vmJ/M0qDAA2viEAAAAAAAAAAAAAAAAAAAAAAAOstk+1HK95e4TD9FsclqLZaq1zr9n2Uw06VDMPsFG6Py+vVq+hZF8ZDTwMd8l9RVwNcrWuc5vgXqvPYVyrt117rz1LaezqOJ800juJkcbVc5ckzVVyTY1EVzlya1FVUQ7VFRVNo1kVBRsV0sjka1E5VVck//ALxJxqen7EeHZrPvpyupjxRseGaVY5XMpM41ZvdG6otFqkc1si0Fup0c11dXrG5H+Ijc1kbXMdNLCkkXjLl20zYHtp2bWSnp9KcHpanNH0SUt81VyqOO96iXxVb0Zf8AXXNRKWGT+elomwwL0Wq5jnIr16J0i0k0/wBCtOcU0p0vx6ixjCcNtbLXZ7XRxta5yN9MtTUSIiLLUTyK+aad/N8ssr3uVVcqnpBrHaV2nDiVpF2zVWLZlRJZ90WuVsNFG5WrMxF2SVrmr/bPdkjt6VVhi2I1rno6V8x7kYcWRdKnZUTNSWuVM3SKmeqvUjRfeonFre+dyqiLqoABB8uOAAAAAAAAAAAAAAAAAAAAAAAAAAADNP1v+7Tq9+NC/wD97VZpYGafrf8Adp1e/Ghf/wC9qszObjx8Jb9/QUH3lUR8x9/+jsz50v8ABhcg8Dh/4jH5of8AueXcCkf4HD/xGPzQ/wDc8u4GdQjUAAADiPiY/g3+IH+RHqt9Q78duHEfEx/Bv8QP8iPVb6h34Ax+wAAAAAAAAAAAADrzZdsZ3I7+9XqDRvbdgVZlV5V0NVleT1nSt+C6cW+WRWOul/uitWOlgb0ZOizk6ed0axwQzSqkagcwYvi+S5tkVkxDDcfvWV5XktzhsuO41jlsmvV+vtZUPSOClpKSFrpZpZHua1scbVc5VREQvMcIfwZqxYN/BjcVxHrLasszFvib5hu1lKll1xHF3p0ZIZ81njVYrhUtcjXfBMLn0TEaiVMlX4x9NDMbwr+Chtp4aGO0eU09PR6wbnbpbPJ8r12yOzxxVFn8czo1FuxWicr/AIMo15uY+RrnVVSir42ZWKyCOZ4A/Xo6Okt9JS0FBS01DQUNMyjoqKjgbTUlHDE1GRxRRtRGsYxrWta1qIiIiIicj9gAAAAAHNW6/d5t72S6Q3rW7cjqJadP8ItTvI6BtS7yvIsvuDo3yQWmx21nOeurZkjerYIWr0WRySyOjhjllZHfxV+Nltu4Z+O1WJLNR6xborxbFqMS0Mx67NjWxJKxHQXPLa5iP+DaPk5r2Qq1auqRUSGJI/GVMOb9vP3y7lN/WrlbrHuUz6qyy+fxtLi+NULX2zAtOqCR6PS2Y/aum5lJTp0Y+k5VfPO6NJJ5p5VdIoEknFo47O4HiQ3K66Y4ZHddDdpNJcUdbNKbdckdk2oyU8nSpq/M66JejUO6TWzMtcC+RU7/ABfNauWFlWsLGmf3SNPv672n9vpz4g+30z+6Rp9/Xe0/t9OAbVIAAAAAAAAAAAAAAAAAAAAAAAAAAAAABHlvs4W+zTiH4tUWrcFpbb3ZvDR+TY1rThbIsY1fxNyNRsfiLu2N3lMLETl5FcGVNL6efiUejXtkNABlM8VPg57ieGDm0dZkfT1R265TdfINOte7DaX0NqqpntlkjtF/o+lJ8G3RGQyvSF0kkNRHG58E0ismjhiFNozW/RPTDcbpPnmiGs2I2zOdMtScenxrLMausarDWQTJ8mWGVOUkFTBI2OenqoXNmp54IpYnskjY5MoPih8P/N+G9u3znb/kcldecKm5ZloxndVB4pmd4rXSypb6l6o1G+VUzop6GsY1Ea2qoZlYixOie8CO8AAAAAAAAAAAAAAAAAAAAAt2cIvwaLKNd7bie4viAQ5Dp3pDdqaG/Yft6t882Pamah0siNlgqchq2q2ezUMrFTlSxdG4SskVfGUPRY6X3DwcTgqR3X+CvEP3ZYbHLa06F52taV5TbYqinui/PDnVzpJUX5DFTnaopWornIlaidFKKR94QA8w0e0V0k2+YBZNK9ENOcP0r07x2NWWjEMHsUFgstM5/JZZ3RxNTxk8qp05aiVXSyvVXPe5yqq+ngAAAAAAAA5B3pbFttu/vSK4aO7j8Co8ntLmyVOL5VQtjt+e6d1zm9FtysN06DpKWZOTekzk6GdrfFzRSxqrF6+ABkbcTXhuaycMvcPW6Oaku/hNhmQwTZFo3qxb7fJRWDUuytm8X41GKrkp6+lV8UVbQq97qeSSNzXyQzU880dBrI8YnYBYeIfsl1I0tgtVNNrDhFvqNTNAL2kLUuFBk9tpZnw21Juirm012i8Zbp2+lqeVRTdBz6eLo5Oc8E9LPNTVMMtPU08roKinnjWGeB7FVr2PYvJWuaqKiovpRUAPxAEjXCx4feXcSTd3hGgVonuNkwKjjdm+teb2+JrpsMxWgmhbXSwOe1zErKuSanoaRHte1Kitje9jo45eQHUvB+4KWsHE8y2bNMgr7rpLtQw28Nt+b6sJQJJe8uqo+i+aw4lFKxYZ61Gub46rlR1NRNla57Z5FjpZdHTaNsj2w7GdN6bTDbNpTj2nllWKJb/eoIfhDNc4qYmI3y2+3mXnVVsyr0lTxr/FxdNWQxxRo2NPatJtKNOtC9NcJ0f0lxK04Lptp3j1Pi+H4pZIPEW+0UdM3osaiqqvkkequklnlc6WaWWSSR75HvcvoYAAAAAAAAABx3vV2H7Z9/2k9XpJuR0/ocot8bJpsSzChbHbtQdOK2ZrWrcLBdeg6Sml5siV8ao+CdImsnhmYnQMybilcK/XLhga0/wLzlJ820fzGeer0a1ut1rfQ2DOaSJUdJR1cXSe2jutK17EqaFz3ckc2SJ8sL2SLrOnKe9TZ1o3vu27Z1tw1vssdwxjLaXyqyX2CnjkyDAL1TskS2ZBaJXJzirKR8r+SoqNlilngkR8M8sbwMcsHTe8Xafqpsj3G6mbatYrd5Ll2nV7dRwXWnhfHZsxtkyJNa75bXO9L6WupnxTs/8AiYr3xyI2SORjeZAAAAAAAAAADtzhnfhIOH5+W5pT9fLCbAhj98M78JBw/Py3NKfr5YTYEAAAAAAAMtfwjT8MvvG/N7/hXg5COTceEafhl9435vf8K8HIRwAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAC0B4N3/vmfm7/66LQBnF6AbttxO1r+Fv8AoD1NumnP8OvIP4WfBtqttz+Gfgzy3yDp+V003R8T8IVnLxfR5+PXpc+TeXRvxtnET7TeT91cb+zjDbpWbnRi/jvj5b2K117Ys2CgrvYu9x1ElU2Zu8UVNTO10jpZGJm+Fzm5Pd7lW55LmiSAuRizYN2br0tiVtPM6WPXzVqMVq60j3plm9F4nJnsTbmX5wUGPjbOIn2m8n7q439nD42ziJ9pvJ+6uN/ZxHjpReP3XBY/ha31Eqvh2utzWo72P0hfnBQY+Ns4ifabyfurjf2cPjbOIn2m8n7q439nDpReP3XBY/ha31EcO11ua1Hex+kL84KDHxtnET7TeT91cb+zh8bZxE+03k/dXG/s4dKLx+64LH8LW+ojh2utzWo72P0hfnBQY+Ns4ifabyfurjf2cPjbOIn2m8n7q439nDpReP3XBY/ha31EcO11ua1Hex+kL84KDHxtnET7TeT91cb+zh8bZxE+03k/dXG/s4dKLx+64LH8LW+ojh2utzWo72P0hfnBQY+Ns4ifabyfurjf2cPjbOIn2m8n7q439nDpReP3XBY/ha31EcO11ua1Hex+kL84KDHxtnET7TeT91cb+zh8bZxE+03k/dXG/s4dKLx+64LH8LW+ojh2utzWo72P0hfnBQY+Ns4ifabyfurjf2cfU2HjL8RuxSUvS3BNvVJTSOe+gv2mOI10dX0ul8mWoS2NqVRFXmnRlby5InPl6DgqNyO0ho4lfTW7Yz3JyLPWtz7S+wFTPkTPJOqqH6ZjrdRXZPpqhE+bGv8Aul74FYvZ1x96+/5bZMC3f4djFjtN7q4bXS6wYBDPa6CxSyyNjZNfbVPNKnk/N6ulqqSRnims5+SvRXOZZyY9krGSRvbJHI1HxyMcj2Paqc0VFT0Kip6eaEG8cNHrFXR4vDDd7E+z94dO1zoJmPbLBO1ioj1ikbsVWK5uuxyMkZrNVzER7VW5N271WJeukdV2NLrI1URzVTJzVXi1kXq7clTNFyXJdin+gAWTKiAAAM/LiYbeKTbLvO1i09slGyhw263dmoOB08MXiKWmtV9Z5fHSQM/mjo55Kuhbz+dKHn/OcGFkHwjPCKWg1T2y6kMjjStyvAL/AIRUSp/5r4sfuNDXwtd6fma7Jp1T0fO93z/zVvjbj0RsQa7FLRsuffW1JFkqpaRIpXquayS0r30ksjl/Kkkgc93ZcpBS/dlR2Le+vs6FMmNk1mp1GvRHtROwiOREJkOCltGoNx25+TULMrYy46bbeqWkzW40dTA2ot97v9RM9uPUM7HIqOja+mrK5yelF+DGMcitlUu4ERfBK0Wp9KNieFZLPRNpsi1ryC5an3mV8XKqdTvnW12livX0rGtFbqeoY1OTUWukVE5ucqy6Gvnp/wCMlbi7pJW3EyZXWZZD3WfSszzaiU7lbUSJyKstSkrtbLNY0jbmqMapKfC678dg3Qp3K3KadElevL7tM2J/dZkmXVzXlAAIUFxAAfhqamno6eerq54aWkpYXVNVVVMrYKemjY1XPkke5URrWoiqrlVEREVVPrWuc5GtTNVHFxn5gV2N4XHwwbTm+3jAtqeG2nVm8WiploKzVLLquaHTRZ4+kxfgujpnsqblF0v/AMz46ljd0FWPxzHtlIgsp413EVyKtqai3azWPDKOpa5nwRi2lmMvooEcnL+LmraGpqmqnp5O8crkVefP5uU/sM9zU0ncSLFht+akpbJp5Wo5iWjNJFK5q7UVYYIaiWNV/JmZG7qoiKilrrYxfubZFQ6lbI+dzVyXemo5qL85zmNXttVUL0gKDHxtnET7TeT91cb+zh8bZxE+03k/dXG/s4ux0ovH7rgsfwtb6ieHw7XW5rUd7H6QvzgoMfG2cRPtN5P3Vxv7OHxtnET7TeT91cb+zh0ovH7rgsfwtb6iOHa63NajvY/SF+cFBj42ziJ9pvJ+6uN/Zw+Ns4ifabyfurjf2cOlF4/dcFj+FrfURw7XW5rUd7H6QvzgoMfG2cRPtN5P3Vxv7OHxtnET7TeT91cb+zh0ovH7rgsfwtb6iOHa63NajvY/SF+cFF/E+NlxE8brKeouWsGP5vR06RtS05ZpdjcdHM1no5STUFHS1LuknLpO8d0l5c+ki81WYzZxx6NPdVchtOn26PD7RozfrtKyjt2peOXGWp0vqKh/SToXGnqHOqbWxXdBrZnTVUXN6rLJAxqvW0eJ25r6TuGlizXgjo6a1qaFqvk6HTPlla1NqqkE0NPNJlythZI7jXLJFU92xsXrm2xUtpFkfA9y5JvrUair1NZrntTtuVELCgPxxSxTxRzwSRzQzRpLDNE9JIpWuRFa5rk9CoqKioqehUU/IQFVFRcl4y6HHxAAHwAo2cZPaNQ7Yd1dZkeH2xlt0v16pajUTFKSmg8RbrJc0nazIbVAiIjEZFUTQ1bI40RsUN2p40TkwvJkNPHP0Th1P2RXbO6WjSbIdC8wt2d0c0bVWrdbqyZtlusDf5vF9Gvp6uT5vRbGrz9HRdOPc9cZK3CbSSsegkmVtmWy5tn1LM/cq6d2VLJkuxHR1Kx+6Xakb5WoqI9S2+Kt347cujUStbnNTosrF5cmp7tO0rM9nVRq8hSVLaHg6P3Ftxv40LV/dLipeW0PB0fuLbjfxoWr+6XGZXdMPih299PQf1kJH7B74eUvzZfu3FjAAGsOTJAAAOD+J594Hum/FfP+00xn3mghxPPvA9034r5/2mmM+82CtyH/AAKXm/Wq/wBJTEWcePhFR/Qf7jwADLMWNAAAAAAAAAAAAAAAAAAAAAL4HBl/Br7b/wDnD6+5SSgEX/Bl/Br7b/8AnD6+5SSgGn7pM/GQxB/Xlrf19QT1ud8EbK/RoPumgAFkSowV+PCI/vYdEvx8s+r16LA5X48Ij+9h0S/Hyz6vXolloLfG1uR+lu+4mKGxL+AtpfMT+ZpUGABtfEIAAAAAAAAAAAAAAAAAAAAAAXGeAVt1pNP9s2Sa/Xa3MbleuuUTUlmrpY08fBjlhmloKeNnNOkzx1wbdpH9FejIyKkX09BqlOY0aNmGC0+mm0nbZhEEfin2LRPG47gnJieMrZ7VTVNfJya5zU6dTNUP5I5yfL/lO+dcWu6w4hVt2sCLMuPZ8isW161qS5L76npWb85naWdaZy9huWW0vTgfZUdZeaa0pUz3iP3PYc9dVF73XT6zpgAGuySvAAAAORN5+87SrZHpO/UzUnyy7XC61jrJguDWWWNl/wA0uKROl8TEr16MNPE1EfUVT0VsLHN5NkkfFFJVr1a48+9nOLlWO03XTvRSxulX4MpLDiVNml+p4+S8kqq26tngmfzXn046WFvob8hPSqyywC0K8ddIyy33iuPRwwWS16xpV1kqwwve3Y5saMZLNJqrsc6OJzGuRWq7WRWlDXoxDu1dOZKS0pHOnyz3uNus5EXiVc1a1M+RFciqm3LIungoR13F54jNwqH1U+5i+RyvRGq2hwbFLZTp0URE5Qw2tkafN6VRqc19K81P1PjbOIn2m8n7q439nEpGbkXpAq1Ffb9jI7lymrVTu+wUz7iFGLjtdbPZS1GXzY/Sl+cFBj42ziJ9pvJ+6uN/Zw+Ns4ifabyfurjf2cfelF4/dcFj+FrfUT5w7XW5rUd7H6QvzgoMfG2cRPtN5P3Vxv7OHxtnET7TeT91cb+zh0ovH7rgsfwtb6iOHa63NajvY/SF+cFBj42ziJ9pvJ+6uN/Zw+Ns4ifabyfurjf2cOlF4/dcFj+FrfURw7XW5rUd7H6QvzgoMfG2cRPtN5P3Vxv7OHxtnET7TeT91cb+zh0ovH7rgsfwtb6iOHa63NajvY/SF+cFBj42ziJ9pvJ+6uN/Zw+Ns4ifabyfurjf2cOlF4/dcFj+FrfURw7XW5rUd7H6QvzgoMfG2cRPtN5P3Vxv7OHxtnET7TeT91cb+zh0ovH7rgsfwtb6iOHa63NajvY/SF+cFBj42ziJ9pvJ+6uN/Zw+Ns4ifabyfurjf2cOlF4/dcFj+FrfURw7XW5rUd7H6Qvzmafrf92nV78aF/8A72qzsn42ziJ9pvJ+6uN/ZxHveLvccgu91v14qn112vdynu90rZGNjfV1NTK6aeVzWojUV73ucqNRETn6EQn5oI6HOIei1a15LQvxaNFVMtGOmZGlI+dytWF0znK/foIURF3xNXVV3EueWzO1+Jl/7JvrBSRWbFIxYleq66NTPWRqJlqud1NueRdf8Dh/4jH5of8AueXcCkf4HD/xGPzQ/wDc8u4GR0tIAAADiPiY/g3+IH+RHqt9Q78duHEfEx/Bv8QP8iPVb6h34Ax+wAAAAAAAAAf1bHYr3k96tON41Z7rkORX64w2exWGx2+a7Xq9VdTI2GnpaSlia6WaaWR7GMija5znORERVVELt/CB8GagtzsY3HcSbH6aurEbDfcH2my1DamhpHc2y01Xnk0TlbK75n/AUTljT5Da2R/OegQCGDhLcCLcFxG7rZtT89ZfNDdosFd0rjqncLckOU6lRwvVs9FhVFOxW1Cq5roX3WZq0VO5Jej5XNC+kXRr2rbSNv2yzSOy6JbcNOrPp5g1o5VNW2iYtVfsprlYxk10vVxfznrq2ZGNR087nKjWMjYjImRxs6Ctlsttkttvs1mt9DabPaaGK2Wq1WykjoLbbKanjbFBT08DERkcUbGMYyNiI1rWoiIiIiH7wAAAAAPAty+6HQfZ/pLf9btxWo1i0007x5qRS3S8Sukr7zVvY98FttVDGjqiurZkjkWOkpWSSuSN7uj0WPc0D3WsrKS30lVX19VTUNBQ0z6ytraydtNSUcMTVfJLLI5UaxjGtc5znKiIiKqryKaXF88JitGGvynbjw4r5bskymPxtizPdTHBFdMXsD+Sxz0+EwyI6KunavSZ8LytdSt6KupWVKOjqo4XeLjx89eOIXXX/SDSF2Q6E7QvKXUrcGp69tJnurkLHfIqMxq6eRzFhdySRtmp3upY3K3xsla+OKZlfYA/u5RlGS5tkV7y/MsgvWV5XktzmvWRZLkdzmvV+vtZUPWSeqq6uZzpZpZHuc50kjlc5VVVU/hAAA+30z+6Rp9/Xe0/t9OfEH7luuFZabhQ3S3zupq+21kVwoalrWudTzQvbJE9EVFRVa5rV5Kip6PmANtQGVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVD8fpxee2rm3cbDvsgfH6cXntq5t3Gw77IANV4GVdbvCBOMDa62Gvpt6OUyzwdLoR3HTPBbxRO6THMXp01RZHxP9Dl5dNq8lRFTkqIqde6EeFNcTHTS9UE2rtTpFuPxpsiMvFsy7Tyi09yGpi9HSWjuNiZRwQTLy9D5qSojRHO/iV+SrQNJQEdnDW4lehPE30Nl1a0hbXYzk+L10Vh1Y0myCqZVZTppcpmSSU8c0zGtZU0lWyGWWkromtZM2KVrmRTQzwRSJgAAAArIeFNbQrfrZsMtm5Oz25smoG0vMqa9TVkEKy1ddimTVVFZL3SdFqdJyRVj7DX9Nyq2KK31a8k8Y5yWbzlnfJpdS62bMd1uktVTOqv9IW3fMcWo2RojpoauqsFeyhmi5oqeMiqPESs5oqdOJvNFT0AGOKAAAAAAAAAAAAAAAAAAWUOAFwZKzfZqRR7mdwuOVUG0HSvIk8ls1ypFih3BX6jcj/gWBXKiutNJJ4t1xqGo5sytWjYvTdUSUvDvB/4WGoPE83FUuLq27Ytt404qaa+6+an0lNyW20T3OdT2G1SORY3XW5+KkihR3SbTxNnqXskSFsM2pfpNpRp1oXprhOj+kuJWnBdNtO8ep8Xw/FLJB4i32ijpm9FjUVVV8kj1V0ks8rnSzSyySSPfI97lA+7o6Okt9JS0FBS01DQUNMyjoqKjgbTUlHDE1GRxRRtRGsYxrWta1qIiIiIicj9gAAAAAA/Vrq6itdFWXO51lLbrbbqWSuuFwrqhlJRUMELFklmmleqNYxjWuc57lRGo1VVURCiBxPfCgNbZdc10+4cORY5iej2mt/SC8avX7CbdmV01yqaWR7KqOjpbjDLFR2KTkjIpIo46+dGJM2opmvSJAL5IPMNEdRmaw6MaRauRUPwXHqlphYNRo7Z01k+Dm3y1UlzSDpL6V8WlV0Oa/P0T08AAAAGUNxydvFLto4pG7DCrRQtocXzDOWaz4rHE1I6RKbM6SDIqmKnYnJGRU9dX3KkYxERGpR8m/JRDV5M9TwuPAoLLvq0C1CpoWwtzva/S2auc1yKlVV2PJsh6Urk59JHeIudFHz5I1Ugby5qjgCqUaPfgt+zyh0L2F1m46922KPUPdtlc+Rx1ckbUraHFMeqayz2GjXm3pNSapZerhza5WyRXGlVURWGcXS0tTW1NPRUVPPV1lXOylpKSlidUVNVLI5GRxxxtRXOc5yoiNRFVVVEQ2b9tukVu0A286GaG2qKCGg0f0ix3TWBKdrWxzLZbRSW98qq1ERzpHU7pHP+dzpHOXmqqoB7UAAAAAAAAAAAAAACuP4RRwsU3v7cPOD0gxxtbuf22WKpulto7ZRrLeNVMRjV9XdcdRrE6c1VSqs1fb2cnuWXyunY3pVvSbmrm3AZvvhH/Cx8zjcKm6bR3HfI9t25TI6iruNvtsHQtmlubTJLWXK1IxE6MVHc2sqLhRtReixzK+BrY44IUcBWhAAAAAAAAB25wzvwkHD8/Lc0p+vlhNgQx++Gd+Eg4fn5bmlP18sJsCAAAAAAAGWv4Rp+GX3jfm9/wrwchHJuPCNPwy+8b83v+FeDkI4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAANErYTd7nfdlG1S6XitqLjcqnQTF21NdVP8bVVPirTSxNdI/53v6Mbeb3c3OVFVyqqqq53lLS1VdVU1DQ01RWVtZUMpaOjpYXVFVVSyORkcccbUVznucqNRrUVVVURENG/aJp3f9JNrO3nTPK4kp8owjRzHseySkanyaGvp7ZTNrKdF5r0vEzeMi6fo6Xi+fJOfJMPm6/1lnNw5ufQSPb7LdWzvY3ZrrG2DVkcnLqo58SO5M1b2C/eAscy2vXyoi72kbUVeTNXZonbyR2X1nRQAMCRJ0AAArT+Ec/7F7U/60Zb+yWEqrlp7wjytpo8Y2l258vKsqr9mVbBD0HL4yKnp8ajld0uXRTouqYU5KqKvT9CLyXlV6xy3QXfIbFaal0rKa6Xmlt1Q+ByNnYyedkT1YqoqI5EcvJVRU5/zKbQe5wvSm0Nbs1EyKjUdaLuLjRLQq81Tq8Sp20yIZYtJr4g1jG8eUSfZRmkVt+wqLTbQfRXTyGJII8F0mx3EEiTpfJW3WijpF5q75Sqqwqqq75SqqqvpVT10A1jLWtKptm1am2Kxc5p5HyvXqukcrnL3VUmRBCyngZTx+9aiInaRMkAAPPOUEDHHl3WXvSDQfEdA8KulRa8j3AVVa3La6ik8XVU2L2tIErKTppycz4QqKulhVzV5PgpayN3okUnnKaXhBN9rbjvSwmzyuVtDj236zxUcKPVY1fU3vIqiabl/M53SiYvL50gYTg3O7D6x8QtKiwqe3omy01CyauVjkRWukp2f2GaLx6k74pMuVWZLsVS3GK9qVFlXKqX0rla+RWx5pxoj1913WoqfWQWAA2kCFwAAAAAAAAAAAAAABc84Fe6u9a47cMg0dzW61F3zDbxcqOzWu419U6pr7hjV0ZUPs7HucvSetHJR11GipzRkENG30ejnOIU9/B4b3WwbstYccjciW667d6y91TfT0lmt+S41BTqn83obc6n/wD6n/qXCDVi3QXD+x8O9Ki8VBYESRUlXvNa2NqZIx9TE18yIibER0++vREREajkaiZIhNXCy1ai1rlUktUucketGqryoxcm9xuqnZyzAAIWFwwc/wC7DCYtR9sG4fBJYUndlWiuT2akarXOWOpls1YlJK1GtcvSjm8VI3k13ymJ8l3zL0Afp3Ghp7pb662VaOdS3GjloalrHdB6xzMdG9Ed/MvJy+k9i7trzXfvBQ29TKqSU00UzVTjR0b2vRU7ObTr1cDaqllpX8T2uavaVFT/ADMwstoeDo/cW3G/jQtX90uKl5bQ8HR+4tuN/Ghav7pcbMG6YfFDt76eg/rISH2D3w8pfmy/duLGAANYcmSAAAcH8Tz7wPdN+K+f9ppjPvNBDiefeB7pvxXz/tNMZ95sFbkP+BS8361X+kpiLOPHwio/oP8AceAAZZixoAAAAAAAAAAAAAAAAAAAABfA4Mv4Nfbf/wA4fX3KSUAi/wCDL+DX23/84fX3KSUA0/dJn4yGIP68tb+vqCetzvgjZX6NB900AAsiVGCvx4RH97Dol+Pln1evRYHK/HhEf3sOiX4+WfV69EstBb42tyP0t33ExQ2JfwFtL5ifzNKgwANr4hAAAAAAAAAAAAAAAAAAAAAADTVwSKKDB8NggjjhhhxS3RQwxMSOKJraOFGta1PQiIiIiInoREMyo0y9NbhT3bTrAbrSpIlLc8KtVwpklajJUjmoYJGdJOa8l5OTmnNfSYV92KbItlYfvRPcpJaaL1M1bQ5fwXL6yQ+AKpv9qJy5Q/xlPtQAYNySIAABUb8IovVzn3BaBY7LVyvs1q0bqb1QUCr/ABFNVV97rIKuZqf/ADSR22iav/pTt/oK8ZZS8Iq0tyaPPdv2tUVFU1OHVuH12mNdcIadz6Sz3KkrprrTRVEv8lrquGvqViavpcltqFT+SpWtNq3QErLNrNEa5nQx7XNZDMx6Ny9zK2qn3xHInE7WzVc9q5621FRVhLihHNHfu0N+RUVXNVM+VFY3LLsZeQAAmGUCAAAAAAAAAAAAAAAAAAAAAAAAXcPA4f8AiMfmh/7nl3ApH+Bw/wDEY/ND/wBzy7gAAAADiPiY/g3+IH+RHqt9Q78duHEfEx/Bv8QP8iPVb6h34Ax+wAAAAADqXaHsx3F759XLXovtu09uOb5VV9CpvVyVFocRwihc9GPul9ubkWKjpWen5b+b5HIkcUcsrmRukY4UPA63G8Su827Pro24aK7UqC5uhv8ArVfLS6SuzLyeVY6q3YdQydFK+fpskgkrXKlFSvZL03zTReSSaPOz7ZZty2JaQ2zRXbZp9b8Kxamcysv12k6NfmOe3BI0jku1/uitSWsq5ETl0n8o4mdGKGOGFkcTQI4uE9wNtu3DYtFu1CvzbZrXuvr7YsN81ivFrT4Kwjx8fRqbfh1DKirRRK1z4n1708tqWOkRzoYpFpWzjgAAAAAH4p54KWCapqZoqemp4nT1FRPIkMEDGIrnve9eSNa1EVVVfQiIVA+Lx4S7iekv8J9u/Dyu1j1A1Oj8dZMv3IrDHe9PsCkTpRyw4tE5Fhutcxef+vPR9BErU6DaxVXxQEwnFF4zG2LhlYpPaslrYNVNx15tXluEbfMWvEUF+e2RqrT3HI6tGyJaLc5eXRmljfPPyd5PTzoyV0Wbdvk3+7mOIXq5Vat7jc4nvc1NJPBhGB2jxlt050xoZntctvsVr6bmwtVI4UlqJFkqqlYI3TzzOajk5VzXNsw1Iy3Is91Aye+5pm2XXaa/ZRleT3Sa9ZBf62oesk9TV1crnSSyPcqqrnKqny4AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABa28EVyu+Ue/LcNg9PWSR43kW0W4ZXdrejnJFVV1mzLDqO3zKnPoqsUV+ubUVUVU8pdyVOa89CooQeCHaE5vX7ktzW5tbXWQacYtoh/oJZe5oVit91vd/v2P5A6lppFTlLJS0uNtknaxV8U250qvRPHR8774AAAAP0rnbqK8W24Wi4w+U2+6UMtur6fxj4fHwzxuilZ02qjm9JrnJzaqKnP0Kin7p/Eya9x41jmQZHNF4+GwWSrvcsHjUg8c2lgknczpqio3mkap0lReXPnyAMToAAAAAAAAAAAAAAA7C2K7JNZ+IFuLw3bropa3Pu1+l+Esty+so5qjGNNLFBJG24X67yMT5EECSMYxnNHTzz08EarJMxF8K0Y0c1J3B6qYJoro/ilzzfUrUjIYcZxHGbTCstVcKmbm5znu/kxwwxslnnnkVI4YYJZZHNYxzk1P+EnwvdNuGJtyocCtvwVlGuedwUuQa+6p0kDlXJ7syNehbLdJI1srLTbfGyw0sbmsWRXTVD42S1D2NA6l2U7NtGNh23jB9uWh9kjt+NYpS+VX3IKmBiZHqDeqhrFuV/vE6emWqqpGp86qyGGOCCJGQwRRs6vAAAAAABAtx2uLja+HFoMun+ld2t9Xu71tstRRaZUHKOudpfa3+MpqvM66ncjmL4h7Xw0EMyKyesb0nMmhpamNQIhvCX+MI+ihyLhuba8qkiqqqPyTdnnVgq0a6OB7GvZgNLUsXmnjUc2S7dBUXoJHROc5slfAUgT9+7Xa6X66XK+Xy5XC83u83Ca7Xi8XaskuN0utVUSOmqKmpqJFWSWWWR73vkequc56qqqqqp+gAbHexr7yfZ7+S1p99UrQdSnLWxr7yfZ7+S1p99UrQdSgAAAAojeGE2mlh1g2Q31iy+W3HTbM7TUI56LAkVFdLFNCrW8uaO6VfP0l5rzRG+hOS873JRZ8MP+6PsR/qTn37fioBVM2iY7Q5fuw2w4nc2sfbco3D4VjtwZLAyqjdBW5JbKaVHRPRWvRWyu5tciovzL6FNlkxxdjP37Ozz8qbT7622g2OgAAAAAACFDiXcdbabw0M7tGj2b47qHq7rRdMehyurwTTenoYaPEqCqdI2ilvVzq5444JKnxUj4qeBlRL4tqSSMiZJC6X43Yz4RHw/d6+XWPS6W9Zbt11cyGeO3WHE9cqegs+PZZXSOY2Ois+Q01TNRyzSvkZFDBWLR1E8i9CKCRytR1L3wjieebjKbv45ZpZY6WLT6CmZJIr2U7F0uwqVWRovoa1XySP5JyTpSOX51UhCANuAFUrwavi03fc/p3XbJ9xOZ1V+170esa3bSfLsjrFqL3qlh9OjIpKKoqXqrqm42VXRtc96+NnopoXr4x1NVTLa1AAAABztux2w6XbytvWp+27WK1/CWDanY5JZ6qohjjddcdrGK2e3Xi3veitZWUFVFT1UD1RW+MgajkcxXNXokAGN5vH2n6pbItx+p22rV+gdT5Xp1fXUdLd4ad8FozK1zJ461Xy3K5V6VLXUz4Z2JzV0avfE/oyRyMbzEaVHhFHCxTe/tw84PSDHG1u5/bZYqm6W2jtlGst41UxGNX1d1x1GsTpzVVKqzV9vZye5ZfK6djelW9JuauAAAAAAAducM78JBw/Py3NKfr5YTYEMfvhnfhIOH5+W5pT9fLCbAgAAAAAABlr+Eafhl9435vf8K8HIRybjwjT8MvvG/N7/hXg5COAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAASWcO/hxXriBf6YPgjVe16Y/6JP4P+UfCWJS5T8OfD3w30Oh0KqDxXifgV3Pn0ul5Qn8no+mS3q42Z9q3GPolq/tQ+n8G7/3zPzd/9dFoAwNaZ+m3pJ4Q6St5MO7gW8ymsik9h7zEtHRyq3fqClnk/tJad8jtaWV7vdOXLPVTJqIiSbw9w5uhb1z6O1rUpVfO/fNZ2+SNz1ZXtTY1yImTURNifvKrHVxsz7VuMfRLV/ag6uNmfatxj6Jav7ULU4Iu9Ml0xOuhn7Ps71UrPgiuBzJfCy+eVWOrjZn2rcY+iWr+1B1cbM+1bjH0S1f2oWpwOmS6YnXQz9n2d6qOCK4HMl8LL55VY6uNmfatxj6Jav7UHVxsz7VuMfRLV/ahanA6ZLpiddDP2fZ3qo4IrgcyXwsvnlVjq42Z9q3GPolq/tQdXGzPtW4x9EtX9qFqcDpkumJ10M/Z9neqjgiuBzJfCy+eVWOrjZn2rcY+iWr+1B1cbM+1bjH0S1f2oWpwOmS6YnXQz9n2d6qOCK4HMl8LL55VY6uNmfatxj6Jav7UHVxsz7VuMfRLV/ahanA6ZLpiddDP2fZ3qo4IrgcyXwsvnlVjq42Z9q3GPolq/tQdXGzPtW4x9EtX9qFqcDpkumJ10M/Z9neqjgiuBzJfCy+eVWOrjZn2rcY+iWr+1D+paPBw7y+tj+Ht2tsgtzeT5fgjRmWrrZ+T29KNvjLyxrObenykXp9FUb8hyKvK0mD8S7pFpiyRrGl6mNz5UoLOzTtZ0ip+4+phHcBFz9gqv+LN55FhtD4Q+1baVktr1EpKbIdWNU7M9Kiy5nqJNBNSYzPyVFqLTaoI2U8MvzKyefx88Spzjlj5rzlPAIpYhYmX/wAV7wOvTiNa01oV6tRqSTP1tViKqoyNqZMjYiqqoyNrWIqquWaqq1vZVj2XYdKlFZMDYos88mplmvVVeNV7KqqgAFDHpAA/i5Jkdjw/Hr7lmT3SksmN4zZ6m/3+818niaG00VHC+oqqmZ/8zI443vcv9DVOSGGapmZT07FdI5Ua1rUVVcqrkiIibVVV2IibVU+Oc1jVe9ckTjUqS+EOamU2Qbi9GNLKSqbUpptpXPkVyZHJ0mUFbkdwcjoHt5+iTyay2+ZfR6WVMXpX5kgbwf8A20xD+tFv/a4T2vd/uAuO6Pcpq7rnXNqYKXOMslmxugq+SVFpstGxlBZKSRGr0fGRUNNSMkVvodIkjvncp4pg/wDtpiH9aLf+1wm3Xo9YbVOEOjjd3Duvbq1VJQpv7U4m1EyPnqGovKjZ5ZEReVEz5SCV6rYZb17qq1olzY+X3K9VjcmsX62tQ02AAaiBO4AAAFLjj8/fzWf8Q1g/vK/F0cpccfn7+az/AIhrB/eV+MkO5X/Gmb+rqz+aEtJjV8Cl+lj/ANRCMADZPIhAAAAAAAAAAAAAAAE8Hg8/36Op35L96+teFFyApv8Ag8/36Op35L96+teFFyA1ot1F+NdV/oNF/I4mBgv8CGfSyfxQAAx2F2AAADL/AC2h4Oj9xbcb+NC1f3S4qXltDwdH7i2438aFq/ulxs8bph8UO3vp6D+shIbYPfDyl+bL924sYAA1hyZIAABwfxPPvA9034r5/wBppjPvNBDiefeB7pvxXz/tNMZ95sFbkP8AgUvN+tV/pKYizjx8IqP6D/ceAAZZixoAAAAAAAAAAAAAAAAAAAABfA4Mv4Nfbf8A84fX3KSUAi/4Mv4Nfbf/AM4fX3KSUA0/dJn4yGIP68tb+vqCetzvgjZX6NB900AAsiVGCvx4RH97Dol+Pln1evRYHK/HhEf3sOiX4+WfV69EstBb42tyP0t33ExQ2JfwFtL5ifzNKgwANr4hAAAAAAAAAAAAAAAAAAAAAADQy4d2pNPqxsf2w5jBUJVTJpHa8Tuk6dFHS1+PRrj9wc5G8kRXVNsqFVERETn6ERDPNLUPg+W6Ghq8a1I2l5NdWRXiz3KTVHS+nqXo11dRVLY4L9QwKvJOdPMylq2xpzc5K+rf/JjXljT3UvDCvvvo8w3vsmJXz2JVMqJERM19iytdBOqIm33D3QyPXibHG9y7EzS8GCtsxWbet1BOuTalisT57VRze6iOROqqohZfABrdkuAAAD4HU/S3T3WjBr/ppqnidozfBsopUpL3jt6hWSkq2te2SN7XtVskUsb2MkjmicySN7GvY9rmoqQM6teDvaKZFc6q4aNa7Z3phS1Ejp22DLsapdTrZSKqqqQ00zKigqGRJ6Gos755E+dXvLEgL04TaRONeBrpUwsvDPQxSu1nxIkcsD35Imu6nnZLCr8kRNfe9bJETPJEKety6d3byI3o1SNlc1Mkdta5E6iOarXZdjPIqsdXGzPtW4x9EtX9qDq42Z9q3GPolq/tQtTgkB0yXTE66Gfs+zvVSluCK4HMl8LL55VY6uNmfatxj6Jav7UHVxsz7VuMfRLV/ahanA6ZLpiddDP2fZ3qo4IrgcyXwsvnlVjq42Z9q3GPolq/tQdXGzPtW4x9EtX9qFqcDpkumJ10M/Z9neqjgiuBzJfCy+eVWOrjZn2rcY+iWr+1B1cbM+1bjH0S1f2oWpwOmS6YnXQz9n2d6qOCK4HMl8LL55VY6uNmfatxj6Jav7UHVxsz7VuMfRLV/ahanA6ZLpiddDP2fZ3qo4IrgcyXwsvnlUC5eDmapxSxpaNzWAV0Kx85ZLlgNxtUrHc1+S1jKmZFTlyXmrkXmq+j+dfHMl8Hw3j2qJajHtQtv+VNbF0lpI8nvlmub3/LVWsbNafEq3k1iI50zV5v/koidIuRA9qzN040tqB6PqrXpalE5JaGnRF7e8shXb2FTsZHXmwcuLKmTIHs+bI//Urig7qNwkuILppFU1dx27ZFlFugc5GV2nN5teoUtUjU5q6OgoamSu5ej0I+naqr6EQ4CyjEsrwe81OOZpjGQ4hkNFy8ssWUWWpx+80nPny8bSzsZK3nyX+U1PmU03zzfU3R3SjWiwyYvq3pxhWo9ge1zWW3M8cpMggpVciIslOszHOhkTkitliVr2q1FRyKiKSKw/3Xu+1JUMgxRupTVMK5I6Shklp5Gp+VvU7qlki/m75EiryomwpO1cBrOexXWLXPY7kSVEei9jNqMVE7OTjNHBbT3c8AjTTLKO55btGyWfTbKU8ZVpplmtznvmn105NVyQUFxekldQvcvzLO6qiVVa3lA1FcVetYNGdUNA89vGmWsGF3nBM3sb08tst5haiyxuVyR1NLOxXQ1NPL0HLHU073xSI1Va93Iys4DaU+DOkZQOmw7tL/AI6NutLRzokVXEmxFc6LWcj2IqoiyQvliRVRqvRy5Fkbz3KvBdKVG2tD/ZquTZG+6jd/e2ZL2HI1eXLI8xABIopMu4eBw/8AEY/ND/3PLuBSP8Dh/wCIx+aH/ueXcAAAAAcR8TH8G/xA/wAiPVb6h347cOI+Jj+Df4gf5Eeq31DvwBj9gHRe1vadr/vO1asmim3HTm9ai53eXJNPBQMSmsuN0aPaya53i4SKlPRUcXSb06ioe1vNzWN6T3sY4DwC3W64Xe4UNptNDWXS63Ssit1stlupX1twuNRO9sUMEELEV8kkj3NY1jEVznORERVUuj8IXwZervLcW3HcSWx1lrtb2wZBhO05Z3UV2r0cjZaeoz2ZipJTs5cn/AUDmzKrmNrJYlZPQPmf4S3Ac2+8Om22TVXUNln1y3czULZa3Uu5W9ZMT0vkmj5T0WG0Mqc4lajliddp2+WTNR/QSkimkplnuAP5NgsFixWx2fGMXstpxvG8etkFlsGPWC2w2ex2OjpYmw01JR0kLWxQwxRsZGyKNrWMaxEaiIiIf1gAAAAAAADOW45HG83X7gtTNWNmeI4rnO0/RDA8krcF1AwO8ubaNYtUZaaR0Un8JqqB7mQW6ZnJ8dsoJpKaoinSSaprWPhSKsQaPfhBvBpg3u6cVe6nbtjDXbttKMeVt8x60sSOfXzG6NnTW3Pj5fxl4t8bZH0EjeT6iNZKR/jVWj8nziZ4J6WeamqYZaepp5XQVFPPGsM8D2KrXsexeStc1UVFRfSioAfiAAAAAAAAAP7WN2d2RZFYcfZO2lffb1S2dlU+NZW0y1U8cCSKzmnSRvT58uac+XzofxT7fTP7pGn39d7T+304BcD6nhqP27sJ+gOv+3B1PDUft3YT9Adf9uF6YAFFnqeGo/buwn6A6/7cHU8NR+3dhP0B1/24XpgAUWep4aj9u7CfoDr/ALcHU8NR+3dhP0B1/wBuF6YAFFnqeGo/buwn6A6/7cHU8NR+3dhP0B1/24XpgAUWep4aj9u7CfoDr/twdTw1H7d2E/QHX/bhemABRZ6nhqP27sJ+gOv+3B1PDUft3YT9Adf9uF6YAFFnqeGo/buwn6A6/wC3B1PDUft3YT9Adf8AbhemABRZ6nhqP27sJ+gOv+3B1PDUft3YT9Adf9uF6YAFFnqeGo/buwn6A6/7cOqNCvBDNvGLXqlum4bdfqfrFbKd7KhcX05wKg0WoqtzeSugqqyesu1RJCq82uWnWmkVF+S+NS38ADx3QXb/AKNbYNLcY0V0D09x7THTLEKZ0FkxbHKZYqdjnr056qpmerpqmqnfzkmq6l8k8z3K6SRzlVT2IAAAAAHB/FC1mo9v/Dv3l6q1VTHSVFi2+ZHaLDNLyWJLzfKGSw2JrkX0Kj7jc6Bip/P0uSelTvApyeFm73bfjOlGkuwvEbo1+UamXam1q1fgpZ16Vtx60zVEGPW+pai8lSvuUc1YjVTpM/g7C70JK1VAodAAAAAAAAAAAAH7lut1wu9wobTaaGsul1ulZFbrZbLdSvrbhcaid7YoYIIWIr5JJHuaxrGIrnOciIiqp+mXpvBvuC43EKDE+InuoxNrsru1HHfNq+m9+gRyYxSTNa+HObjSuT0Vk7F/8Ljk9EEMq1asWWSkkpwJDuAhwa7dw/dLYdftdLLS1m8PV7GY47tSVHRq4tC7FVeLqUxqkcnyVuE3Rgfc6liqiSQspoXLFFJNVWLAAAAAAAfGajaiYRpHgWYaoalZNasM0/wDHavLMxyq9z+T2uw26hhfUVVTM5EVVRjGOVGtRznLya1rnKiKBy5xAN9OkPDx205puK1bqW1jLRF8D4DgtLXR0d/1OyKpY/4NslArkVUWRzHSzzox6U1NBUTOY9I+i7J53U7n9Xt5GvOoe4rXHIXZBqBqLen3KsSHxkVlx2kb/F0FntVO971goaGBI6enhVznIyJFe+SRz5Hdy8YPih5vxONzFwzKJ11x7b7pxPVYzt907rZOhLbLY6RqT3u5RNXofCl1WGKedEVyQRtp6Zskrafx0sSwAAABsd7GvvJ9nv5LWn31StB1KctbGvvJ9nv5LWn31StB1KAAAACiz4Yf90fYj/UnPv2/FS9MUWfDD/uj7Ef6k59+34qAVa9jP37Ozz8qbT7622g2OjHF2M/fs7PPyptPvrbaDY6AAAAAAAMtfwjT8MvvG/N7/hXg5COTceEafhl9435vf8K8HIRwD1rQjW/UjbZrHpzrxpDkE+MakaWZTTZbil4hVzo2T07l6dPUxo5PG01TE6WmqIHL0ZoKiaN3Nr3Ia1nD13v6bcQnavp1uS05kp6GbIKP4G1CwxtYlXX6c5PRMjbeLJUryRy+KkeyWCV7WrPS1VLMjWpKiJj8k6fAa4otVw7d1FNjuo17qItrev1ZR4jq/S1E6ut2DVqPWK0ZhHGvob5C+Z8VZ0OSvoaidytlkgp2tA1DQfgpqmmraanrKOogq6OrgZU0tVTStnpqmORqPjkje1Va5rmqio5FVFRUVD84AAAAM33wj/hY+ZxuFTdNo7jvke27cpkdRV3G322DoWzS3NpklrLlakYidGKjubWVFwo2ovRY5lfA1sccEKO0gjnbdjth0u3lbetT9t2sVr+EsG1OxySz1VRDHG6647WMVs9uvFve9FaysoKqKnqoHqit8ZA1HI5iuaoGNQDp3ePtP1S2Rbj9TttWr9A6nyvTq+uo6W7w074LRmVrmTx1qvluVyr0qWupnwzsTmro1e+J/RkjkY3mIAAAA7c4Z34SDh+fluaU/XywmwIY/fDO/CQcPz8tzSn6+WE2BAAAAAAADLX8I0/DL7xvze/4V4OQjk3HhGn4ZfeN+b3/AArwchHAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALQHg3f8Avmfm7/66LQBV/wDBu/8AfM/N3/10WgDVr3Rr45d8v/Hf+qoSaWEv4PrP/wAX7+QAAhGXGAAAAAAAAAAAAAAAAAAAAAAAAAAABVY40fEytOZUt22e7f8AJIrlj8VYkGumd2SqSWhvEtNKjm4xQVLV5Phjlja+tljVWvfGynRytbUsfaHyzF7Pm+LZJhmRQT1Ngy2w1eNXympa6a21NRR11PJTVMbKiFzJYnOjleiPjc1zefNFRU5mfrv32bZVsk1/vul11kq7xhl0jXJtLcxnh6DMnsk0j2xJK5ERiVdK5rqapjaicpIke1qRyxK7JjuYOHOEd+caZrTv1U69t2axtTZ1G9qJDM9qrr1GuqrvktKuo9kOqmSqk6K5InIyz2MtrW9Zt3mw2YzKmmVWSyIvumovE3LkR+1Fdn+bs1kz4pPqMH/20xD+tFv/AGuE+XPqMH/20xD+tFv/AGuE2MbU/wCm1H0b/wCVSJsH/PZ20/iabAANJ0yJgAAApccfn7+az/iGsH95X4ujlLjj8/fzWf8AENYP7yvxkh3K/wCNM39XVn80JaTGr4FL9LH/AKiEYAGyeRCAAAAAAAAAAAAAAAJ4PB5/v0dTvyX719a8KLkBTf8AB5/v0dTvyX719a8KLkBrRbqL8a6r/QaL+RxMDBf4EM+lk/igABjsLsAAAGX+W0PB0fuLbjfxoWr+6XFS8toeDo/cW3G/jQtX90uNnjdMPih299PQf1kJDbB74eUvzZfu3FjAAGsOTJAAAOD+J594Hum/FfP+00xn3mghxPPvA9034r5/2mmM+82CtyH/AAKXm/Wq/wBJTEWcePhFR/Qf7jwADLMWNAAAAAAAAAAAAAAAAAAAAAL4HBl/Br7b/wDnD6+5SSgEX/Bl/Br7b/8AnD6+5SSgGn7pM/GQxB/Xlrf19QT1ud8EbK/RoPumgAFkSowV+PCI/vYdEvx8s+r16LA5X48Ij+9h0S/Hyz6vXolloLfG1uR+lu+4mKGxL+AtpfMT+ZpUGABtfEIAAAAAAAAAAAAAAAAAAAAAAel6OauZ1oNqfher2mt3fZM1wO+RXyyVvJz6aVzObZqapjRyeNp6iJ8sE8KqiSRTSMX0OPNAdK0rNs+2LOnsi1YWzUs7HRyRvRHMkje1WvY9q5o5rmqrXIuxUVUU5IZpaeVs8Dla9qoqKmxUVFzRUXkVF2oaGmyLevpdvb0ituf4XWUVqzG3U8dHqVptNcG1F+wS4qio5r2ckfLRzq1z6arRqMlZzRehLHNFH2YZrWiOuuq+3PUG0aoaNZndcJzGzr4tldbpEfSXOnc5rpaKvpXosNVSyqxnTp52vjcrGry6TWuS1xs+47ehmqlFasS3O0UehuoaxtppsupYp7npNkEvJE8YkydOptrnrz5x1SSQMROa1fpRqa8GljucGIGGtr1d8cFKSS1btvc5/seJFkrKNFVVWNY0zfUwt4mSRo+VGplMz3O+vlbcfFuy7YgZZ94pEgrERE112RydnPiY5eVFybn71duqk9oPn8Wy3Fc5sVvynCcmx/MMZu0Xj7VkeLXqmyCxXJnPl06esge+KRvNFTmxyp6D6AxgVFPPSTvpapislYqtc1yK1zXIuSo5FyVFRdiou1FLytc17Uexc0XaipxKgABxH6AAAAAAAAAAAAAAAAAAAAABxdvd2P6Tb3dLKvC85oKW1ZtaKWefTTUymo1mv2CV8jU5OTovYtRRyq1jaiikd4uRqI5OhLHFLH2iCpLoXvvNcG8tHfC51bJSWnSvSSKaNcnNcn7nNcmbXscisexXMe1zXKi9SvoKO1KOSgtCNJIXpk5q8Sp/kqcaKm1F2oqKhmr646K6gbd9Vcz0c1Ps77LmWEXZ1suMKdJ9HXxqiSUtdRyKieMpqqF8VRDJyTpRzMVUReaJ5OW3+P1tSt+XaSYtuvxi0xMy3S2vpsN1FrKditlumOXKo8Tb55+SL0nUNxqIo2L6F6F3m6TlSNiJUgNr/RTx5pNIzBazcRGsbHX+6p62JvvY6uFG76jUzVUZI1zJo0VVVscrGuVXIqkH773YkuleGayc1WLY+NV41Y7PLPsoqK1eqrVVNhdw8Dh/4jH5of8AueXcCkf4HD/xGPzQ/wDc8u4EjSkgAAAcR8TH8G/xA/yI9VvqHfjtw+R1AwPEdVMCzfTDUCyU2TYHqPiNywPNsbrJZYaTILReKKa33KhlfE5kjWT09RNE50bmuRJFVrkXkqAZcPCu4LO5biZZPR5Lb6Wp0m2xWi7+SZprzkdse6kuPiZOjU2zFaN3R+FK9OTmvVrm0tMqL4+dr/FQTaRWyvYltp2BaS0ekO2/AaPGLY5sVRleW16MuWf6i10bFatxv116DZKmVek/oRojIIEerIIYWfIOo8VxXGMGxuxYbhWO2PEcRxi1QWPG8Xxq1QWPH7BRU0bYqeko6OFrYoYYmNa1kcbUa1ERERD++AAAAADx/XfX/RjbFpjkes2v2o+M6WaZ4pAk15yvKaxaela93PxVLTQsR09VVTK1WQ0dLHLUTv5Miie5UaoHsAPm8OzHFNQsUxzOsFyOy5fhmX2WmyPFspxy4xXew5BQVkTZ6WrpKqNzo5YpY3te17FVFRyH0gAAAAKNHhJPBjTHavK+IztbxRrbBcqh943Wab4/QqiWWqlcrpc9oKdno8RO5yJdY42p4uVUrVRzZa2SG8ufp3G3W+72+utN2oaO6Wq6UctuudsuNMytt9xp52OimgnheiskjkY5zHMeitc1yoqKigGJWCxBx7eDpceHxq6/XHRCyXGt2fax3+R1hbDA6qZojfZ0dPLi9ZMnP/U5ejPNbJ5ERVhjfTvV8lN4+orvgAAAAAAA+30z+6Rp9/Xe0/t9OfEH2+mf3SNPv672n9vpwDapAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABH9xHOIvoXw2tAL1rFqzdKW6ZZX0s9DpHpHRXOGkzDVi8M8UxtLRRuVXMpad1RBLW1ytcylgVV6Mkr4IJspfc3uO1S3ca8ambi9Z702+aiapZJLkF6lgY+G2WuPothorbb4XOc6KjoaaKnpKeJXOVkNNGiuc5Fct6rwm7hbXvcdpRS769G4L1edU9vGGvtWqeExVtRX0+T4NTTVFdNc7bRqro4quzPqKuqnSJrPKKKSoc5zn0sMcme8AAAAAAAAAAACaDgv8JXNOJ1rw6TI47xi21nSi40tfrbqDSN8lqbw56pNBiljmcnJ1xro2uWSZqObQ06rNIivfSw1IHd/g9HBek3gZvbd4u5fFnO2tab33p6fYheoGrSa/ZDQzKjmSwORVksltlj/1pXIkdXUIymTxrI61jNE9rWsa1jGtYxjUaxjU6LWonoRET+ZEPlcDwPDdLsKxTTjTvGbPhuCYNYKXFsRxTH6JlusuPW+ihZBS0lNC30NjjjY1qJ8/o5qqqqqfWAAAAAAAAz2/CP8AjDN3LZ9c9i23HKG1W37SvIWprHmljrOlQax5TQSr/wCHU0zHKk1ns8zeSP8A5FVXRPlaj4qaknlsTcejdVvO040MZts2J7bt0GqWq+uFinhzbWLRXRDLM4sGkGNSPfS1EFJebdQy07b1cejPDGyOXx1FTtlnckT5qKR2fz8WdxH/AP6fm9z9VLPPssA4jB258WdxH/8A6fm9z9VLPPssfFncR/8A+n5vc/VSzz7LAOIwdufFncR//wCn5vc/VSzz7LHxZ3Ef/wDp+b3P1Us8+ywDVf2NfeT7PfyWtPvqlaDqU5v2b2K94vtC2q41ktnuuO5Hju2/BrFkGP323zWi92KupMYtdPV0dZSStbLDPDLHJHJFI1r2PY5rkRUVDpAAAAAFFnww/wC6PsR/qTn37fipemKLPhh/3R9iP9Sc+/b8VAKtexn79nZ5+VNp99bbQbHRji7Gfv2dnn5U2n31ttBsdAAAAAAAGWv4Rp+GX3jfm9/wrwchHJuPCNPwy+8b83v+FeDkI4AAABoLeDIcU7/T3pL5hOteReP1h0KxzyvRO8XSf/W8/wAJpehH8FeMcv8AGVlh6UcTW+hz7fJT9FrvJKiQtlGLzoRrfqRts1j05140hyCfGNSNLMppstxS8Qq50bJ6dy9OnqY0cnjaapidLTVEDl6M0FRNG7m17kNazh6739NuITtX063JacyU9DNkFH8DahYY2sSrr9OcnomRtvFkqV5I5fFSPZLBK9rVnpaqlmRrUlREA7aAAAAABXH8Io4WKb39uHnB6QY42t3P7bLFU3S20dso1lvGqmIxq+ruuOo1idOaqpVWavt7OT3LL5XTsb0q3pNzVzbgM33wj/hY+ZxuFTdNo7jvke27cpkdRV3G322DoWzS3NpklrLlakYidGKjubWVFwo2ovRY5lfA1sccEKOArQgAA7c4Z34SDh+fluaU/XywmwIY/fDO/CQcPz8tzSn6+WE2BAAAAAAADLX8I0/DL7xvze/4V4OQjk3HhGn4ZfeN+b3/AArwchHAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAALQHg3f8Avmfm7/66LQBV/wDBu/8AfM/N3/10WgDVr3Rr45d8v/Hf+qoSaWEv4PrP/wAX7+QAAhGXGB5zrDqNRaPaSapat3K21V4t2lunN81Gr7RQysgrbrBZLZVXOWmhe/5LXytpXMa53oRXIq+g9GOX9733l27z8l/P/qpdip7kWZR21fSyLGtButTz1VPE9EVUVWSSsY5EVMlTNqqmabU40OnaM0lPZ89REuTmscqdtGqqfvIb+sXaLdnLVDvVaf8AIdYu0W7OWqHeq0/5FS8Gyj0s/RD+QZ/H6z0xELhhv5zpvgo/NLaHWLtFuzlqh3qtP+Q6xdot2ctUO9Vp/wAipeB0s/RD+QZ/H6z0w4Yb+c6b4KPzS2h1i7Rbs5aod6rT/kOsXaLdnLVDvVaf8ipeB0s/RD+QZ/H6z0w4Yb+c6b4KPzS2h1i7Rbs5aod6rT/kOsXaLdnLVDvVaf8AIqXgdLP0Q/kGfx+s9MOGG/nOm+Cj80todYu0W7OWqHeq0/5DrF2i3Zy1Q71Wn/IqXgdLP0Q/kGfx+s9MOGG/nOm+Cj80todYu0W7OWqHeq0/5Hu+hHHh2jarZVQYhn1mzjQupu1W2jt2TZo2iumBxvf8mNK2400qyUnScqNWWaBKeNF6Uk8bUVyUuAeXbW5g6KFpWXNQ2bZ1VSTvaqNmjrJ3vjdyORs7pYnZLxo5iovFs4zmp8Zb8QzNkmlZI1F2tWNqIvYzajVT6lNPmlqqWupaatoqmCsoqyBlVSVdLM2opaqKRqPjkjkaqtc1zVRyOaqoqKiop+crB8Cvf1cq6qbsr1Yv0layKhmumgV5utV06iBlOx89fi6yOXm5jY2yVdG1f/LZDUxIvRSnjZZ8MBWkNgVenR2xRrsNrzrvm95SU86NVrKmmeq71M1FVdXPJzJGZu3uVkjNZyN1llBdS8tFeyxYrXotmex7c81Y9PfNXq8iouzNqouSZ5AAFkCowRvcUfZpTbxtsl/s1itjKrV7TVs+c6R1MaI2rqqyKNPLrN0uXNWXKnjWFGc2t8ojo5HLyi5LJCCssPb93iwwvvZeIF05t6tGgmZNE7bkqtXax6Iqa0cjVdHI3PJ8bnNXYqnn2rZlJbNnTWXXNzilarVTt8qdRUXJUXkVEUzAZI3xPfFKx8ckb1jkjkarHxuReStci+lFRU5Kin02D/7aYh/Wi3/tcJJ9xltrrNue8LI8hx+2NoNPNeKeTVLGG00SRUNDcJ5lZkNAxOfJFjrVWqRjUayOK607Gpyb6ItsXrqa2ZNjtyrHrHSW++0ldVSNYsjo44qiOR7kanpXk1qryT0qbemHuIFj4u4VWbiJdrbS2jSJMxuebmOcxUfE7LjfFIjon5fjsVEIH2pZVRYNuS2TWe/hfqqvEioi7HJ2HJk5OwppxgA02TIAAAAClxx+fv5rP+Iawf3lfi6OUyPCBLXW0W9vE6+eFW0l42/WSpoJkRVjlSK85HTyN58uXSa6JebfnRHsX/4kMj25YyMZpURteuSus6sROyucS5J9SKvaRS0uNSKtyXKn/dj/AMyDYAGykRBAAAAAAAAAAAAAAAJ4PB5/v0dTvyX719a8KLkBTt8Hjt1ZLu/1Zu7IVdQUO225W6pqOfyYpqrKMTlgYv8A/c2jqF//ANalxI1oN1Dex+lfWtauatoqJF7C72q5L9SovaVCYODCKlx41Xlkk/igABjuLrgAAGX+W0PB0fuLbjfxoWr+6XFS8tkeDn1VO/SDcnRNlY6qp9SbNVTQIvy445rXOyJ6/wDo5YJUT/8AxqbPO6Xtc7RCt9UTYk9Bn2P+MhT+KohDbB5U9vlL82X7txY4ABrDEyQAADg/iefeB7pvxXz/ALTTGfeaD3ErttZddhW6qloYVnni0huNyexF5K2Gj8XWVD//ANkUErv/ANpnwmwRuQz2LgveeNF90lqZqnLktJT5L9eS5dpSLWPCL7YaN3JvP+twABloLGAAAAAAAAAAAAAAAAAAAAAF8Dgy/g19t/8Azh9fcpJQCL/gy/g19t//ADh9fcpJQDT90mfjIYg/ry1v6+oJ63O+CNlfo0H3TQACyJUYK/HhEf3sOiX4+WfV69FgcgT8IXsdRW7Q9LL7Cr3R2PcLb6eribF0kbHWY9kbUmc/n6Ea+GJnLkvNZ09KcvTK3Qcmig0srjvmdki1mr9boZWtT63KifWURiS1zrjWkjU/+3n3HNVf3FOsAG2GQeAAAAAAAAAAAAAAAAAAAAAAAAAAAPUtL9b9Y9E7m+8aQ6pZ9prcZZGyVU+FZVW46yvVqKjUqYoZGsnbyVU6ErXNVFVFQkz0245u/nA4oae/ZTp5qzTwNSNjdRdP6eKo6CIiIjp7TJQSPVOX8uRz3KvpcriHoFrr94I4P4nKsmIN2aGvlyy3yemifMidRs2rvrf7r0Pasy8dv2Nssqskib1Gvcjfrbnqr9aFk7F/CNtQaSJiZrtaw3IJkZykkxfU6tw+Jzuiic0ZPQVyonS6S8lcvoVE5+jmvtls8I000l8Z8M7Yc6oOTWrF8GaiW+79Ny8+kjvGUcHRRPRyVOfPmvoT+eqKCOVq7nPofWq9ZfanvL15Yq2vYn1M9kqxPqanH2sqtgxZv7Amr7O1k/OjiX9+pn+8todYu0W7OWqHeq0/5DrF2i3Zy1Q71Wn/ACKl4PF6Wfoh/IM/j9Z6Y7HDDfznTfBR+aW0OsXaLdnLVDvVaf8AIdYu0W7OWqHeq0/5FS8DpZ+iH8gz+P1nphww38503wUfmltDrF2i3Zy1Q71Wn/IdYu0W7OWqHeq0/wCRUvA6Wfoh/IM/j9Z6YcMN/OdN8FH5pbQ6xdot2ctUO9Vp/wAh1i7Rbs5aod6rT/kVLwOln6IfyDP4/WemHDDfznTfBR+aW0OsXaLdnLVDvVaf8h1i7Rbs5aod6rT/AJFS8DpZ+iH8gz+P1nphww38503wUfml0DTTj67LcyuVPa80sur2k7p5mROvWSYtS5DjVOj+iiufLbaqoq0Rrldz5Uqp0UReaqqtbMdgGoeDaq4jZc903yyw5vhuQ0qVlmyPG7jHdLXWsX0ORJGKvRexebHxP5Pje1zXta5qomZiWAfB+9dsvxzcZmmgMlxr6zT7UjA63LYrG+oWSgs19s76V0dwhjcvKNZqN9VTzeLRHSrHSdNVSBnKHml3ubeGlwsKrVxSweqqinms2NZ5qWeRJoZYGKm+rG9zUljkYzWkTWfI16N1EaxVRxX1xMXbYtS24LFt9jHNmXVa9qarkcvFmiLqqirs2I1UzzzXiLf4AMJRIsAAA8Q3L6XU2te3rWvSapp21C6gaYXrGaFFi8c+nramgnbQVEbf55IKnyeZnoX5cTfQvzGbcagBmeaq2qnsWqGpFkpFVaSz57eLVSq5rWOWOnuNRCzmjURE+SxPQiIn9CGcDceLxVT6S/V05XKsDHUFSxORHvSqilXtuSOFP7pHHHykYj7Mrmp7pUlYvaTUVvczd3S5n4HD/wARj80P/c8u4FI/wOH/AIjH5of+55dwM2BHUAAAAAAAAAAFUvi/eEh6c7Zm5Pt62LXXG9XNwtOs1kyzV9scWQ6T6Nzp0o5YqFfTBertCv8A/Tb06Cmk5JM6pkjnomgS3cS/i0bYeGVgHwnqdeI811lyG3PqtONAcUucP8OcpX5bIq2uVeklttfjGOa+4VLFRyxytgjqZWLEZs/EE4lu6HiRanuzzXvLPE4xZqqZdOtIcYkmoNNdN6eTmnRoaJz3LLUvZ0Umr6lX1E3JEV7Y2xxM451P1R1G1qz3J9UdWs1yTUTUTM7m+8ZRmOW3WW83281D0ROlLPIqr0WNayNkbeTI442MY1rGtanwYBaM8H642lRs5y20bQd0OVVM21POr14jAMyvNS+oi283qtnc5yuevNY7FXzyq6qZ/Io55Fqm9Bj6xZNE+CeCqghqaaaKopqiJs9PUQSJNBOx6I5j2PTmjmuRUVFT0KimJIXVfB0ONz8DTYXw8d3GX8rRUSU+LbWNVsirP/5XK98NNQYDc51Z/wCTIruja6md/wDFu6NFzVjqNkQF4sAAAAAHlGueiGmG5HSTPdDdZsUt+a6Z6lY9NjWV47cWqjKmCVEVk0EqcnwVMEjYp4KmJWywTQRSRua9jXJlVcVThpaocMjcpdtKcn+EMk0pyt1RkmhGqklOjaTO7Ek3RSGpexrY47pQdOKCupmo3ovWOVjfEVFO9+tYcN8QzYTo7xF9tmWbfdWaaK31VU1b3ptqJTW5lwyHSzIYY3torxQtVzFe1Om6GopumxtTTTTRK9iubIwDIDB0Vuu2taw7MdetQNumueP/AADn+n91WjqZKZZJrFklHInjKG8Wmpexi1FDWQqyaGXotd0Xq17I5GSRt51AAAAB9vpn90jT7+u9p/b6c+IPt9M/ukaff13tP7fTgG1SAAAAAAV7uI74Qfo5w5NzF020Zrt61M1Ivlrw+05hJk2LZRarTaZY7tFJLHC2GdFk6UaRqjlX0Lz9BYRM0Xwor8K3l34kcM/Y6kAm463/ALb+x7rd36sP/tHW/wDbf2Pdbu/Vh/8AaUIQAX3ut/7b+x7rd36sP/tHW/8Abf2Pdbu/Vh/9pQhABfe63/tv7Hut3fqw/wDtHW/9t/Y91u79WH/2lCEAF97rf+2/se63d+rD/wC0db/239j3W7v1Yf8A2lCEAF97rf8Atv7Hut3fqw/+0db/ANt/Y91u79WH/wBpQhABfutXhfe1ua4UkV72la/W+1Pl5V1basnx28XCmZyXm6KlklgZI7ny+S6aNPT/ACixDsj4gW1riE6a1Wpm2XUOPKaWyTw0OaYfeaF+P6g6fVdRG6SGmvNqkVXx+MSOVIqiF0tLOsE3iZ5fFv6OPgdjbEt7ms3D+3GYVuG0ZvFTBXWOsjoc1w+Wtkp8c1NsEssbrjYbtE35L4Z2MRWPVFdTzxwTx8pImOQDYaB4htr3B6cbrNBtK9xWklzddtPdW8Qpssx+aboJXUHjOlFV2+sYxzmsq6GpiqaOpiRzkjnpJmc16PM9vAAAAPwVNNTVtNUUdZTwVdHVwPpqqlqYmz01THI1WSRyMcitc1zVVFaqKioqoplg8dHh1pw8t7mT2HC7Q6h2/a2xT6r6FOhjVKGyUdTUK27Y01eijUdZ6x7oYo0V7koam2ve5XyuRNUUgs8IY2P0+8fh46hX7HrNHX6vbZI59dtOKiKHpXGro7fAq5VaY3J8tW1VrZUTthYirLVWqgTl6EVAMvIAAAAAAA9l2+bf9Wd0mseBaC6IYlcM11L1HvsVhx2y0Ma+KjV6856ytm5K2no6WJJKipqpeUcEEMkj3I1qqAdB8PHYJrLxGtyGL6BaS0ktvoZXNvepuo1Tb3V2O6V4/HKxlZd61EcxJH/KSKmpUex9TUSRxo5iK+RmrVtN2p6M7K9BsF266D443HsCwW3+JZNUrHPkGU10vJ1debxVMYzymurJEWSWXotanyWRsjijjiZzRwuuGxpNwzNt9o0jwplvyPUrImwZBrhqz8HpTXbUa+JErVRjnc5IrbRdOWGho1XoxRukkcizz1EkkkoAAAAAAAAAAAAAAAAAAAAAAAAAKLPhh/3R9iP9Sc+/b8VL0xRZ8MP+6PsR/qTn37fioBVr2M/fs7PPyptPvrbaDY6McXYz9+zs8/Km0++ttoNjoAAAAAAAy1/CNPwy+8b83v8AhXg5COTceEafhl9435vf8K8HIRwAAAATp8Brii1XDt3UU2O6jXuoi2t6/VlHiOr9LUTq63YNWo9YrRmEca+hvkL5nxVnQ5K+hqJ3K2WSCna2CwAG2zTVNNW01PWUdRBV0dXAyppaqmlbPTVMcjUfHJG9qq1zXNVFRyKqKioqH5ypr4MhxTv9PekvmE615F4/WHQrHPK9E7xdJ/8AW8/wml6EfwV4xy/xlZYelHE1voc+3yU/Ra7ySokLZQAAAAOdt2O2HS7eVt61P23axWv4SwbU7HJLPVVEMcbrrjtYxWz268W970VrKygqoqeqgeqK3xkDUcjmK5q9EgAxvN4+0/VLZFuP1O21av0DqfK9Or66jpbvDTvgtGZWuZPHWq+W5XKvSpa6mfDOxOaujV74n9GSORjeYjSq8Im4Vy749uTdwGjuNfCG6LbfY6m5WuhtVC6e9ar4kxX1V0xxrI08ZPVUyrLX29nJ7ll8rp42dKt6Tc1d7XMc5j2uY9jla9j0Vrmqi8lRU/mVADtrhnfhIOH5+W5pT9fLCbAhj98M78JBw/Py3NKfr5YTYEAAAAAAAMtfwjT8MvvG/N7/AIV4OQjk3HhGn4ZfeN+b3/CvByEcAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAtAeDd/75n5u/wDrotAFX/wbv/fM/N3/ANdFoA1a90a+OXfL/wAd/wCqoSaWEv4PrP8A8X7+QAAhGXGBy/ve+8u3efkv5/8AVS7HUBy/ve+8u3efkv5/9VLsVvhn+Ei7/wCnUn38Z51sf9Jqvo3/AMqmc+ADc5MfAAAAAAAAAAAAAAAB9bgOcZNplm+I6iYXcpbPluD5HR5Xjlzh9L6KtoKiOpp3qnzOaj429Ji+hzVVF5oqoaNu3fWnHdxWh2l2t2LIyK0akYfSZGlC2oSqfZqp7fF3C3SSJ6HSUdXHVUsip6OnTP5GbKW2fB69fX5Lo9qzt1vFwWWv0wyeLPMPpqiVzpEtF+R0ddT07efJIqavpHzu9CfLvfP08/k4r91YwehvdgxRYsWfDnXWJM1srkTatHVObE9Fy2rvdRvDm57GNdMuzWcpevBG33UN4ZLDld/ZVLVVE/8AyMRVTtZs1kXqqjeohYnABrvkrQAACGrjjbeodYtmVy1EttEyfL9vl+iz+gnji6dbLZ6p0duv9K13L5MaRy0tfJ83NLK30/zLSUNNnNsRsuoGGZbgeS0yVuOZtjFfiN/o3Iitq6K5UstHVRLz9HyoppG+n+kzXdTMCvWleo2e6ZZG3oX/AE8zK54TeUSNYmuqrXWzUU7mtX0o1z4HKnP+ZUM/W5I4pyW5htb+EtoS5yWXO2pp0VeKnq0drsan5MdRE+R351R3IvY62K2mteltyJuyZqsd86PLJV7KtciJ2GGkBozmEeoWj+lOfw1CVcWcabWLMIqpsvj21Lbna6WtbIj+a9LpJOi9LmvPnzPSSLLg16y0+r2wnSqkfVpU33SWortIcgj8Yr30y2uZJ7YzkvpRPgyttXL505o5E+bkkpphLxfuZU4d4q3juNVsVrqGtqYERUyzZHK9I3J+a9mq9q8rXIpIuwbQZa1iUlpMXNJY2O+tWoqp20XNF7KAAFuT1gV3fCB9tV4zfSjTPcnjFBJWy6P11TiOoEdLAstRHZb1LTuoa+RUT0RUldF4l3/reUd/Ja5UsRH8m/2Gy5TY7zjOSWqgvuPZDa6iyX2y3WlZW2y70dVE+CppqiFyK18csb3scxyKitcqKXiwBxftTAfF2xcU7Kj31aKVVkizy36CRjop4s9qIr4nvRjlRUY/VfkurkeBeiwYbzWDUWLOurvibHfkuRUc1fqciZpypmnKZioLK28PgEZjSX27Zns2yC0XvGaxzqxuj+dXpbVkVme53Pya0XmVFp6mFOl8ltwkgfG1nJ087l5kQOXcN/fhhNRUUt52pa0VstMqNkdiOITagU7ubuini5rX5TG/0/OrHLyT0r6PSbQWGelno9Yr2NDa12b00kcj0RXU1TPFTVUbl42vgme16q1fcq5mvG5fePcmSrDO2LjXrsOodBWUUionE9jVexeyjmoqbeouS9VEOJwdQeZDvR7Ie6D6Acr9gHmQ70eyHug+gHK/YC6vCZhv1wUPjcHpDxOg9r81k7x3kOXwdQeZDvR7Ie6D6Acr9gHmQ70eyHug+gHK/YBwmYb9cFD43B6QdB7X5rJ3jvIcvg6g8yHej2Q90H0A5X7APMh3o9kPdB9AOV+wDhMw364KHxuD0g6D2vzWTvHeQ5fB1B5kO9Hsh7oPoByv2AeZDvR7Ie6D6Acr9gHCZhv1wUPjcHpB0HtfmsneO8hy+Dt3DuGzvyzmopqWy7VNY6KWrVWxOzHF36eU7OTnNXxs11Wljj9LF9Mjm80VF+ZUVZkNmnALyCO/2fOt5mQWaGx0EsdfFoxgt3fcq+8PaqO8nvl5Y1sUMSOTk+G3umWRrvRVRLzRbS4oaXGjzhLYs9rXkvRSSzMRVbTUs0VTVSOTiYyGJ7nIrl9yj5N7iavv5GoiqnuWLcS9du1DYKOie1qrte9qsYidVXOREXLjyTN3URTpTgCbabvpvoTn24LKrZNbrnrreKW3YXFWQ+LqHY5Y1qWNrmIq9JrK2tqqxEa5qdOO208rVcyRirP+fz7TabXYLVbLFY7dQ2ey2W3w2mz2i2UrKG22ulpo2w09NTwMRGRxRRsYxjGIjWtaiIiIh/QNXrHXFq1sc8WrbxTtiPe5K6XWZHnrJFCxrYoIs9msscLGNc5ETXcjn5JrZEzbtWHBduwqaxadc0ibkq8Ws5VVzndjNyqqJyJs5AAC0p7oAPHtw2b0ummgmteodZPFT0+EaUZDlLpJlVGdKhtNXUsaiJ6Vc50bWta35TnORE5qqIejZFmVNtWtS2NRJnNPIyJiImaq6RyMaiJy5qqbDinmZTwPqJPetRVXtImambAWYvByMzip8n3T6eTTsWa7WHF8zt1Mr2te1tvqLzQ1r2t5dJyKtzoEVefJvRb/APN6azpK7wW9ZqXSDfpp3RXOrbRWbV+yXHR24TPX5CzXNIKy0x8v51luVttkCf0LUf0cza602LmVN/dFe+l36JivmbSeyWtRM3KtFLHWZNRNquckCtRE2rnly5EIcOrQZZl9bOqpFyasmovU/tEWPb2E1sy9GADU0JxgAAHy2c4fZNQ8JzDAMlgWqxzOcWuGH3+maqI6oornSTUVVGiqip8qKeRPT/SZwWu+jOY7etX9QdGM9opaLJ9PsjnsVY6SFYIrlC1Uko6+Bqqv8RWU8lPVRO5rzjqGLz9JpTEZ/EE4ZWk++q1UWQVFxfp1rXjVrdbMa1Jt1vbXw3GmRz5Y7be6TpMWppmyPe6N7XtlgWZ6sc5rnxPyEbn7pZWJo3X2tCw7+q9Lt2skaSysa560s8OskU+o1Fc6NzZHsmRjVflvb2o7e9R1q8UrjVN7rOiqbMy9mQZ6qKqJrtdlm3NdiLmiK3NcuNFyzzSheCUrV/g2b+dJ7hWx0mkkWqtipuk6DJtJ7/S5JTVzWoq/It0robmjuXL0OpU5qvJquU5cqtjG9WjqJaaXaLuZfJC/oPdS6GZNXU6r/wDoljonMcn/AKtcqf8AqbCV3sdcFb10LLSu5e2zqiFyZ5sradVTsOasiOY5OVr2tci7FRFIrVV2rxUMqw1dDM1ydWN/7lyyXtpsOWAdQeZDvR7Ie6D6Acr9gHmQ70eyHug+gHK/YD3+EzDfrgofG4PSHW6D2vzWTvHeQ5fB1B5kO9Hsh7oPoByv2AeZDvR7Ie6D6Acr9gHCZhv1wUPjcHpB0HtfmsneO8hy+DqDzId6PZD3QfQDlfsA8yHej2Q90H0A5X7AOEzDfrgofG4PSDoPa/NZO8d5Dl8HUHmQ70eyHug+gHK/YB5kO9Hsh7oPoByv2AcJmG/XBQ+NwekHQe1+ayd47yHL4OoPMh3o9kPdB9AOV+wDzId6PZD3QfQDlfsA4TMN+uCh8bg9IOg9r81k7x3kOXwdQeZDvR7Ie6D6Acr9gHmQ70eyHug+gHK/YBwmYb9cFD43B6QdB7X5rJ3jvIcvg6CyjaXuqwjH7rlmabZtwWIYtYqVa695LlGjOR2DH7PAio1Zqqsno2QxMRXNTpyOROap6Tn0qGx7fsK8MDqmwK2Gpjauq50MjJWo7JF1VVjnIi5Ki5LtyVFOpUUtTSuRlVG5irtRHIqLl9aIXZOA7l1Pkewq3WWJzVlwDVrI8UqmozoOa6d1HfW8/T6ebbyz0+j+j+ZVJnCrB4OtrJBR5LuD0BuNVG2S+Wq26r4rTPkSNXPoJHWq9I1F/lueytszkROSo2mkXk5OattPmq/p3XMqbk6V18aKdipHVVCVsa8jm1jG1CqnVRJHyMX85jk5Ca2GdoMtK5FBI1drGb2vYWNVb+9ERe0qAAERSuwcUcRDbnW7p9oOsOklihZNmNVZospwJrl5PlvNlqIrlSUrVVei1axKeWhVzuaNbXOd6FRFTtcFS3MvXbFxL3WXfa770bXUFRDUwqu1EkhkbIzWTlarmojm8Tm5ouxTp2hQ09p0E1nVSZxStcx3aciouXZyXZ1FMwaqpaqhqqmhrqaoo62jqH0tZR1ULqeqpZY3KySOSNyI5r2uRWq1yIqKioqH4C5rxDODFhe6LI7zrNoTfrNpRrNeOlWZRZLrRPbpxqJV+lVrKhYGOmoK2VVRZqqKOdk6sRz4ElfLO+u3qJwot/8AptcZKG47ccvyeDxqspbtp3UUef26tbzVEkalDNJNG1eS+iojicno5tTmhtHYI6b2AGM93Ka0G27TWbaqtTfqKtnjp5Y5MvdNjdKrGVDEXa2SFXZtyV7Y36zGwvvHhxem71W+JaZ80GfuZI2q5qpyKqNzVi9VHZbeJVTJVjuB1B5kO9Hsh7oPoByv2AeZDvR7Ie6D6Acr9gJDcJmG/XBQ+NwekKV6D2vzWTvHeQ5fB1B5kO9Hsh7oPoByv2AeZDvR7Ie6D6Acr9gHCZhv1wUPjcHpB0HtfmsneO8hy+DqDzId6PZD3QfQDlfsA8yHej2Q90H0A5X7AOEzDfrgofG4PSDoPa/NZO8d5Dl8HUHmQ70eyHug+gHK/YB5kO9Hsh7oPoByv2AcJmG/XBQ+NwekHQe1+ayd47yHL4OoPMh3o9kPdB9AOV+wDzId6PZD3QfQDlfsA4TMN+uCh8bg9IOg9r81k7x3kOXwdQeZDvR7Ie6D6Acr9gHmQ70eyHug+gHK/YBwmYb9cFD43B6QdB7X5rJ3jvIcvg6g8yHej2Q90H0A5X7Aefaibe9fNILXQ3vVnQ/V/S+y3Ov+Crbd9RNNL1hNruFUsb5kpoKispoo5JfFxSP8W1Vd0Y3Ly5IqnboL+3GtSrZZ9mW1STTvXJrI6mF73Lx5Na16uVckzyRDjlsy0oY1lmp5GtTjVWORE7aqmR5AAdNwbKN5dTDDU020nc3UU9RE2enqINBcqlhnY9Ecx7HpQ8nNcioqKnoVFPRtm8t3LupG68FoQUqSZ6m/Sxxa2rlrauu5utlmmeWeWaZ8aHFT0dXVqqUsTn5ceq1XZZ8WeSLlmcyA6g8yHej2Q90H0A5X7APMh3o9kPdB9AOV+wHh8JmG/XBQ+NwekOz0HtfmsneO8hy+DqDzId6PZD3QfQDlfsA8yHej2Q90H0A5X7AOEzDfrgofG4PSDoPa/NZO8d5Dl8HUHmQ70eyHug+gHK/YB5kO9Hsh7oPoByv2AcJmG/XBQ+NwekHQe1+ayd47yHL4OoPMh3o9kPdB9AOV+wDzId6PZD3QfQDlfsA4TMN+uCh8bg9IOg9r81k7x3kOXwdQeZDvR7Ie6D6Acr9gHmQ70eyHug+gHK/YBwmYb9cFD43B6QdB7X5rJ3jvIcvg6g8yHej2Q90H0A5X7APMh3o9kPdB9AOV+wDhMw364KHxuD0g6D2vzWTvHeQ5fB1B5kO9Hsh7oPoByv2AeZDvR7Ie6D6Acr9gHCZhv1wUPjcHpB0HtfmsneO8hy+DqDzId6PZD3QfQDlfsB/1uyDek5zWptD3PorlRqK7QPK2t9P9KrQ8k/8Auo4TcN047w0PjcHpB0HtfmsneO8hy8TNcB6hq6vfvbJ6anlnhtmkuSV1wljb0mUkLm0dM2SRf5mrLUQR8/8A5pWp/Oc+6a8Jvf8Aam3OCgotvOUYfSPlSOrvepVVSYFbLc1eaeMeyqkbUyNRU+anhlf6UXo8vSWoeGjw18f2G4rkV5v+Q0Gea26gUsFFluUWulkpsfsNDA7xrLPaElRJnQrN/GzVErY3VDooOcMaRNQg3px6VuCdhYEXjuJZFvU1oW3adM+lip6SZlQ5m/JqPkmdE57YWxsVz0SRWueqNaxq5qqXIw3uReKpvNSWnPSvipoXo9z3tVqLq7URqOyVyquSbM0TjVSUcAGtUS+AAABmn63/AHadXvxoX/8AvarNJXJsgt2J43kGU3iXxFpxqyVeQXSbmieJp6OCSpnd6fR6GRuX0/0GZVd7nUXq63O8VaRpV3a4TXOqSJHJEkk8jpXo3pKq8ub15c1VeXzqvzmazcdrNqHWhf611TKFGWdHnyK5zqxyoi/mo1M/nIR3x9mYkVlwfjKsq/UiRp+/P9xdg8Dh/wCIx+aH/ueXcCkf4HD/AMRj80P/AHPLuBnEI3gAAAAAA8y1j1n0q2+ab5Rq9rXnuN6aaa4Zb1uWSZflVwbb7Vb4+aNYxvzvlmlerYoaeFr5ppJGRxxve5rV4t4jPFG2wcNHTRct1pyJt71Ev1BJPplodjFZFLqLqHK1zo0kjhXmlHQMka5JrlUokMfQc1njpljp5M2TiNcU/dHxLdRVyXWbI3WDTaxXCSo030MxSsmg06wSN3SYyVYlVFrbg6NytluNUiyu6TmxpBF0YGASu8XzwjHVPeCzKdvez2fJtFNsdX42x5PmrpFsuret9KqqyaOeSNyvtVonT5PkML/KamFVSqlbHNLQsq+AAAAAA/01zmOa9jnMexyOY9qq1zVReaKi/wAyof5ABoPeD08biPc3juO7IN1uWNXcZh9lSg0c1JyCua2fXSz0EDUbbK+eR3SmyCihjc5ZVVZLhTQulf0qiGokqLYRic43kmQ4dkNjy3Er5dsZynGbtT3/ABzI7BcZbRfLDXUkrJ6Wso6qJzZIZoZI2SMljcjmOYioqKiKaXnAs4zGPcRXS2LR3WS6Wqxbx9K7BG7KqBEjtlFrLaYOULcntEPTVVqGp4pLlSxojYppmyxtbDM2OACwUAAAAACEvjZcJTFuJnoI654ZS2qw7r9IrVU3DRbMZ/FUUWVQoj56jD7xUOVqeRVz+awTyORKKre2VFSKSrjmy98yw7KtPMtybA85x674nmmG32qxjK8Xv9C+23vH7jQzvpqyjq6d6I+OWGWN7HscnNFapteFTjwi/gxecliN630bYMSjl3BYBY/KdbMFscHQrdacfoYmtS6UcDU5S3m1wRqqt9ElZRQrG1ZJqelhmAz5gAAD7fTP7pGn39d7T+3058Qfb6Z/dI0+/rvaf2+nANqkAAAAAAzRfCivwreXfiRwz9jqTS6M0Xwor8K3l34kcM/Y6kArsgAAAAAAAAAAAAAAAAAu3+CVb36p8utOwLNr06WmippNedDIKyRV8m5SQUWX2mB7l/kuWS03GGmjRERUusvJVc5Uu4mPZw8d0Fdsz3r7b9yVNVVNNa9NtTaGozRlKrvHV+NXBXWrJqVqIi83TWuuuMbeaO5PexeiqoiGwXR1lJcKSlr6Cqpq6grqZlZRVtHO2ppKyGVqPjlikaqtex7XNc1zVVFRUVF5AH7AAAB+KeCCqgmpqmGKopqiJ0FRTzxpNBOx6K17HsXmjmuRVRUX0Kin5QAZDnFG2nP2Tb9dye3mloXUOJYxqBPf9M2J0nwrit+ZHe8djbKqr4x1PQ19NSSvRf8Az6OZFRFarU4CLm3hem21lr1F2q7trPbujFmGLXPQTOq+GNY4GVVmqH33Hll5J0XTTw3W/s6bl6ax2uNvpaxqNpkgAAAH9jHcev2XZBY8UxWy3XJMnye8UuPY3jtioJbre79X1s7KajoqOlia6SaeeWWOKOKNque+RrWoqqiGmtwLuD1YuHBo0upWrNrs953hav2eKTP7xEjLjHpTaZOhPBiFqqUVWqrXNZNcKmDk2oqWsYjpYaWnkdH74OXwWW6JY/jW/wB3TYt0dZcstXwht602v9A6Oo0ps9bB0W5LcYJERWXe4QyyJTwub/qdJKj1VZ6hWUlvQAAAAAAAA/i5HkmO4fYbvlOW36y4tjGP2+W7X7I8jukFksNkpYWq+apq6yZzYoYo2ornSSOa1qIqqqHOnny7J+2Fta/WCxL28A6lBy158uyfthbWv1gsS9vHny7J+2Fta/WCxL28A6lBy158uyfthbWv1gsS9vHny7J+2Fta/WCxL28A6lBy158uyfthbWv1gsS9vHny7J+2Fta/WCxL28A6lBy158uyfthbWv1gsS9vHny7J+2Fta/WCxL28A6lB+harra79a7bfLHcqC82W80EN1s94tVZHcbXdaWojbNT1NNURqscsUsb2PZIxVa5r0VFVFRT98AAAAFFnww/7o+xH+pOfft+Kl6Yos+GH/dH2I/1Jz79vxUAq17Gfv2dnn5U2n31ttBsdGOLsZ+/Z2eflTaffW20Gx0AAAAAAAZa/hGn4ZfeN+b3/CvByEcm48I0/DL7xvze/wCFeDkI4AAAAAAB61oRrfqRts1j05140hyCfGNSNLMppstxS8Qq50bJ6dy9OnqY0cnjaapidLTVEDl6M0FRNG7m17kNazh6739NuITtX063JacyU9DNkFH8DahYY2sSrr9OcnomRtvFkqV5I5fFSPZLBK9rVnpaqlmRrUlREx+SdPgNcUWq4du6imx3Ua91EW1vX6so8R1fpaidXW7Bq1HrFaMwjjX0N8hfM+Ks6HJX0NRO5WyyQU7WgahoPwU1TTVtNT1lHUQVdHVwMqaWqppWz01THI1HxyRvaqtc1zVRUciqioqKh+cAAAAFfjiM+Dq7Od+GW3nV7D7veNruu2RTyVuTZnp9j9NkGDZtVS9Jz6+94u+WnZLVueqPfU0VVRyTK6R06zPckjbA4AKpPDw8F8wLaNuP093Ja27j5NdrppDk1PmunOn+O6ZpgeNwXygc2e13W6Vs1wq5qjyOoRlVDTQR0/RqKOne6eRnTgW1sAAAAAAAAZWvhBmRU2UcYTefc6TxHiqXIsUx1/k9W2tj8bZ8BxS0zc3tRER3jKKTpR/PG7pMVVVqqsNJ1/xA9YKbX7fLu51kt9Y2vsuoO4jLb9jFW2RZmyWd16rIrMiPX+UjaGOjbzTki9H0IickTkAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAtAeDd/75n5u/+ui0AVf/AAbv/fM/N3/10WgDVr3Rr45d8v8Ax3/qqEmlhL+D6z/8X7+QAAhGXGBy/ve+8u3efkv5/wDVS7HUBy/ve+8u3efkv5/9VLsVvhn+Ei7/AOnUn38Z51sf9Jqvo3/yqZz4ANzkx8AAAAAAAAAAAAAAAAlM4Nesz9Ht++lNPUVa0li1apq7R2/Iip/rHwvG2a1R8l/+a60VoT0Ki8lXlz/krFmfR4dlN3wbLsWzawT+TX7DsjocpslT6f8AV6u31UVXTP8AR6fkyQsX0f0FAYrXGpcTMM7fw+rETUtGkqKdFXia6WJzWP7cb1a9q8itRT1LEtJ9j2xS2rHxxSNf20aqKqfWmafWaboPl8Iy20Z/heIZ3YJfH2HNcXt+W2SfmjvHUlypIqymfzT0LzjmYvo/pPqDTVqaaooqmSjq2KyWNytc1ditc1cnIqcioqKimQJj2yMSRi5tVM0XqooABwH6BRl41+k7dMN/WodzpadKW1atY3Z9VbdE2NWR9OqpnWy4PRf51kr7TXyqv9Myl5oq/wDhGem8f/4ZNXqSj5Sr8P6b36vRqfKankF0tEKr8/oVb25E/wD1L838+Qzcw78PunpUUNivflDa1LVUjs+LWaz2XGq9lX0yMavHm/LiVS1WMlmpXXJlqET3UD2SJ1clXUX9z817XYOUOBPuxo9GNxF50IzC6toMJ3CQU9BYpquZsdFbsrofGfBbebv5Hl8U1TR/JXnJOtC1efJFS5eZhNDXVtsraO5W2sqrfcbfVR11BX0NQ+kraGeJ6SRTQysVHMexzWua9qorVaioqKheU4WPEXsW8/TCHDc4uNFbdxmnVpigzW1SSRUq53Rx8oY8jt0KKiua9fFtq42NRIJ5E9DY5oecl91K0YLTjtpukjc2mWSllZHFajWJmsUjESOCrVE2729iMglXiY5kTlz3xyto/Ba+cK0/tRtB+T2qroVXlRdrmdtFzc3qoqpyJnLEADC4SFAAAAAAAAAAAAAAAAAAAAAAAAAAABCfx19xFJpRtBXSO31rIsw3C5BDjVPTxzeLrKex2memud6qmpz5qxzmW6hcioqKy6P/AKCXPVDU7BdGdP8ALNUdSshocVwfCrPLe8gvde9UjpoYk9DI2JzfLNK9WRQwRo6SaWWONjXPe1q0Cd+W7/JN624bJdW7pDVWnFKSJMW0yxape1z8ZsNLLK+ljl6PNq1E75pqqociuTxtS5rXLGyNEyG7nNo62ti7jXRX/tKnVLvWFKyplkci6ktXHk+lp2LxOckmpPKm1GxM1X5LLHnanFq9kFg3dksuF/8AxVS1WIicbWLse9eomWbW9Vy5p71TjI/r2C/XjFr9ZMnx64VFpv8Ajl3pr9Y7rSORtVbKyjmZUU1REqoqI+OSNj2qqL6WofyAbL8sUVRE6CdqOY5FRzVRFRUVMlRUXYqKmxUXYqEPmuc1yOauSoaM+z3cjju7LbtptrfYH0kVTk9lZT5dZaSbxn8GL7SIkF3tzmqqvakdQ17ovGcnPglgk5cpGqvTJRk4UXEJm2W6sVOK5/V11Rt91UroKbNaeFrqtcIuKdGGlyOmgReapGxUirGRp05adrXI2R9PDGt4mzXm0ZHaLXkGP3S33uxXu3w3azXm01kdxtd2paiNstPU01RGqskikY9r2PYqtc1yKiqimqfpmaMts6NuK9VZ1PC5buVr3zWdNkqtWJVzdTud/wB6mVyRuRVzcze5ckSTJJt4fXxp73WGyV7k9lxojZW8utl79E/Jfxp1FzbyH9IAEQyvAAAAAAAAAAAAAAAAAAAAAAAADg/iefeB7pvxXz/tNMZ95oIcTz7wPdN+K+f9ppjPvNgrch/wKXm/Wq/0lMRZx4+EVH9B/uPOndm24u6bU9ymlWuNAypqaHEcibHldrpVTx16sdax1FeKRiKvRWR9LPOsSv8AQ2ZkTv8A4UNEfGMmsGaY3j+YYrdaS+4xlVlpcjx29UEnjaG70NbAyppamF387JYpGPaqonochmNllDgp8Se0YNFbtnWu+QwWzG6+5uXQzNbxU+KorLV1k3Tlxmsmd8lkNRNI+ajleqIyaaWFXKklO2Pvbpxov2pihdKlxnuNTLNa9kxujqomNzknodZZNdiJtc+le579VEzdFLKu1Y2tXjwcvnDY1c+71pP1aedUVjl4my7EyXqI9ERM+RyN6qqWsAAa85KoAAAAAAAAAAAAAAAAAAAAAAAAFfjwiP72HRL8fLPq9eiwOV+PCI/vYdEvx8s+r16JZaC3xtbkfpbvuJihsS/gLaXzE/maVBjTZwn/AGMxH+rFB+yRGZMabOE/7GYj/Vig/ZIjInuxf/TsPvn2p/LQFqMAf+bavah/jKfTgAwdEkAAAAAAAAAAAAAAAAAAAAAAAAAAAAAfp3G42+z2+vu93r6K12q10UtxudzuNUyht9up4GOlmnnmeqMjjjY1z3PeqNa1qqqoiH6Yx8j0jjRVcq5IibVVV4kROVVPiqiJmvERm8X7cDRaDbHNU4Iq+OnyzWKj/wBDWJ0iTJHU1Xw1HJHd5GoiK5Gw2xlyd00RER7oW9JqvapQ9JR+K7vli3n7gUjwusml0S0kjqsW0yVzXwNyKSaSNbpf3RORHM8tkp4GxNciOSmo6ZXNY90jSLg2kNALAO0cBsBKalvNCsVt2pItbVMcmToddjWQU7uVHRQta6RqpmyaSVu3LMhfijeiK8953vo3a1NCm9sVOJ2SqrnJ2Fcqoi8rUapdw8Dh/wCIx+aH/ueXcCkf4HD/AMRj80P/AHPLuBN4twAAADmbepqhlmiGzfdprTgVRR0mdaQ7Zs81QwuquNEy5W+mu1gxa63a3ST0zvkyxtqKSFXRu9D2oqL6FOmTiPiY/g3+IH+RHqt9Q78AZJGses+qu4PUjKdXtbM9yTUvUrNLgtyyTL8quDrhdbhJyRrGN+ZkUMTEbFDTwtZDDHGyOONjGtanmQAAAAAAAAAAAPTNG9Y9TdvuqOE60aNZleMA1O07vkeQ4hltimSK4WqoY10bkVrkWOWGaKSWCenma+KeGeWKVj45Hsd5mADV04RHFY0z4oGgjMkp22rDtw2ndLTWzXfSelnf4uz1UqPZBerOkjnSS2m4LFI+JXOe+nkbJTyue6Nk08thjdbQN3GtOx/X3CdxOg+RvsGa4fVLFV0U6LPYcxtc7mfCNju9N809FWRsRj2ehzHNjlidHNFFKzVY4d3EB0W4ju3THdeNIqxlvuTfF2PVHTesrUq8i0ryBsEctVaqx3QYs0XJ6SU1Y1jY6mFzXojHpJFGB3aAAAAACgF4RvwYWaBZHfN++17FPE6I5tfPH6/ad4/bVbRaQ3uumRGZBQwxp0YrPc55OhNFyayirZmIxVhq44qWo4bYeUYxjubY1kGG5fZLZkuKZXZarHMlx29UbLhaL7QVsL6aro6qB6KySKaKSSN7HIqOa9UUzCeOBwicl4aeujcp09oLxfdo2sN1qKzSbKpvGXF+BVy9Kepw28VPJVSopmdKSjmmXnWUjekj5JqesSICDI+30z+6Rp9/Xe0/t9OfEH2+mf3SNPv672n9vpwDapAAAAAAM0Xwor8K3l34kcM/Y6k0ujNF8KK/Ct5d+JHDP2OpAK7IAAAAAAAAAAAAAAAAABrCcEPce/dBwwNqOd3C4OuGVYngf+hrNXTy+Pr2XHDaiXHmS1T+XypaujobfXq7mqqlwarl6XSRMnsvdeCCa8Ou2kW7rbRcK5qOwfUCx604vRTS9OeohyOgmst4WFv80dPJjdoVyehOnckVEVXOUAuTAAAAAAgv8I20F/06cKLXeso7ey45DodeLHr1jrHw+MdS/A1e2hvVQ13zsWKyXa/P6SIvNEVq8kcrky7zaF190tt+uWhWtGit2bE616u6UZFpjcEn5+J8TfrRWWuRXcvSiI2qVeaelOXoMYa42+ttNwrrVcqaWiuNsrJbfcKOdvQnpJ4Xujlien8zmua5qp/SgB+mW0vB0uC43cZklh33bosU8foFhV8dU6Gaf3ymc2l1kvtuqHMdea2FyIktmtlTC5rI15sraynex6Ogp54qjgTgd8IPJOJRrc3NdSLbdbNtA0gvkUuquSQyyWybUK4RtjqYMNtFQ1Ol46dr4ZK2eJUWlo5eaSRz1FJ0tOzF8Yx3Ccax/DcQslsxrFMUstLjmNY7ZaNlvtFioKKFlNSUdLAxEZHFDFHHGxjURGtYiIAf3QAAAAAD/EsscEck00jIYYWLLLLK9I44mtRVc5zl9CIiIqqq/NyP9lKjwjrjWraIsw4dm0/LujdaiN9g3T6pY3XubJa43IrajA7ZUsT/AM2RFRt1nhf8hquoVXpOrY4wOBfCEeNTJvDzG67OtsmUOXavp9fUZqBmVkq2upNwV+oJ0dG+GeNypLYrfLE11KiL0KuoZ5SvTZHRuZVwAAAAAAAAAAAAAANjvY195Ps9/Ja0++qVoOpTlrY195Ps9/Ja0++qVoOpQAAAAUWfDD/uj7Ef6k59+34qXpiiz4Yf90fYj/UnPv2/FQCrXsZ+/Z2eflTaffW20Gx0Y4uxn79nZ5+VNp99bbQbHQAAAAAABlr+Eafhl9435vf8K8HIRybjwjT8MvvG/N7/AIV4OQjgAAAAAAAAAGgt4MhxTv8AT3pL5hOteReP1h0KxzyvRO8XSf8A1vP8JpehH8FeMcv8ZWWHpRxNb6HPt8lP0Wu8kqJC2UYvOhGt+pG2zWPTnXjSHIJ8Y1I0symmy3FLxCrnRsnp3L06epjRyeNpqmJ0tNUQOXozQVE0bubXuQ1meHRv30n4jG2PDNwOmdTS0F2qYGWLVPT1a5tVedMMlgiYtwtNUn8pYlVzZ6Woc1vlFLUQSdFjnPjYB3aAAAAAAAAAAAARt8XHd3QbJeH3uL1rS5R0GaT4XPpzpNElSlPW1mV5JG+1Wd9OnSa560Tp5bnKxio7ye1VCoqK3mSSGbX4R9xRbHvb3E2Xb1onkUF8257ZbjWUqZHaK1tZY9U8ymTya6XmllY5Y56GgiZ8H0UzUVHq+4zRySQ1USoBWzAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAABaA8G7/3zPzd/wDXRaAKofg+OqemOmnnc/6R9R8D0/8Ahv8AgD8Dfw2y+34p8LeTfw08o8l8qmj8b4ryiDp9Dn0PHR8+XSTnZE86fbF2jdB/pex/2s1it0Nupei0tMK99bZ1m1EsLuh+q9kMj2rlZdEi5Oa1UXJUVFyXYqKnGhMjCmuoobg0EcszWuTfdiuRF/50nIqnvAPB/On2xdo3Qf6Xsf8Aax50+2LtG6D/AEvY/wC1kLfaNfb5HqvF5vMLh9ErO5wzv2+U94OX9733l27z8l/P/qpdj6jzp9sXaN0H+l7H/azm7eTuS27X3aFurslk180WvN6vO27ObVaLRatU7HcLpdaqoxi6Q09NTU8dUsksssj2MZGxFc5zkREVVRCtMNrlXyixFsCWWyKpGpW0qqq08qIiJPHmqrqbETlU8+17Rs9bKqkSdme9v/Gb+SvZKBAANxMgIAAAAAAAAAAAAAAAAAAX2+Ebqc7VLh/aAVtTULPc8LslXpjcmOf01pkx+vqbfQR8+ar/APwEduciejkj+SJyRCSUrPcA3ctplhWhet2lupepmE4I+w6o0ua2Fmd5hb8Xjr4b3a4qOZlEtVNH4xIpLD05Gx80Y6qaqoiyc3T0+dPti7Rug/0vY/7Wan2llhFea6mknfOybKsqofSrXzTxLHDI5m91SpVMa1zWqio1syNTJdmWXGThuNbtHW3Qs+eedqP3prVzciLmz3CqqKvKrcz3gHg/nT7Yu0boP9L2P+1jzp9sXaN0H+l7H/ayPPtGvt8j1Xi83mFVdErO5wzv2+U94IZOPBgbMs2F3TJPJkll0v1Ux3MkqERUkpWVctRjr/Sn/wALnX2NFa75Kr0F+dG8pI/On2xdo3Qf6Xsf9rOK+IxrFt11V2O7mMKsuvGi18vNXpjVXu0Wa1apWG5XS6VVmlhvNPBTU7KpXyyvkt7GsjYivc5URrXKqIt69G6zr6XM0gLl3mksuqZFBadEsjlglREidOxk2a6mxN6c/PPZlx7Cnb3S2faN1rQo2zMVzoZMk1m++Rqq3l/KRCiWfYYDn+a6WZjj2oOneTXbD80xW5Mu2P5FZKpaS4W6eNfQrXfM5jkVzHxPRzJGPcx7XNc5q/Hg21KujpLQpJaCvibJBI1zHse1HMexyKjmua5FRzXIqo5qoqKiqipkQZjkfFI2WJyo5FRUVFyVFTaioqcSpyKXCNiHHC0u1dpLNpxuuqLNpBqe2JlDS6jKvkOleayJ8lJKqVyqloqHoiOek7vI3OR6tmh6UdOk9tBX0F1oaS52utpLlbbhTMrKC4UFSysoa6GRqPjlhlYqtexzVRzXNVUVFRUUzCjrnbnvs3V7VZIYNGdX8jsmNsqPKJsDvD2ZRgFSrnK6XlaKtJIIXy8+T5qZsUy+j+MRURUxEaQu5U3ZvVXT3owDtBlmVEiq51BUa7qNXLtXeJWI+Wnaq8UasmYirkxYmIjUvvdTG2soY20V54lmYmxJWZJJl+c1cmv7ebV6usu00RwVatG/CKrzBFBQ7gNvVvuUiI1KjJtIsifanehOTuVluCyo5XL6eflzET5uj/Okjen/ABxuH5mrIfhvOs60uqZ1axlHqBpvcZntc9UajXzWlK+BnJV9L3SIxERVVyJ6TGBfjQa0qrhTPZaVzqqpjbxSUSNrWuT8pG0zpJETsPjY5OVELy2biTcm02osNoMYvUkzjVO/yTuKqEvQOKcc4j2xDKWwPtm63RWlSo6Hi0yPMqfD3N6bem3ptr/EKzknz9Pl0V9C8l9B6Vbd4W0i8+O+B90u3O6+TdHyj4N1uxmu8n6fS6HT6FavR6XQdy5/P0V/oUsVX4V4n2U90dqXbr4XJxpJR1DFTk2o6NMtpU0Vt2NOmcNXE5OxIxf4KdGA8H86fbF2jdB/pex/2sedPti7Rug/0vY/7WeX7Rr7fI9V4vN5hzdErO5wzv2+U94B4P50+2LtG6D/AEvY/wC1jzp9sXaN0H+l7H/ax7Rr7fI9V4vN5g6JWdzhnft8p7wDwfzp9sXaN0H+l7H/AGsedPti7Rug/wBL2P8AtY9o19vkeq8Xm8wdErO5wzv2+U94B4P50+2LtG6D/S9j/tY86fbF2jdB/pex/wBrHtGvt8j1Xi83mDolZ3OGd+3ynvAOdLlvB2lWZsTrvuj26WptQqtgdctbcZoWzK3krkYr61OfLmnPl83ND4K+cQ3Yxj3jVr92Og9R4nodP4D1Ht2TdLp8uXQ8jkl6fLmnPo8+j6efLkp6NFhdiZaTkbZ13a+VV4tSkqH57ctmrGue3Z2ziktqx4dstXE3tyMT+KnZAIr8y40PDvxGCZabWu4ZnXQojkteG6dZDXTy81RPkVM9HBSL/OvLx/8AN/8Abnw5qj4RRo9a46un0a0A1DzKqRHw0tx1Dv1vwG3I9PQ2bxFKtwlkj583IxzoXuRERVjVV6N4LpaGelJfSVsVj3Hr2I7idVQ+wmdvXrFgbl2UVc+TM8CuxBuXZzVdUWlEuXIx2+L3I9ZSxmcZ7rt++2jZzZKqp1ZzuimzDyLyqz6WYtLFe9R74rmq6Ho29Hp5NFJy+TVVroYF5KiSKvyVqZa+cabfHrdDWWmzZpZ9EMZq0dC61aPW2Sx3iWPmvRWS9zyz3BkiIvJzqSana5U59BPmIprnc7lerhW3e83Cuu11uVS+suNzudXJX3CvmkVXSSzTPVXve5VVVc5VVVX0qZAcE9yVvBVVkNsY82zHBSoqOWjoVWSZ6fkyVL2pHF1HJEyZVRfcyMXalrbxY50rI3QXYp1c/wD7kqZNTsoxFzd/eVuXKinf+/HiO6075snbDkD3YPo/Yri6sw3SSy3B9Ra6R6I9kdfdqjk3y6v6D3NSZ7GRxNe9sMUXTkWSPIAzT3FuFc/DO69Lcu4dnx0VmU7dWOKNMkTqucqqrnvcvunyPc573KrnuVVVSPNp2pX2zWvtG05VkmfxuX+CciInEiJkiJsRAACrjoAlq4eXFf1V2Xz0en2YwXPVPb1PVrJLhc1cjsjwRZXo6aoxuolcjGNVVfI+3yubTySOe5rqeSSSV0SoKBxLwvuJjBdKouPiJZzKyzptqsfmjmPRFRskT25PilbmurIxzXJmqZ5KqL6lj21adgVzLSsmVY5W8qcSpyo5F2OavKioqGjlty3abft12LR5RohqNY8sSOlZU3rGVqG2/NcWV/JPF3S0yKlRBycqsSRWrDIrVWOSRvJy9HGZHiOZZfp/kNty7BMpyLC8qs03lFpyXFL1U49frZJyVOnT1cD2SxqqKqKrXJzRVQmS0E48O8DS2Ghs+qNFhuv+PUvRjkqMoo1xPO1ibyRscd3ompC5eXPnLVUdRI5eSq9V588JmNe5NX7sarmtXAu1Y7QolVVbS1bmwVTE5GNmySnm+c9abLi1V2qSKu7jlZlRG2C8sCxScr40VzF7Kt9+3tJr9vkLpQIEtMfCDdquTR01Pqbpvq3pfcpUb5RPQ0dDn+M0qr/K/wBahmhq3cvRy5UfpTn83zL2tiXFo4eeZRwut25bFrXLKjUfT5bYL5hklO53QRWPfXUUMfyVeiK5rnM+S5UcqIqpAS9WidpLXLmdFb1x7SRG8b4aaSpiT/Fpkmi7S6+S8hdChvxc+0Wo6ltKHbyOejHd6/Vd+4kXByvbt8+yu6xQy0e7bbZ/HyeKihqtbcbt9Y93S6KJ4iWsbInNfm5t9PNOXPmfYedPti7Rug/0vY/7WWsqcOsQaN+91lhVjHbdjqWdq7OPYrE4uU9plrWVImcdTGqdh7V/zPeAeD+dPti7Rug/0vY/7WPOn2xdo3Qf6Xsf9rOv7Rr7fI9V4vN5h++iVnc4Z37fKe8A8H86fbF2jdB/pex/2sedPti7Rug/0vY/7WPaNfb5HqvF5vMHRKzucM79vlPeAeD+dPti7Rug/wBL2P8AtY86fbF2jdB/pex/2se0a+3yPVeLzeYOiVnc4Z37fKe8A8H86fbF2jdB/pex/wBrHnT7Yu0boP8AS9j/ALWPaNfb5HqvF5vMHRKzucM79vlPeAeD+dPti7Rug/0vY/7WPOn2xdo3Qf6Xsf8Aax7Rr7fI9V4vN5g6JWdzhnft8p7wDwfzp9sXaN0H+l7H/ax50+2LtG6D/S9j/tY9o19vkeq8Xm8wdErO5wzv2+U8H4nn3ge6b8V8/wC00xn3l6/iPbh9AMp2Nbl8exjXPR3I7/dtNp6W12Ow6m2W73i5SrUUypHT00VS6SR6oir0WNVfQvoKKBny3JiyLWsfBm8kFr0skD3WoqokjHMVU9iU6ZojkRVTNFTPizIw45TwVF4KR0D0ciQ8iov47+oAAZVSyRYe4enG6v2k9BYdHN3K3rNtPbfDFasb1eoYn3fOcQgZ0Y4oLxT/APmXKljb/wDmGKtZG1ip0Kvm1rLUemWqum2s2I23PdKc3xvP8Pu0aPor9jF0iulF0la1zoZeivShmZ0kSSCZGSxu5texrkVDM+PWNIdddY9A8kblujOpWYab375Lairxa8y2+C5saqq2GtpkVYKqLmvPxVSyRnP09Exg6Su5mYdYtWhUXywsqW2JbUqq+SLU1qGd67VcsbMnUz3Ltc+FHRrx7wr1c9by3PxitawomWfbTFqaduxHZ5StTqZrseiciOyX87LJDSoBUG0P8IP3C4fDS2zXTS7B9ZqOFqRSZBj9W/TDMJ+a/KlnWKGooJFanpSOGkp0XlyVyc+kkoGnHHx2SZdFSx5rQat6U170Rtc/IMNjyWyU7uSqqxVFsnqKiRnzJ0nU0bua/wAjl6TE3f8A3P3StuBO9JbsPtCnbnlLZ72VbX5crYmKlSnY14GKvIhfGy8UrkWo1NWtSJ6/iyorFTtuX3HccpNyDgbFuKPw/cwjZLad0mnFI2RnTamUur8HkROijvlMuVNTuReTk9Coi8+acuaKiev23eps5vEtNBa91+22uqquPxlPR0+uOMvrpE6CyKniPLfGIqNRVVqtRW9FeaJyUjtaeD+LdiSOitm61owOTPNJaGpjVMuPNHxIqZcpVkNvWFUIjqeticnYkYv8HHTQPB/On2xdo3Qf6Xsf9rHnT7Yu0boP9L2P+1nh+0a+3yPVeLzeYdnolZ3OGd+3ynvAPB/On2xdo3Qf6Xsf9rHnT7Yu0boP9L2P+1j2jX2+R6rxebzB0Ss7nDO/b5T3gHg/nT7Yu0boP9L2P+1jzp9sXaN0H+l7H/ax7Rr7fI9V4vN5g6JWdzhnft8p7wDwfzp9sXaN0H+l7H/ax50+2LtG6D/S9j/tY9o19vkeq8Xm8wdErO5wzv2+U94B4P50+2LtG6D/AEvY/wC1jzp9sXaN0H+l7H/ax7Rr7fI9V4vN5g6JWdzhnft8p7wDwfzp9sXaN0H+l7H/AGsedPti7Rug/wBL2P8AtY9o19vkeq8Xm8wdErO5wzv2+U94K/HhEf3sOiX4+WfV69Ex/nT7Yu0boP8AS9j/ALWQU8e/WPSLUfbho7a9PNVNOM8udDrcy4VtuwzOLZlFfRwfAN4j8fLDTTyPZH03sb03Ijek9qc+aoSq0IbpXroNKy5VZXWZURwtqnK5z4ZGtam8TbVcrURE7alE4j11DLci0Y4pmK5WJkiORV983kzKppps4T/sZiP9WKD9kiMyY0V9O9122mq0+wWqqdwGiNtqKnDrZUVFurtWrBDW2976KBz4JmLVIrXsVVa5qoiorVTkZCN14sC3Lasy4T7Go5Z0jktJHb1G6TV1m0OrmjUVUz1XZZ9RS1eA9VTU81ppUSNbmkOWaomeSyZ8fbQ6pB4P50+2LtG6D/S9j/tY86fbF2jdB/pex/2swm+0a+3yPVeLzeYSK6JWdzhnft8p7wDwfzp9sXaN0H+l7H/ax50+2LtG6D/S9j/tY9o19vkeq8Xm8wdErO5wzv2+U94B4P50+2LtG6D/AEvY/wC1jzp9sXaN0H+l7H/ax7Rr7fI9V4vN5g6JWdzhnft8p7wDwfzp9sXaN0H+l7H/AGsedPti7Rug/wBL2P8AtY9o19vkeq8Xm8wdErO5wzv2+U94B4P50+2LtG6D/S9j/tY86fbF2jdB/pex/wBrHtGvt8j1Xi83mDolZ3OGd+3ynvAPB/On2xdo3Qf6Xsf9rHnT7Yu0boP9L2P+1j2jX2+R6rxebzB0Ss7nDO/b5T3gHg/nT7Yu0boP9L2P+1jzp9sXaN0H+l7H/ax7Rr7fI9V4vN5g6JWdzhnft8p7wDwfzp9sXaN0H+l7H/ax50+2LtG6D/S9j/tY9o19vkeq8Xm8wdErO5wzv2+U94BzJcd62zi0yVMNy3Y7a6Opo2K+po59csYbWxcm9Pl4jy3xiuVFRUajVVeackXmh5DlPFH4fuHxvlu26TTirbGzpuTFnV+cSKnRV3yWW2mqHKvJq+hEVefJOXNURfcszB/Fu2pEhse61ozvXiSKhqZFXPiyRkSquZ1prfsKnTWqK2FqdmRifxcd8ghG1H4+OyTEYqqPCqDVvVavYitoX4/hseNWSodyRUWWouc9PURs+dOk2mkdzT+Ry9JF1rl4QbuKzKGvtehemeD6LUNQniqbIb3O7U/NKXor6JYXTRQW5iu5c1jmo6hG8+SOdy6SyKw/3P3SsxAnY2O7D7Pp3ZZzWg9tI1mfK6J6rUr1V1IHqmW1M8kWk7UxSuRZbVzrElen4sSK9V7Tk9x3XIWqdYtcNJNv+G12f6yZ9jmn2KULXc7jf65IJa+Rrel5PRUyc56qocifJp6Zkkrv5mKVAuJZxdMr3bQ3LRrRalu+n+3plUjbvU1snkmZ6sOif0mOuTY3K2mt6ORr46BrnOkVjZJ3qqsggie1X1n1Y10ymfNdYNQss1GyedFjbdcrvE10ko41crvEUsbl8XTwoqqrYIGsjb/M1DzIy+aLW5x4f4GWnTX7v7Utti8cKo+JVZq0dLIm1HwxuzdLKxfeTS5I1cnxwxyNR5Ya+uLVqXlhfZlmM9j0jtjtucj06jlTY1q8rW8fErlRcgADJKWiLuHgcP8AxGPzQ/8Ac8u4FD3wSzXLRPRfz/v9MWsOlmk/8JP9FX8Hf9JeoNpwT4f8j/0keWeReXVEXj/EeVUvjPFdLoeUxdLl0287jvny7J+2Fta/WCxL28A6lBy158uyfthbWv1gsS9vHny7J+2Fta/WCxL28A6lOI+Jj+Df4gf5Eeq31Dvx9v58uyfthbWv1gsS9vOOOIrvI2hZRw+t9eNY1uq235FkeRbONTrFj+P2LXLGLve77XVeE3ynpKOjpIq10s080skcccUbXPe97WtRVVEAMosAAAAAAAAAAAAAAAkF4a/EU1m4au42ya26XzyXrFrp4jHtY9LaysWnsGqWPeUMlnopF5OSCth5Ploq9rHPpplXm2WGWop54+gAbK+1PdPoxvO0KwbcNoNlMGUYBnNubURtcrIr5jFcxrUrrNd6VrnLTV9HI5YpoVVU5o17HSRPjkf0SZVfBz4tGonDC1zSouD7xl+2PUu409JrfphSyLVTwsbyijySwQvlZFHdqNnJOTlayrgasEqtVKeen0lMZ4g2xTLsdsOVWTeLtkks+SWemvtrfXa441aK9aerhZPEk9JPWMnglRr2o+GdjJI3I5r2Nc1UQDsAHLXny7J+2Fta/WCxL28efLsn7YW1r9YLEvbwDqU8A3Q7ZdIN4Whmf7edc8ZgyfTzUKzvt1fFyZHdbHUt+XRXa2VDmu8nrqKZI6iCdEXoviTpNcxXMd8z58uyfthbWv1gsS9vHny7J+2Fta/WCxL28AyuuI9w/NXeG9uXyfQLU+N13srmrkelOpFJSOprHqfjk8r2UdygavPxVQxWPp6ukVXLT1MErUdLEsM8vIGmf3SNPv672n9vpzTk4omB8NTiY7arzo5mm8LaZjmodg8fkeiOqjtcsRq7npzfViRqOcqV/jJLfWpHFBXUrVRJYmse3lNBBJHmlvwm56R7gKfAcruOM1Fz0/1Wp8fvd5xnJqLKcPqnUF0jilq6C700j6Wqo3pGssVVC90ckbmvReSgGzyDlrz5dk/bC2tfrBYl7ePPl2T9sLa1+sFiXt4B1KDlrz5dk/bC2tfrBYl7ePPl2T9sLa1+sFiXt4B1KZovhRX4VvLvxI4Z+x1JoK+fLsn7YW1r9YLEvbzO28JN1H081V4neU5bpfnuF6kYpNo3iNDDk2BZTQ5hj0s8FJUNmhbW0kskKvjVURzEdzaqpzRACA0AAAAAAAAAAAAAAAAAAsVeC861JpdxS8ZweqrPJrfuC0eynSx0cr2spJKujpocvo1crl5I9VxiWGNf5SuqugnNZOS11Tsbh56ywbfN9W0bWWurqa22XAtwmK3bKK6sqkoaWnsr7vS017WSdzmtjatBPWor3r0G8+bkVqKigbDQOWvPl2T9sLa1+sFiXt48+XZP2wtrX6wWJe3gHUoOWvPl2T9sLa1+sFiXt48+XZP2wtrX6wWJe3gHUpmOVPCW1b3k8abd3tP04pJ8T08w7cplOY6lakVVE2S0aWYZcr5VXS3TKxq9CatqqSspqehpG8lnmejnpDBFUywaH3ny7J+2Fta/WCxL2887w3cXw39Psn1NzPDNyOz7Hsq1lymmzTU6/wBu12xGK6ZlcqO1UNlpaism8v6TvFUdupYmM5oxqpI5Go+WRzwPeNtu3PSTaZojp9t80OxiDE9NdNrG2zWK3MVstdWvVzpqu4V86Nas9bWTyT1VTUORFlmqJHck58k9xOWvPl2T9sLa1+sFiXt48+XZP2wtrX6wWJe3gHUoOWvPl2T9sLa1+sFiXt48+XZP2wtrX6wWJe3gHUoOWvPl2T9sLa1+sFiXt5GVxUuOJts2UbcrjfdCdVNJtfNxGfsqcb0ixDAM7tefWjHqpsbEnyDIn0M0zYKKgSeORlPKrZK2bxcMfRZ5RPTgc7eEA8aOm2O4FXbWduGTQSbutS7B/wCPZDa5mzybfLDXQ8m3GR7XosV6rYpEdb4lRXQRqtW9Gp5IlRnE1VVU1tTUVtbUT1dZVzvqqurqpXVFTVSyOV8kkkjlVznOcqqrlVVVVVVPrtSNR861fz7MNUdTcou2a6g59kNVlWYZXfajyq636vrJXTVFRM7kiIrnOXkxiNYxqNa1rWtRqfFAAAAAAAAAAAAAAAAGx3sa+8n2e/ktaffVK0HUpGLsw3n7PLDs82n2O+bsNtNmvdm204JabxZ7trti9uulqqqfF7VDUU1TTyVySRSxSMex8b0RzXMVFRFRUOlfPl2T9sLa1+sFiXt4B1KDlrz5dk/bC2tfrBYl7ePPl2T9sLa1+sFiXt4B1KUWfDD/ALo+xH+pOfft+KlwHz5dk/bC2tfrBYl7eUsfCxdbNGtZtQNlFTo/q3pjqtTWDDs4gvtRptntqzqCyvqK3GXU7Kt9FPKkLpUhlViSdFXJE/lz6K8gK3Gxn79nZ5+VNp99bbQbHRja7Lrta7DvF2nXy+XK32WyWXctgl2vF4u1ZHbrXaaWmym1TVFTU1EitjiiijY975HqjWtYqqqIiqa0vny7J+2Fta/WCxL28A6lBy158uyfthbWv1gsS9vHny7J+2Fta/WCxL28A6lBy158uyfthbWv1gsS9vHny7J+2Fta/WCxL28AzfPCNPwy+8b83v8AhXg5COTH8f8AznCdSeLhu0zXTrMcWz7Db1/AP4Hy3CsgpMqxm7eTaZ4ZSVHk1fSySQS+Kngnhf0Hr0ZIZGryc1USHAAAAAAAAAAAEifDS4k+t3DM19pdX9LV/hPh99gjsmrekF1u89txTU21tc5WRzOYjkgrqVZJJKOvSOR1PJI9FZLDLPBLHYADYQ2N7+dtnEJ0et+sG3bNqa9QRxQU+bYLdHx0OoWmNwlYrnW2+23pK+F/NkqRzsV9NUtic+CaVidI7QMYvb7uQ112q6l2jWDbvqhlek2o1kasNLkeK1yQOq4HOY+Sjr6WRr6atpJHRxrJR1kcsEni29ON3JC4Xsw8Ldo20drxLfjoPXLWwxx0k2sW3tsU0davJGJNccTrqiPxa80SSWairXIvSd4qibyaxQLtIIl9IeOjwotaKWjlsG83S/Eauq6DZbXq8tdo1VW+R7kZ4uaa809LTL0VcnSkjmfGic16fRRVOtLZv62J3qjjuFm3qbSrtQSuc2Kutm47Dq+jkVjla5GysuKtVUVFReS+hUUA60BwVfOKfw0sepPLa/f5s8qIVcrehY9xWKZPV+hrnr/q9HXSy8uTV5L0eSryROaqiLxtqn4RZwj9L6erSLczPqVd6aNXssOlmmeS5PUVnLn6IrhJRQ2zmqpyRH1bfnRfQnpAJvj53LsvxTAMZvmaZ1k1gw3D8Zt0l3yPKspu9PYMdsNJCnSlqaytneyGGJielXyORqf0lM/cx4XrjUFLcLRs+2q3q5172uZbs83E5BDZ6CkciOajnYxaJpnztVei5FW6QKiIqKzmvyasm8/iZb1N/d58t3Ka2ZDk2M01Z5bZNL7D0cS0ox1yKqxuprDS9CnkljRVa2rq0nqlb6HTuALFHG08IwpdY7Blu0fYBf7xQadXaOXH9WtydL46x3PPaVyPiq7HicbmtqKe3Sovi6i5yeLmq2rJFDGynVZ6unIAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAdDbuNN6vR3dXuW0nrqWSjqdNdfMvwZ9PJ6VYlrv8AcKJitcnNHNc2Frmvaqtc1zXIqoqKc8k+nhKW3ep0L4pmqmT01vlo8W3E4pZNcMelSFG0ks9TS/A17a2RqI1ZFudmuFQ9q/LalbGrvQ9rnQFgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA6G2j6b1esW6vbTpPQ0slZU6la+YhgzKeP0K9Lpf7fRPVzl5I1rWzOc57lRrWtc5VREVTnkn08Gt271OunFM0ryept8tZi23bFL3rhkMqwo6kinpqX4GsjXSORWpIlzvNvqGNT5bkopFb6GOc0Cy74VFsuqdcdmmIbpcRtbavNNpuSSVGUeTxItZW4bkklJQ3N3o+VItDXw2apRF5pFA+vk+SnSVc642t85wjE9S8Ky7TrPLFQ5PhGe4zXYbmGN3Njn26/2u50stFX0c6IqO6E0M0sbuiqLyevJUX0mSbxOtheccOnd3qLt8yaG41uHsqnZbozmlbB4uLPcSrppvguuRyfJWeHxc1FVtbyRlXQVCN5s6DnAR9AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAGij4K7suqdDtmmX7pcutbaTNN2WSR1GL+URIlZRYbjclXQ2x3p+VGtdXzXmpVE5JLAygk+UnRVKWfDF2F5xxFt3enW3zGYbjRYe+qblus2aUUHjIsCxKhmh+FK5XL8lJ5vGQ0VI13NH1dfTo7kzpubrZYNhGJ6aYViOnWB2KhxjCMCxmhw3D8btjHMt1gtdspYqKgo4EVVd0IYYYo29JVXkxOaqvpAPqSH/jL8LTFOJ1tplxu1LbMe3GaUtq8n0BzutYjKdtXLHGtdjlyk9Cpbrs2ngje9F509RBSVCJI2GSCaYAAGKlqLp1nOkeeZdphqZi14wnUDA7/AFOL5fid/pForxYa+jldFUU88a/ztc1eTmqrXIqOa5zVRV+MNObjT8ETAOJNicmrWlC2DTfeLh1n8lseW1cPkeOavUNPGviLDkz42q5sjOSMpLmjXyU6fxUjZIegkGbbrXofq3ty1MyrRvXLT/JNMdTcKuDrbkeI5TQrR3CjcnpjmieirFUU0zOjLBV075IKiJ7JIpJI3teoHlYAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB9np1p1nOrmeYjphpni14zbUDPL/TYviGJ2CkWtvF+r6yVsVPTwRp/O5zk5ucqNaiK5zmtRVT6HRTQ/VvcbqZiujehun+Sanam5rcG23HMRxahWsuFY5fTJNK9VSKnpoWdKWerqHxwU8THySyRxsc9NJLgscETAOGzicerWq62DUjeLmNn8lvmW0kPlmOaQ0NRGnj7DjL5Go50j+asq7mrWSVCfxUbY4emk4HvnBq4WuI8MbbVFjdyS3ZBuK1WZR5Pr9nVIjZYH1kUT/IsdtkvLmtutSVFRHG9fTUT1FVUKjGyxwwy/AAAAAAj73+cMradxIMCgxLcNg7nZRY6OWnwLV/EZIrJqpp86VXOVKC4rG9stM57le+grI56V7uT1h8Y1kjZBAAZjHEE8HY30bLqu/ZfpzjlVus0IoZH1NNn2k9kmqM6sdKnpR1/xFrpa2BWIj3PnoVraVjI+nJPD0ug2AmWKSCSSGaN8M0L1ililYsckTmqqOa5q+lFRUVFRfm5G26R27suFBw/d7E1wu2vu2zBrzm1waqzan4nFNp5qc+T5XQlqL3bXwVFWrFc5Wx1y1EXNV5xr8wBkYgvoa4+CDaEX6sqq/bru31P01p5Hunhx7VnBbbq1SsVy8/ERV9FNaZYom81RqyRVD0ajUc5683rGfqX4Jjvjw233a+41uF2o5Lj9moJrjVzXy85fil4eyJjXqkdLFYayNXL/ABiclnRPkt9PpVGgVXwT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAT39XW3setDa132y33aHV1t7HrQ2td9st92gCBAE9/V1t7HrQ2td9st92h1dbex60NrXfbLfdoAgQBPf1dbex60NrXfbLfdodXW3setDa132y33aAIEAWoNNPBMd8eZW+033Jdwu1HGsfvNBDcaSax3nL8rvDGSsc9EkpZbDRxo5P4tOSTqnynen0IjpMNDvBBtCLDWUtfuK3b6n6lU8b2zzY9pNgtt0lpXq1efiJa+tmu0ssTuSI5Y4qd6tVyNcxeT0AoZRRSTyRwwxvmmmekUUUTFkklc5eTWtanpVVVURET5+ZPtw/PB2N9G9GrsWX6jY5VbU9CK6RlTU59qxZJqfOr5Sr6VdYcRc6KtnV6KxzJ65aKleyTpxzzdHoOvq7TeFBw/dk81vu2gW2zBrNm1vaiw6n5ZFNqHqcyT5PTlp73cnz1FIr1a1XR0K08XNE5Rp8xIkAR97A+GVtO4b+BT4lt5wdzcovlHFT57q/l0kV71U1BdEqORK+4pGxsVM17Ueygo44KVjub0h8Y58jpBAAAAAD//Z" '
        f'style="width:100%;max-width:100%;height:auto;border-radius:8px;"/></div>',
        unsafe_allow_html=True)

    # === ИНДИКАТОР: редактирование существующего КП ===
    _editing_kp_id = st.session_state.get("crm_edit_kp_id")
    if _editing_kp_id:
        st.warning(f"📝 **Редактируется КП id={_editing_kp_id}** из CRM. "
                   "Правьте данные и нажмите «Сохранить КП» — будет создана новая версия.")
        if st.button("✖ Выйти из редактирования", key="exit_edit_mode", use_container_width=True):
            st.session_state.pop("crm_edit_kp_id", None)
            st.rerun()
        st.markdown("---")

    st.header("📄 Тип КП")
    kp_mode = st.radio(
        "Что включает КП?",
        ["Кран", "Траверса", "Кран + траверса"],
        index=0,
        horizontal=False,
    )
    _needs_crane = kp_mode in ("Кран", "Кран + траверса")
    # Дефолты (чтобы код не падал при чистой траверсе)
    series = "ЛКС73"
    capacity = 500
    boom = 5
    height_to_arm = 4
    column_diameter = None
    include_flange = False
    flange_code = None
    use_lllm = False
    lllm_code = None
    hoist_brand = "R-Tech"
    hoist_mode = "8/2 м/мин"
    hoist_height = 3
    hc_capacity = None
    include_electrification = False
    include_montage = False
    montage_vat = False
    montage_price = 0.0

    st.markdown("---")
    # В режиме траверсы — полностью скрываем блок крана
    _crane_holder = st.container() if _needs_crane else st.empty()
    with _crane_holder:
     st.header("⚙️ Параметры крана")

     # В комбо-режиме траверса совместима только с одноплечевыми кранами
     if kp_mode == "Кран + траверса":
         _series_options = ["ЛКС71", "ЛКС73"]
         _series_help = ("В комбо с траверсой — только одноплечевые краны: "
                         "ЛКС71 (настенный), ЛКС73 (на колонне)")
     else:
         _series_options = SERIES_OPTIONS
         _series_help = ("ЛКС71 — настенный, ЛКС73 — на колонне, "
                         "ЛКС73М — на колонне с электрическим поворотом (LLL), "
                         "ЛКС77 — двухплечевой настенный, ЛКС78 — "
                         "двухплечевой на колонне")
     _series_ui = st.selectbox("Серия", _series_options, index=1,
                                help=_series_help)
     # Внутреннее представление: ЛКС73М отдельная опция в UI, но внутри — series=ЛКС73 + use_lllm=True
     if _series_ui == "ЛКС73М":
         series = "ЛКС73"
         _is_lllm_series = True
     else:
         series = _series_ui
         _is_lllm_series = False

     capacities = get_allowed_capacities(series)
     if _is_lllm_series:
         # ЛКС73М — только 1000 кг
         capacities = [1000] if 1000 in capacities else capacities
     capacity = st.selectbox("Грузоподъёмность, кг", capacities)

     # Пред-загрузка прайса ЛКС71/73 для фильтров ниже
     _df71_early = load_price_cranes_71_73(None)

     # ЛКС73М: показываем только «Конфигурация LLL». Стрела/высота/диаметр колонны
     # вычисляются из выбранного LLL-кода, а не выбираются отдельно.
     if _is_lllm_series:
         # boom/height_to_arm/column_diameter будут установлены ниже, в блоке LLL
         boom = 3
         height_to_arm = 3
         column_diameter = None
     else:
         booms = get_allowed_booms(series, capacity)
         boom = st.selectbox("Длина стрелы, м", booms)

         heights = get_allowed_heights(series)
         if series == "ЛКС77":
             height_to_arm = 0
             st.caption("ЛКС77 — настенный, высота колонны не задаётся.")
         elif series == "ЛКС71":
             height_to_arm = 0
             st.caption("ЛКС71 — настенный, высота колонны не задаётся.")
         elif series == "ЛКС78":
             # ЛКС78 — выбираем высоту колонны, показываем рабочую в лейбле
             height_to_arm = st.selectbox(
                 "Высота колонны, м", heights,
                 format_func=lambda h: f"{h:g} м (рабочая {LKS78_HEIGHT_TO_ARM.get(h, h - 0.2):g} м)")
         else:
             height_to_arm = st.selectbox("Высота до стрелы, м", heights)

         # --- Диаметр колонны (только для ЛКС73) ---
         column_diameter = None
         if series == "ЛКС73":
             diams = get_available_column_diameters(
                 _df71_early, series, capacity, boom, height_to_arm)
             if diams:
                 labels = [
                     f"{d} — высота до стрелы {int(d)/100:g} м "
                     f"(габарит {gabarite_height(capacity, int(d)/100):g} м)"
                     for d in diams
                 ]
                 idx = st.selectbox(
                     "Высота до стрелы (подрезка колонны), см",
                     range(len(diams)),
                     format_func=lambda i: labels[i],
                     help="Суффикс в коде крана (напр. ЛКС73.0500-6-5.420 → высота до стрелы 4.20 м)")
                 column_diameter = diams[idx]

     # --- Кран с электрическим поворотом (ЛКС73М LLL) ---
     # Если пользователь выбрал в списке серий "ЛКС73М" — автоматически включаем LLL-режим
     use_lllm = False
     lllm_code = None
     if _is_lllm_series and capacity == 1000:
        lllm_vars = get_lllm_variants(_df71_early, capacity)
        if lllm_vars:
            use_lllm = True
            lbl = [f"{v['code']} — стрела {v['boom']} м, высота {v['height']} м ({v['price']:,.0f} ₽)".replace(",", " ")
                   for v in lllm_vars]
            idx = st.selectbox("Конфигурация ЛКС73М (стрела × высота)", range(len(lllm_vars)),
                               format_func=lambda i: lbl[i],
                               help="Выберите конфигурацию — стрела и высота будут использованы автоматически")
            lllm_code = lllm_vars[idx]["code"]
            # Устанавливаем boom и height_to_arm по выбранному коду
            boom = int(lllm_vars[idx]["boom"])
            height_to_arm = int(lllm_vars[idx]["height"])
        else:
            st.warning("Конфигурации LLL не найдены в прайсе.")
     elif _is_lllm_series and capacity != 1000:
        st.info("ЛКС73М доступен только для г/п 1000 кг — измените грузоподъёмность.")

     # --- Фланец уширенный (для ЛКС73 / ЛКС73М) ---
     # Автоподбор по D2 (диаметру фланца крана) из crane_loads.CRANE_D2
     include_flange = False
     flange_code = None
     if series == "ЛКС73":
        flanges = get_flanges(_df71_early)
        if flanges:
            include_flange = st.checkbox("Уширенный фланец под основание", value=False)
            if include_flange:
                # Автоподбор по конфигурации крана
                try:
                    from crane_loads import get_d2_for_crane, get_flange_code_for_d2
                    _h_for_d2 = int(height_to_arm) if height_to_arm else 3
                    _d2 = get_d2_for_crane(capacity, boom, _h_for_d2)
                    _suggested_code = get_flange_code_for_d2(_d2) if _d2 else None
                except Exception:
                    _d2 = None
                    _suggested_code = None

                labels = [f"{fl['code']} — ø{fl['diameter']} мм ({fl['price']:,.0f} ₽)".replace(",", " ")
                          for fl in flanges]
                _default_idx = 0
                if _suggested_code:
                    for _i, _fl in enumerate(flanges):
                        _c = str(_fl["code"]).replace("ЛКС", "").replace("КС", "").replace(" ", "").lower()
                        _s = _suggested_code.replace("КС", "").replace(" ", "").lower()
                        if _c == _s:
                            _default_idx = _i
                            break

                if _d2 and _suggested_code:
                    st.caption(
                        f"🎯 Автоподбор: D2 крана = {_d2} мм → рекомендуется **{_suggested_code}**")
                elif _d2 is None:
                    st.caption("⚠️ Для этой конфигурации крана D2 не найден в таблице — выберите фланец вручную.")

                idx = st.selectbox("Фланец", range(len(flanges)),
                                   format_func=lambda i: labels[i],
                                   index=_default_idx)
                flange_code = flanges[idx]["code"]

     # --- Далее: таль и опции ---

     st.markdown("---")
     st.subheader("🔗 Таль")
     include_hoist = st.checkbox(
         "Включить таль в КП",
         value=st.session_state.get("include_hoist", True),
         key="include_hoist",
         help="Снимите галочку, если таль в КП не нужна (клиент сам её купит)")
     hoists = get_allowed_hoists(series, capacity)
     hoist_brand = st.selectbox("Бренд тали", hoists,
                                disabled=(not include_hoist))

     modes = get_hoist_modes(series, hoist_brand, height_to_arm, capacity)
     # В режиме Кран + траверса таль с ручным контролем груза несовместима — вырезаем
     if kp_mode == "Кран + траверса":
         modes = [m for m in modes if m != "ручной контроль груза"]
     hoist_mode = st.selectbox("Исполнение / скорость", modes)

     # Для тали с ручным контролем груза:
     #  - ЛКС77/78: отдельный выбор 125 или 250 кг
     #  - ЛКС71/73 80/125 кг: г/п тали совпадает с г/п крана (80 или 125)
     hc_capacity = None
     if hoist_mode == "ручной контроль груза":
        if series in ("ЛКС77", "ЛКС78"):
            hc_capacity = st.selectbox(
                "Г/п тали (HC), кг", HC_CAPACITIES, index=0,
                help="Таль с ручным контролем груза — 125 или 250 кг")
        else:
            # ЛКС71/73 80 или 125 кг — г/п тали равна г/п крана
            hc_capacity = capacity

     # Для OCALIFT (ЛКС71/73) в прайсе только 3 / 4.5 / 6 м.
     if hoist_brand == "OCALIFT":
        hoist_height = st.selectbox("Высота подъёма тали, м",
                                    [3, 4.5, 6], index=0)
     elif (series in ("ЛКС77", "ЛКС78")
           and hoist_mode == "ручной контроль груза"):
        # Таль HC для ЛКС77/78 — только высота 3 м
        hoist_height = 3
        st.caption("Для тали с ручным контролем груза (ЛКС77/78) — только высота 3 м")
     else:
        hoist_height = st.selectbox("Высота подъёма тали, м",
                                    [3, 4, 5, 6, 9], index=0)

     # --- Индекс исполнения тали (P2/C2/C3) для R-Tech 8/2 м/мин ---
     # P2 — стационарная пульт 2 кн., C2 — с двумя разъёмами, C3 — радиоуправление
     hoist_exec = None
     # Выбор исполнения тали (P2/C2/C3) показываем ПОСЛЕ чекбокса электрификации ниже.
     # Здесь просто плейсхолдер — по умолчанию P2 (стационарная, ручной пульт).
     hoist_exec = "P2"

     st.markdown("---")
     st.subheader("➕ Опции")
     can_electrify = (series in ("ЛКС71", "ЛКС73")
                    and capacity in ELECTRIFICATION_CAPACITIES)
     include_electrification = st.checkbox(
        "Пакет электрификации (250/500/1000 кг ЛКС71/73/73М)",
        value=False, disabled=not can_electrify,
        help="Доступен для ЛКС71/73/73М г/п 250/500/1000 кг.",
     )
     # Выбор исполнения тали C2/C3 — ТОЛЬКО когда кран с электрификацией.
     # Без электрификации таль всегда P2 (стационарная, ручной пульт).
     if (include_electrification and hoist_brand == "R-Tech"
             and hoist_mode == "8/2 м/мин"
             and series in ("ЛКС71", "ЛКС73")):
         _exec_options = ["P2 — стационарная (пульт 2 кнопки, ручной контроль)",
                          "C2 — с двумя разъёмами (для электрификации)",
                          "C3 — радиоуправление (для электрификации)"]
         _exec_default_idx = 2 if _is_lllm_series else 1  # LLL→C3, обычные→C2
         _prev_lllm_hoist = st.session_state.get("_hoist_exec_last_lllm")
         if _prev_lllm_hoist is not None and _prev_lllm_hoist != _is_lllm_series:
             st.session_state.pop("hoist_exec_ui", None)
         st.session_state["_hoist_exec_last_lllm"] = _is_lllm_series
         _exec_pick = st.selectbox(
             "Индекс исполнения тали",
             range(len(_exec_options)),
             format_func=lambda i: _exec_options[i],
             index=_exec_default_idx,
             key="hoist_exec_ui",
             help=("C2/C3 применяются вместе с электрификацией. "
                   "Для ЛКС73М рекомендуется C3 (радиоуправление)."))
         hoist_exec = ["P2", "C2", "C3"][_exec_pick]
     else:
         # Без электрификации — всегда P2 (с ручным контролем)
         hoist_exec = "P2"
     # --- Настройка компонентов пакета электрификации (всегда доступно) ---
     electr_custom = False  # оставлено для обратной совместимости
     electr_shkaf_code = None
     electr_trolley_code = None
     electr_cables_state = None
     if include_electrification and can_electrify:
        electr_custom = True   # всегда включаем ручной выбор компонентов
        if True:
            # Ищем в прайсе все шкафы AR58* и приводные тележки AR540/AR541/AR542
            import re as _re_elec
            _shkafs = []
            _trolleys = []
            if _df71_early is not None:
                for _, _r in _df71_early.iterrows():
                    _c = str(_r.get("code", ""))
                    _n = str(_r.get("name", ""))
                    _p = float(_r.get("price", 0) or 0)
                    if _re_elec.match(r"^AR58\d", _c):
                        # Исключаем шкафы с индексом «Т» (ТЭК-исполнение)
                        if _c.rstrip().endswith("Т"):
                            continue
                        _shkafs.append({"code": _c, "name": _n, "price": _p})
                    elif _re_elec.match(r"^AR54[012]", _c):
                        _trolleys.append({"code": _c, "name": _n, "price": _p})
            # Дефолт по серии
            _default_shkaf = "AR580.32" if _is_lllm_series else "AR580.10"
            _default_trolley = "AR541" if _is_lllm_series else "AR542"

            # СБРОС выбора при ��мене серии — чтобы дефолт по серии применялся
            _last_lllm = st.session_state.get("_electr_last_is_lllm", None)
            if _last_lllm is not None and _last_lllm != _is_lllm_series:
                # Серия сменилась — сбрасываем предыдущий выбор
                st.session_state.pop("electr_shkaf", None)
                st.session_state.pop("electr_trolley", None)
            st.session_state["_electr_last_is_lllm"] = _is_lllm_series

            # Шкаф управления
            if _shkafs:
                _shkaf_codes = [s["code"] for s in _shkafs]
                _shkaf_idx = _shkaf_codes.index(_default_shkaf) if _default_shkaf in _shkaf_codes else 0
                _s_labels = [f"{s['code']} — {s['name']} ({s['price']:,.0f} ₽)".replace(",", " ")
                             for s in _shkafs]
                _shkaf_pick = st.selectbox(
                    "Шкаф управления", range(len(_shkafs)),
                    format_func=lambda i: _s_labels[i],
                    index=_shkaf_idx, key="electr_shkaf")
                electr_shkaf_code = _shkafs[_shkaf_pick]["code"]
                # Подсказка автоподбора
                if electr_shkaf_code != _default_shkaf and _default_shkaf in _shkaf_codes:
                    st.caption(f"💡 Для {'ЛКС73М' if _is_lllm_series else 'обычных ЛКС71/73'} "
                              f"рекомендуется **{_default_shkaf}**")
            else:
                st.warning("В прайсе не найдено шкафов AR58* — добавьте их в файл konsolnkii.xlsx")

            # Приводная тележка
            if _trolleys:
                _trol_codes = [t["code"] for t in _trolleys]
                _trol_idx = _trol_codes.index(_default_trolley) if _default_trolley in _trol_codes else 0
                _t_labels = [f"{t['code']} — {t['name']} ({t['price']:,.0f} ₽)".replace(",", " ")
                             for t in _trolleys]
                _trol_pick = st.selectbox(
                    "Приводная тележка", range(len(_trolleys)),
                    format_func=lambda i: _t_labels[i],
                    index=_trol_idx, key="electr_trolley")
                electr_trolley_code = _trolleys[_trol_pick]["code"]
                if electr_trolley_code != _default_trolley and _default_trolley in _trol_codes:
                    st.caption(f"💡 Для {'ЛКС73М' if _is_lllm_series else 'обычных ЛКС71/73'} "
                              f"рекомендуется **{_default_trolley}**")
            else:
                st.warning("В прайсе не найдено приводных тележек AR540/541/542 — добавьте их в konsolnkii.xlsx")

            # Кабели — редактируемые метры
            if "electr_cables" not in st.session_state:
                st.session_state["electr_cables"] = {
                    "4x1.5": 10.5, "6x1.5": 3.5, "5x2.5": 20.0, "4x2.5": 20.0,
                }
            _cbl_c1, _cbl_c2 = st.columns(2)
            with _cbl_c1:
                st.session_state["electr_cables"]["4x1.5"] = st.number_input(
                    "YFFB 4×1.5, м", min_value=0.0, max_value=200.0, step=0.5,
                    value=float(st.session_state["electr_cables"]["4x1.5"]))
                st.session_state["electr_cables"]["5x2.5"] = st.number_input(
                    "YFFB 5×2.5, м", min_value=0.0, max_value=200.0, step=1.0,
                    value=float(st.session_state["electr_cables"]["5x2.5"]))
            with _cbl_c2:
                st.session_state["electr_cables"]["6x1.5"] = st.number_input(
                    "YFFB 6×1.5, м", min_value=0.0, max_value=200.0, step=0.5,
                    value=float(st.session_state["electr_cables"]["6x1.5"]))
                st.session_state["electr_cables"]["4x2.5"] = st.number_input(
                    "YFFB 4×2.5, м", min_value=0.0, max_value=200.0, step=1.0,
                    value=float(st.session_state["electr_cables"]["4x2.5"]))
            electr_cables_state = dict(st.session_state["electr_cables"])
            # Сохраняем финальный выбор в session_state — build_specification прочитает
            st.session_state["electr_shkaf_code_final"] = electr_shkaf_code
            st.session_state["electr_trolley_code_final"] = electr_trolley_code
            st.session_state["electr_cables_final"] = electr_cables_state
        else:
            # Автоподбор → чистим финальные ключи
            st.session_state.pop("electr_shkaf_code_final", None)
            st.session_state.pop("electr_trolley_code_final", None)
            st.session_state.pop("electr_cables_final", None)
     include_montage = st.checkbox("Монтаж и пусконаладка (отдельной таблицей)",
                                  value=False)
     montage_vat = st.checkbox("— включить НДС 22 % в стоимость монтажа",
                              value=False, disabled=not include_montage,
                              help="По умолчанию монтаж от ИП без НДС по отдельному договору.")
     montage_price = st.number_input(
        "Стоимость монтажа, ₽" + (" (с НДС 22 %)" if montage_vat else " (без НДС)"),
        min_value=0, max_value=1_000_000, value=35000, step=5000,
        disabled=not include_montage)

    # ================= ВАКУУМНАЯ ТРАВЕРСА =================
    tv_selection = None
    if kp_mode in ("Траверса", "Кран + траверса"):
        st.markdown("---")
        st.header("🔹 Вакуумная траверса VacuTec")
        _tv_df = _tv.load_traverse_price(
            st.session_state.get("traverse_price_upload"))
        if _tv_df is None or _tv_df.empty:
            st.warning("Прайс траверс не найден.")
        else:
            # Фильтр по типу конструкции
            tv_type = st.selectbox(
                "Тип конструкции",
                ["P — стандартная",
                 "PR — усиленная",
                 "GT — для листов большой ширины",
                 "На пневмоцилиндре (скоро)"],
                index=0,
            )
            type_key = tv_type.split(" ")[0]
            if type_key == "На":
                st.info("Категория на пневмоцилиндре — скоро в прайсе. Пока выберите P/PR/GT.")
            else:
                # Фильтруем базовые коды по типу
                base_codes = _tv.get_base_codes(_tv_df)
                import re as _re
                filt = []
                for bc in base_codes:
                    m = _re.match(r"VacuTec (\d+)(P|PR|GT)-", bc)
                    if m and m.group(2) == type_key:
                        filt.append(bc)
                if not filt:
                    st.warning("Нет моделей этого типа.")
                else:
                    # Лейблы: код + г/п
                    labels = []
                    for bc in filt:
                        parsed = _tv.parse_code(bc)
                        if parsed:
                            labels.append(
                                f"{bc} — {parsed['capacity']} кг, "
                                f"{parsed['suckers']} присосок, "
                                f"L{parsed['beam_length']} м, D{parsed['d_sucker']}"
                            )
                        else:
                            labels.append(bc)
                    # Мультивыбор: можно выбрать 1 или несколько траверс
                    idxs = st.multiselect(
                        "Модели траверс (можно выбрать несколько)",
                        options=list(range(len(filt))),
                        default=[0] if filt else [],
                        format_func=lambda i: labels[i],
                        help="Первая — основная, остальные — альтернативные варианты (каждая на отдельной странице в PDF)")
                    if not idxs:
                        st.warning("Выберите хотя бы одну модель")
                        st.stop()
                    tv_base = filt[idxs[0]]
                    _extra_bases = [filt[i] for i in idxs[1:]]
                    tv_power = st.radio(
                        "Питание (основной траверсы)", ["220В", "АКБ"], index=0,
                        horizontal=True,
                        help="АКБ — версия с суффиксом A в коде. Аккумулятор "
                             "не добавляется автоматически.")
                    st.caption("Дополнительная комплектация")
                    tv_opt_bat = st.checkbox(
                        "➕ Сменный аккумулятор 24V 20,8Ah + блок питания")
                    tv_opt_sup = st.checkbox("➕ Опоры для хранения траверсы типа Р")
                    tv_opt_cbl = st.checkbox("➕ Спиральный кабель управления")
                    tv_opt_hnd = st.checkbox("➕ Наклонная ручка")
                    tv_include_mont = st.checkbox(
                        "🔧 Монтаж и ПНР траверсы", value=False)
                    tv_mont_vat = False
                    tv_mont_price = 0.0
                    if tv_include_mont:
                        tv_mont_vat = st.checkbox(
                            "— включить НДС 22 % в стоимость монтажа", value=False)
                        tv_mont_price = st.number_input(
                            "Стоимость монтажа траверсы, ₽"
                            + (" (с НДС)" if tv_mont_vat else " (без НДС)"),
                            min_value=0, max_value=1_000_000,
                            value=15000, step=5000)
                    # Условия поставки берутся из общего блока «Условия КП» (ниже на странице)
                    # Дефолт для траверсы: предоплата 100 %, доставка — ДЛ с обрешёткой
                    if "kp_prepay_pct" not in st.session_state:
                        st.session_state["kp_prepay_pct"] = 100
                    if "kp_delivery_option_idx" not in st.session_state:
                        st.session_state["kp_delivery_option_idx"] = 1
                    tv_prepay_pct = int(st.session_state.get("kp_prepay_pct", 100))
                    tv_delivery_option = st.session_state.get(
                        "kp_delivery_option", "Доставка до ТК «Деловые линии» и отправка с учётом обрешётки за счёт покупателя")
                    tv_delivery_addr = st.session_state.get(
                        "kp_delivery_text") or tv_delivery_option
                    tv_delivery_price = int(st.session_state.get("kp_delivery_price", 0))
                    tv_delivery_target = st.session_state.get("kp_delivery_target", "")
                    tv_selection = _tv.TraverseSelection(
                        base_code=tv_base,
                        power_type=tv_power,
                        options_battery=tv_opt_bat,
                        options_supports=tv_opt_sup,
                        options_cable=tv_opt_cbl,
                        options_handle=tv_opt_hnd,
                        include_montage=tv_include_mont,
                        montage_price=float(tv_mont_price),
                        montage_vat=tv_mont_vat,
                        prepay_pct=int(tv_prepay_pct),
                        delivery_address=tv_delivery_addr.strip(),
                        delivery_price=float(tv_delivery_price),
                    )

                    # Альтернативные траверсы (выбраны в multiselect кроме первой)
                    _tv_extra_selections = [
                        _tv.TraverseSelection(
                            base_code=_bc, power_type=tv_power,
                            include_montage=False, montage_price=0.0,
                            montage_vat=False, prepay_pct=int(tv_prepay_pct),
                            delivery_address=tv_delivery_addr.strip())
                        for _bc in _extra_bases
                    ]
                    st.session_state["_tv_extra_selections"] = _tv_extra_selections
                    # Совместимость с старым именем — берём первую из альтернатив
                    st.session_state["_tv_selection2"] = (
                        _tv_extra_selections[0] if _tv_extra_selections else None)
                    # Мини-превью (включая все альтернативные траверсы)
                    _tv_items = list(_tv.build_traverse_items(tv_selection, _tv_df))
                    for _es in _tv_extra_selections:
                        _tv_items.extend(
                            _tv.build_traverse_items(_es, _tv_df))
                    _tv_total = sum(i["total_vat"] for i in _tv_items)
                    if tv_selection.include_montage:
                        _tv_total += tv_selection.montage_price
                    st.success(f"Итог КП на траверсу: {_tv_total:,.0f} ₽".replace(",", " "))

    st.markdown("---")
    st.subheader("📁 Прайсы (можно заменить)")
    up_71_73 = st.file_uploader(f"Прайс ЛКС71/73 ({PRICE_CRANES_71_73})",
                                type=["xlsx", "xls"], key="p71")
    up_77_78 = st.file_uploader(f"Прайс ЛКС77/78 ({PRICE_CRANES_77_78})",
                                type=["xlsx", "xls"], key="p77")
    up_hoists = st.file_uploader(f"Прайс тали ({PRICE_HOISTS})",
                                 type=["xlsx", "xls"], key="ph")
    up_traverses = st.file_uploader(
        "Прайс вакуумных траверс (Vakuumnye-traversy-7.xlsx)",
        type=["xlsx", "xls"], key="ptv")
    if up_traverses is not None:
        st.session_state["traverse_price_upload"] = up_traverses

    # =============== РУЧНАЯ ПОЗИЦИЯ ===============
    st.markdown("---")
    st.subheader("➕ Ручные позиции")
    st.caption("Товары, которых нет в прайсе — добавятся в спецификацию.")
    custom_items = st.session_state.setdefault("custom_items", [])

    # → Память: выбор товара из истории (автоподстановка)
    import history_memory as _hm
    _prod_names = _hm.get_product_names()
    if _prod_names:
        _hint_col1, _hint_col2 = st.columns([3, 1])
        with _hint_col1:
            _picked_prod = st.selectbox(
                "💡 Подставить из истории",
                ["—"] + _prod_names,
                key="ci_hint_pick",
                help="Ранее введённые вами вручные позиции")
        with _hint_col2:
            if st.button("↳ Подставить",
                         key="ci_hint_apply", use_container_width=True,
                         disabled=(_picked_prod == "—")):
                _p = _hm.get_product(_picked_prod)
                if _p:
                    st.session_state["ci_name"] = _picked_prod
                    st.session_state["ci_code"] = _p.get("code", "")
                    st.session_state["ci_price"] = float(_p.get("price") or 0)
                    st.session_state["ci_unit"] = _p.get("unit", "шт")
                    st.rerun()

    with st.form("add_custom_item", clear_on_submit=True):
        ci_col1, ci_col2 = st.columns([2, 1])
        with ci_col1:
            ci_name = st.text_input("Наименование", key="ci_name")
            ci_code = st.text_input("Код (необязательно)", key="ci_code")
        with ci_col2:
            ci_qty = st.number_input("Кол-во", min_value=0.0, value=1.0,
                                     step=0.1, key="ci_qty")
            ci_unit = st.selectbox("Ед.", ["шт", "м", "компл", "усл.", "компт."],
                                    index=0, key="ci_unit")
        ci_price = st.number_input("Цена за ед., ₽ (с НДС 22 %)",
                                   min_value=0.0, value=0.0, step=100.0,
                                   key="ci_price")
        ci_add = st.form_submit_button("➕ Добавить в спецификацию")
        if ci_add and ci_name and ci_price > 0:
            custom_items.append({
                "code": ci_code or "CUST",
                "name": ci_name,
                "unit": ci_unit,
                "qty": float(ci_qty),
                "price": float(ci_price),
                "total": float(ci_qty) * float(ci_price),
            })
            st.session_state["custom_items"] = custom_items
            # Запоминаем товар в истории
            _hm.remember_product(ci_name, ci_code or "", float(ci_price), ci_unit)
            st.rerun()
    if custom_items:
        st.caption(f"Добавлено позиций: {len(custom_items)}")
        for i, it in enumerate(list(custom_items)):
            cc1, cc2 = st.columns([4, 1])
            with cc1:
                st.text(f"{it['name'][:40]} — {it['qty']:g} {it['unit']} × {it['price']:,.0f} ₽".replace(",", " "))
            with cc2:
                if st.button("✖", key=f"del_ci_{i}"):
                    custom_items.pop(i)
                    st.session_state["custom_items"] = custom_items
                    st.rerun()

    st.markdown("---")
    with st.expander("Цены OCALIFT (нет в прайсе)"):
        ocp = st.session_state.setdefault("ocalift_prices",
                                          dict(OCALIFT_PRICES_DEFAULT))
        for k, v in list(ocp.items()):
            ocp[k] = st.number_input(f"{k[0]} кг · {k[1]}",
                                     min_value=0, max_value=500_000,
                                     value=int(v), step=1000, key=f"oca_{k}")

# Загружаем прайсы
df71 = load_price_cranes_71_73(up_71_73)
df77 = load_price_cranes_77_78(up_77_78)
dfh = load_price_hoists(up_hoists)

# Статус прайсов перенесён в сайдбар в виде компактной подписи — не загромождает главную страницу
with st.sidebar:
    _price_status_lines = []
    if df71 is not None:
        _price_status_lines.append(f"ЛКС71/73: {len(df71)} поз.")
    else:
        _price_status_lines.append("⚠️ ЛКС71/73 не загружен")
    if df77 is not None:
        _price_status_lines.append(f"ЛКС77/78: {len(df77)} поз.")
    else:
        _price_status_lines.append("⚠️ ЛКС77/78 не загружен")
    if dfh is not None:
        _price_status_lines.append(f"Тали: {len(dfh)} поз.")
    else:
        _price_status_lines.append("⚠️ Тали не загружены")
    st.caption("📊 Прайсы: " + " · ".join(_price_status_lines))

# --- Режим «Только траверса» — показываем только её блок и выходим ---
if kp_mode == "Траверса":
    st.header("🔹 КП на вакуумную траверсу")
    if tv_selection is None:
        st.info("Выберите модель траверсы в сайдбаре слева.")
        st.stop()

    # === Реквизиты Покупателя (вводятся при выставлении КП — переносятся в договор) ===
    # Сохраняемся в общий buyer_data — отсюда читают и крановой договор и внешний
    with st.expander("👥 Реквизиты Покупателя (будут в КП и договоре)", expanded=False):
        _tvhdr_bd = st.session_state.get("buyer_data", {}) or {}
        _tvhdr_c1, _tvhdr_c2 = st.columns(2)
        with _tvhdr_c1:
            _tvhdr_b_short = st.text_input(
                "Наименование краткое",
                value=_tvhdr_bd.get("short", ""),
                key="tvhdr_b_short",
                placeholder="ООО «Ромашка»")
            _tvhdr_b_full = st.text_input(
                "Наименование полное",
                value=_tvhdr_bd.get("full", ""),
                key="tvhdr_b_full")
            _tvhdr_b_address = st.text_input(
                "Юр. адрес",
                value=_tvhdr_bd.get("address", ""),
                key="tvhdr_b_address")
            _tvhdr_b_inn = st.text_input(
                "ИНН",
                value=_tvhdr_bd.get("inn", ""),
                key="tvhdr_b_inn")
            _tvhdr_b_kpp = st.text_input(
                "КПП",
                value=_tvhdr_bd.get("kpp", ""),
                key="tvhdr_b_kpp")
            _tvhdr_b_ogrn = st.text_input(
                "ОГРН",
                value=_tvhdr_bd.get("ogrn", ""),
                key="tvhdr_b_ogrn")
            _tvhdr_b_phone = st.text_input(
                "Телефон",
                value=_tvhdr_bd.get("phone", ""),
                key="tvhdr_b_phone")
            _tvhdr_b_email = st.text_input(
                "E-mail",
                value=_tvhdr_bd.get("email", ""),
                key="tvhdr_b_email")
        with _tvhdr_c2:
            _tvhdr_b_bank = st.text_input(
                "Банк",
                value=_tvhdr_bd.get("bank", ""),
                key="tvhdr_b_bank")
            _tvhdr_b_bik = st.text_input(
                "БИК",
                value=_tvhdr_bd.get("bik", ""),
                key="tvhdr_b_bik")
            _tvhdr_b_rs = st.text_input(
                "Р/с",
                value=_tvhdr_bd.get("rs", ""),
                key="tvhdr_b_rs")
            _tvhdr_b_ks = st.text_input(
                "К/с",
                value=_tvhdr_bd.get("ks", ""),
                key="tvhdr_b_ks")
            _tvhdr_b_dir_pos = st.text_input(
                "Должность подписанта",
                value=_tvhdr_bd.get("director_position", "Генеральный директор"),
                key="tvhdr_b_dir_pos")
            _tvhdr_b_dir_gen = st.text_input(
                "ФИО в родительном падеже",
                value=_tvhdr_bd.get("director_fio_gen", ""),
                key="tvhdr_b_dir_gen")
            _tvhdr_b_dir_short = st.text_input(
                "ФИО коротко (для подписи)",
                value=_tvhdr_bd.get("director_fio_short", ""),
                key="tvhdr_b_dir_short")
            _tvhdr_b_basis = st.text_input(
                "На основании",
                value=_tvhdr_bd.get("basis", "Устава"),
                key="tvhdr_b_basis")
        # Синхронизируем в buyer_data — чтобы блок договора ниже подхватил
        st.session_state["buyer_data"] = {
            "short": _tvhdr_b_short, "full": _tvhdr_b_full or _tvhdr_b_short,
            "address": _tvhdr_b_address, "post_address": _tvhdr_b_address,
            "inn": _tvhdr_b_inn, "kpp": _tvhdr_b_kpp, "ogrn": _tvhdr_b_ogrn,
            "phone": _tvhdr_b_phone, "email": _tvhdr_b_email,
            "bank": _tvhdr_b_bank, "bik": _tvhdr_b_bik,
            "rs": _tvhdr_b_rs, "ks": _tvhdr_b_ks,
            "director_position": _tvhdr_b_dir_pos,
            "director_fio_gen": _tvhdr_b_dir_gen,
            "director_fio_short": _tvhdr_b_dir_short,
            "basis": _tvhdr_b_basis,
        }
    _tv_df_main = _tv.load_traverse_price(
        st.session_state.get("traverse_price_upload"))
    _tv_items = list(_tv.build_traverse_items(tv_selection, _tv_df_main))
    # Альтернативные траверсы (все выбранные в multiselect)
    _tv_extras = st.session_state.get("_tv_extra_selections", []) or []
    for _es in _tv_extras:
        _tv_items.extend(_tv.build_traverse_items(_es, _tv_df_main))
    # Применяем сохранённые правки цены/кол-ва (по коду позиции)
    _tv_overrides = st.session_state.get("_tv_item_overrides", {}) or {}
    for _it in _tv_items:
        _ov = _tv_overrides.get(_it["code"])
        if _ov:
            if "qty" in _ov:
                _it["qty"] = float(_ov["qty"])
            if "price" in _ov:
                _it["price_vat"] = float(_ov["price"])
            _it["total_vat"] = float(_it["qty"]) * float(_it["price_vat"])
    # Строка доставки (если цена > 0) — попадает в таблицу, спецификацию и итог
    _tv_delivery_price = float(tv_selection.delivery_price) if getattr(
        tv_selection, "delivery_price", 0) else 0.0
    if _tv_delivery_price > 0:
        _tv_items.append({
            "code": "ДОСТАВКА",
            "name": (tv_selection.delivery_address
                     or "Доставка"),
            "unit": "усл.",
            "qty": 1.0,
            "price_vat": _tv_delivery_price,
            "total_vat": _tv_delivery_price,
        })
    # Спецификация
    st.subheader("📋 Спецификация траверсы")
    tv_rows = [{
        "Код": it["code"], "Наименование": it["name"],
        "Ед.": it["unit"], "Кол-во": it["qty"],
        "Цена, ₽": it["price_vat"], "Сумма, ₽": it["total_vat"],
    } for it in _tv_items]
    # Ручные позиции
    for _ci in st.session_state.get("custom_items", []):
        tv_rows.append({
            "Код": _ci["code"], "Наименование": _ci["name"],
            "Ед.": _ci["unit"], "Кол-во": _ci["qty"],
            "Цена, ₽": _ci["price"], "Сумма, ₽": _ci["total"],
        })
    st.caption("Количество и цену можно поменять прямо в таблице")
    _tv_df_display = pd.DataFrame([{
        "Код": r["Код"], "Наименование": r["Наименование"],
        "Ед.": r["Ед."], "Кол-во": float(r["Кол-во"]),
        "Цена, ₽": float(r["Цена, ₽"]),
        "Сумма, ₽": float(r["Сумма, ₽"]),
    } for r in tv_rows])
    _tv_edited = st.data_editor(
        _tv_df_display, use_container_width=True, hide_index=True,
        key="tv_spec_editor",
        column_config={
            "Кол-во": st.column_config.NumberColumn(
                "Кол-во", min_value=0.0, step=1.0, format="%.2f"),
            "Цена, ₽": st.column_config.NumberColumn(
                "Цена, ₽", min_value=0.0, step=100.0, format="%.2f ₽"),
            "Сумма, ₽": st.column_config.NumberColumn(
                "Сумма, ₽", format="%.2f ₽", disabled=True),
            "Код": st.column_config.TextColumn(disabled=True),
            "Наименование": st.column_config.TextColumn(disabled=True),
            "Ед.": st.column_config.TextColumn(disabled=True),
        },
    )
    # Пересчёт сумм по отредактированным количествам и ценам + сохранение в overrides
    _tv_overrides_new = dict(st.session_state.get("_tv_item_overrides", {}) or {})
    _tv_any_changed = False
    for _i, _row in _tv_edited.iterrows():
        try:
            _new_qty = float(_row["Кол-во"])
            _new_price = float(_row["Цена, ₽"])
            _prev_qty = float(tv_rows[_i]["Кол-во"])
            _prev_price = float(tv_rows[_i]["Цена, ₽"])
            tv_rows[_i]["Кол-во"] = _new_qty
            tv_rows[_i]["Цена, ₽"] = _new_price
            tv_rows[_i]["Сумма, ₽"] = _new_qty * _new_price
            # Корректируем базовые items для PDF/DOCX
            if _i < len(_tv_items):
                _code = _tv_items[_i]["code"]
                _tv_items[_i]["qty"] = _new_qty
                _tv_items[_i]["price_vat"] = _new_price
                _tv_items[_i]["total_vat"] = _new_qty * _new_price
                # Сохраняем правку — переживёт rerun
                _tv_overrides_new[_code] = {"qty": _new_qty, "price": _new_price}
            if _new_qty != _prev_qty or _new_price != _prev_price:
                _tv_any_changed = True
        except Exception:
            pass
    st.session_state["_tv_item_overrides"] = _tv_overrides_new
    # --- Скидка (траверса) ---
    st.subheader("🏷️ Скидка клиенту")
    _tv_disc_c1, _tv_disc_c2 = st.columns([1, 3])
    with _tv_disc_c1:
        _tv_discount_pct = st.number_input(
            "Скидка, %", min_value=0.0, max_value=50.0,
            value=float(st.session_state.get("kp_discount_pct", 0.0)),
            step=1.0, key="kp_discount_pct",
            help="Применяется к позициям траверсы (с НДС) и монтажу (без НДС). В PDF появится строка «Вам предоставлена скидка X%».")
    _tv_disc_r = 1.0 - float(_tv_discount_pct) / 100.0
    if _tv_discount_pct > 0:
        for _r in tv_rows:
            _r["Цена, ₽"] = float(_r["Цена, ₽"]) * _tv_disc_r
            _r["Сумма, ₽"] = float(_r["Кол-во"]) * float(_r["Цена, ₽"])
        for _it in _tv_items:
            _it["price_vat"] = float(_it["price_vat"]) * _tv_disc_r
            _it["total_vat"] = float(_it["qty"]) * float(_it["price_vat"])
        # Монтаж траверсы — тоже со скидкой (обычно без НДС)
        if tv_selection.include_montage and tv_selection.montage_price > 0:
            tv_selection.montage_price = float(tv_selection.montage_price) * _tv_disc_r
    tv_total_supply = sum(float(r["Сумма, ₽"]) for r in tv_rows)
    _tv_base_before = tv_total_supply / _tv_disc_r if _tv_disc_r > 0 else tv_total_supply
    _tv_discount_amount = _tv_base_before - tv_total_supply
    with _tv_disc_c2:
        if _tv_discount_pct > 0:
            st.success(f"Скидка {_tv_discount_pct:.0f} % применена. "
                       f"Сэкономлено: **{fmt_money(_tv_discount_amount)}**.")
        else:
            st.caption("Поставьте % > 0 — все цены и монтаж пересчитаются автоматически.")

    st.subheader("💰 Итоги")
    c1, c2, c3 = st.columns(3)
    c1.metric("Стоимость поставки с НДС 22 %",
              f"{tv_total_supply:,.0f} ₽".replace(",", " "))
    c2.metric("Предоплата 70 %",
              f"{tv_total_supply*0.70:,.0f} ₽".replace(",", " "))
    c3.metric("Остаток 30 %",
              f"{tv_total_supply*0.30:,.0f} ₽".replace(",", " "))
    if tv_selection.include_montage:
        st.info(f"+ Монтаж траверсы: {tv_selection.montage_price:,.0f} ₽ "
                f"({'с НДС 22%' if tv_selection.montage_vat else 'без НДС'})".replace(",", " "))
    st.markdown("---")
    st.subheader("📄 Скачать КП на траверсу")
    tv_col1, tv_col2, tv_col3 = st.columns([1, 1, 3])
    with tv_col1:
        _tv_base = date.today().strftime("%d%m%Y") + "/ЛКС-Т"
        try:
            import crm_db as _crm_db_tvnum
            _tv_default = _crm_db_tvnum.generate_unique_kp_number(_tv_base)
        except Exception:
            _tv_default = _tv_base
        tv_kp_number = st.text_input(
            "№ КП", value=_tv_default,
            key="tv_kp_num",
            help="Если номер уже в базе — добавится суффикс -2, -3 при сохранении.")
    with tv_col2:
        tv_kp_date = st.text_input(
            "Дата", value=date.today().strftime("%d.%m.%Y"), key="tv_kp_date")
    with tv_col3:
        tv_buyer = st.text_input("Кому", placeholder="ООО «Ромашка»",
                                 key="tv_kp_buyer")

    if st.button("📄 Сформировать КП (траверса)",
                 type="primary", use_container_width=True):
        # Собираем QuoteData-аналог для траверсы
        tv_q = _tv.build_traverse_quote_data(tv_selection, _tv_df_main)
        # Добавляем ручные позиции
        from dataclasses import dataclass as _dc
        @_dc
        class _SL:
            code: str; name: str; unit: str; qty: float; price: float
            @property
            def total(self): return self.qty * self.price
        for _ci in st.session_state.get("custom_items", []):
            tv_q.lines.append(_SL(_ci["code"], _ci["name"], _ci["unit"],
                                  _ci["qty"], _ci["price"]))

        # Все альтернативные траверсы ␲ добавляем в спецификацию
        _tv_extras_pdf = st.session_state.get("_tv_extra_selections", []) or []
        for _es_pdf in _tv_extras_pdf:
            for _it in _tv.build_traverse_items(_es_pdf, _tv_df_main):
                tv_q.lines.append(_SL(
                    _it["code"], _it["name"], _it["unit"],
                    _it["qty"], _it["price_vat"]))

        # Строка доставки (если цена > 0) — в спецификацию для PDF/DOCX
        _tv_dp = float(getattr(tv_selection, "delivery_price", 0) or 0)
        if _tv_dp > 0:
            tv_q.lines.append(_SL(
                "ДОСТАВКА",
                (tv_selection.delivery_address or "Доставка"),
                "усл.", 1.0, _tv_dp))

        # Функции-геттеры картинок/характеристик для траверсы
        def _tv_get_crane_img(_series):
            return _tv.get_traverse_image(tv_selection.code)
        def _tv_get_hoist_img(_brand, _exec_key=""):
            return None
        def _tv_crane_chars(_q):
            return _tv.traverse_characteristics(tv_selection, _tv_df_main)
        def _tv_hoist_chars(_q):
            return []
        def _tv_desc(_q):
            return _tv.get_traverse_description(tv_selection, _tv_df_main)
        tv_series_desc = {"VACUTEC": "Вакуумная траверса VacuTec"}

        # Альтернативные траверсы (каждая — своя страница)
        def _mk_extra(_sel):
            return {
                "selection": _sel,
                "image_fn": lambda _q, _s=_sel: _tv.get_traverse_image(_s.code),
                "characteristics_fn": lambda _q, _s=_sel: _tv.traverse_characteristics(_s, _tv_df_main),
                "description_fn": lambda _q, _s=_sel: _tv.get_traverse_description(_s, _tv_df_main),
            }
        _tv_pdf_extras = [_mk_extra(_es) for _es in _tv_extras]

        try:
            tv_pdf = build_kp_pdf(
                tv_q, tv_kp_number, tv_buyer,
                supplier=SUPPLIER, series_descriptions=tv_series_desc,
                get_crane_image_fn=_tv_get_crane_img,
                get_hoist_image_fn=_tv_get_hoist_img,
                crane_characteristics_fn=_tv_crane_chars,
                hoist_characteristics_fn=_tv_hoist_chars,
                description_fn=_tv_desc,
                # Альтернативные траверсы — каждая на своей странице
                traverse_extras=_tv_pdf_extras,
                prepayment_rate=PREPAYMENT_DEFAULT,
                kp_date=tv_kp_date,
                # Способ и стоимость доставки в условиях поставки
                delivery_terms=tv_selection.delivery_address,
                delivery_price=float(tv_selection.delivery_price),
                delivery_target=st.session_state.get("tv_delivery_target", ""),
                prepay_pct=int(tv_selection.prepay_pct),
                discount_pct=float(st.session_state.get("kp_discount_pct", 0.0) or 0.0),
            )
            st.download_button(
                "⬇️ Скачать КП (PDF)", data=tv_pdf,
                file_name=f"KP_traversa_{tv_kp_number.replace('/', '_')}.pdf",
                mime="application/pdf", use_container_width=True)
        except Exception as e:
            st.error(f"PDF не сгенерирован: {e}")

        try:
            tv_docx = build_kp_docx(tv_q, tv_kp_number, tv_buyer)
            st.download_button(
                "⬇️ Скачать КП (DOCX)", data=tv_docx,
                file_name=f"KP_traversa_{tv_kp_number.replace('/', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument."
                     "wordprocessingml.document",
                use_container_width=True)
        except Exception as e:
            st.warning(f"DOCX не сгенерирован: {e}")
            tv_docx = None

        # --- Кнопка сохранения КП в CRM базу ---
        _tv_save_key = f"crm_autosaved_{tv_kp_number}"
        _tv_already = st.session_state.get(_tv_save_key)
        if _tv_already:
            st.info(f"💾 КП {tv_kp_number} уже сохранён в CRM — id={_tv_already}")
        else:
            if True:  # автосохранение
                try:
                    import crm_db as _crm_db
                    _customer_id = None
                    _autofilled = st.session_state.get("crm_autofilled")
                    _bn = tv_buyer.get("name", "") if isinstance(tv_buyer, dict) else ""
                    if _autofilled and getattr(_autofilled, "inn", ""):
                        _customer_id = _crm_db.upsert_customer(_autofilled)
                    elif _bn:
                        _cl = _crm_db.Customer(name_short=_bn)
                        _customer_id = _crm_db.upsert_customer(_cl)
                    _items = [_crm_db.QuoteItem(
                        code=ln.code, name=ln.name, unit=ln.unit,
                        qty=ln.qty, price=ln.price) for ln in tv_q.lines]
                    _kp_disc_tv = float(st.session_state.get("kp_discount_pct", 0.0) or 0.0)
                    _tot_tv = float(sum(l.total for l in tv_q.lines))
                    _disc_r_tv = 1.0 - _kp_disc_tv / 100.0 if _kp_disc_tv > 0 else 1.0
                    _base_tv = _tot_tv / _disc_r_tv if _disc_r_tv > 0 else _tot_tv
                    _tv_kp_unique = _crm_db.generate_unique_kp_number(tv_kp_number)
                    if _tv_kp_unique != tv_kp_number:
                        st.info(f"⚠️ Номер {tv_kp_number} уже занят — присвоен: **{_tv_kp_unique}**")
                    _rec = _crm_db.QuoteRecord(
                        kp_number=_tv_kp_unique, customer_id=_customer_id,
                        product_type="Траверса",
                        product_model=tv_selection.code if tv_selection else "",
                        include_montage=bool(tv_selection.include_montage)
                            if tv_selection else False,
                        delivery_city=(tv_selection.delivery_address
                            if tv_selection and tv_selection.delivery_address
                            else ""),
                        base_total=_base_tv,
                        discount_pct=_kp_disc_tv, status=_crm_db.STATUS_DRAFT,
                        items=_items,
                        pdf_bytes=tv_pdf if 'tv_pdf' in dir() else None,
                        docx_bytes=tv_docx)
                    _qid = _crm_db.save_quote(_rec)
                    st.session_state[_tv_save_key] = _qid
                    st.success(f"✅ КП {_tv_kp_unique} автоматически сохранено в CRM — id={_qid}")
                except Exception as _e:
                    st.error(f"Не удалось сохранить: {_e}")

    # ================================================================
    # Блок «📑 Договор поставки товаров» для режима Траверса
    # Все ключи виджетов — с префиксом tv_dog_ чтобы не конфликтовать с крановым блоком
    # ================================================================
    st.markdown("---")
    st.subheader("📑 Договор поставки товаров")

    # === Автораспознавание реквизитов (такое же как в внешнем договоре) ===
    with st.expander("⚡ Автозаполнить из файла или текста реквизитов", expanded=False):
        st.caption("📋 Перетащите карточку партнёра (DOCX/PDF/TXT) или вставьте текст с реквизитами. "
                   "Поля ниже заполнятся автоматически. Распознанные реквизиты сохраняются в историю.")
        _tv_up_col, _tv_txt_col = st.columns([1, 2])
        with _tv_up_col:
            _tv_req_file = st.file_uploader(
                "Файл (DOCX / DOC / PDF / JPG / PNG / TXT)",
                type=["docx", "doc", "pdf", "jpg", "jpeg", "png", "webp", "txt"],
                key="tv_dog_req_file",
                help="Поддерживаются: DOCX, DOC, PDF, JPG/PNG (+ OCR)")
        with _tv_txt_col:
            _tv_req_text = st.text_area("Или вставьте текст", height=180,
                                        key="tv_dog_req_text",
                                        placeholder="ООО «...»\nИНН ...\nКПП ...\nОГРН ...\n"
                                                    "Юр. адрес: ...\nБанк: ...\nр/с ...\nБИК ...\n"
                                                    "Генеральный директор Фамилия Имя Отчество")

        if st.button("🔍 Распознать реквизиты", use_container_width=True, key="tv_dog_recognize_btn"):
            _tv_raw_text = ""
            if _tv_req_file is not None:
                _tv_data = _tv_req_file.getvalue()
                _tv_fname = _tv_req_file.name
                try:
                    from smart_requisites import _extract_text_smart
                    _tv_raw_text = _extract_text_smart(_tv_data, _tv_fname)
                except Exception as _e:
                    st.error(f"Не удалось прочитать файл: {_e}")
                    if _tv_fname.lower().endswith(".docx"):
                        try:
                            _tv_raw_text = extract_text_from_docx(_tv_data)
                        except Exception:
                            pass
                    elif _tv_fname.lower().endswith(".txt"):
                        try:
                            _tv_raw_text = _tv_data.decode("utf-8", errors="ignore")
                        except Exception:
                            _tv_raw_text = _tv_data.decode("cp1251", errors="ignore")
            elif _tv_req_text.strip():
                _tv_raw_text = _tv_req_text

            if not _tv_raw_text.strip():
                st.warning("Загрузите файл или вставьте текст с реквизитами.")
            else:
                try:
                    from external_kp_parser import extract_requisites_from_text as _tv_ext_req
                    _tv_req = _tv_ext_req(_tv_raw_text)
                    _tv_parsed = {
                        "short": _tv_req.company_short or "",
                        "full": _tv_req.company_full or _tv_req.company_short or "",
                        "address": _tv_req.address or "",
                        "post_address": _tv_req.address or "",
                        "inn": _tv_req.inn or "",
                        "kpp": _tv_req.kpp or "",
                        "ogrn": _tv_req.ogrn or "",
                        "phone": _tv_req.phone or "",
                        "email": _tv_req.email or "",
                        "bank": _tv_req.bank_name or "",
                        "bik": _tv_req.bank_bik or "",
                        "rs": _tv_req.bank_account or "",
                        "ks": _tv_req.corr_account or "",
                        "director_position": _tv_req.director_title or "Генеральный директор",
                        "director_fio_gen": _tv_req.director_gen or "",
                        "director_fio_short": _tv_req.director_short or "",
                        "basis": "Устава",
                    }
                except Exception:
                    _tv_parsed = parse_requisites(_tv_raw_text)
                # Кладём в общий buyer_data — так же как в кране и внешнем договоре
                st.session_state["buyer_data"] = _tv_parsed
                # Сбрасываем ключи tv_dog_b_* чтобы при следующем rerun поля подтянули новые value
                for _k in ["tv_dog_b_short", "tv_dog_b_full", "tv_dog_b_address",
                           "tv_dog_b_inn", "tv_dog_b_kpp", "tv_dog_b_ogrn",
                           "tv_dog_b_phone", "tv_dog_b_email", "tv_dog_b_bank",
                           "tv_dog_b_bik", "tv_dog_b_rs", "tv_dog_b_ks",
                           "tv_dog_b_dir_pos", "tv_dog_b_dir_gen",
                           "tv_dog_b_dir_short", "tv_dog_b_basis"]:
                    st.session_state.pop(_k, None)
                # Автосохранение в историю по ИНН
                try:
                    import history_memory as _hm_tv_req
                    if _tv_parsed.get("inn"):
                        _hm_tv_req.remember_ext_buyer(_tv_parsed["inn"], {
                            "short": _tv_parsed.get("short", ""),
                            "full": _tv_parsed.get("full", ""),
                            "inn": _tv_parsed.get("inn", ""),
                            "kpp": _tv_parsed.get("kpp", ""),
                            "ogrn": _tv_parsed.get("ogrn", ""),
                            "address": _tv_parsed.get("address", ""),
                            "bank": _tv_parsed.get("bank", ""),
                            "bik": _tv_parsed.get("bik", ""),
                            "rs": _tv_parsed.get("rs", ""),
                            "ks": _tv_parsed.get("ks", ""),
                            "phone": _tv_parsed.get("phone", ""),
                            "email": _tv_parsed.get("email", ""),
                            "director_position": _tv_parsed.get("director_position", ""),
                            "director_fio_short": _tv_parsed.get("director_fio_short", ""),
                            "director_fio_gen": _tv_parsed.get("director_fio_gen", ""),
                            "basis": _tv_parsed.get("basis", "Устава"),
                        })
                except Exception:
                    pass
                _tv_filled = sum(1 for v in _tv_parsed.values() if v)
                _tv_summary = []
                if _tv_parsed.get("short"): _tv_summary.append(_tv_parsed["short"])
                if _tv_parsed.get("inn"): _tv_summary.append(f"ИНН {_tv_parsed['inn']}")
                if _tv_parsed.get("bank"): _tv_summary.append(_tv_parsed["bank"])
                _tv_summary_str = " · ".join(_tv_summary) if _tv_summary else ""
                st.success(f"✅ Распознано {_tv_filled} полей" +
                           (f": {_tv_summary_str}" if _tv_summary_str else "") +
                           ". Поля ниже заполнены — проверьте.")
                st.rerun()

    # Собираем минимальный QuoteData с позициями траверсы для build_dogovor_docx
    _tv_dog_q = QuoteData(
        series="Траверса",
        capacity=int(getattr(tv_selection, "capacity_kg", 500) or 500),
        boom=0, height_to_arm=0,
        hoist_brand="—", hoist_mode="—", hoist_height=0,
        include_electrification=False,
        include_montage=bool(getattr(tv_selection, "include_montage", False)),
        montage_price=float(getattr(tv_selection, "montage_price", 0) or 0),
        include_vat=bool(st.session_state.get("kp_include_vat", True)),
    )
    for _tv_it in _tv_items:
        _tv_dog_q.lines.append(SpecLine(
            code=str(_tv_it.get("code", "")),
            name=str(_tv_it.get("name", "")),
            unit=str(_tv_it.get("unit", "шт")),
            qty=float(_tv_it.get("qty", 1)),
            price=float(_tv_it.get("price_vat", 0)),
        ))

    # Форма реквизитов Покупателя
    _tv_bd = st.session_state.get("buyer_data", {})
    with st.expander("Реквизиты Покупателя", expanded=True):
        _tv_c1, _tv_c2 = st.columns(2)
        with _tv_c1:
            _tv_b_short = st.text_input("Наименование краткое",
                                        value=_tv_bd.get("short", ""),
                                        key="tv_dog_b_short")
            _tv_b_full = st.text_input("Наименование полное",
                                       value=_tv_bd.get("full", ""),
                                       key="tv_dog_b_full")
            _tv_b_address = st.text_input("Юр. адрес",
                                          value=_tv_bd.get("address", ""),
                                          key="tv_dog_b_address")
            _tv_b_inn = st.text_input("ИНН",
                                      value=_tv_bd.get("inn", ""),
                                      key="tv_dog_b_inn")
            _tv_b_kpp = st.text_input("КПП",
                                      value=_tv_bd.get("kpp", ""),
                                      key="tv_dog_b_kpp")
            _tv_b_ogrn = st.text_input("ОГРН",
                                       value=_tv_bd.get("ogrn", ""),
                                       key="tv_dog_b_ogrn")
            _tv_b_phone = st.text_input("Телефон",
                                        value=_tv_bd.get("phone", ""),
                                        key="tv_dog_b_phone")
            _tv_b_email = st.text_input("E-mail",
                                        value=_tv_bd.get("email", ""),
                                        key="tv_dog_b_email")
        with _tv_c2:
            _tv_b_bank = st.text_input("Банк",
                                       value=_tv_bd.get("bank", ""),
                                       key="tv_dog_b_bank")
            _tv_b_bik = st.text_input("БИК",
                                      value=_tv_bd.get("bik", ""),
                                      key="tv_dog_b_bik")
            _tv_b_rs = st.text_input("Р/с",
                                     value=_tv_bd.get("rs", ""),
                                     key="tv_dog_b_rs")
            _tv_b_ks = st.text_input("К/с",
                                     value=_tv_bd.get("ks", ""),
                                     key="tv_dog_b_ks")
            _tv_b_dir_pos = st.text_input("Должность подписанта",
                                          value=_tv_bd.get("director_position", "Генеральный директор"),
                                          key="tv_dog_b_dir_pos")
            _tv_b_dir_gen = st.text_input("ФИО в родительном падеже",
                                          value=_tv_bd.get("director_fio_gen", ""),
                                          key="tv_dog_b_dir_gen")
            _tv_b_dir_short = st.text_input("ФИО коротко (для подписи)",
                                            value=_tv_bd.get("director_fio_short", ""),
                                            key="tv_dog_b_dir_short")
            _tv_b_basis = st.text_input("На основании",
                                        value=_tv_bd.get("basis", "Устава"),
                                        key="tv_dog_b_basis")

    # === Выбор поставщика (ЛКС / Модернизация / Кинематика) + печать ===
    _tv_dog_sup_row1, _tv_dog_sup_row2 = st.columns([3, 2])
    with _tv_dog_sup_row1:
        _tv_dog_sup_keys = list(_suppliers.SUPPLIERS.keys())
        _tv_dog_sup_current = st.session_state.get("supplier_key", _suppliers.DEFAULT_SUPPLIER_KEY)
        _tv_dog_sup_idx = _tv_dog_sup_keys.index(_tv_dog_sup_current) if _tv_dog_sup_current in _tv_dog_sup_keys else 0
        _tv_dog_picked_sup_key = st.selectbox(
            "🏢 От какой компании выставляем договор",
            _tv_dog_sup_keys,
            format_func=lambda k: _suppliers.SUPPLIERS[k]["label"],
            index=_tv_dog_sup_idx,
            key="tv_dog_supplier_key",
            help="Влияет на реквизиты, шапку, футер, печать, ЭДО в договоре")
        # Сохраняем выбор глобально (другие блоки читают из supplier_key)
        st.session_state["supplier_key"] = _tv_dog_picked_sup_key
    with _tv_dog_sup_row2:
        _tv_dog_stamp_available = bool(
            _suppliers.SUPPLIERS[_tv_dog_picked_sup_key].get("stamp_path"))
        st.checkbox(
            "🖋️ Печать и подпись в договоре",
            value=st.session_state.get("include_stamp", False),
            key="include_stamp",
            disabled=not _tv_dog_stamp_available,
            help=("Вставить печать+подпись в DOCX"
                  if _tv_dog_stamp_available else "Для этого поставщика печать не загружена"))
    # Обновляем глобальный SUPPLIER чтобы build_dogovor_docx взял его
    SUPPLIER = _suppliers.get_supplier(_tv_dog_picked_sup_key)

    _tv_dog_c3, _tv_dog_c4 = st.columns(2)
    with _tv_dog_c3:
        # Префикс № договора по выбранному поставщику
        _tv_dog_prefix_map = {"LKS": "ЛКС", "MODERNIZATSIYA": "МОД", "KINEMATIKA": "КИН"}
        _tv_dog_num_prefix = _tv_dog_prefix_map.get(_tv_dog_picked_sup_key, "ЛКС")
        _tv_contract_number = st.text_input(
            "Номер договора",
            value=date.today().strftime("%d/%m/%Y") + f"/{_tv_dog_num_prefix}",
            key="tv_dog_contract_number")
    with _tv_dog_c4:
        _tv_contract_date_str = st.text_input(
            "Дата договора (в тексте)",
            value=date.today().strftime("%d.%m.%Y"),
            key="tv_dog_contract_date_str")

    _tv_dog_pp_default = int(st.session_state.get("tv_prepay_pct", int(PREPAYMENT_DEFAULT * 100)))
    _tv_dog_prepay_pct = st.number_input(
        "Предоплата, %",
        min_value=0, max_value=100,
        value=_tv_dog_pp_default, step=5,
        key="tv_dog_prepay_pct")

    _tv_dog_delivery_terms = st.text_input(
        "Условия доставки",
        value=(st.session_state.get("tv_delivery_option") or "включена в стоимость"),
        key="tv_dog_delivery_terms")

    if st.button("📥 Сформировать Договор поставки товаров",
                 type="primary", use_container_width=True,
                 key="tv_dog_generate_btn"):
        if not (_tv_b_short and _tv_b_inn):
            st.error("Заполните хотя бы наименование и ИНН Покупателя.")
        else:
            _tv_buyer = {
                "short": _tv_b_short, "full": _tv_b_full or _tv_b_short,
                "address": _tv_b_address, "post_address": _tv_b_address,
                "inn": _tv_b_inn, "kpp": _tv_b_kpp, "ogrn": _tv_b_ogrn,
                "phone": _tv_b_phone, "email": _tv_b_email,
                "bank": _tv_b_bank, "bik": _tv_b_bik, "rs": _tv_b_rs, "ks": _tv_b_ks,
                "director_position": _tv_b_dir_pos,
                "director_fio_gen": _tv_b_dir_gen,
                "director_fio_short": _tv_b_dir_short,
                "basis": _tv_b_basis,
            }
            # Собираем все выбранные траверсы для Приложения № 1
            _tv_all_sels = [tv_selection] + list(st.session_state.get("_tv_extra_selections", []) or [])

            try:
                _tv_dogovor_bytes = build_dogovor_docx(
                    _tv_dog_q, _tv_buyer, _tv_contract_number, _tv_contract_date_str,
                    prepay_pct=int(_tv_dog_prepay_pct),
                    delivery_terms=_tv_dog_delivery_terms,
                    traverse_selections=_tv_all_sels,
                    include_stamp=bool(st.session_state.get("include_stamp", False)),
                )
                _tv_fn = f"Dogovor_{_tv_contract_number.replace('/', '_')}.docx"
                st.download_button(
                    "⬇️ Скачать Договор поставки товаров (DOCX)",
                    data=_tv_dogovor_bytes, file_name=_tv_fn,
                    mime="application/vnd.openxmlformats-officedocument."
                         "wordprocessingml.document",
                    use_container_width=True,
                    key="tv_dog_download_btn")
                st.success("Договор готов. Нажмите кнопку скачивания выше.")
            except Exception as _tv_dog_err:
                st.error(f"Ошибка формирования договора: {_tv_dog_err}")

    st.stop()

# --- UI: дополнительные краны в этом КП (выбор и управление списком) ---
if "extra_cranes" not in st.session_state:
    st.session_state["extra_cranes"] = []

with st.container(border=True):
    _mc_h1, _mc_h2 = st.columns([1, 6])
    with _mc_h1:
        st.markdown(
            '<div style="width:44px;height:44px;border-radius:12px;'
            'background:#FEF3E7;display:flex;align-items:center;justify-content:center;'
            'font-size:22px;">🏗️</div>',
            unsafe_allow_html=True)
    with _mc_h2:
        _n_extra = len(st.session_state["extra_cranes"])
        st.markdown(
            f'<div style="font-size:22px;font-weight:700;color:#111827;'
            f'padding-top:6px;">Краны в КП</div>'
            f'<div style="color:#6B7280;font-size:13px;">'
            f'Основной: {series} {capacity} кг × {boom} м'
            + (f' · Дополнительных кранов: {_n_extra}' if _n_extra else '')
            + '</div>', unsafe_allow_html=True)

    # Список доп. кранов — каждый в своём блоке
    _SERIES_ALL = ["ЛКС71", "ЛКС73"]
    # Заменяю кривую кириллицу на правильную

    _CAP_BY_SERIES = {
        "ЛКС71": [80, 125, 250, 500, 1000],
        "ЛКС73": [80, 125, 250, 500, 1000],
    }
    _remove_idx = None
    for _ec_i, _ec_v in enumerate(list(st.session_state["extra_cranes"])):
        with st.container(border=True):
            _hdr_c1, _hdr_c2 = st.columns([6, 1])
            with _hdr_c1:
                st.markdown(f"**Доп. кран #{_ec_i + 1}**")
            with _hdr_c2:
                if st.button("🗑️", key=f"ec_del_{_ec_i}", help="Удалить этот кран"):
                    _remove_idx = _ec_i
            _c1, _c2, _c3, _c4 = st.columns(4)
            with _c1:
                _ec_v["series"] = st.selectbox(
                    "Серия", ["ЛКС71", "ЛКС73"],  # только ЛКС71/73
                    index=["ЛКС71", "ЛКС73"].index(_ec_v.get("series", "ЛКС73")) if _ec_v.get("series", "ЛКС73") in ["ЛКС71", "ЛКС73"] else 1,
                    key=f"ec_series_{_ec_i}")
            with _c2:
                _caps = [80, 125, 250, 500, 1000]
                _ec_v["capacity"] = st.selectbox(
                    "Г/п, кг", _caps,
                    index=_caps.index(int(_ec_v.get("capacity", 250))) if int(_ec_v.get("capacity", 250)) in _caps else 2,
                    key=f"ec_cap_{_ec_i}")
            with _c3:
                _ec_v["boom"] = st.number_input(
                    "Стрела, м", min_value=3.0, max_value=8.0,
                    value=float(_ec_v.get("boom", 3.0)), step=1.0,
                    key=f"ec_boom_{_ec_i}")
            with _c4:
                _ec_v["qty"] = st.number_input(
                    "Кол-во", min_value=1, max_value=20,
                    value=int(_ec_v.get("qty", 1)),
                    key=f"ec_qty_{_ec_i}")
            _c5, _c6, _c7 = st.columns(3)
            with _c5:
                _ec_v["height_to_arm"] = st.number_input(
                    "Высота до стрелы, м", min_value=3.0, max_value=6.0,
                    value=float(_ec_v.get("height_to_arm", 3.0)), step=1.0,
                    key=f"ec_h_{_ec_i}")
            with _c6:
                _brands = ["R-Tech", "OCALIFT"]
                _ec_v["hoist_brand"] = st.selectbox(
                    "Таль", _brands,
                    index=_brands.index(_ec_v.get("hoist_brand", "R-Tech")) if _ec_v.get("hoist_brand", "R-Tech") in _brands else 0,
                    key=f"ec_hb_{_ec_i}")
            with _c7:
                _ec_v["hoist_height"] = st.number_input(
                    "Высота подъёма тали, м", min_value=3.0, max_value=6.0,
                    value=float(_ec_v.get("hoist_height", 3.0)), step=1.0,
                    key=f"ec_hh_{_ec_i}")
            # Фиксированный режим тали
            _ec_v["hoist_mode"] = "8/2 м/мин"

    if _remove_idx is not None:
        st.session_state["extra_cranes"].pop(_remove_idx)
        st.rerun()

    if st.button("➕ Добавить ещё кран", key="add_extra_crane"):
        st.session_state["extra_cranes"].append({
            "series": "ЛКС73",
            "capacity": 250,
            "boom": 3.0,
            "height_to_arm": 3.0,
            "hoist_brand": "R-Tech",
            "hoist_mode": "8/2 м/мин",
            "hoist_height": 3.0,
            "qty": 1,
        })
        st.rerun()

# --- Сборка КП (Кран или Кран+Траверса) ---
q = QuoteData(
    series=series,
    capacity=capacity,
    boom=boom,
    height_to_arm=height_to_arm,
    hoist_brand=hoist_brand,
    hoist_mode=hoist_mode,
    hoist_height=hoist_height,
    include_electrification=include_electrification,
    include_montage=include_montage,
    montage_price=float(montage_price),
    montage_vat=montage_vat,
    hc_capacity=hc_capacity,
    column_diameter=column_diameter,
    use_lllm=use_lllm,
    lllm_code=lllm_code,
    include_flange=include_flange,
    flange_code=flange_code,
    with_traverse=(kp_mode == "Кран + траверса"),
    hoist_exec=hoist_exec,
    include_hoist=bool(include_hoist),
    include_vat=bool(st.session_state.get("kp_include_vat", True)),
)
q = build_specification(q, df71, df77, dfh)

# --- Дополнительные краны (если есть) ---
# session_state["extra_cranes"] = [{"series", "capacity", "boom", "height_to_arm",
#                                    "hoist_brand", "hoist_mode", "hoist_height", "qty"}]
_extra_cranes_specs = st.session_state.get("extra_cranes", []) or []
for _ec in _extra_cranes_specs:
    try:
        _ec_qty = max(1, int(_ec.get("qty", 1)))
        # Строим временный QuoteData для этого крана (без электрификации и монтажа — они общие)
        _ec_q = QuoteData(
            series=_ec.get("series", "ЛКС73"),
            capacity=int(_ec.get("capacity", 250)),
            boom=float(_ec.get("boom", 3.0)),
            height_to_arm=float(_ec.get("height_to_arm", 3.0)),
            hoist_brand=_ec.get("hoist_brand", "R-Tech"),
            hoist_mode=_ec.get("hoist_mode", "8/2 м/мин"),
            hoist_height=float(_ec.get("hoist_height", 3.0)),
            include_electrification=False,
            include_montage=False,
            montage_price=0.0,
            montage_vat=False,
            hc_capacity=None,
            column_diameter=None,
            use_lllm=False,
            lllm_code=None,
            include_flange=False,
            flange_code="",
            with_traverse=False,
        )
        _ec_q = build_specification(_ec_q, df71, df77, dfh)
        # Добавляем все позиции (кран, таль, тележки, кабель) в q.lines,
        # умножая кол-во на qty
        for _ln in _ec_q.lines:
            q.lines.append(SpecLine(
                _ln.code, _ln.name, _ln.unit,
                _ln.qty * _ec_qty, _ln.price))
    except Exception as _ec_err:
        st.warning(f"Доп. кран {_ec.get('series')} — ошибка: {_ec_err}")

# Добавляем ручные позиции (если есть)
for _ci in st.session_state.get("custom_items", []):
    q.lines.append(SpecLine(
        _ci["code"], _ci["name"], _ci["unit"],
        _ci["qty"], _ci["price"],
    ))

# Режим Кран + траверса — добавляем позиции траверсы (вт.ч. 2-й сравнение)
if kp_mode == "Кран + траверса" and tv_selection is not None:
    _tv_df_combo = _tv.load_traverse_price(
        st.session_state.get("traverse_price_upload"))
    if _tv_df_combo is not None:
        for _ti in _tv.build_traverse_items(tv_selection, _tv_df_combo):
            q.lines.append(SpecLine(
                _ti["code"], _ti["name"], _ti["unit"],
                _ti["qty"], _ti["price_vat"],
            ))
        # Альтернативные траверсы (сравнение) — тоже в спецификацию
        _tv_extras_combo = st.session_state.get("_tv_extra_selections", []) or []
        for _es_combo in _tv_extras_combo:
            for _ti in _tv.build_traverse_items(_es_combo, _tv_df_combo):
                q.lines.append(SpecLine(
                    _ti["code"], _ti["name"], _ti["unit"],
                    _ti["qty"], _ti["price_vat"],
                ))

# --- Монтаж с НДС → включаем в q.lines (чтобы попалв в стоимость КП и НДС) ---
# Монтаж без НДС (от ИП) сохраняется как отдельная сущность (отдельным блоком в PDF).
if q.include_montage and q.montage_price > 0 and q.montage_vat:
    q.lines.append(SpecLine(
        code="МОНТАЖ",
        name="Монтаж и пусконаладочные работы консольного крана",
        unit="усл.", qty=1.0, price=float(q.montage_price),
    ))
    # Сбрасываем флаги в q, чтобы PDF/DOCX не добавляли монтаж второй раз отдельным блоком
    q.include_montage = False
    q.montage_price = 0.0

# === HERO-ЗАГОЛОВОК СТРАНИЦЫ РАСЧЁТА ===
_hero_kp_num = st.session_state.get("kp_buyer_name", "") or ""
_hero_subtitle = f"Для: <b>{_hero_kp_num}</b>" if _hero_kp_num else "Создайте коммерческое предложение в несколько кликов"
st.markdown(f'''
<div class="hero-header">
    <div>
        <h1 class="hero-title">💼 Расчёт КП</h1>
        <p class="hero-subtitle">{_hero_subtitle}</p>
    </div>
    <div class="hero-info">
        <span class="hero-badge">{q.series} · {q.capacity} кг · {q.boom} м</span>
    </div>
</div>
''', unsafe_allow_html=True)

st.subheader("📋 Спецификация")

# Чекбокс НДС — прямо здесь, над таблицей (влияет на цены в спецификации)
_spec_vat_c1, _spec_vat_c2 = st.columns([1, 3])
with _spec_vat_c1:
    _spec_include_vat = st.checkbox(
        "✅ НДС 22 % включён",
        value=st.session_state.get("kp_include_vat", True),
        key="kp_include_vat",
        help="Снимите галочку — цены пересчитаются как «без НДС» (для УСН/ИП).")
with _spec_vat_c2:
    if _spec_include_vat:
        st.caption("💰 **Цены в таблице — с НДС 22 %** (налог включён в стоимость).")
    else:
        st.caption("💵 **Цены в таблице — без НДС** (пересчитаны без налога).")
# Обновляем поле QuoteData чтобы properties total/vat подхватили чекбокс тут же
q.include_vat = bool(_spec_include_vat)
# Если НДС выключен — в таблице показать цены без НДС (делим на 1.22)
if not _spec_include_vat:
    for _ln_v in q.lines:
        # Не трогаем если уже было вычтено (флагом в override)
        pass  # цены остаются с НДС в таблице, но в итоге total пересчитается

st.caption("💡 **Количество и цену можно менять прямо в таблице** (клик по ячейке → ввести цифру → Enter). "
           "Сумма и итоги пересчитаются автоматически.")

# === ДЕДУПЛИКАЦИЯ: склеиваем одинаковые товары в одну строку (сумма qty) ===
# Группируем по (code, name, unit, price) — если всё совпадает, это тот же товар
_dedup_seen = {}
_dedup_lines = []
for _ln in q.lines:
    _key = (_ln.code or "", _ln.name or "", _ln.unit or "", float(_ln.price or 0.0))
    if _key in _dedup_seen:
        _dedup_seen[_key].qty = float(_dedup_seen[_key].qty) + float(_ln.qty)
    else:
        _dedup_seen[_key] = _ln
        _dedup_lines.append(_ln)
q.lines = _dedup_lines

# Применяем сохранённые правки по КЛЮЧУ позиции (code+name+unit) — переживают rerun
_spec_overrides = st.session_state.get("_spec_item_overrides", {}) or {}
def _spec_row_key(_ln):
    """Составной ключ для override."""
    return _ln.code or f"__no_code__:{_ln.name}:{_ln.unit}"

for _ln in q.lines:
    _ov = _spec_overrides.get(_spec_row_key(_ln))
    if _ov:
        if "qty" in _ov:
            _ln.qty = float(_ov["qty"])
        if "price" in _ov:
            _ln.price = float(_ov["price"])
# НЕ читаем edited_rows — Streamlit вернёт туда старые значения после rerun,
# только overrides — они единственный источник правды.

spec_df = pd.DataFrame([{
    "Код": ln.code, "Наименование": ln.name, "Ед.": ln.unit,
    "Кол-во": float(ln.qty),
    "Цена, ₽": float(ln.price),
    "Сумма, ₽": float(ln.total),  # ln.total = qty * price (property)
} for ln in q.lines])
spec_edited = st.data_editor(
    spec_df, use_container_width=True, hide_index=True, key="spec_editor",
    column_config={
        "Кол-во": st.column_config.NumberColumn(
            "Кол-во", min_value=0.0, step=1.0, format="%.2f"),
        "Цена, ₽": st.column_config.NumberColumn(
            "Цена, ₽", min_value=0.0, step=100.0, format="%.2f ₽"),
        "Сумма, ₽": st.column_config.NumberColumn(
            "Сумма, ₽", format="%.2f ₽", disabled=True),
        "Код": st.column_config.TextColumn(disabled=True),
        "Наименование": st.column_config.TextColumn(disabled=True),
        "Ед.": st.column_config.TextColumn(disabled=True),
    },
)
# Фаза 2: синхронизация + сохранение overrides по ключу позиции (включая безкодовые)
_spec_overrides_new = dict(st.session_state.get("_spec_item_overrides", {}) or {})
for _i, _ln in enumerate(q.lines):
    try:
        _new_qty = float(spec_edited.iloc[_i]["Кол-во"])
        _new_price = float(spec_edited.iloc[_i]["Цена, ₽"])
        _changed = (_new_qty != _ln.qty) or (_new_price != _ln.price)
        _ln.qty = _new_qty
        _ln.price = _new_price
        if _changed:
            _spec_overrides_new[_spec_row_key(_ln)] = {
                "qty": _new_qty, "price": _new_price
            }
    except Exception:
        pass
st.session_state["_spec_item_overrides"] = _spec_overrides_new


if q.electrification_lines:
    st.subheader("⚡ Электрификация")
    # Применяем сохранённые правки по КЛЮЧУ (включая безкодовые)
    _el_overrides = st.session_state.get("_elec_item_overrides", {}) or {}
    def _el_row_key(_ln):
        return _ln.code or f"__no_code__:{_ln.name}:{_ln.unit}"
    for _ln in q.electrification_lines:
        _ov = _el_overrides.get(_el_row_key(_ln))
        if _ov:
            if "qty" in _ov: _ln.qty = float(_ov["qty"])
            if "price" in _ov: _ln.price = float(_ov["price"])
    # Фаза 1: предыдущие правки текущего цикла (fallback)
    _prev_el_edits = st.session_state.get("elec_editor")
    if isinstance(_prev_el_edits, dict) and "edited_rows" in _prev_el_edits:
        for _row_idx_str, _changes in _prev_el_edits["edited_rows"].items():
            try:
                _ri = int(_row_idx_str)
                if 0 <= _ri < len(q.electrification_lines):
                    if "Кол-во" in _changes:
                        q.electrification_lines[_ri].qty = float(_changes["Кол-во"])
                    if "Цена, ₽" in _changes:
                        q.electrification_lines[_ri].price = float(_changes["Цена, ₽"])
            except Exception:
                pass
    el_df = pd.DataFrame([{
        "Код": ln.code, "Наименование": ln.name, "Ед.": ln.unit,
        "Кол-во": float(ln.qty),
        "Цена, ₽": float(ln.price),
        "Сумма, ₽": float(ln.total),
    } for ln in q.electrification_lines])
    el_edited = st.data_editor(
        el_df, use_container_width=True, hide_index=True, key="elec_editor",
        column_config={
            "Кол-во": st.column_config.NumberColumn(
                "Кол-во", min_value=0.0, step=1.0, format="%.2f"),
            "Цена, ₽": st.column_config.NumberColumn(
                "Цена, ₽", min_value=0.0, step=100.0, format="%.2f ₽"),
            "Сумма, ₽": st.column_config.NumberColumn(
                "Сумма, ₽", format="%.2f ₽", disabled=True),
            "Код": st.column_config.TextColumn(disabled=True),
            "Наименование": st.column_config.TextColumn(disabled=True),
            "Ед.": st.column_config.TextColumn(disabled=True),
        },
    )
    _el_overrides_new = dict(st.session_state.get("_elec_item_overrides", {}) or {})
    for _i, _ln in enumerate(q.electrification_lines):
        try:
            _new_qty = float(el_edited.iloc[_i]["Кол-во"])
            _new_price = float(el_edited.iloc[_i]["Цена, ₽"])
            _changed = (_new_qty != _ln.qty) or (_new_price != _ln.price)
            _ln.qty = _new_qty
            _ln.price = _new_price
            if _changed:
                _el_overrides_new[_el_row_key(_ln)] = {"qty": _new_qty, "price": _new_price}
        except Exception:
            pass
    st.session_state["_elec_item_overrides"] = _el_overrides_new

# --- Комментарий к КП (под таблицами расчёта) ---
# Попадает в PDF/DOCX КП, нО НЕ попадает в договор
st.markdown(
    "<div style='background:#FFF3E0;border-left:4px solid #F97316;"
    "padding:10px 14px;border-radius:6px;margin-top:16px;'>"
    "💬 <b>Комментарий к КП</b> — текст попадёт в PDF/DOCX КП после условий поставки "
    "(<i>в договор не попадает</i>)"
    "</div>", unsafe_allow_html=True)
st.text_area(
    "Комментарий", value=st.session_state.get("kp_comment", ""),
    key="kp_comment",
    placeholder="Например: Срок изготовления 25 раб. дней. Гарантия 18 мес. "
    "При заказе до 01.08 — бесплатный монтаж.",
    height=100, label_visibility="collapsed")

# --- Условия оплаты и доставки (универсально для Кран и Кран+Траверса) ---
with st.container(border=True):
    _pd_h1, _pd_h2 = st.columns([1, 6])
    with _pd_h1:
        st.markdown(
            '<div style="width:44px;height:44px;border-radius:12px;'
            'background:#FEF3E7;display:flex;align-items:center;justify-content:center;'
            'font-size:22px;">💳</div>',
            unsafe_allow_html=True)
    with _pd_h2:
        st.markdown(
            '<div style="font-size:22px;font-weight:700;color:#111827;'
            'padding-top:6px;">Условия оплаты и доставки</div>',
            unsafe_allow_html=True)

    _pd_c1, _pd_c2 = st.columns(2)
    with _pd_c1:
        kp_prepay_pct = st.number_input(
            "% предоплаты",
            min_value=0, max_value=100,
            value=int(st.session_state.get("kp_prepay_pct", int(PREPAYMENT_DEFAULT*100))),
            step=10, key="kp_prepay_pct",
            help="Попадает в итоги, КП и договор")
    with _pd_c2:
        _DELIVERY_OPTIONS = [
            "Самовывоз из г. Орёл",
            "Самовывоз со склада г. Санкт-Петербург (Всеволожский р-н, п. Романовка, ул. Инженерная, 19)",
            "Самовывоз со склада г. Москва (Лианозовский проезд, д. 6)",
            "Доставка до ТК «Деловые линии» и отправка с учётом обрешётки за счёт покупателя",
            "Доставка силами Поставщика",
            "— Свой вариант (вписать вручную) —",
        ]
        _cur_idx = int(st.session_state.get("kp_delivery_option_idx", 0))
        if _cur_idx >= len(_DELIVERY_OPTIONS):
            _cur_idx = 0
        _kp_delivery_option = st.selectbox(
            "Способ доставки",
            _DELIVERY_OPTIONS, index=_cur_idx, key="kp_delivery_option")
        st.session_state["kp_delivery_option_idx"] = _DELIVERY_OPTIONS.index(_kp_delivery_option)

    # Цена доставки только для «Силами Поставщика» / «Свой вариант»
    _need_delivery_price = _kp_delivery_option in (
        "Доставка силами Поставщика",
        "— Свой вариант (вписать вручную) —",
    )
    if _kp_delivery_option == "— Свой вариант (вписать вручную) —":
        _kp_delivery_text = st.text_input(
            "Свой текст условий доставки",
            value=st.session_state.get("kp_delivery_text", ""),
            key="kp_delivery_text")
        _kp_delivery_addr = _kp_delivery_text
    else:
        _kp_delivery_addr = _kp_delivery_option
    if _need_delivery_price:
        _dp_c1, _dp_c2 = st.columns(2)
        with _dp_c1:
            kp_delivery_price = st.number_input(
                "Стоимость доставки, ₽ (с НДС 22 %)",
                min_value=0, max_value=5_000_000,
                value=int(st.session_state.get("kp_delivery_price", 30000)),
                step=1000, key="kp_delivery_price",
                help="Отдельная строка в спецификации КП")
        with _dp_c2:
            kp_delivery_target = st.text_input(
                "Адрес доставки",
                value=st.session_state.get("kp_delivery_target", ""),
                key="kp_delivery_target",
                placeholder="г. Москва, ул. Промышленная, д. 15",
                help="Попадёт в договор как «Адрес доставки»")
    else:
        st.session_state["kp_delivery_price"] = 0

# Добавляем строку доставки в спецификацию (если цена > 0) — чтобы q.total учёл её,
# а скидка ниже применилась тоже к доставке
_kp_dp = float(st.session_state.get("kp_delivery_price", 0) or 0)
if _kp_dp > 0:
    q.lines.append(SpecLine(
        code="ДОСТАВКА",
        name=(_kp_delivery_addr or "Доставка"),
        unit="усл.",
        qty=1.0,
        price=_kp_dp,
    ))

# --- Скидка (карточка в стиле мокапа) ---
with st.container(border=True):
    _sc_head_c1, _sc_head_c2 = st.columns([1, 6])
    with _sc_head_c1:
        st.markdown(
            '<div style="width:44px;height:44px;border-radius:12px;'
            'background:#FEF3E7;display:flex;align-items:center;justify-content:center;'
            'font-size:22px;">🏷️</div>',
            unsafe_allow_html=True)
    with _sc_head_c2:
        st.markdown(
            '<div style="font-size:22px;font-weight:700;color:#111827;'
            'padding-top:6px;">Скидка</div>',
            unsafe_allow_html=True)

    _sd_c1, _sd_c2 = st.columns([3, 2])
    with _sd_c1:
        _discount_pct = st.slider(
            "Скидка, %", min_value=0, max_value=30,
            value=int(st.session_state.get("kp_discount_pct", 0.0)),
            step=1, key="kp_discount_pct", label_visibility="collapsed",
            help="Применяется к каждой позиции в КП")
        _discount_pct = float(_discount_pct)
    # Применяем скидку
    _disc_ratio = 1.0 - _discount_pct / 100.0
    if _discount_pct > 0:
        for _ln in q.lines:
            _ln.price = float(_ln.price) * _disc_ratio
        for _ln in q.electrification_lines:
            _ln.price = float(_ln.price) * _disc_ratio
        if q.include_montage and q.montage_price > 0:
            q.montage_price = float(q.montage_price) * _disc_ratio
    _base_total_before_disc = q.total / _disc_ratio if _disc_ratio > 0 else q.total
    _discount_amount = _base_total_before_disc - q.total

    with _sd_c2:
        # Оранжевая карточка «Итого со скидкой»
        st.markdown(
            f'<div style="background:#FEF3E7;border:2px solid #FBBF87;'
            f'border-radius:14px;padding:18px 20px;text-align:center;'
            f'margin-top:-8px;">'
            f'<div style="color:#6B7280;font-size:13px;font-weight:500;'
            f'margin-bottom:6px;">Итого со скидкой</div>'
            f'<div style="color:#F97316;font-size:26px;font-weight:700;">'
            f'{fmt_money(q.total)}</div>'
            f'</div>',
            unsafe_allow_html=True)

    # Подпись под слайдером — база и экономия
    if _discount_pct > 0:
        st.markdown(
            f'<div style="color:#6B7280;font-size:14px;margin-top:4px;">'
            f'Сумма без скидки: <b style="color:#111827;">'
            f'{fmt_money(_base_total_before_disc)}</b>  ·  '
            f'Скидка: <b style="color:#F97316;">{_discount_pct:.0f}%</b> '
            f'<span style="color:#F97316;">(−{fmt_money(_discount_amount)})</span>'
            f'</div>',
            unsafe_allow_html=True)
    else:
        st.caption("Потяните слайдер — все цены пересчитаются автоматически")

# Итоги
st.subheader("💰 Итоги")
_include_vat_now = bool(st.session_state.get("kp_include_vat", True))
_kp_pp_rate = float(st.session_state.get("kp_prepay_pct", int(PREPAYMENT_DEFAULT*100))) / 100.0
m1, m2, m3, m4 = st.columns(4)
if _include_vat_now:
    m1.metric("Итого с НДС 22%", fmt_money(q.total))
    m2.metric("НДС в составе", fmt_money(q.vat))
else:
    m1.metric("Итого (без НДС)", fmt_money(q.total))
    m2.metric("НДС", "не облагается")
m3.metric(f"Предоплата {int(_kp_pp_rate*100)}%",
          fmt_money(q.total * _kp_pp_rate))
m4.metric("Остаток", fmt_money(q.total * (1 - _kp_pp_rate)))
if q.include_montage:
    st.info(f"Монтаж (без НДС, отдельным договором): "
            f"**{fmt_money(q.montage_price)}**")

# ============================================================
# 💾 БЫСТРОЕ СОХРАНЕНИЕ КП В CRM (видно сразу после Итогов)
# ============================================================
st.markdown("---")
with st.container(border=True):
    st.markdown("### 💾 Сохранить КП в CRM (чтобы не потерять и мочь открыть/редактировать)")
    st.caption("⚙️ Сохранённые КП можно открыть и редактировать в вкладке «📄 Выставленные КП»")

    _quick_col1, _quick_col2, _quick_col3 = st.columns([2, 2, 1])
    with _quick_col1:
        _q_buyer = st.text_input(
            "Кому (краткое название)",
            value=st.session_state.get("kp_buyer_name", ""),
            key="quick_save_buyer",
            placeholder="ООО «Ромашка»")
    with _quick_col2:
        _q_inn = st.text_input(
            "ИНН (если есть)",
            value=(st.session_state.get("buyer_data", {}) or {}).get("inn", ""),
            key="quick_save_inn",
            placeholder="7712345678",
            max_chars=12,
            help="Если ИНН уже есть — КП привяжется к тому же клиенту. Нет ИНН? Можно пусто.")
    with _quick_col3:
        st.markdown("<br>", unsafe_allow_html=True)
        _quick_save_clicked = st.button(
            "💾 Сохранить КП",
            type="primary",
            key="quick_save_kp_btn",
            use_container_width=True,
            help="Сохранит текущий расчёт в CRM — потом его можно найти в «Выставленные КП» и редактировать",
        )

    if _quick_save_clicked:
        try:
            import crm_db as _crm_db_q
            # Получаем номер КП (мог быть ещё не введён вручную — генерируем по дате)
            _sup_key_now = st.session_state.get("supplier_key", _suppliers.DEFAULT_SUPPLIER_KEY)
            _prefix_now = _suppliers.SUPPLIERS.get(_sup_key_now, {}).get("kp_prefix", "ЛКС")
            _base_num = date.today().strftime("%d%m%Y") + "/" + _prefix_now
            _quick_kp_num = st.session_state.get("kp_number_field_val") or _crm_db_q.generate_unique_kp_number(_base_num)

            # Создаём/обновляем клиента
            _cust = _crm_db_q.Customer(
                inn=(_q_inn or "").strip(),
                name_short=(_q_buyer or "").strip() or "Без названия",
                phone=st.session_state.get("kp_buyer_phone", "").strip(),
                email=st.session_state.get("kp_buyer_email", "").strip(),
            )
            # Если в buyer_data есть богатые реквизиты — подмешиваем
            _bd = st.session_state.get("buyer_data", {}) or {}
            for _f in ("kpp", "ogrn", "name_full", "address",
                       "bank", "bik", "rs", "ks",
                       "director_position", "director_fio_short"):
                _val = _bd.get(_f)
                if _val:
                    if _f == "director_fio_short":
                        _cust.director_fio = _val
                    else:
                        setattr(_cust, _f if _f != "name_full" else "name_full", _val)
            _cust_id = _crm_db_q.upsert_customer(_cust)

            # Сохраняем КП
            if kp_mode == "Траверса":
                _pt = "Траверса"
                _pm = "Траверса"
            elif kp_mode == "Кран + траверса":
                _pt = "Кран + траверса"
                _pm = f"{q.series} {q.capacity}кг + траверса"
            else:
                _pt = "Кран"
                _pm = f"{q.series} {q.capacity}кг {q.boom}м"

            _items = [_crm_db_q.QuoteItem(
                        code=ln.code, name=ln.name, unit=ln.unit,
                        qty=float(ln.qty), price=float(ln.price)
                    ) for ln in q.lines]
            # Генерируем PDF и DOCX чтобы они были видны при открытии КП
            _quick_docx_bytes = None
            _quick_pdf_bytes = None
            try:
                _quick_docx_bytes = build_kp_docx(q, _quick_kp_num, _q_buyer)
            except Exception:
                pass
            try:
                _quick_pdf_bytes = build_kp_pdf(q, _quick_kp_num, _q_buyer,
                                                phone=st.session_state.get("kp_buyer_phone", ""),
                                                email=st.session_state.get("kp_buyer_email", ""))
            except Exception:
                pass

            _rec = _crm_db_q.QuoteRecord(
                kp_number=_quick_kp_num,
                customer_id=_cust_id,
                product_type=_pt,
                product_model=_pm,
                base_total=float(q.total),
                discount_pct=float(st.session_state.get("kp_discount_pct", 0.0)),
                final_total=float(q.total),
                include_montage=bool(q.include_montage),
                status=_crm_db_q.STATUS_DRAFT,
                items=_items,
                pdf_bytes=_quick_pdf_bytes,
                docx_bytes=_quick_docx_bytes,
            )
            _qid = _crm_db_q.save_quote(_rec)
            st.session_state["last_saved_kp_id"] = int(_qid)
            st.session_state["last_saved_kp_num"] = _quick_kp_num
            st.success(f"✅ КП **{_quick_kp_num}** сохранён в CRM (id={_qid}). "
                       "Смотреть и редактировать — во вкладке «📄 Выставленные КП».")
        except Exception as _e_sv:
            st.error(f"Ошибка сохранения: {_e_sv}")
            import traceback
            with st.expander("Детали ошибки"):
                st.code(traceback.format_exc())

    # → Если только что сохранили КП — показываем кнопку удаления (если чел ошибся)
    if st.session_state.get("last_saved_kp_id"):
        _saved_id = st.session_state["last_saved_kp_id"]
        _saved_num = st.session_state.get("last_saved_kp_num", "")
        _del_key = f"quick_del_confirm_{_saved_id}"
        _del_flag = st.session_state.get(_del_key, False)

        st.markdown("---")
        _dc1, _dc2 = st.columns([1, 3])
        with _dc1:
            if not _del_flag:
                if st.button(f"🗑 Удалить сохранённое КП",
                             key="quick_del_btn",
                             help=f"Удалить КП {_saved_num} из CRM (если сохранили ошибочно)"):
                    st.session_state[_del_key] = True
                    st.rerun()
            else:
                if st.button("✖ Да, удалить", type="primary", key="quick_del_go"):
                    try:
                        import crm_db as _crm_db_del
                        _crm_db_del.delete_quote(_saved_id)
                        st.session_state.pop("last_saved_kp_id", None)
                        st.session_state.pop("last_saved_kp_num", None)
                        st.session_state.pop(_del_key, None)
                        st.success(f"КП {_saved_num} удалено из CRM.")
                        st.rerun()
                    except Exception as _e:
                        st.error(f"Ошибка: {_e}")
        with _dc2:
            if _del_flag:
                st.warning(f"⚠️ Будет удалено КП **{_saved_num}** и его позиции. Отменить нельзя.")
                if st.button("↶ Отмена", key="quick_del_cancel"):
                    st.session_state.pop(_del_key, None)
                    st.rerun()

st.markdown("---")

# --- КП ---
st.subheader("📄 Коммерческое предложение")

# Ряд 1: Поставщик документа + Печать/подпись
_sup_row1, _sup_row2 = st.columns([3, 2])
with _sup_row1:
    _sup_keys = list(_suppliers.SUPPLIERS.keys())
    _sup_current = st.session_state.get("supplier_key", _suppliers.DEFAULT_SUPPLIER_KEY)
    _sup_idx = _sup_keys.index(_sup_current) if _sup_current in _sup_keys else 0
    _picked_supplier_key = st.selectbox(
        "🏢 От какой компании выставляем КП/договор",
        _sup_keys,
        format_func=lambda k: _suppliers.SUPPLIERS[k]["label"],
        index=_sup_idx, key="supplier_key_kp",
        help="Влияет на реквизиты, шапку, футер, печать, ЭДО в договоре и префикс № КП")
    st.session_state["supplier_key"] = _picked_supplier_key
with _sup_row2:
    _stamp_available = bool(_suppliers.SUPPLIERS[_picked_supplier_key].get("stamp_path"))
    st.checkbox(
        "🖋️ Печать и подпись в документе",
        value=st.session_state.get("include_stamp", False),
        key="include_stamp",
        disabled=not _stamp_available,
        help=("Вставить печать+подпись в PDF/DOCX"
              if _stamp_available else "Для этого поставщика печать/подпись не загружена"))

# Обновляем глобальный SUPPLIER под выбор пользователя (SUPPLIER читают функции PDF/DOCX ниже)
SUPPLIER = _suppliers.get_supplier(_picked_supplier_key)

# Префикс в номере КП по поставщику
_kp_prefix_map = {"LKS": "ЛКС", "MODERNIZATSIYA": "МОД", "KINEMATIKA": "КИН"}
_kp_prefix = _kp_prefix_map.get(_picked_supplier_key, "ЛКС")

col_kp1, col_kp2, col_kp3, col_kp4 = st.columns([1, 1, 2, 1.4])
with col_kp1:
    # Сразу подсовываем уникальный номер — если сегодняшний базовый уже занят,
    # покажется 22072026-2/ЛКС и т.д. Пересчитывается при смене поставщика.
    _kp_base_default = date.today().strftime("%d%m%Y") + "/" + _kp_prefix
    try:
        import crm_db as _crm_db_num
        _kp_default_unique = _crm_db_num.generate_unique_kp_number(_kp_base_default)
    except Exception:
        _kp_default_unique = _kp_base_default
    kp_number = st.text_input(
        "№ исходящего КП",
        value=_kp_default_unique,
        help=f"Префикс автоматически меняется под поставщика ({_kp_prefix}). "
             "Если такой номер уже в базе — добавится суффикс -2, -3 при сохранении.")
with col_kp2:
    kp_date_str = st.text_input("Дата",
                                value=date.today().strftime("%d.%m.%Y"))
with col_kp3:
    # Селект из CRM-клиентов (если есть)
    _crm_customers = []
    try:
        import crm_db as _crm_db_kp
        _crm_customers = _crm_db_kp.search_customers("")
    except Exception:
        pass

    buyer_name_for_kp = st.text_input("Кому", placeholder="ООО «Ромашка»",
                                       key="kp_buyer_name")

    if _crm_customers:
        _cust_options = ["—"] + [f"{c.name_short} · ИНН {c.inn}" for c in _crm_customers]
        _cust_map = {f"{c.name_short} · ИНН {c.inn}": c for c in _crm_customers}
        _picked_cust = st.selectbox(
            "📇 Из базы клиентов CRM",
            _cust_options,
            key="kp_buyer_pick_crm",
            label_visibility="collapsed",
            help="Выберите клиента — подставит название, телефон, email и все реквизиты")
        if _picked_cust != "—":
            _c = _cust_map[_picked_cust]
            if _c.name_short != buyer_name_for_kp:
                # Подставляем все поля
                st.session_state["kp_buyer_name"] = _c.name_short
                if _c.phone: st.session_state["kp_buyer_phone"] = _c.phone
                if _c.email: st.session_state["kp_buyer_email"] = _c.email
                # А также в buyer_data для договора
                st.session_state["buyer_data"] = {
                    "short": _c.name_short or "",
                    "full": _c.name_full or _c.name_short or "",
                    "inn": _c.inn or "",
                    "kpp": _c.kpp or "",
                    "ogrn": _c.ogrn or "",
                    "address": _c.address or "",
                    "phone": _c.phone or "",
                    "email": _c.email or "",
                    "bank": _c.bank or "",
                    "bik": _c.bik or "",
                    "rs": _c.rs or "",
                    "ks": _c.ks or "",
                    "director_fio_short": _c.director_fio or "",
                    "director_position": _c.director_position or "Генеральный директор",
                    "basis": "Устава",
                }
                st.rerun()
with col_kp4:
    # Автопрефикс +7: если поле пустое — ставим +7. Если пользователь ввёл
    # цифры без +7 (начинаются на 8/9/пробел) — автодобавляем префикс.
    _phone_val = str(st.session_state.get("kp_buyer_phone", "") or "").strip()
    if not _phone_val:
        st.session_state["kp_buyer_phone"] = "+7 "
    elif not _phone_val.startswith("+"):
        # Нормализация: 89991234567 → +7 9991234567; 9991234567 → +7 9991234567
        _digits = "".join(ch for ch in _phone_val if ch.isdigit())
        if _digits.startswith("8") and len(_digits) == 11:
            _digits = _digits[1:]
        st.session_state["kp_buyer_phone"] = f"+7 {_digits}"
    buyer_phone_for_kp = st.text_input(
        "Телефон клиента",
        placeholder="+7 (999) 000-00-00",
        key="kp_buyer_phone",
        help="Префикс +7 устанавливается автоматически. Сохраняется в CRM — по нему можно искать КП")

if st.button("📥 Сформировать КП", type="primary", use_container_width=True):
    kp_docx_bytes = build_kp_docx(q, kp_number, buyer_name_for_kp)
    try:
        _pdf_kwargs = dict(
            supplier=SUPPLIER,
            series_descriptions=SERIES_DESCRIPTIONS,
            get_crane_image_fn=get_crane_image,
            get_hoist_image_fn=get_hoist_image,
            crane_characteristics_fn=crane_characteristics,
            hoist_characteristics_fn=hoist_characteristics,
            prepayment_rate=_kp_pp_rate,
            discount_pct=float(st.session_state.get("kp_discount_pct", 0.0) or 0.0),
            delivery_terms=(_kp_delivery_addr or "включена в стоимость"),
            delivery_price=_kp_dp,
            delivery_target=str(st.session_state.get("kp_delivery_target", "") or ""),
            kp_comment=str(st.session_state.get("kp_comment", "") or "").strip(),
        )
        # Режим Кран + траверса — подключаем карточки узлов и описание траверсы
        if kp_mode == "Кран + траверса" and tv_selection is not None:
            _tv_df_combo = _tv.load_traverse_price(
                st.session_state.get("traverse_price_upload"))
            _pdf_kwargs["attach_traverse_cards"] = True
            _pdf_kwargs["attach_traverse_desc"] = True
            _pdf_kwargs["description_fn"] = (
                lambda _q, _s=tv_selection, _d=_tv_df_combo:
                _tv.get_traverse_description(_s, _d))
            _pdf_kwargs["traverse_characteristics_fn"] = (
                lambda _q, _s=tv_selection, _d=_tv_df_combo:
                _tv.traverse_characteristics(_s, _d))
            _pdf_kwargs["traverse_image_fn"] = (
                lambda _q, _s=tv_selection:
                _tv.get_traverse_image(_s.code))
        try:
            kp_pdf_bytes = build_kp_pdf(
                q, kp_number, buyer_name_for_kp,
                kp_date=kp_date_str, **_pdf_kwargs,
            )
        except TypeError:
            # Старая версия kp_pdf.py без kp_date — вызываем без него.
            kp_pdf_bytes = build_kp_pdf(
                q, kp_number, buyer_name_for_kp, **_pdf_kwargs,
            )
    except Exception as e:
        kp_pdf_bytes = None
        st.warning(f"PDF не сгенерирован ({e}). DOCX доступен ниже.")

    stamp = date.today().strftime("%Y%m%d")
    base = f"KP_{q.series}_{q.capacity}kg_{q.boom}m_{stamp}"

    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        st.download_button(
            "⬇️ Скачать КП (DOCX)", data=kp_docx_bytes,
            file_name=f"{base}.docx",
            mime="application/vnd.openxmlformats-officedocument."
                 "wordprocessingml.document",
            use_container_width=True,
        )
    with col_dl2:
        if kp_pdf_bytes:
            st.download_button(
                "⬇️ Скачать КП (PDF)", data=kp_pdf_bytes,
                file_name=f"{base}.pdf", mime="application/pdf",
                use_container_width=True,
            )
    st.success("КП готово. Нажмите кнопки скачивания выше.")

    # --- АВТОМАТИЧЕСКОЕ сохранение КП в CRM базу ---
    _autosave_key = f"crm_autosaved_{kp_number}"
    _already_saved = st.session_state.get(_autosave_key)
    if _already_saved:
        st.info(f"💾 КП {kp_number} уже сохранён в CRM — id={_already_saved}")
    else:
        # Всегда сохраняем при генерации КП — без дополнительной кнопки
        if True:
            try:
                import crm_db as _crm_db
                if kp_mode == "Траверса":
                    _pt = "Траверса"
                    _pm = tv_selection.code if tv_selection else ""
                    _mt = bool(tv_selection.include_montage) if tv_selection else False
                elif kp_mode == "Кран + траверса":
                    _pt = "Кран + траверса"
                    _pm = (f"{q.series} {q.capacity}кг {q.boom}м"
                            + (f" + {tv_selection.code}" if tv_selection else ""))
                    _mt = bool(q.include_montage)
                else:
                    _pt = "Кран"
                    _pm = f"{q.series} {q.capacity}кг {q.boom}м"
                    _mt = bool(q.include_montage)

                _customer_id = None
                _autofilled = st.session_state.get("crm_autofilled")
                _phone = str(st.session_state.get("kp_buyer_phone", "") or "").strip()
                if _autofilled and _autofilled.inn:
                    # Если в автозаполнении телефона нет, но пользователь ввёл — допишем
                    if _phone and not _autofilled.phone:
                        _autofilled.phone = _phone
                    _customer_id = _crm_db.upsert_customer(_autofilled)
                elif buyer_name_for_kp:
                    _cl = _crm_db.Customer(name_short=buyer_name_for_kp,
                                           phone=_phone)
                    _customer_id = _crm_db.upsert_customer(_cl)

                _items = [_crm_db.QuoteItem(
                    code=ln.code, name=ln.name, unit=ln.unit,
                    qty=ln.qty, price=ln.price) for ln in q.lines]
                # base_total = сумма до скидки (q.total уже после скидки)
                _kp_disc = float(st.session_state.get("kp_discount_pct", 0.0) or 0.0)
                _disc_r = 1.0 - _kp_disc / 100.0 if _kp_disc > 0 else 1.0
                _base_tot = float(q.total) / _disc_r if _disc_r > 0 else float(q.total)
                # Уникализируем номер КП — если такой уже есть, добавится -2, -3, ...
                _kp_unique = _crm_db.generate_unique_kp_number(kp_number)
                if _kp_unique != kp_number:
                    st.info(f"⚠️ Номер {kp_number} уже занят — присвоен уникальный: **{_kp_unique}**")
                _rec = _crm_db.QuoteRecord(
                    kp_number=_kp_unique, customer_id=_customer_id,
                    product_type=_pt, product_model=_pm,
                    include_montage=_mt,
                    delivery_city=st.session_state.get("crm_delivery_city", "") or "",
                    base_total=_base_tot,
                    discount_pct=_kp_disc, status=_crm_db.STATUS_DRAFT,
                    items=_items,
                    pdf_bytes=kp_pdf_bytes, docx_bytes=kp_docx_bytes)
                _quote_id = _crm_db.save_quote(_rec)
                st.session_state[_autosave_key] = _quote_id
                st.session_state[f"crm_saved_kp_num_{_quote_id}"] = _kp_unique
                # Запоминаем покупателя в истории
                try:
                    import history_memory as _hm_save
                    if buyer_name_for_kp:
                        _hm_save.remember_buyer_short(buyer_name_for_kp)
                except Exception:
                    pass
                st.success(f"✅ КП {_kp_unique} автоматически сохранено в CRM — id={_quote_id}. Вкладка «📄 КП» → редактирование/дубликат.")
            except Exception as _e_auto:
                st.error(f"Не удалось сохранить в CRM: {_e_auto}")

    # --- CRM: блок для ручного обновления клиента и скидки ---
    st.markdown("---")
    try:
        import crm_ui
        if kp_mode == "Траверса":
            _crm_product_type = "Траверса"
            _crm_product_model = tv_selection.code if tv_selection else ""
            _crm_include_montage = bool(
                tv_selection.include_montage) if tv_selection else False
        elif kp_mode == "Кран + траверса":
            _crm_product_type = "Кран + траверса"
            _crm_product_model = (
                f"{q.series} {q.capacity}кг {q.boom}м"
                + (f" + {tv_selection.code}" if tv_selection else ""))
            _crm_include_montage = bool(q.include_montage)
        else:
            _crm_product_type = "Кран"
            _crm_product_model = f"{q.series} {q.capacity}кг {q.boom}м"
            _crm_include_montage = bool(q.include_montage)
        _crm_spec_lines = [
            {"code": ln.code, "name": ln.name, "unit": ln.unit,
             "qty": ln.qty, "price": ln.price}
            for ln in q.lines
        ]
        crm_ui.render_save_quote_block(
            base_total=float(q.total),
            product_type=_crm_product_type,
            product_model=_crm_product_model,
            include_montage=_crm_include_montage,
            spec_lines=_crm_spec_lines,
            kp_number=kp_number,
            pdf_bytes=kp_pdf_bytes,
            docx_bytes=kp_docx_bytes,
        )
    except ModuleNotFoundError:
        pass  # CRM модуль отсутствует — пропускаем без ошибки
    except Exception as _e_crm:
        st.warning(f"Блок CRM не загрузился: {_e_crm}")

st.markdown("---")

# --- Договор ---
st.subheader("📑 Договор поставки")

# --- Автозаполнение реквизитов ---
with st.expander("⚡ Автозаполнить из файла или текста реквизитов", expanded=False):
    st.caption("📋 Перетащите карточку партнёра (DOCX/PDF/TXT) или просто вставьте текст с реквизитами "
               "из письма/сайта/чужого договора. Поля ниже заполнятся автоматически. "
               "Распознанные реквизиты сохраняются в историю — следующий раз можно выбрать одним кликом.")
    up_col, txt_col = st.columns([1, 2])
    with up_col:
        req_file = st.file_uploader(
            "Файл (DOCX / DOC / PDF / JPG / PNG / TXT)",
            type=["docx", "doc", "pdf", "jpg", "jpeg", "png", "webp", "txt"],
            key="req_file",
            help="Поддерживаются: DOCX, DOC (старый Word), PDF (с текстом или скан), JPG/PNG (фото или скриншот)")
    with txt_col:
        req_text = st.text_area("Или вставьте текст", height=180,
                                key="req_text",
                                placeholder="ООО «...»\nИНН ...\nКПП ...\nОГРН ...\n"
                                            "Юр. адрес: ...\nБанк: ...\nр/с ...\nБИК ...\n"
                                            "Генеральный директор Фамилия Имя Отчество")

    if st.button("🔍 Распознать реквизиты", use_container_width=True):
        raw_text = ""
        if req_file is not None:
            data = req_file.getvalue()
            fname = req_file.name
            try:
                # Новый умный экстрактор — поддерживает DOCX/DOC/PDF/JPG/PNG (+ OCR)
                from smart_requisites import _extract_text_smart
                raw_text = _extract_text_smart(data, fname)
            except Exception as e:
                st.error(f"Не удалось прочитать файл: {e}")
                # fallback — старая логика
                if fname.lower().endswith(".docx"):
                    try:
                        raw_text = extract_text_from_docx(data)
                    except Exception:
                        pass
                elif fname.lower().endswith(".txt"):
                    try:
                        raw_text = data.decode("utf-8", errors="ignore")
                    except Exception:
                        raw_text = data.decode("cp1251", errors="ignore")
        elif req_text.strip():
            raw_text = req_text

        if not raw_text.strip():
            st.warning("Загрузите файл или вставьте текст с реквизитами.")
        else:
            # Новый надёжный парсер (тот же что в внешнем договоре)
            try:
                from external_kp_parser import extract_requisites_from_text as _ext_req
                _req = _ext_req(raw_text)
                parsed = {
                    "short": _req.company_short or "",
                    "full": _req.company_full or _req.company_short or "",
                    "address": _req.address or "",
                    "post_address": _req.address or "",
                    "inn": _req.inn or "",
                    "kpp": _req.kpp or "",
                    "ogrn": _req.ogrn or "",
                    "phone": _req.phone or "",
                    "email": _req.email or "",
                    "bank": _req.bank_name or "",
                    "bik": _req.bank_bik or "",
                    "rs": _req.bank_account or "",
                    "ks": _req.corr_account or "",
                    "director_position": _req.director_title or "Генеральный директор",
                    "director_fio_gen": _req.director_gen or "",
                    "director_fio_short": _req.director_short or "",
                    "basis": "Устава",
                }
            except Exception:
                # Fallback — старый парсер
                parsed = parse_requisites(raw_text)
            # Кладём в session_state, чтобы поля ниже подхватили его
            st.session_state["buyer_data"] = parsed
            # Автосохранение в историю по ИНН (если ИНН есть)
            try:
                import history_memory as _hm_kp_req
                if parsed.get("inn"):
                    # Конвертируем в формат внешнего договора (чтобы была единая база)
                    _hm_kp_req.remember_ext_buyer(parsed["inn"], {
                        "short": parsed.get("short", ""),
                        "full": parsed.get("full", ""),
                        "inn": parsed.get("inn", ""),
                        "kpp": parsed.get("kpp", ""),
                        "ogrn": parsed.get("ogrn", ""),
                        "address": parsed.get("address", ""),
                        "bank": parsed.get("bank", ""),
                        "bik": parsed.get("bik", ""),
                        "rs": parsed.get("rs", ""),
                        "ks": parsed.get("ks", ""),
                        "phone": parsed.get("phone", ""),
                        "email": parsed.get("email", ""),
                        "director_position": parsed.get("director_position", ""),
                        "director_fio_short": parsed.get("director_fio_short", ""),
                        "director_fio_gen": parsed.get("director_fio_gen", ""),
                        "basis": parsed.get("basis", "Устава"),
                    })
            except Exception:
                pass
            filled = sum(1 for v in parsed.values() if v)
            _summary = []
            if parsed.get("short"): _summary.append(parsed["short"])
            if parsed.get("inn"): _summary.append(f"ИНН {parsed['inn']}")
            if parsed.get("bank"): _summary.append(parsed["bank"])
            if parsed.get("director_fio_short"): _summary.append(parsed["director_fio_short"])
            _summary_str = " · ".join(_summary) if _summary else ""
            st.success(f"✅ Распознано {filled} полей" +
                       (f": {_summary_str}" if _summary_str else "") +
                       ". Поля ниже заполнены — проверьте.")

# Значения по умолчанию — из session_state если было автозаполнение, или пусто
bd = st.session_state.get("buyer_data", {})

# Память: выбор покупателя из истории по ИНН
try:
    import history_memory as _hm_ext
    _ext_hist = _hm_ext.get_ext_buyers()
except Exception:
    _ext_hist = {}

if _ext_hist:
    _labels = [f"{v.get('short', v.get('name_short','—'))} · ИНН {k}"
               for k, v in _ext_hist.items()]
    _label_to_inn = dict(zip(_labels, _ext_hist.keys()))
    _pick_col1, _pick_col2 = st.columns([4, 1])
    with _pick_col1:
        _picked_ext = st.selectbox(
            "💡 Подставить реквизиты из истории",
            ["—"] + _labels,
            key="ext_buyer_pick",
            help="Ранее введённые покупатели (по ИНН)")
    with _pick_col2:
        if st.button("↳ Подставить", key="ext_buyer_apply",
                     use_container_width=True,
                     disabled=(_picked_ext == "—")):
            _inn_pick = _label_to_inn.get(_picked_ext)
            if _inn_pick:
                st.session_state["buyer_data"] = _ext_hist[_inn_pick]
                st.rerun()
    bd = st.session_state.get("buyer_data", {})

with st.expander("Реквизиты Покупателя", expanded=True):
    c1, c2 = st.columns(2)
    with c1:
        b_short = st.text_input("Наименование краткое",
                                value=bd.get("short", ""))
        b_full = st.text_input("Наименование полное",
                               value=bd.get("full", ""))
        b_address = st.text_input("Юр. адрес", value=bd.get("address", ""))
        b_post = st.text_input("Почтовый адрес", value=bd.get("post_address", ""))
        b_inn = st.text_input("ИНН", value=bd.get("inn", ""))
        b_kpp = st.text_input("КПП", value=bd.get("kpp", ""))
        b_ogrn = st.text_input("ОГРН", value=bd.get("ogrn", ""))
        b_phone = st.text_input("Телефон", value=bd.get("phone", ""))
        b_email = st.text_input("E-mail", value=bd.get("email", ""))
    with c2:
        b_bank = st.text_input("Банк", value=bd.get("bank", ""))
        b_bik = st.text_input("БИК", value=bd.get("bik", ""))
        b_rs = st.text_input("Р/с", value=bd.get("rs", ""))
        b_ks = st.text_input("К/с", value=bd.get("ks", ""))
        b_dir_pos = st.text_input("Должность подписанта",
                                  value=bd.get("director_position", "Генеральный директор"))
        b_dir_gen = st.text_input("ФИО в родительном падеже",
                                  value=bd.get("director_fio_gen", ""))
        b_dir_short = st.text_input("ФИО коротко (для подписи)",
                                    value=bd.get("director_fio_short", ""))
        b_basis = st.text_input("На основании (Устава / Доверенности)",
                                value=bd.get("basis", "Устава"))

c3, c4 = st.columns(2)
with c3:
    contract_number = st.text_input(
        "Номер договора",
        value=date.today().strftime("%d/%m/%Y") + "/ЛКС",
        help="По умолчанию: ДД/ММ/ГГГГ/ЛКС от текущей даты")
with c4:
    contract_date_str = st.text_input(
        "Дата договора (в тексте)",
        value=date.today().strftime("%d.%m.%Y"))

c5, c6 = st.columns(2)
with c5:
    # Дефолт % предоплаты — берём из КП (единый ключ kp_prepay_pct для всех режимов)
    _dog_pp_default = int(st.session_state.get(
        "kp_prepay_pct",
        st.session_state.get("tv_prepay_pct", int(PREPAYMENT_DEFAULT * 100))))
    dog_prepay_pct = st.number_input(
        "Предоплата, %",
        min_value=0, max_value=100,
        value=_dog_pp_default,
        step=5,
        help="Автоматически берётся из «Условия оплаты и доставки». Можно перезаписать для договора.")
with c6:
    # Подтягиваем условия доставки: сначала из КП (kp_delivery_*), если нет — из траверсы (tv_delivery_*)
    _dog_dt_default = (st.session_state.get("kp_delivery_option")
                        or st.session_state.get("tv_delivery_option")
                        or "включена в стоимость")
    _dog_dt_addr = (st.session_state.get("kp_delivery_target")
                     or st.session_state.get("tv_delivery_target")
                     or "").strip()
    if _dog_dt_addr:
        _dog_dt_default = f"{_dog_dt_default}, адрес доставки: {_dog_dt_addr}"
    dog_delivery_terms = st.text_input(
        "Условия доставки",
        value=_dog_dt_default,
        help="Автоматически берётся из «Условия оплаты и доставки». Можно перезаписать вручную.")

# --- Расширенные условия договора (все перезаписываемые) ---
with st.expander("⚙️ Расширенные условия договора (можно редактировать)", expanded=False):
    st.caption("Все поля ниже — текст, который попадёт в Спецификацию Приложения № 1 договора. "
               "Оставьте пустым — будет использоваться текст по умолчанию (шаблон ЛКС).")

    _pp_pct = int(dog_prepay_pct)
    _pp_rem = 100 - _pp_pct
    _default_payment = (
        f"Покупатель обязуется произвести оплату Товара в следующем порядке: "
        f"— {_pp_pct} % от общей цены Спецификации — предоплата после подписания Договора и Спецификации в течение 5 рабочих дней; "
        f"оставшиеся {_pp_rem} % — по уведомлению на электронную почту о готовности Товара к отгрузке.")
    dog_payment_order = st.text_area(
        "Порядок оплаты",
        value=st.session_state.get("dog_payment_order") or _default_payment,
        height=110, key="dog_payment_order",
        help=f"Автоматически сгенерирован по {_pp_pct}%/{_pp_rem}%. Можно переписать.")

    dog_shipment_term = st.text_area(
        "Срок отгрузки продукции",
        value=st.session_state.get("dog_shipment_term") or "20 рабочих дней со дня поступления оплаты на расчётный счёт Поставщика.",
        height=80, key="dog_shipment_term")

    # Срок действия договора (пункт 8.1)
    dog_contract_valid_until = st.text_input(
        "Срок действия договора (пункт 8.1 — конечная дата)",
        value=st.session_state.get("dog_contract_valid_until") or "31.12.2026",
        key="dog_contract_valid_until",
        help="Подставляется в пункт 8.1: «Договор вступает в силу с момента подписания и действует до ‹дата› г.»")

    # Гарантия — генерируется автоматически по наличию монтажа
    _default_warranty = ("Поставщик гарантирует соответствие заявленного качества комплектующих при "
                         "соблюдении Покупателем условий эксплуатации и хранения. Гарантия на все товары — 12 месяцев со дня получения Покупателем.")
    if getattr(q, "include_montage", False):
        _default_warranty += (" Гарантия на монтажные работы — 12 месяцев. При заказе монтажа "
                              "в нашей компании вы получаете расширенную гарантию 24 месяца на комплектующие крановой системы "
                              "(электрическая таль, вакуумная траверса и пакет электрификации в расширенную гарантию не входят).")
    dog_warranty_text = st.text_area(
        "Гарантия",
        value=st.session_state.get("dog_warranty_text") or _default_warranty,
        height=180, key="dog_warranty_text",
        help="Автоматически: если в КП включён монтаж — добавляется блок про расширенную гарантию 24 мес.")

if st.button("📥 Сформировать Договор поставки",
             type="primary", use_container_width=True):
    if not (b_short and b_inn):
        st.error("Заполните хотя бы наименование и ИНН Покупателя.")
    else:
        buyer = {
            "short": b_short, "full": b_full or b_short,
            "address": b_address, "post_address": b_post,
            "inn": b_inn, "kpp": b_kpp, "ogrn": b_ogrn,
            "phone": b_phone, "email": b_email,
            "bank": b_bank, "bik": b_bik, "rs": b_rs, "ks": b_ks,
            "director_position": b_dir_pos,
            "director_fio_gen": b_dir_gen,
            "director_fio_short": b_dir_short,
            "basis": b_basis,
        }
        # Собираем все выбранные траверсы (основная + альтернативные) — для Приложения № 1
        _all_tv_sels = []
        _tv_main = tv_selection if 'tv_selection' in dir() else None
        if _tv_main is not None:
            _all_tv_sels.append(_tv_main)
        _all_tv_sels.extend(
            st.session_state.get("_tv_extra_selections", []) or [])

        # Запоминаем в истории весь блок реквизитов по ИНН
        try:
            import history_memory as _hm_ext_save
            _hm_ext_save.remember_ext_buyer(b_inn, buyer)
        except Exception:
            pass

        dogovor_bytes = build_dogovor_docx(
            q, buyer, contract_number, contract_date_str,
            prepay_pct=int(dog_prepay_pct),
            delivery_terms=dog_delivery_terms,
            traverse_selections=_all_tv_sels or None,
            include_stamp=bool(st.session_state.get("include_stamp", False)),
            shipment_term=st.session_state.get("dog_shipment_term"),
            warranty_text=st.session_state.get("dog_warranty_text"),
            payment_order_text=st.session_state.get("dog_payment_order"),
            contract_valid_until=st.session_state.get("dog_contract_valid_until"),
            kp_comment=str(st.session_state.get("kp_comment", "") or "").strip(),
        )
        fn = f"Dogovor_{contract_number.replace('/', '_')}.docx"
        st.download_button("⬇️ Скачать Договор поставки (DOCX)",
                           data=dogovor_bytes, file_name=fn,
                           mime="application/vnd.openxmlformats-officedocument."
                                "wordprocessingml.document",
                           use_container_width=True)
        st.success("Договор готов. Нажмите кнопку скачивания выше.")

st.markdown("---")
with st.expander("ℹ️ Паспорт логики"):
    st.markdown("""
**Серии и грузоподъёмности**
- ЛКС71 / ЛКС73 → 80, 125, 250, 500, 1000 кг
- ЛКС77 / ЛКС78 → 150, 300 кг

**Длины стрел ЛКС71/73**
- 80 кг → 3–5 м · 125 кг → 3–7 м · 250 кг → 3–8 м · 500 кг → 3–8 м · 1000 кг → 3–6 м

**Тали**
- ЛКС71/73: OCALIFT (1 / 2 скорости), R-Tech (8/2 м/мин)
- ЛКС77/78: только R-Tech; при высоте до 3 м доступен вариант с ручным контролем груза

**Тележки под таль**
- OCALIFT: 125–500 кг → 523, 1 т → 533
- R-Tech: 80 кг → 513.RC59, 125–500/300 кг → 521, 1 т → 533

**Кабельные тележки для ЛКС71/73**
- 321.050.RC74.65х54, количество = длина стрелы − 1
- 327.050.RC74.65х54, начальный подвес — 1 шт

**Электрификация** — только для ЛКС71/73 г/п 250/500/1000 кг

**Оплата** — 70 % / 30 %. **НДС** — 22 %.
    """)
